from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"F:\filesFires\base_import")
OUTPUT_PATH = ROOT / "documents" / "fire_map_related_files.docx"


def ru(value: str) -> str:
    return value.encode("ascii").decode("unicode_escape")


def text(*parts: str) -> str:
    return ru("".join(parts))


def configure_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)


def set_font(run, name: str, size: int) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def add_paragraph(document: Document, body: str, *, bold: bool = False, size: int = 12) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(body)
    run.bold = bold
    set_font(run, "Times New Roman", size)


def add_file_list(document: Document, files: list[Path]) -> None:
    for path in files:
        relative = path.relative_to(ROOT)
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(str(relative))
        set_font(run, "Consolas", 10)


def resolve(paths: list[str]) -> list[Path]:
    result: list[Path] = []
    for item in paths:
        path = ROOT / item
        if path.exists():
            result.append(path)
    return result


def main() -> None:
    forecast_risk_module = sorted((ROOT / "app" / "services" / "forecast_risk").glob("*.py"))
    mapping_module = sorted((ROOT / "core" / "mapping").rglob("*.py"))

    categories: list[tuple[str, list[Path], str]] = [
        (
            text(
                r"\u041c\u0430\u0440\u0448\u0440\u0443\u0442\u044b \u0438 ",
                r"\u0441\u0435\u0440\u0432\u0435\u0440\u043d\u0430\u044f \u0441\u0431\u043e\u0440\u043a\u0430 ",
                r"\u0441\u0442\u0440\u0430\u043d\u0438\u0446\u044b",
            ),
            resolve(
                [
                    "app/routes/pages.py",
                    "app/routes/page_common.py",
                ]
            ),
            text(
                r"\u0424\u0430\u0439\u043b\u044b, \u043a\u043e\u0442\u043e\u0440\u044b\u0435 ",
                r"\u043e\u0442\u0432\u0435\u0447\u0430\u044e\u0442 \u0437\u0430 URL ",
                r"\u0440\u0430\u0437\u0434\u0435\u043b\u0430 Fire-map, ",
                r"\u043e\u0442\u0440\u0438\u0441\u043e\u0432\u043a\u0443 ",
                r"\u0441\u0442\u0440\u0430\u043d\u0438\u0446\u044b \u043a\u0430\u0440\u0442\u044b ",
                r"\u0438 embed-\u0440\u0435\u0436\u0438\u043c.",
            ),
        ),
        (
            text(
                r"\u0421\u0435\u0440\u0432\u0438\u0441 \u043a\u0430\u0440\u0442\u044b ",
                r"\u0438 \u043e\u0431\u0449\u0438\u0435 ",
                r"\u0430\u043d\u0430\u043b\u0438\u0442\u0438\u0447\u0435\u0441\u043a\u0438\u0435 ",
                r"\u0441\u0435\u0440\u0432\u0438\u0441\u044b",
            ),
            resolve(["app/services/fire_map_service.py", "app/services/executive_brief.py"]) + forecast_risk_module,
            text(
                r"\u0421\u0435\u0440\u0432\u0438\u0441\u043d\u044b\u0439 \u0441\u043b\u043e\u0439, ",
                r"\u043a\u043e\u0442\u043e\u0440\u044b\u0439 \u0444\u043e\u0440\u043c\u0438\u0440\u0443\u0435\u0442 ",
                r"\u043a\u043e\u043d\u0442\u0435\u043a\u0441\u0442 \u0441\u0442\u0440\u0430\u043d\u0438\u0446\u044b, ",
                r"\u0441\u043e\u0431\u0438\u0440\u0430\u0435\u0442 HTML-\u043a\u0430\u0440\u0442\u0443, ",
                r"\u0441\u0442\u0440\u043e\u0438\u0442 \u043a\u0440\u0430\u0442\u043a\u0438\u0439 ",
                r"\u0431\u0440\u0438\u0444 \u0438 \u0442\u0435\u0440\u0440\u0438\u0442\u043e\u0440\u0438\u0430\u043b\u044c\u043d\u044b\u0439 ",
                r"\u043f\u0440\u0438\u043e\u0440\u0438\u0442\u0435\u0442 \u0434\u043b\u044f ",
                r"\u043e\u0442\u043e\u0431\u0440\u0430\u0436\u0435\u043d\u0438\u044f \u043d\u0430 \u043a\u0430\u0440\u0442\u0435.",
            ),
        ),
        (
            text(
                r"\u041f\u0430\u0439\u043f\u043b\u0430\u0439\u043d ",
                r"\u043f\u043e\u0441\u0442\u0440\u043e\u0435\u043d\u0438\u044f ",
                r"\u043a\u0430\u0440\u0442\u044b",
            ),
            resolve(
                [
                    "core/processing/steps/create_fire_map.py",
                    "core/processing/steps/fire_map_loader.py",
                ]
            )
            + mapping_module,
            text(
                r"\u0424\u0430\u0439\u043b\u044b, \u043e\u0442\u0432\u0435\u0447\u0430\u044e\u0449\u0438\u0435 ",
                r"\u0437\u0430 \u0437\u0430\u0433\u0440\u0443\u0437\u043a\u0443 ",
                r"\u0434\u0430\u043d\u043d\u044b\u0445, \u043f\u0440\u0435\u043e\u0431\u0440\u0430\u0437\u043e\u0432\u0430\u043d\u0438\u0435 ",
                r"\u0442\u043e\u0447\u0435\u043a \u0432 \u0432\u0435\u0431-\u043a\u0430\u0440\u0442\u0443, ",
                r"\u0430\u043d\u0430\u043b\u0438\u0442\u0438\u043a\u0443 \u0433\u043e\u0440\u044f\u0447\u0438\u0445 ",
                r"\u0437\u043e\u043d, \u043f\u0440\u0438\u043e\u0440\u0438\u0442\u0435\u0442\u043e\u0432 ",
                r"\u0438 HTML-\u0444\u0440\u0430\u0433\u043c\u0435\u043d\u0442\u043e\u0432 \u043a\u0430\u0440\u0442\u044b.",
            ),
        ),
        (
            text(
                r"\u0428\u0430\u0431\u043b\u043e\u043d\u044b \u0438 ",
                r"\u043e\u0431\u0449\u0438\u0439 layout",
            ),
            resolve(
                [
                    "app/templates/base.html",
                    "app/templates/fire_map.html",
                    "app/templates/fire_map_error.html",
                    "app/templates/includes/sidebar_nav.html",
                ]
            ),
            text(
                r"\u0428\u0430\u0431\u043b\u043e\u043d\u044b, \u0438\u0437 ",
                r"\u043a\u043e\u0442\u043e\u0440\u044b\u0445 \u0441\u043e\u0431\u0438\u0440\u0430\u0435\u0442\u0441\u044f ",
                r"\u0441\u0442\u0440\u0430\u043d\u0438\u0446\u0430 Fire-map, ",
                r"\u0435\u0435 \u043e\u0448\u0438\u0431\u043e\u0447\u043d\u044b\u0439 ",
                r"\u0441\u0446\u0435\u043d\u0430\u0440\u0438\u0439 \u0438 \u0431\u0430\u0437\u043e\u0432\u044b\u0439 ",
                r"layout.",
            ),
        ),
        (
            text(r"\u0421\u0442\u0438\u043b\u0438 \u0438 supporting JavaScript"),
            resolve(
                [
                    "app/static/css/base.css",
                    "app/static/css/layout.css",
                    "app/static/css/shared-components.css",
                    "app/static/css/dashboard.css",
                    "app/static/css/fire_map.css",
                    "app/static/js/sidebar.js",
                ]
            ),
            text(
                r"\u0421\u0442\u0438\u043b\u0438 \u0438 \u043c\u0438\u043d\u0438\u043c\u0430\u043b\u044c\u043d\u044b\u0439 ",
                r"frontend-\u0441\u043b\u043e\u0439, \u043d\u0435\u043e\u0431\u0445\u043e\u0434\u0438\u043c\u044b\u0439 ",
                r"\u0434\u043b\u044f \u043e\u0442\u0440\u0438\u0441\u043e\u0432\u043a\u0438 ",
                r"\u0441\u0442\u0440\u0430\u043d\u0438\u0446\u044b \u043a\u0430\u0440\u0442\u044b ",
                r"\u0438 \u0431\u043e\u043a\u043e\u0432\u043e\u0439 \u043d\u0430\u0432\u0438\u0433\u0430\u0446\u0438\u0438.",
            ),
        ),
        (
            text(
                r"\u0422\u0435\u0441\u0442\u044b, \u043e\u0442\u043d\u043e\u0441\u044f\u0449\u0438\u0435\u0441\u044f ",
                r"\u043a Fire-map",
            ),
            resolve(
                [
                    "tests/test_mapping_analytics.py",
                    "tests/test_mapping_template_fragments.py",
                ]
            ),
            text(
                r"\u0424\u0430\u0439\u043b\u044b \u0442\u0435\u0441\u0442\u043e\u0432, ",
                r"\u043f\u0440\u043e\u0432\u0435\u0440\u044f\u044e\u0449\u0438\u0435 ",
                r"\u0430\u043d\u0430\u043b\u0438\u0442\u0438\u043a\u0443 \u043a\u0430\u0440\u0442\u044b ",
                r"\u0438 \u0448\u0430\u0431\u043b\u043e\u043d\u043d\u044b\u0435 ",
                r"\u0444\u0440\u0430\u0433\u043c\u0435\u043d\u0442\u044b \u0432\u0435\u0431-\u043a\u0430\u0440\u0442\u044b.",
            ),
        ),
    ]

    total_files = sum(len(files) for _, files, _ in categories)

    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title_run = title.add_run(text(r"\u0424\u0430\u0439\u043b\u044b, \u043e\u0442\u043d\u043e\u0441\u044f\u0449\u0438\u0435\u0441\u044f \u043a \u0431\u043b\u043e\u043a\u0443 Fire-map"))
    title_run.bold = True
    set_font(title_run, "Times New Roman", 16)

    add_paragraph(
        document,
        text(
            r"\u041a\u043e\u0440\u043d\u0435\u0432\u043e\u0439 ",
            r"\u043a\u0430\u0442\u0430\u043b\u043e\u0433 ",
            r"\u043f\u0440\u043e\u0435\u043a\u0442\u0430: F:\\filesFires\\base_import",
        ),
    )
    add_paragraph(
        document,
        text(
            r"\u041f\u0440\u0438\u043d\u0446\u0438\u043f \u043e\u0442\u0431\u043e\u0440\u0430: ",
            r"\u0432 \u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442 \u0432\u043a\u043b\u044e\u0447\u0435\u043d\u044b ",
            r"\u043f\u0440\u044f\u043c\u044b\u0435 \u0444\u0430\u0439\u043b\u044b, ",
            r"\u043e\u043f\u0440\u0435\u0434\u0435\u043b\u044f\u044e\u0449\u0438\u0435 ",
            r"\u0441\u0442\u0440\u0430\u043d\u0438\u0446\u0443 Fire-map, ",
            r"\u0433\u0435\u043d\u0435\u0440\u0430\u0446\u0438\u044e HTML-\u043a\u0430\u0440\u0442\u044b, ",
            r"\u0433\u0435\u043e\u0430\u043d\u0430\u043b\u0438\u0442\u0438\u043a\u0443, ",
            r"\u0441\u0432\u043e\u0434\u043d\u044b\u0439 \u0431\u0440\u0438\u0444 ",
            r"\u0438 \u0431\u043b\u043e\u043a \u0442\u0435\u0440\u0440\u0438\u0442\u043e\u0440\u0438\u0430\u043b\u044c\u043d\u043e\u0433\u043e ",
            r"\u043f\u0440\u0438\u043e\u0440\u0438\u0442\u0435\u0442\u0430. ",
            r"\u041e\u0431\u0449\u0435\u0441\u0438\u0441\u0442\u0435\u043c\u043d\u044b\u0435 ",
            r"\u0444\u0430\u0439\u043b\u044b \u0431\u0435\u0437 \u043f\u0440\u044f\u043c\u043e\u0439 ",
            r"\u0441\u0432\u044f\u0437\u0438 \u0441 \u0431\u043b\u043e\u043a\u043e\u043c ",
            r"\u043d\u0435 \u0432\u043a\u043b\u044e\u0447\u0430\u043b\u0438\u0441\u044c.",
        ),
    )
    add_paragraph(document, text(r"\u0418\u0442\u043e\u0433\u043e \u043e\u0442\u043e\u0431\u0440\u0430\u043d\u043e ", str(total_files), r" \u0444\u0430\u0439\u043b\u043e\u0432."))

    for heading, files, description in categories:
        add_paragraph(document, heading, bold=True)
        add_paragraph(document, description)
        add_paragraph(document, text(r"\u041a\u043e\u043b\u0438\u0447\u0435\u0441\u0442\u0432\u043e \u0444\u0430\u0439\u043b\u043e\u0432: ", str(len(files))))
        add_file_list(document, files)

    note = document.add_paragraph()
    note_head = note.add_run(text(r"\u041f\u0440\u0438\u043c\u0435\u0447\u0430\u043d\u0438\u0435. "))
    note_head.bold = True
    set_font(note_head, "Times New Roman", 12)
    note_text = note.add_run(
        text(
            r"\u0415\u0441\u043b\u0438 \u043d\u0443\u0436\u043d\u043e, \u043d\u0430 ",
            r"\u043e\u0441\u043d\u043e\u0432\u0435 \u044d\u0442\u043e\u0433\u043e \u0441\u043f\u0438\u0441\u043a\u0430 ",
            r"\u043c\u043e\u0436\u043d\u043e \u043e\u0442\u0434\u0435\u043b\u044c\u043d\u043e ",
            r"\u0441\u043e\u0441\u0442\u0430\u0432\u0438\u0442\u044c \u0441\u0445\u0435\u043c\u0443 ",
            r"\u0432\u0437\u0430\u0438\u043c\u043e\u0434\u0435\u0439\u0441\u0442\u0432\u0438\u044f ",
            r"\u0444\u0430\u0439\u043b\u043e\u0432 \u0431\u043b\u043e\u043a\u0430 Fire-map ",
            r"\u0434\u043b\u044f \u0434\u0438\u0441\u0441\u0435\u0440\u0442\u0430\u0446\u0438\u0438.",
        )
    )
    set_font(note_text, "Times New Roman", 12)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
