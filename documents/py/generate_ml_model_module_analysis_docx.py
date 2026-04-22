from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parents[1]
DOCUMENTS_DIR = BASE_DIR / "documents"
OUTPUT_PATH = DOCUMENTS_DIR / "ml_model_module_analysis.docx"


def ru(value: str) -> str:
    return value.encode("ascii").decode("unicode_escape")


def set_font(run, *, name: str = "Times New Roman", size: int = 12, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)


def add_paragraph(document: Document, text: str, *, bold: bool = False, size: int = 12, align=None) -> None:
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_font(run)


def add_section(document: Document, title: str, paragraphs: list[str], bullets: list[str] | None = None) -> None:
    add_paragraph(document, title, bold=True, size=14)
    for item in paragraphs:
        add_paragraph(document, item)
    for item in bullets or []:
        add_bullet(document, item)


def main() -> None:
    document = Document()
    configure_styles(document)

    add_paragraph(
        document,
        ru("\\u0422\\u0435\\u0445\\u043d\\u0438\\u0447\\u0435\\u0441\\u043a\\u0438\\u0439 \\u0430\\u043d\\u0430\\u043b\\u0438\\u0437 \\u043c\\u043e\\u0434\\u0443\\u043b\\u044f `app/services/ml_model/`"),
        bold=True,
        size=16,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        document,
        ru(
            "\\u0414\\u043e\\u043a\\u0443\\u043c\\u0435\\u043d\\u0442 \\u043f\\u043e\\u0441\\u0442\\u0440\\u043e\\u0435\\u043d \\u043f\\u043e \\u0444\\u0430\\u043a\\u0442\\u0438\\u0447\\u0435\\u0441\\u043a\\u043e\\u043c\\u0443 "
            "\\u043a\\u043e\\u0434\\u0443 \\u043c\\u043e\\u0434\\u0443\\u043b\\u0435\\u0439 `core.py`, `jobs.py`, `training/*`, "
            "`backtesting/*`, `payloads.py`, `ml_model_config_types.py`, \\u0430 \\u0442\\u0430\\u043a\\u0436\\u0435 "
            "\\u0442\\u043e\\u0447\\u0435\\u043a \\u0438\\u043d\\u0442\\u0435\\u0433\\u0440\\u0430\\u0446\\u0438\\u0438 \\u0441 `forecasting` "
            "\\u0438 `forecast_risk`."
        ),
    )

    add_section(
        document,
        ru("1. \\u0410\\u043b\\u0433\\u043e\\u0440\\u0438\\u0442\\u043c\\u044b \\u043c\\u0430\\u0448\\u0438\\u043d\\u043d\\u043e\\u0433\\u043e \\u043e\\u0431\\u0443\\u0447\\u0435\\u043d\\u0438\\u044f, \\u0442\\u0438\\u043f \\u0437\\u0430\\u0434\\u0430\\u0447\\u0438 \\u0438 \\u0446\\u0435\\u043b\\u0435\\u0432\\u044b\\u0435 \\u043f\\u0435\\u0440\\u0435\\u043c\\u0435\\u043d\\u043d\\u044b\\u0435"),
        [
            ru(
                "\\u0412 \\u043c\\u043e\\u0434\\u0443\\u043b\\u0435 `app/services/ml_model/training/training_models.py:12-331` "
                "\\u0438\\u0441\\u043f\\u043e\\u043b\\u044c\\u0437\\u0443\\u044e\\u0442\\u0441\\u044f \\u0434\\u0432\\u0435 \\u043e\\u0441\\u043d\\u043e\\u0432\\u043d\\u044b\\u0435 ML-\\u0437\\u0430\\u0434\\u0430\\u0447\\u0438: "
                "\\u0440\\u0435\\u0433\\u0440\\u0435\\u0441\\u0441\\u0438\\u044f \\u043f\\u043e \\u0441\\u0447\\u0451\\u0442\\u043d\\u043e\\u0439 \\u0446\\u0435\\u043b\\u0438 \\u0438 \\u0431\\u0438\\u043d\\u0430\\u0440\\u043d\\u0430\\u044f "
                "\\u043a\\u043b\\u0430\\u0441\\u0441\\u0438\\u0444\\u0438\\u043a\\u0430\\u0446\\u0438\\u044f \\u0441\\u043e\\u0431\\u044b\\u0442\\u0438\\u044f \\u043f\\u043e\\u0436\\u0430\\u0440\\u0430."
            ),
            ru(
                "\\u0414\\u043b\\u044f count-\\u043f\\u0440\\u043e\\u0433\\u043d\\u043e\\u0437\\u0430 \\u0438\\u0441\\u043f\\u043e\\u043b\\u044c\\u0437\\u0443\\u0435\\u0442\\u0441\\u044f "
                "`sklearn.linear_model.PoissonRegressor` \\u0441 \\u043f\\u0430\\u0440\\u0430\\u043c\\u0435\\u0442\\u0440\\u0430\\u043c\\u0438 "
                "`alpha=0.40`, `max_iter=2000` (`config/constants.py:187-190`). \\u042d\\u0442\\u043e "
                "\\u0440\\u0435\\u0433\\u0440\\u0435\\u0441\\u0441\\u0438\\u043e\\u043d\\u043d\\u0430\\u044f \\u043c\\u043e\\u0434\\u0435\\u043b\\u044c, "
                "\\u0446\\u0435\\u043b\\u0435\\u0432\\u0430\\u044f \\u043f\\u0435\\u0440\\u0435\\u043c\\u0435\\u043d\\u043d\\u0430\\u044f - `count`, "
                "\\u0442\\u043e \\u0435\\u0441\\u0442\\u044c \\u0434\\u043d\\u0435\\u0432\\u043d\\u043e\\u0435 \\u0447\\u0438\\u0441\\u043b\\u043e \\u043f\\u043e\\u0436\\u0430\\u0440\\u043e\\u0432."
            ),
            ru(
                "\\u0414\\u043e\\u043f\\u043e\\u043b\\u043d\\u0438\\u0442\\u0435\\u043b\\u044c\\u043d\\u043e \\u043f\\u0440\\u0438 "
                "\\u043f\\u0435\\u0440\\u0435\\u0434\\u0438\\u0441\\u043f\\u0435\\u0440\\u0441\\u0438\\u0438 \\u0440\\u044f\\u0434\\u0430 "
                "(`variance / mean >= 1.35`, `config/constants.py:157`) \\u043c\\u043e\\u0434\\u0443\\u043b\\u044c "
                "\\u043c\\u043e\\u0436\\u0435\\u0442 \\u043e\\u0431\\u0443\\u0447\\u0430\\u0442\\u044c `statsmodels.api.GLM` "
                "\\u0441 \\u0441\\u0435\\u043c\\u0435\\u0439\\u0441\\u0442\\u0432\\u043e\\u043c `NegativeBinomial` "
                "(`training_models.py:160-197`). \\u042d\\u0442\\u0430 \\u043c\\u043e\\u0434\\u0435\\u043b\\u044c "
                "\\u043d\\u0435 \\u043e\\u0442\\u043d\\u043e\\u0441\\u0438\\u0442\\u0441\\u044f \\u043a sklearn, \\u043d\\u043e "
                "\\u0432\\u0445\\u043e\\u0434\\u0438\\u0442 \\u0432 \\u0441\\u043e\\u0441\\u0442\\u0430\\u0432 count-\\u043a\\u0430\\u043d\\u0434\\u0438\\u0434\\u0430\\u0442\\u043e\\u0432."
            ),
            ru(
                "\\u0414\\u043b\\u044f \\u0437\\u0430\\u0434\\u0430\\u0447\\u0438 \\u0432\\u0435\\u0440\\u043e\\u044f\\u0442\\u043d\\u043e\\u0441\\u0442\\u0438 "
                "P(>=1 \\u043f\\u043e\\u0436\\u0430\\u0440 \\u0432 \\u0434\\u0435\\u043d\\u044c) \\u0438\\u0441\\u043f\\u043e\\u043b\\u044c\\u0437\\u0443\\u0435\\u0442\\u0441\\u044f "
                "`sklearn.linear_model.LogisticRegression` \\u0441 `solver=\"liblinear\"`, `max_iter=500`, "
                "`class_weight=\"balanced\"`, `random_state=42` (`config/constants.py:192-197`). "
                "\\u042d\\u0442\\u043e \\u0431\\u0438\\u043d\\u0430\\u0440\\u043d\\u0430\\u044f \\u043a\\u043b\\u0430\\u0441\\u0441\\u0438\\u0444\\u0438\\u043a\\u0430\\u0446\\u0438\\u044f, "
                "\\u0430 \\u0446\\u0435\\u043b\\u0435\\u0432\\u0430\\u044f \\u043f\\u0435\\u0440\\u0435\\u043c\\u0435\\u043d\\u043d\\u0430\\u044f - `event`, "
                "\\u0433\\u0434\\u0435 `event = 1`, \\u0435\\u0441\\u043b\\u0438 `count > 0`, \\u0438 `0` \\u0438\\u043d\\u0430\\u0447\\u0435 "
                "(`training_dataset.py:20-28`)."
            ),
            ru(
                "\\u0412\\u043e\\u043a\\u0440\\u0443\\u0433 \\u043c\\u043e\\u0434\\u0435\\u043b\\u0435\\u0439 \\u0441\\u043e\\u0431\\u0440\\u0430\\u043d "
                "sklearn-\\u043f\\u0430\\u0439\\u043f\\u043b\\u0430\\u0439\\u043d: `ColumnTransformer`, `StandardScaler`, `Pipeline` "
                "(`training_models.py:20-30`, `66-87`). \\u041c\\u0430\\u0441\\u0448\\u0442\\u0430\\u0431\\u0438\\u0440\\u043e\\u0432\\u0430\\u043d\\u0438\\u0435 "
                "\\u043f\\u0440\\u0438\\u043c\\u0435\\u043d\\u044f\\u0435\\u0442\\u0441\\u044f \\u043a \\u043d\\u0435\\u043f\\u0440\\u0435\\u0440\\u044b\\u0432\\u043d\\u044b\\u043c "
                "\\u043f\\u0440\\u0438\\u0437\\u043d\\u0430\\u043a\\u0430\\u043c count-\\u043c\\u043e\\u0434\\u0435\\u043b\\u0438."
            ),
        ],
    )

    add_section(
        document,
        ru("2. \\u041f\\u043e\\u043b\\u043d\\u044b\\u0439 \\u043f\\u0430\\u0439\\u043f\\u043b\\u0430\\u0439\\u043d \\u043e\\u0431\\u0443\\u0447\\u0435\\u043d\\u0438\\u044f: \\u0437\\u0430\\u0433\\u0440\\u0443\\u0437\\u043a\\u0430 \\u0434\\u0430\\u043d\\u043d\\u044b\\u0445 -> \\u043f\\u0440\\u0438\\u0437\\u043d\\u0430\\u043a\\u0438 -> \\u043e\\u0431\\u0443\\u0447\\u0435\\u043d\\u0438\\u0435 -> \\u0432\\u0430\\u043b\\u0438\\u0434\\u0430\\u0446\\u0438\\u044f -> \\u043a\\u044d\\u0448"),
        [
            ru(
                "\\u0412\\u0435\\u0440\\u0445\\u043d\\u0438\\u0439 \\u043e\\u0440\\u043a\\u0435\\u0441\\u0442\\u0440\\u0430\\u0442\\u043e\\u0440 - "
                "`app/services/ml_model/core.py:get_ml_model_data()` (`171-305`). \\u041e\\u043d "
                "\\u0441\\u043e\\u0431\\u0438\\u0440\\u0430\\u0435\\u0442 `MlRequestState`, \\u043d\\u043e\\u0440\\u043c\\u0430\\u043b\\u0438\\u0437\\u0443\\u0435\\u0442 "
                "\\u0444\\u0438\\u043b\\u044c\\u0442\\u0440\\u044b, \\u0432\\u044b\\u0431\\u0438\\u0440\\u0430\\u0435\\u0442 \\u0438\\u0441\\u0445\\u043e\\u0434\\u043d\\u044b\\u0435 "
                "\\u0442\\u0430\\u0431\\u043b\\u0438\\u0446\\u044b \\u0438 \\u0441\\u0442\\u0440\\u043e\\u0438\\u0442 \\u043a\\u043b\\u044e\\u0447 \\u043a\\u044d\\u0448\\u0430."
            ),
            ru(
                "\\u0417\\u0430\\u0433\\u0440\\u0443\\u0437\\u043a\\u0430 \\u0432\\u0445\\u043e\\u0434\\u043d\\u044b\\u0445 \\u0434\\u0430\\u043d\\u043d\\u044b\\u0445 "
                "\\u0438\\u0434\\u0451\\u0442 \\u0447\\u0435\\u0440\\u0435\\u0437 `training/data_access.py`. "
                "`load_ml_filter_bundle()` \\u0438\\u0441\\u043f\\u043e\\u043b\\u044c\\u0437\\u0443\\u0435\\u0442 "
                "\\u0444\\u0443\\u043d\\u043a\\u0446\\u0438\\u0438 \\u0431\\u043b\\u043e\\u043a\\u0430 `forecasting`: "
                "`_collect_forecasting_metadata`, `_build_option_catalog_sql`, "
                "`_resolve_option_value` (`core.py:106-137`). `load_ml_aggregation_inputs()` "
                "\\u0441\\u0442\\u0440\\u043e\\u0438\\u0442 `daily_history` \\u0438 "
                "`filtered_records_count` \\u0447\\u0435\\u0440\\u0435\\u0437 "
                "`forecasting.data._build_daily_history_sql` \\u0438 "
                "`forecasting.data._count_forecasting_records_sql`."
            ),
            ru(
                "\\u0418\\u0441\\u0442\\u043e\\u0440\\u0438\\u0447\\u0435\\u0441\\u043a\\u0438\\u0439 \\u0440\\u044f\\u0434 "
                "\\u043f\\u0440\\u0435\\u0432\\u0440\\u0430\\u0449\\u0430\\u0435\\u0442\\u0441\\u044f \\u0432 `DataFrame` "
                "\\u0444\\u0443\\u043d\\u043a\\u0446\\u0438\\u0435\\u0439 `_build_history_frame()` "
                "(`training_dataset.py:57-66`), \\u0437\\u0430\\u0442\\u0435\\u043c "
                "`_prepare_reference_frame()` \\u0434\\u043e\\u0431\\u0430\\u0432\\u043b\\u044f\\u0435\\u0442 "
                "`weekday`, `event`, `avg_temperature` (`20-28`)."
            ),
            ru(
                "\\u0418\\u043d\\u0436\\u0435\\u043d\\u0435\\u0440\\u0438\\u044f \\u043f\\u0440\\u0438\\u0437\\u043d\\u0430\\u043a\\u043e\\u0432 "
                "\\u0440\\u0435\\u0430\\u043b\\u0438\\u0437\\u043e\\u0432\\u0430\\u043d\\u0430 \\u0432 `_feature_frame()` "
                "(`69-78`): `weekday`, `month`, `lag_1`, `lag_7`, `lag_14`, "
                "`rolling_7`, `rolling_28`, `trend_gap`, \\u0430 \\u0442\\u0430\\u043a\\u0436\\u0435 `temp_value` "
                "\\u0438\\u0437 \\u0431\\u043b\\u043e\\u043a\\u0430 \\u0442\\u0435\\u043c\\u043f\\u0435\\u0440\\u0430\\u0442\\u0443\\u0440\\u044b. "
                "`_build_design_matrix()` \\u0434\\u0435\\u043b\\u0430\\u0435\\u0442 one-hot \\u043a\\u043e\\u0434\\u0438\\u0440\\u043e\\u0432\\u0430\\u043d\\u0438\\u0435 "
                "\\u0434\\u043b\\u044f `weekday` \\u0438 `month` (`81-110`)."
            ),
            ru(
                "\\u0414\\u0430\\u043b\\u044c\\u0448\\u0435 `_train_ml_model()` "
                "(`training.py:359-459`) \\u0434\\u0435\\u043b\\u0430\\u0435\\u0442 \\u0441\\u043b\\u0435\\u0434\\u0443\\u044e\\u0449\\u0443\\u044e "
                "\\u0446\\u0435\\u043f\\u043e\\u0447\\u043a\\u0443: "
                "1) \\u043f\\u0440\\u043e\\u0432\\u0435\\u0440\\u043a\\u0430 `MIN_DAILY_HISTORY = 60`; "
                "2) \\u043f\\u043e\\u0434\\u0433\\u043e\\u0442\\u043e\\u0432\\u043a\\u0430 seed-\\u0434\\u0430\\u0442\\u0430\\u0441\\u0435\\u0442\\u0430; "
                "3) rolling-origin backtesting; 4) \\u0432\\u044b\\u0431\\u043e\\u0440 \\u0440\\u0430\\u0431\\u043e\\u0447\\u0435\\u0439 "
                "count-\\u043c\\u043e\\u0434\\u0435\\u043b\\u0438; 5) \\u0434\\u043e\\u043e\\u0431\\u0443\\u0447\\u0435\\u043d\\u0438\\u0435 "
                "\\u043d\\u0430 \\u043f\\u043e\\u043b\\u043d\\u043e\\u0439 \\u0438\\u0441\\u0442\\u043e\\u0440\\u0438\\u0438; 6) "
                "\\u043f\\u043e\\u0441\\u0442\\u0440\\u043e\\u0435\\u043d\\u0438\\u0435 future forecast rows \\u0441 "
                "\\u0438\\u043d\\u0442\\u0435\\u0440\\u0432\\u0430\\u043b\\u0430\\u043c\\u0438; 7) permutation importance; "
                "8) \\u0441\\u0435\\u0440\\u0438\\u0430\\u043b\\u0438\\u0437\\u0430\\u0446\\u0438\\u044f \\u0440\\u0435\\u0437\\u0443\\u043b\\u044c\\u0442\\u0430\\u0442\\u0430."
            ),
            ru(
                "\\u0418\\u0442\\u043e\\u0433\\u043e\\u0432\\u044b\\u0439 \\u043e\\u0442\\u0447\\u0451\\u0442 "
                "\\u0434\\u043b\\u044f UI \\u0441\\u043e\\u0431\\u0438\\u0440\\u0430\\u044e\\u0442 `_assemble_training_result()` "
                "\\u0438 `_build_ml_payload()` (`training_result.py:252-332`, `payloads.py:33-104`). "
                "\\u0412 \\u043f\\u0430\\u0439\\u043b\\u043e\\u0430\\u0434 \\u0432\\u0445\\u043e\\u0434\\u044f\\u0442 "
                "\\u0441\\u0432\\u043e\\u0434\\u043a\\u0430, \\u0442\\u0430\\u0431\\u043b\\u0438\\u0446\\u044b backtest, "
                "`forecast_rows`, `feature_importance`, charts \\u0438 quality assessment."
            ),
        ],
    )

    add_section(
        document,
        ru("3. \\u0421\\u0442\\u0440\\u0430\\u0442\\u0435\\u0433\\u0438\\u044f backtesting: \\u0440\\u0430\\u0437\\u0431\\u0438\\u0435\\u043d\\u0438\\u0435 \\u0432\\u0440\\u0435\\u043c\\u0435\\u043d\\u043d\\u043e\\u0433\\u043e \\u0440\\u044f\\u0434\\u0430 \\u0438 \\u043c\\u0435\\u0442\\u0440\\u0438\\u043a\\u0438 \\u043a\\u0430\\u0447\\u0435\\u0441\\u0442\\u0432\\u0430"),
        [
            ru(
                "\\u041e\\u0441\\u043d\\u043e\\u0432\\u043d\\u0430\\u044f \\u043b\\u043e\\u0433\\u0438\\u043a\\u0430 \\u043d\\u0430\\u0445\\u043e\\u0434\\u0438\\u0442\\u0441\\u044f "
                "\\u0432 `backtesting/training_backtesting_execution.py:62-478`. "
                "\\u0418\\u0441\\u043f\\u043e\\u043b\\u044c\\u0437\\u0443\\u0435\\u0442\\u0441\\u044f expanding-window "
                "rolling-origin scheme: \\u0434\\u043b\\u044f \\u043a\\u0430\\u0436\\u0434\\u043e\\u0439 origin-date "
                "\\u0431\\u0435\\u0440\\u0451\\u0442\\u0441\\u044f \\u0432\\u0441\\u044f \\u0438\\u0441\\u0442\\u043e\\u0440\\u0438\\u044f \\u0434\\u043e \\u044d\\u0442\\u043e\\u0439 \\u0434\\u0430\\u0442\\u044b, "
                "\\u0430 \\u0437\\u0430\\u0442\\u0435\\u043c \\u043f\\u0440\\u043e\\u0433\\u043d\\u043e\\u0437 \\u0441\\u0442\\u0440\\u043e\\u0438\\u0442\\u0441\\u044f "
                "\\u043d\\u0430 1..H \\u0434\\u043d\\u0435\\u0439 \\u0432\\u043f\\u0435\\u0440\\u0451\\u0434."
            ),
            ru(
                "\\u0414\\u043e\\u0441\\u0442\\u0443\\u043f\\u043d\\u044b\\u0435 origin-\\u0442\\u043e\\u0447\\u043a\\u0438 "
                "\\u0432\\u044b\\u0431\\u0438\\u0440\\u0430\\u044e\\u0442\\u0441\\u044f \\u0444\\u0443\\u043d\\u043a\\u0446\\u0438\\u0435\\u0439 "
                "`_select_backtest_origins()` (`278-295`). \\u0412 \\u0440\\u0430\\u0441\\u0447\\u0451\\u0442 "
                "\\u043f\\u043e\\u043f\\u0430\\u0434\\u0430\\u044e\\u0442 \\u043d\\u0435 \\u0431\\u043e\\u043b\\u0435\\u0435 "
                "`MAX_BACKTEST_POINTS = 45` \\u043e\\u043a\\u043e\\u043d, \\u043f\\u0440\\u0438 "
                "\\u043c\\u0438\\u043d\\u0438\\u043c\\u0443\\u043c\\u0435 `MIN_BACKTEST_POINTS = 8` "
                "\\u0441\\u043e\\u043f\\u043e\\u0441\\u0442\\u0430\\u0432\\u0438\\u043c\\u044b\\u0445 \\u043e\\u043a\\u043e\\u043d "
                "(`config/constants.py:126`, `130`)."
            ),
            ru(
                "\\u041c\\u0438\\u043d\\u0438\\u043c\\u0430\\u043b\\u044c\\u043d\\u044b\\u0439 \\u0440\\u0430\\u0437\\u043c\\u0435\\u0440 "
                "\\u043e\\u0431\\u0443\\u0447\\u0430\\u044e\\u0449\\u0435\\u0433\\u043e \\u043e\\u043a\\u043d\\u0430 "
                "\\u0432 `context.min_train_rows` \\u0441\\u0447\\u0438\\u0442\\u0430\\u0435\\u0442\\u0441\\u044f \\u043a\\u0430\\u043a "
                "`min(max(ROLLING_MIN_TRAIN_ROWS, MIN_FEATURE_ROWS), len(dataset) - MIN_BACKTEST_POINTS)` "
                "(`_prepare_backtest_run_context()`, `321-325`). \\u041f\\u0440\\u0438 "
                "current constants \\u044d\\u0442\\u043e \\u043d\\u0435 \\u043c\\u0435\\u043d\\u0435\\u0435 28 "
                "\\u0441\\u0442\\u0440\\u043e\\u043a \\u043f\\u0440\\u0438\\u0437\\u043d\\u0430\\u043a\\u043e\\u0432."
            ),
            ru(
                "\\u0412\\u0430\\u043b\\u0438\\u0434\\u0430\\u0446\\u0438\\u044f lead-time-aware: "
                "`_lead_time_validation_horizons()` \\u0432\\u043e\\u0437\\u0432\\u0440\\u0430\\u0449\\u0430\\u0435\\u0442 "
                "\\u0432\\u0441\\u0435 \\u0433\\u043e\\u0440\\u0438\\u0437\\u043e\\u043d\\u0442\\u044b `1..H` "
                "(`training_backtesting_support.py:95-96`). \\u0421\\u0442\\u0440\\u043e\\u0438\\u0442\\u0441\\u044f "
                "\\u043e\\u0431\\u0449\\u0438\\u0439 \\u043d\\u0430\\u0431\\u043e\\u0440 `BacktestWindowRow` \\u0434\\u043b\\u044f "
                "\\u043a\\u0430\\u0436\\u0434\\u043e\\u0433\\u043e horizon."
            ),
            ru(
                "\\u041a\\u0430\\u043d\\u0434\\u0438\\u0434\\u0430\\u0442\\u044b \\u0434\\u043b\\u044f \\u0441\\u0440\\u0430\\u0432\\u043d\\u0435\\u043d\\u0438\\u044f: "
                "`seasonal_baseline`, `heuristic_forecast`, `poisson`, `negative_binomial` "
                "(`_simulate_candidate_paths()`, `101-135`). "
                "\\u0412\\u044b\\u0431\\u043e\\u0440 \\u0440\\u0430\\u0431\\u043e\\u0447\\u0435\\u0433\\u043e count-\\u043c\\u0435\\u0442\\u043e\\u0434\\u0430 "
                "\\u0434\\u0435\\u043b\\u0430\\u0435\\u0442 `_select_count_method()` "
                "(`training_selection.py:104-134`) \\u043f\\u043e Poisson deviance, MAE, RMSE "
                "\\u0438 explainability tie-break."
            ),
            ru(
                "\\u041c\\u0435\\u0442\\u0440\\u0438\\u043a\\u0438 count-\\u043a\\u0430\\u0447\\u0435\\u0441\\u0442\\u0432\\u0430 "
                "\\u0441\\u0447\\u0438\\u0442\\u0430\\u044e\\u0442\\u0441\\u044f \\u0432 `app/services/model_quality.py:53-75`. "
                "\\u0422\\u043e\\u0447\\u043d\\u044b\\u0435 \\u0444\\u043e\\u0440\\u043c\\u0443\\u043b\\u044b:"
            ),
        ],
        bullets=[
            ru("MAE = (1 / n) * sum(|y_hat_i - y_i|)."),
            ru("RMSE = sqrt((1 / n) * sum((y_hat_i - y_i)^2))."),
            ru("sMAPE = (100 / n) * sum(2 * |y_hat_i - y_i| / (|y_i| + |y_hat_i|))."),
            ru(
                "Mean Poisson deviance = 2 * mean(y_i * ln(y_i / y_hat_i) - (y_i - y_hat_i)), "
                "\\u043f\\u0440\\u0438 `y_hat_i >= 1e-6`."
            ),
            ru(
                "\\u041a\\u043b\\u0430\\u0441\\u0441\\u0438\\u0447\\u0435\\u0441\\u043a\\u0438\\u0439 MAPE \\u0432 "
                "\\u043c\\u043e\\u0434\\u0443\\u043b\\u0435 \\u043d\\u0435 \\u0432\\u044b\\u0447\\u0438\\u0441\\u043b\\u044f\\u0435\\u0442\\u0441\\u044f; "
                "\\u0432\\u043c\\u0435\\u0441\\u0442\\u043e \\u043d\\u0435\\u0433\\u043e \\u0438\\u0441\\u043f\\u043e\\u043b\\u044c\\u0437\\u0443\\u0435\\u0442\\u0441\\u044f sMAPE, "
                "\\u043a\\u043e\\u0442\\u043e\\u0440\\u044b\\u0439 \\u0441\\u0442\\u0430\\u0431\\u0438\\u043b\\u044c\\u043d\\u0435\\u0435 \\u043f\\u0440\\u0438 \\u043d\\u0443\\u043b\\u0435\\u0432\\u044b\\u0445 \\u0434\\u043d\\u044f\\u0445."
            ),
        ],
    )

    add_section(
        document,
        ru("4. \\u041a\\u0430\\u043b\\u0438\\u0431\\u0440\\u043e\\u0432\\u043a\\u0430 \\u0438\\u043d\\u0442\\u0435\\u0440\\u0432\\u0430\\u043b\\u043e\\u0432 \\u043f\\u0440\\u0435\\u0434\\u0441\\u043a\\u0430\\u0437\\u0430\\u043d\\u0438\\u044f"),
        [
            ru(
                "\\u0418\\u043d\\u0442\\u0435\\u0440\\u0432\\u0430\\u043b\\u044c\\u043d\\u0430\\u044f \\u043b\\u043e\\u0433\\u0438\\u043a\\u0430 "
                "\\u0440\\u0430\\u0437\\u043d\\u0435\\u0441\\u0435\\u043d\\u0430 \\u043f\\u043e `calibration_compute.py`, "
                "`calibration_output.py`, `forecast_intervals.py`, "
                "`training_backtesting_horizons.py`. \\u0411\\u0430\\u0437\\u043e\\u0432\\u044b\\u0439 "
                "\\u0443\\u0440\\u043e\\u0432\\u0435\\u043d\\u044c \\u043f\\u043e\\u043a\\u0440\\u044b\\u0442\\u0438\\u044f - "
                "`PREDICTION_INTERVAL_LEVEL = 0.8` (`config/constants.py:162`)."
            ),
            ru(
                "\\u041a\\u0430\\u043b\\u0438\\u0431\\u0440\\u043e\\u0432\\u043a\\u0430 "
                "`_build_prediction_interval_calibration()` "
                "(`calibration_compute.py:127-150`) \\u0431\\u0435\\u0440\\u0451\\u0442 \\u043e\\u0441\\u0442\\u0430\\u0442\\u043a\\u0438 "
                "`|actual - prediction|`, \\u0432\\u044b\\u0447\\u0438\\u0441\\u043b\\u044f\\u0435\\u0442 "
                "\\u043a\\u0432\\u0430\\u043d\\u0442\\u0438\\u043b\\u044c \\u044d\\u0442\\u0438\\u0445 "
                "\\u0430\\u0431\\u0441\\u043e\\u043b\\u044e\\u0442\\u043d\\u044b\\u0445 \\u043e\\u0448\\u0438\\u0431\\u043e\\u043a "
                "\\u0438 \\u0444\\u043e\\u0440\\u043c\\u0438\\u0440\\u0443\\u0435\\u0442 "
                "adaptive bins \\u043f\\u043e \\u043a\\u0432\\u0430\\u043d\\u0442\\u0438\\u043b\\u044f\\u043c "
                "\\u043f\\u0440\\u0435\\u0434\\u0441\\u043a\\u0430\\u0437\\u0430\\u043d\\u043d\\u043e\\u0433\\u043e count "
                "(`_build_prediction_interval_bins()`, `54-124`)."
            ),
            ru(
                "\\u0410\\u0431\\u0441\\u043e\\u043b\\u044e\\u0442\\u043d\\u044b\\u0439 error quantile "
                "\\u0431\\u0435\\u0440\\u0451\\u0442\\u0441\\u044f \\u043a\\u0430\\u043a \\u044d\\u043b\\u0435\\u043c\\u0435\\u043d\\u0442 "
                "\\u0441 \\u0440\\u0430\\u043d\\u0433\\u043e\\u043c `ceil((n + 1) * level)` "
                "(`_prediction_interval_absolute_error_quantile()`, `45-51`). "
                "\\u0417\\u0430\\u0442\\u0435\\u043c \\u0433\\u0440\\u0430\\u043d\\u0438\\u0446\\u044b \\u0441\\u0442\\u0440\\u043e\\u044f\\u0442\\u0441\\u044f "
                "\\u043a\\u0430\\u043a `lower = y_hat - q`, `upper = y_hat + q` "
                "\\u0447\\u0435\\u0440\\u0435\\u0437 `_count_interval()`."
            ),
            ru(
                "\\u0414\\u043b\\u044f \\u0432\\u0430\\u043b\\u0438\\u0434\\u0430\\u0446\\u0438\\u0438 "
                "\\u0438\\u043d\\u0442\\u0435\\u0440\\u0432\\u0430\\u043b\\u043e\\u0432 "
                "`_evaluate_prediction_interval_backtest()` "
                "(`calibration_output.py:74-207`) \\u0441\\u0442\\u0440\\u043e\\u0438\\u0442 "
                "\\u0442\\u0440\\u0438 \\u0441\\u0445\\u0435\\u043c\\u044b:"
            ),
        ],
        bullets=[
            ru(
                "\\u0424\\u0438\\u043a\\u0441\\u0438\\u0440\\u043e\\u0432\\u0430\\u043d\\u043d\\u044b\\u0439 chrono split: "
                "\\u043f\\u0435\\u0440\\u0432\\u044b\\u0435 60% \\u043e\\u043a\\u043e\\u043d - calibration, "
                "\\u043e\\u0441\\u0442\\u0430\\u0432\\u0448\\u0438\\u0435\\u0441\\u044f 40% - evaluation "
                "(`PREDICTION_INTERVAL_CALIBRATION_FRACTION = 0.6`, `config/constants.py:163`)."
            ),
            ru(
                "\\u0411\\u043b\\u043e\\u0447\\u043d\\u0430\\u044f \\u0441\\u0445\\u0435\\u043c\\u0430 `blocked_forward_cv`: "
                "\\u043f\\u043e\\u0437\\u0434\\u043d\\u0438\\u0435 \\u043e\\u043a\\u043d\\u0430 \\u0434\\u0435\\u043b\\u044f\\u0442\\u0441\\u044f "
                "\\u043d\\u0430 2, 3 \\u0438\\u043b\\u0438 4 \\u0431\\u043b\\u043e\\u043a\\u0430 \\u0432 "
                "\\u0437\\u0430\\u0432\\u0438\\u0441\\u0438\\u043c\\u043e\\u0441\\u0442\\u0438 \\u043e\\u0442 \\u0447\\u0438\\u0441\\u043b\\u0430 "
                "evaluation windows (`calibration_compute.py:199-210`)."
            ),
            ru(
                "\\u0421\\u043a\\u043e\\u043b\\u044c\\u0437\\u044f\\u0449\\u0430\\u044f \\u0441\\u0445\\u0435\\u043c\\u0430 "
                "`rolling_split_conformal`: \\u043f\\u0435\\u0440\\u0435\\u043a\\u0430\\u043b\\u0438\\u0431\\u0440\\u043e\\u0432\\u043a\\u0430 "
                "\\u043f\\u0435\\u0440\\u0435\\u0434 \\u043a\\u0430\\u0436\\u0434\\u044b\\u043c \\u0441\\u043b\\u0435\\u0434\\u0443\\u044e\\u0449\\u0438\\u043c "
                "evaluation point (`407-446`)."
            ),
            ru(
                "\\u041f\\u043e\\u0434\\u0434\\u0435\\u0440\\u0436\\u0438\\u0432\\u0430\\u0435\\u0442\\u0441\\u044f "
                "\\u0438 \\u044f\\u0440\\u043b\\u044b\\u043a `Jackknife+`, \\u043d\\u043e "
                "\\u0432 `calibration_output.py:66-71` \\u043f\\u0440\\u044f\\u043c\\u043e "
                "\\u0437\\u0430\\u0444\\u0438\\u043a\\u0441\\u0438\\u0440\\u043e\\u0432\\u0430\\u043d\\u043e, \\u0447\\u0442\\u043e "
                "\\u044d\\u0442\\u0430 \\u0441\\u0445\\u0435\\u043c\\u0430 \\u043d\\u0435 \\u043f\\u0440\\u0438\\u043c\\u0435\\u043d\\u0435\\u043d\\u0430, "
                "\\u043f\\u043e\\u0442\\u043e\\u043c\\u0443 \\u0447\\u0442\\u043e \\u0434\\u043b\\u044f honest time-series "
                "\\u0432\\u0430\\u0440\\u0438\\u0430\\u043d\\u0442\\u0430 \\u043d\\u0443\\u0436\\u043d\\u044b leave-one-block-out refits."
            ),
        ],
    )

    add_section(
        document,
        ru("5. \\u041e\\u0446\\u0435\\u043d\\u043a\\u0430 \\u0432\\u0430\\u0436\\u043d\\u043e\\u0441\\u0442\\u0438 \\u043f\\u0440\\u0438\\u0437\\u043d\\u0430\\u043a\\u043e\\u0432 \\u0438 \\u0438\\u0441\\u043f\\u043e\\u043b\\u044c\\u0437\\u043e\\u0432\\u0430\\u043d\\u0438\\u0435 \\u0432 UI"),
        [
            ru(
                "\\u041c\\u0435\\u0442\\u043e\\u0434 \\u0432\\u0430\\u0436\\u043d\\u043e\\u0441\\u0442\\u0438 - "
                "`sklearn.inspection.permutation_importance` "
                "(`training_importance.py:18-77`). \\u041e\\u0446\\u0435\\u043d\\u043a\\u0430 \\u0432\\u0435\\u0434\\u0451\\u0442\\u0441\\u044f "
                "\\u043d\\u0430 holdout-\\u0441\\u0440\\u0435\\u0437\\u0435 \\u0432 \\u043a\\u043e\\u043d\\u0446\\u0435 \\u0440\\u044f\\u0434\\u0430: "
                "`IMPORTANCE_TRAIN_SPLIT_RATIO = 0.7`, "
                "`IMPORTANCE_MAX_SAMPLE_SIZE = 120`, `PERMUTATION_REPEATS = 8`."
            ),
            ru(
                "\\u0426\\u0435\\u043b\\u0435\\u0432\\u0430\\u044f \\u0444\\u0443\\u043d\\u043a\\u0446\\u0438\\u044f "
                "\\u0434\\u043b\\u044f permutation importance - `neg_mean_absolute_error` "
                "(`training_importance.py:55-73`). \\u0421\\u043b\\u0435\\u0434\\u043e\\u0432\\u0430\\u0442\\u0435\\u043b\\u044c\\u043d\\u043e, "
                "\\u0432\\u0430\\u0436\\u043d\\u043e\\u0441\\u0442\\u044c \\u043f\\u043e\\u043a\\u0430\\u0437\\u044b\\u0432\\u0430\\u0435\\u0442, "
                "\\u043d\\u0430\\u0441\\u043a\\u043e\\u043b\\u044c\\u043a\\u043e \\u0440\\u0430\\u0441\\u0442\\u0451\\u0442 MAE "
                "\\u043f\\u0440\\u0438 \\u043f\\u0435\\u0440\\u0435\\u043c\\u0435\\u0448\\u0438\\u0432\\u0430\\u043d\\u0438\\u0438 "
                "\\u043f\\u0440\\u0438\\u0437\\u043d\\u0430\\u043a\\u0430."
            ),
            ru(
                "\\u0415\\u0441\\u043b\\u0438 permutation importance \\u043d\\u0435\\u0434\\u043e\\u0441\\u0442\\u0443\\u043f\\u0435\\u043d "
                "\\u0438\\u043b\\u0438 \\u0443\\u043f\\u0430\\u043b, \\u0432\\u043a\\u043b\\u044e\\u0447\\u0430\\u0435\\u0442\\u0441\\u044f "
                "fallback: `feature_importances_`, `coef_` \\u0438\\u043b\\u0438 `params` "
                "(`_fallback_feature_importance()`, `103-118`). "
                "\\u0414\\u043b\\u044f \\u043b\\u0438\\u043d\\u0435\\u0439\\u043d\\u044b\\u0445 \\u043c\\u043e\\u0434\\u0435\\u043b\\u0435\\u0439 "
                "\\u0431\\u0435\\u0440\\u0451\\u0442\\u0441\\u044f `abs(coef_)`."
            ),
            ru(
                "\\u0412\\u0430\\u0436\\u043d\\u043e\\u0441\\u0442\\u0438 \\u0430\\u0433\\u0440\\u0435\\u0433\\u0438\\u0440\\u0443\\u044e\\u0442\\u0441\\u044f "
                "\\u043f\\u043e \\u0441\\u043c\\u044b\\u0441\\u043b\\u043e\\u0432\\u044b\\u043c \\u043f\\u0440\\u0438\\u0437\\u043d\\u0430\\u043a\\u0430\\u043c: "
                "\\u0432\\u0441\\u0435 dummy-\\u043a\\u043e\\u043b\\u043e\\u043d\\u043a\\u0438 `weekday_*` "
                "\\u0441\\u0432\\u043e\\u0434\\u044f\\u0442\\u0441\\u044f \\u0432 `weekday`, "
                "`month_*` - \\u0432 `month` (`121-126`)."
            ),
            ru(
                "\\u0412 UI \\u0440\\u0435\\u0437\\u0443\\u043b\\u044c\\u0442\\u0430\\u0442 "
                "\\u0438\\u0434\\u0451\\u0442 \\u0432 `payloads.py:69-78`: "
                "`charts.importance`, `feature_importance`, `top_feature_label`. "
                "`presentation_training.py:103-123` \\u0441\\u0442\\u0440\\u043e\\u0438\\u0442 "
                "bar-chart \\u043f\\u043e \\u0442\\u043e\\u043f-8 \\u0434\\u0440\\u0430\\u0439\\u0432\\u0435\\u0440\\u0430\\u043c, "
                "\\u0430 `feature_importance_note` \\u043e\\u0431\\u044a\\u044f\\u0441\\u043d\\u044f\\u0435\\u0442, "
                "\\u0435\\u0441\\u043b\\u0438 \\u0434\\u0440\\u0430\\u0439\\u0432\\u0435\\u0440\\u044b \\u043f\\u043e\\u0441\\u0442\\u0440\\u043e\\u0435\\u043d\\u044b "
                "\\u043d\\u0435 \\u043f\\u043e \\u0442\\u043e\\u0439 \\u0436\\u0435 \\u043c\\u043e\\u0434\\u0435\\u043b\\u0438, "
                "\\u0447\\u0442\\u043e \\u0432\\u044b\\u0431\\u0440\\u0430\\u043d\\u0430 \\u0434\\u043b\\u044f production forecast "
                "(`training.py:281-328`)."
            ),
        ],
    )

    add_section(
        document,
        ru("6. \\u0421\\u0438\\u0441\\u0442\\u0435\\u043c\\u0430 \\u043a\\u044d\\u0448\\u0435\\u0439 ML-\\u0430\\u0440\\u0442\\u0435\\u0444\\u0430\\u043a\\u0442\\u043e\\u0432"),
        [
            ru(
                "\\u0411\\u0430\\u0437\\u043e\\u0432\\u044b\\u0439 \\u043a\\u043e\\u043d\\u0442\\u0435\\u0439\\u043d\\u0435\\u0440 "
                "`MLModelCaches` \\u043e\\u043f\\u0438\\u0441\\u0430\\u043d \\u0432 `core.py:38-48`. "
                "\\u041e\\u043d \\u0441\\u043e\\u0434\\u0435\\u0440\\u0436\\u0438\\u0442 \\u0434\\u0432\\u0430 "
                "\\u0443\\u0440\\u043e\\u0432\\u043d\\u044f \\u043a\\u044d\\u0448\\u0430: `ml_cache` "
                "\\u0438 `artifact_cache`."
            ),
            ru(
                "`ml_cache` \\u0441\\u0442\\u0440\\u043e\\u0438\\u0442\\u0441\\u044f \\u0447\\u0435\\u0440\\u0435\\u0437 "
                "`build_immutable_payload_lru_cache(max_size=128)` "
                "(`core.py:44-48`, `config/constants.py:184-185`). "
                "\\u042d\\u0442\\u043e \\u0432\\u0435\\u0440\\u0445\\u043d\\u0438\\u0439 LRU-\\u043a\\u044d\\u0448 "
                "\\u0433\\u043e\\u0442\\u043e\\u0432\\u044b\\u0445 UI payload."
            ),
            ru(
                "\\u0415\\u0433\\u043e \\u043a\\u043b\\u044e\\u0447 \\u0444\\u043e\\u0440\\u043c\\u0438\\u0440\\u0443\\u0435\\u0442 "
                "`build_ml_cache_key()` (`shared/request_state.py:116-134`): "
                "(`ML_CACHE_SCHEMA_VERSION`, `selected_table`, `cause`, `object_category`, "
                "`temperature`, `days_ahead`, `history_window`). "
                "\\u0412 current project `ML_CACHE_SCHEMA_VERSION = 2` "
                "(`config/constants.py:184`)."
            ),
            ru(
                "`artifact_cache` - \\u044d\\u0442\\u043e `OrderedDict` "
                "\\u0432 \\u043f\\u0430\\u043c\\u044f\\u0442\\u0438 process, \\u043a\\u043e\\u0442\\u043e\\u0440\\u044b\\u0439 "
                "\\u043a\\u044d\\u0448\\u0438\\u0440\\u0443\\u0435\\u0442 \\u0442\\u044f\\u0436\\u0451\\u043b\\u044b\\u0435 "
                "\\u0430\\u0440\\u0442\\u0435\\u0444\\u0430\\u043a\\u0442\\u044b \\u043e\\u0431\\u0443\\u0447\\u0435\\u043d\\u0438\\u044f: "
                "\\u0444\\u0438\\u043d\\u0430\\u043b\\u044c\\u043d\\u0443\\u044e frame, dataset, backtest, fitted models, "
                "feature importance (`training.py:35-47`, `331-356`). "
                "\\u041b\\u0438\\u043c\\u0438\\u0442 - 32 \\u044d\\u043b\\u0435\\u043c\\u0435\\u043d\\u0442\\u0430 "
                "(`_TRAINING_ARTIFACT_CACHE_LIMIT = 32`, `83`)."
            ),
            ru(
                "\\u041a\\u043b\\u044e\\u0447 artifact-\\u043a\\u044d\\u0448\\u0430: "
                "`(forecast_days, daily_history_signature)` "
                "(`training.py:121-126`). "
                "`daily_history_signature` \\u0441\\u043e\\u0441\\u0442\\u043e\\u0438\\u0442 "
                "\\u0438\\u0437 \\u043f\\u043e\\u0441\\u043b\\u0435\\u0434\\u043d\\u0438\\u0445 "
                "\\u0434\\u043e `MAX_HISTORY_POINTS = 900` \\u0437\\u0430\\u043f\\u0438\\u0441\\u0435\\u0439 "
                "(`date.isoformat`, `count`, `avg_temperature`) (`110-118`)."
            ),
            ru(
                "\\u041d\\u0438\\u0436\\u0435 \\u0441\\u0442\\u043e\\u0438\\u0442 TTL-\\u0441\\u043b\\u043e\\u0439 \\u0432 "
                "`training/data_access.py`: `_ML_FILTER_BUNDLE_CACHE` \\u0438 "
                "`_ML_AGGREGATION_INPUT_CACHE`, \\u043e\\u0431\\u0430 \\u0441 TTL 120 \\u0441\\u0435\\u043a\\u0443\\u043d\\u0434 "
                "(`11-12`). \\u041a\\u043b\\u044e\\u0447\\u0438 \\u0432\\u043a\\u043b\\u044e\\u0447\\u0430\\u044e\\u0442 "
                "\\u043d\\u043e\\u0440\\u043c\\u0430\\u043b\\u0438\\u0437\\u043e\\u0432\\u0430\\u043d\\u043d\\u044b\\u0435 tables, "
                "history window, filters \\u0438 `callable_cache_scope(...)`, "
                "\\u0447\\u0442\\u043e\\u0431\\u044b \\u043f\\u0435\\u0440\\u0435\\u0441\\u0447\\u0435\\u0442 "
                "\\u043e\\u0442\\u0434\\u0435\\u043b\\u044f\\u043b\\u0441\\u044f \\u043f\\u0440\\u0438 \\u0441\\u043c\\u0435\\u043d\\u0435 "
                "\\u0444\\u0443\\u043d\\u043a\\u0446\\u0438\\u0439-\\u0438\\u0441\\u0442\\u043e\\u0447\\u043d\\u0438\\u043a\\u043e\\u0432."
            ),
            ru(
                "\\u0418\\u043d\\u0432\\u0430\\u043b\\u0438\\u0434\\u0430\\u0446\\u0438\\u044f "
                "\\u0432\\u044b\\u0437\\u044b\\u0432\\u0430\\u0435\\u0442\\u0441\\u044f "
                "`clear_ml_model_cache()` (`core.py:353-358`). \\u041e\\u043d\\u0430 "
                "\\u043e\\u0447\\u0438\\u0449\\u0430\\u0435\\u0442 `ml_cache`, TTL-\\u043a\\u044d\\u0448\\u0438 "
                "input layer, artifact_cache \\u0438 \\u0434\\u043e\\u043f\\u043e\\u043b\\u043d\\u0438\\u0442\\u0435\\u043b\\u044c\\u043d\\u043e "
                "\\u0441\\u0431\\u0440\\u0430\\u0441\\u044b\\u0432\\u0430\\u0435\\u0442 "
                "`forecasting` SQL cache \\u0447\\u0435\\u0440\\u0435\\u0437 "
                "`clear_forecasting_sql_cache()`."
            ),
        ],
        bullets=[
            ru("\\u0412\\u044b\\u0442\\u0435\\u0441\\u043d\\u0435\\u043d\\u0438\\u0435 `ml_cache`: LRU, \\u043b\\u0438\\u043c\\u0438\\u0442 128 payload."),
            ru("\\u0412\\u044b\\u0442\\u0435\\u0441\\u043d\\u0435\\u043d\\u0438\\u0435 `artifact_cache`: manual LRU \\u0447\\u0435\\u0440\\u0435\\u0437 `move_to_end()` / `popitem(last=False)`, \\u043b\\u0438\\u043c\\u0438\\u0442 32."),
            ru("\\u0412\\u044b\\u0442\\u0435\\u0441\\u043d\\u0435\\u043d\\u0438\\u0435 TTL-\\u043a\\u044d\\u0448\\u0435\\u0439: \\u043f\\u043e \\u0438\\u0441\\u0442\\u0435\\u0447\\u0435\\u043d\\u0438\\u044e 120 \\u0441\\u0435\\u043a\\u0443\\u043d\\u0434 \\u0438\\u043b\\u0438 \\u043f\\u0440\\u0438 manual clear."),
        ],
    )

    add_section(
        document,
        ru("7. \\u0410\\u0441\\u0438\\u043d\\u0445\\u0440\\u043e\\u043d\\u043d\\u044b\\u0439 \\u0444\\u043e\\u043d\\u043e\\u0432\\u044b\\u0439 \\u0437\\u0430\\u043f\\u0443\\u0441\\u043a, \\u0441\\u0442\\u0430\\u0442\\u0443\\u0441\\u044b \\u0438 \\u043b\\u043e\\u0433\\u0438"),
        [
            ru(
                "\\u0424\\u043e\\u043d\\u043e\\u0432\\u043e\\u0435 \\u0432\\u044b\\u043f\\u043e\\u043b\\u043d\\u0435\\u043d\\u0438\\u0435 "
                "\\u0441\\u043e\\u0441\\u0440\\u0435\\u0434\\u043e\\u0442\\u043e\\u0447\\u0435\\u043d\\u043e \\u0432 "
                "`jobs.py:3-379`. \\u0418\\u0441\\u043f\\u043e\\u043b\\u044c\\u0437\\u0443\\u0435\\u0442\\u0441\\u044f "
                "`ThreadPoolExecutor(max_workers=2, thread_name_prefix=\"ml-model\")` "
                "(`23`)."
            ),
            ru(
                "\\u0422\\u043e\\u0447\\u043a\\u0430 \\u0432\\u0445\\u043e\\u0434\\u0430 - "
                "`start_ml_model_job()` (`34-91`), \\u043a\\u043e\\u0442\\u043e\\u0440\\u0430\\u044f "
                "\\u0441\\u043e\\u0431\\u0438\\u0440\\u0430\\u0435\\u0442 request state, \\u0441\\u0435\\u0440\\u0438\\u0430\\u043b\\u0438\\u0437\\u0443\\u0435\\u0442 "
                "cache key, \\u043f\\u0440\\u043e\\u0432\\u0435\\u0440\\u044f\\u0435\\u0442 cache-aware reuse "
                "\\u0438 \\u0437\\u0430\\u043f\\u0443\\u0441\\u043a\\u0430\\u0435\\u0442 background job."
            ),
            ru(
                "\\u0414\\u0435\\u0434\\u0443\\u043f\\u043b\\u0438\\u043a\\u0430\\u0446\\u0438\\u044f \\u0437\\u0430\\u0434\\u0430\\u0447 "
                "\\u0434\\u0435\\u043b\\u0430\\u0435\\u0442\\u0441\\u044f \\u0447\\u0435\\u0440\\u0435\\u0437 "
                "`JobReuseCoordinator`, `_ML_JOB_IDS_BY_CACHE_KEY` \\u0438 `_ML_JOB_LOCK = RLock()` "
                "(`24-25`, `60-65`, `119-124`). "
                "\\u0422\\u0430\\u043a \\u043d\\u0435 \\u0437\\u0430\\u043f\\u0443\\u0441\\u043a\\u0430\\u044e\\u0442\\u0441\\u044f "
                "\\u0434\\u0432\\u0435 \\u0442\\u044f\\u0436\\u0451\\u043b\\u044b\\u0435 \\u0444\\u043e\\u043d\\u043e\\u0432\\u044b\\u0435 "
                "\\u0437\\u0430\\u0434\\u0430\\u0447\\u0438 \\u0434\\u043b\\u044f \\u043e\\u0434\\u043d\\u043e\\u0433\\u043e \\u0438 "
                "\\u0442\\u043e\\u0433\\u043e \\u0436\\u0435 cache key."
            ),
            ru(
                "\\u0420\\u0430\\u0441\\u0447\\u0451\\u0442 \\u0438\\u0434\\u0451\\u0442 "
                "\\u0432 `_run_ml_model_job()` (`98-151`). \\u041f\\u0440\\u043e\\u0433\\u0440\\u0435\\u0441\\u0441 "
                "\\u043f\\u0435\\u0440\\u0435\\u0434\\u0430\\u0451\\u0442\\u0441\\u044f \\u0447\\u0435\\u0440\\u0435\\u0437 "
                "`StageTrackingJobProgressReporter` \\u0438 "
                "`LinkedJobProgressReporter`: "
                "\\u043e\\u0441\\u043d\\u043e\\u0432\\u043d\\u0430\\u044f job \\u043e\\u0442\\u0441\\u043b\\u0435\\u0436\\u0438\\u0432\\u0430\\u0435\\u0442 "
                "phase `ml_model.*`, \\u0430 \\u0441\\u0432\\u044f\\u0437\\u0430\\u043d\\u043d\\u0430\\u044f "
                "backtest-job - `ml_backtest.*`."
            ),
            ru(
                "\\u0418\\u0441\\u043f\\u043e\\u043b\\u044c\\u0437\\u0443\\u044e\\u0442\\u0441\\u044f "
                "\\u0441\\u0442\\u0430\\u0442\\u0443\\u0441\\u044b `pending`, `running`, `completed`, `failed` "
                "(`_ml_backtest_status_resolver()`, `369-379`). "
                "\\u041b\\u043e\\u0433\\u0438 \\u043f\\u0438\\u0448\\u0443\\u0442\\u0441\\u044f \\u0432 `job_store.add_log()` "
                "\\u0438 \\u0437\\u0430\\u0432\\u0435\\u0440\\u0448\\u0430\\u044e\\u0442\\u0441\\u044f \\u0447\\u0435\\u0440\\u0435\\u0437 "
                "`complete_job()` \\u0438\\u043b\\u0438 `fail_job()` (`179-194`, `249-313`)."
            ),
            ru(
                "\\u0415\\u0441\\u043b\\u0438 \\u0434\\u0430\\u043d\\u043d\\u044b\\u0435 \\u0443\\u0436\\u0435 "
                "\\u0435\\u0441\\u0442\\u044c \\u0432 `ml_cache`, \\u0444\\u043e\\u043d\\u043e\\u0432\\u0430\\u044f "
                "\\u0437\\u0430\\u0434\\u0430\\u0447\\u0430 \\u043d\\u0435 \\u0433\\u043e\\u043d\\u044f\\u0435\\u0442 "
                "\\u043f\\u0435\\u0440\\u0435\\u0441\\u0447\\u0451\\u0442: `_handle_cached_ml_payload()` "
                "(`201-229`) \\u0441\\u0440\\u0430\\u0437\\u0443 \\u0437\\u0430\\u043a\\u0440\\u044b\\u0432\\u0430\\u0435\\u0442 "
                "\\u043e\\u0431\\u0435 job \\u043a\\u0430\\u043a completed."
            ),
        ],
    )

    add_section(
        document,
        ru("8. \\u0412\\u0437\\u0430\\u0438\\u043c\\u043e\\u0434\\u0435\\u0439\\u0441\\u0442\\u0432\\u0438\\u0435 \\u0441 `forecasting` \\u0438 `forecast_risk`"),
        [
            ru(
                "\\u0421 `forecasting` \\u0435\\u0441\\u0442\\u044c \\u043f\\u0440\\u044f\\u043c\\u0430\\u044f \\u0438 "
                "\\u043f\\u043b\\u043e\\u0442\\u043d\\u0430\\u044f \\u0438\\u043d\\u0442\\u0435\\u0433\\u0440\\u0430\\u0446\\u0438\\u044f. "
                "ML-\\u043c\\u043e\\u0434\\u0443\\u043b\\u044c \\u0438\\u0441\\u043f\\u043e\\u043b\\u044c\\u0437\\u0443\\u0435:"
            ),
        ],
        bullets=[
            ru("`forecasting.data._build_daily_history_sql`, `_count_forecasting_records_sql`, `_collect_forecasting_metadata`, `_build_option_catalog_sql` (`core.py:11-22`, `106-137`)."),
            ru("`forecasting.utils` \\u0434\\u043b\\u044f parsing/formatting \\u0444\\u0438\\u043b\\u044c\\u0442\\u0440\\u043e\\u0432 (`core.py:23-30`)."),
            ru("`forecasting.data._build_forecast_rows` \\u0432 \\u0440\\u043e\\u043b\\u0438 heuristic scenario-forecast engine (`forecast_intervals.py:11`, `102-126`)."),
            ru("`forecasting.data._temperature_quality_from_daily_history` \\u0434\\u043b\\u044f quality block (`core.py:20-21`, `274`)."),
            ru("`forecasting.presentation._build_feature_cards_with_quality` \\u0434\\u043b\\u044f unified cards \\u0432 UI (`payloads.py:6`, `69`)."),
        ],
    )

    add_paragraph(
        document,
        ru(
            "\\u0421 `forecast_risk` \\u043f\\u0440\\u044f\\u043c\\u043e\\u0439 \\u0441\\u0432\\u044f\\u0437\\u043a\\u0438 "
            "\\u0432\\u043d\\u0443\\u0442\\u0440\\u0438 `app/services/ml_model/` \\u043d\\u0435 "
            "\\u043e\\u0431\\u043d\\u0430\\u0440\\u0443\\u0436\\u0435\\u043d\\u043e: \\u0432 "
            "\\u043a\\u043e\\u0434\\u0435 `ml_model` \\u043d\\u0435\\u0442 \\u0438\\u043c\\u043f\\u043e\\u0440\\u0442\\u043e\\u0432 "
            "\\u0438\\u043b\\u0438 \\u0432\\u044b\\u0437\\u043e\\u0432\\u043e\\u0432 `forecast_risk.core`. "
            "\\u041d\\u0430 \\u0443\\u0440\\u043e\\u0432\\u043d\\u0435 \\u043f\\u0440\\u0438\\u043b\\u043e\\u0436\\u0435\\u043d\\u0438\\u044f "
            "\\u044d\\u0442\\u0438 \\u0431\\u043b\\u043e\\u043a\\u0438 \\u0441\\u043e\\u0441\\u0443\\u0449\\u0435\\u0441\\u0442\\u0432\\u0443\\u044e\\u0442 "
            "\\u043f\\u0430\\u0440\\u0430\\u043b\\u043b\\u0435\\u043b\\u044c\\u043d\\u043e: `forecasting` \\u0441\\u0430\\u043c "
            "\\u043e\\u0431\\u0449\\u0430\\u0435\\u0442\\u0441\\u044f \\u0441 `forecast_risk`, \\u0430 "
            "`ml_model` \\u043f\\u0435\\u0440\\u0435\\u0438\\u0441\\u043f\\u043e\\u043b\\u044c\\u0437\\u0443\\u0435\\u0442 "
            "\\u0438\\u043c\\u0435\\u043d\\u043d\\u043e \\u0431\\u0430\\u0437\\u043e\\u0432\\u044b\\u0439 `forecasting`-\\u0441\\u043b\\u043e\\u0439."
        )
    )

    add_section(
        document,
        ru("9. \\u0424\\u043e\\u0440\\u043c\\u0443\\u043b\\u044b \\u0438 \\u0432\\u0430\\u0436\\u043d\\u044b\\u0435 \\u0442\\u0435\\u0445\\u043d\\u0438\\u0447\\u0435\\u0441\\u043a\\u0438\\u0435 \\u043a\\u043e\\u043d\\u0441\\u0442\\u0430\\u043d\\u0442\\u044b"),
        [
            ru(
                "\\u041a\\u0440\\u043e\\u043c\\u0435 count-\\u043c\\u0435\\u0442\\u0440\\u0438\\u043a, \\u0434\\u043b\\u044f "
                "event-\\u0431\\u043b\\u043e\\u043a\\u0430 \\u0441\\u0447\\u0438\\u0442\\u0430\\u044e\\u0442\\u0441\\u044f "
                "Brier score, ROC AUC, F1, log loss "
                "(`training_backtesting_events.py:269-527`, `model_quality.py:78-128`)."
            ),
        ],
        bullets=[
            ru("Brier = (1 / n) * sum((p_i - y_i)^2)."),
            ru("LogLoss = -(1 / n) * sum(y_i * log(p_i) + (1 - y_i) * log(1 - p_i))."),
            ru("F1 = 2 * Precision * Recall / (Precision + Recall), threshold = 0.5."),
            ru("Coverage = (1 / n) * sum(I(lower_i <= y_i <= upper_i))."),
            ru("Stability score = |coverage - target_level| + std(segment_coverages)."),
            ru("\\u041f\\u043e\\u0440\\u043e\\u0433 \\u0438\\u043d\\u0444\\u043e\\u0440\\u043c\\u0430\\u0442\\u0438\\u0432\\u043d\\u043e\\u0441\\u0442\\u0438 event-\\u0431\\u043b\\u043e\\u043a\\u0430: event rate \\u043d\\u0435 \\u0434\\u043e\\u043b\\u0436\\u0435\\u043d \\u0431\\u044b\\u0442\\u044c \\u043d\\u0430\\u0441\\u044b\\u0449\\u0435\\u043d\\u043d\\u044b\\u043c, margin = 0.05."),
            ru("\\u0414\\u043b\\u044f \\u043e\\u0431\\u0443\\u0447\\u0435\\u043d\\u0438\\u044f logistic-\\u043c\\u043e\\u0434\\u0435\\u043b\\u0438 \\u043d\\u0443\\u0436\\u043d\\u043e \\u043c\\u0438\\u043d\\u0438\\u043c\\u0443\\u043c 8 \\u043f\\u043e\\u043b\\u043e\\u0436\\u0438\\u0442\\u0435\\u043b\\u044c\\u043d\\u044b\\u0445 \\u0438 8 \\u043e\\u0442\\u0440\\u0438\\u0446\\u0430\\u0442\\u0435\\u043b\\u044c\\u043d\\u044b\\u0445 \\u043d\\u0430\\u0431\\u043b\\u044e\\u0434\\u0435\\u043d\\u0438\\u0439 (`MIN_EVENT_CLASS_COUNT = 8`)."),
            ru("\\u0414\\u043b\\u044f negative binomial \\u043d\\u0443\\u0436\\u043d\\u044b \\u043c\\u0438\\u043d\\u0438\\u043c\\u0443\\u043c 56 \\u043e\\u0431\\u0443\\u0447\\u0430\\u044e\\u0449\\u0438\\u0445 \\u0441\\u0442\\u0440\\u043e\\u043a \\u0438 overdispersion >= 1.35."),
        ],
    )

    add_section(
        document,
        ru("10. \\u0418\\u0442\\u043e\\u0433\\u043e\\u0432\\u0430\\u044f \\u0442\\u0435\\u0445\\u043d\\u0438\\u0447\\u0435\\u0441\\u043a\\u0430\\u044f \\u0441\\u0445\\u0435\\u043c\\u0430"),
        [
            ru(
                "\\u041c\\u043e\\u0434\\u0443\\u043b\\u044c `app/services/ml_model/` "
                "\\u043f\\u0440\\u0435\\u0434\\u0441\\u0442\\u0430\\u0432\\u043b\\u044f\\u0435\\u0442 \\u0441\\u043e\\u0431\\u043e\\u0439 "
                "\\u043e\\u0442\\u0434\\u0435\\u043b\\u044c\\u043d\\u044b\\u0439 \\u0430\\u043d\\u0430\\u043b\\u0438\\u0442\\u0438\\u0447\\u0435\\u0441\\u043a\\u0438\\u0439 "
                "\\u043a\\u043e\\u043d\\u0442\\u0443\\u0440, \\u0433\\u0434\\u0435 `forecasting` \\u0438\\u0441\\u043f\\u043e\\u043b\\u044c\\u0437\\u0443\\u0435\\u0442\\u0441\\u044f "
                "\\u043a\\u0430\\u043a data layer \\u0438 heuristic baseline, "
                "\\u0430 ML-\\u0447\\u0430\\u0441\\u0442\\u044c \\u0434\\u043e\\u0431\\u0430\\u0432\\u043b\\u044f\\u0435\\u0442 "
                "lead-time-aware rolling-origin backtesting, "
                "model selection, conformal-style intervals, permutation importance "
                "\\u0438 cache-aware background execution."
            ),
            ru(
                "\\u041f\\u043e \\u0444\\u0430\\u043a\\u0442\\u0443 \\u044d\\u0442\\u043e "
                "\\u0433\\u0438\\u0431\\u0440\\u0438\\u0434\\u043d\\u044b\\u0439 \\u043f\\u0430\\u0439\\u043f\\u043b\\u0430\\u0439\\u043d: "
                "sklearn + statsmodels + pandas/numpy + in-memory caches + background job orchestration."
            ),
        ],
    )

    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
