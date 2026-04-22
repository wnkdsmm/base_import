from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"F:\filesFires\base_import")
OUTPUT_PATH = ROOT / "documents" / "project_structure.docx"
EXCLUDE_DIRS = {".git", "__pycache__"}
EXCLUDE_FILES = {OUTPUT_PATH.name}


def ru(value: str) -> str:
    return value.encode("ascii").decode("unicode_escape")


def text(*parts: str) -> str:
    return ru("".join(parts))


TOP_LEVEL_NOTES = {
    ".claude": text(
        r"\u043b\u043e\u043a\u0430\u043b\u044c\u043d\u0430\u044f ",
        r"\u0441\u043b\u0443\u0436\u0435\u0431\u043d\u0430\u044f ",
        r"\u043a\u043e\u043d\u0444\u0438\u0433\u0443\u0440\u0430\u0446\u0438\u044f ",
        r"\u0430\u0441\u0441\u0438\u0441\u0442\u0438\u0432\u043d\u044b\u0445 ",
        r"\u0438\u043d\u0441\u0442\u0440\u0443\u043c\u0435\u043d\u0442\u043e\u0432",
    ),
    ".vscode": text(
        r"\u043b\u043e\u043a\u0430\u043b\u044c\u043d\u0430\u044f ",
        r"\u043a\u043e\u043d\u0444\u0438\u0433\u0443\u0440\u0430\u0446\u0438\u044f ",
        r"\u0441\u0440\u0435\u0434\u044b \u0440\u0430\u0437\u0440\u0430\u0431\u043e\u0442\u043a\u0438 ",
        r"Visual Studio Code",
    ),
    "app": text(
        r"\u0432\u0435\u0431-\u043f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u0435: ",
        r"\u043c\u0430\u0440\u0448\u0440\u0443\u0442\u044b, \u0441\u0435\u0440\u0432\u0438\u0441\u044b, ",
        r"\u0448\u0430\u0431\u043b\u043e\u043d\u044b \u0438\u043d\u0442\u0435\u0440\u0444\u0435\u0439\u0441\u0430, ",
        r"\u0441\u0442\u0430\u0442\u0438\u0447\u0435\u0441\u043a\u0438\u0435 \u0440\u0435\u0441\u0443\u0440\u0441\u044b ",
        r"\u0438 \u043f\u0440\u0438\u043a\u043b\u0430\u0434\u043d\u0430\u044f \u043b\u043e\u0433\u0438\u043a\u0430 ",
        r"\u0430\u043d\u0430\u043b\u0438\u0442\u0438\u043a\u0438",
    ),
    "config": text(
        r"\u043a\u043e\u043d\u0444\u0438\u0433\u0443\u0440\u0430\u0446\u0438\u044f ",
        r"\u043f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u044f, \u043f\u0443\u0442\u0438, ",
        r"\u043f\u0430\u0440\u0430\u043c\u0435\u0442\u0440\u044b \u0438 ",
        r"\u043f\u043e\u0434\u043a\u043b\u044e\u0447\u0435\u043d\u0438\u0435 \u043a \u0434\u0430\u043d\u043d\u044b\u043c",
    ),
    "core": text(
        r"\u0431\u0430\u0437\u043e\u0432\u044b\u0435 \u043c\u0435\u0445\u0430\u043d\u0438\u0437\u043c\u044b ",
        r"\u043e\u0431\u0440\u0430\u0431\u043e\u0442\u043a\u0438 \u0434\u0430\u043d\u043d\u044b\u0445 ",
        r"\u0438 \u043f\u043e\u0441\u0442\u0440\u043e\u0435\u043d\u0438\u044f ",
        r"\u043a\u0430\u0440\u0442\u043e\u0433\u0440\u0430\u0444\u0438\u0447\u0435\u0441\u043a\u0438\u0445 ",
        r"\u043f\u0440\u0435\u0434\u0441\u0442\u0430\u0432\u043b\u0435\u043d\u0438\u0439",
    ),
    "data": text(
        r"\u0440\u0430\u0431\u043e\u0447\u0438\u0435 \u0434\u0430\u043d\u043d\u044b\u0435, ",
        r"\u0440\u0435\u0437\u0443\u043b\u044c\u0442\u0430\u0442\u044b \u043e\u0431\u0440\u0430\u0431\u043e\u0442\u043a\u0438 ",
        r"\u0438 \u0437\u0430\u0433\u0440\u0443\u0436\u0435\u043d\u043d\u044b\u0435 \u043d\u0430\u0431\u043e\u0440\u044b",
    ),
    "documents": text(
        r"\u043f\u0440\u043e\u0435\u043a\u0442\u043d\u0430\u044f \u0438 ",
        r"\u043f\u043e\u043b\u044c\u0437\u043e\u0432\u0430\u0442\u0435\u043b\u044c\u0441\u043a\u0430\u044f ",
        r"\u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442\u0430\u0446\u0438\u044f",
    ),
    "sample_data": text(
        r"\u043f\u0440\u0438\u043c\u0435\u0440\u043d\u044b\u0435 \u0432\u0445\u043e\u0434\u043d\u044b\u0435 ",
        r"\u0434\u0430\u043d\u043d\u044b\u0435 \u0434\u043b\u044f \u0434\u0435\u043c\u043e\u043d\u0441\u0442\u0440\u0430\u0446\u0438\u0438 ",
        r"\u0438 \u0442\u0435\u0441\u0442\u043e\u0432\u044b\u0445 \u0437\u0430\u043f\u0443\u0441\u043a\u043e\u0432",
    ),
    "scripts": text(
        r"\u0441\u043b\u0443\u0436\u0435\u0431\u043d\u044b\u0435 \u0441\u0446\u0435\u043d\u0430\u0440\u0438\u0438 ",
        r"\u0434\u043b\u044f \u043f\u043e\u0434\u0433\u043e\u0442\u043e\u0432\u043a\u0438 ",
        r"\u0434\u0430\u043d\u043d\u044b\u0445 \u0438 \u043f\u0440\u043e\u0432\u0435\u0440\u043e\u043a",
    ),
    "tests": text(
        r"\u0430\u0432\u0442\u043e\u043c\u0430\u0442\u0438\u0447\u0435\u0441\u043a\u0438\u0435 ",
        r"\u0442\u0435\u0441\u0442\u044b \u0438 \u0444\u0438\u043a\u0441\u0442\u0443\u0440\u044b",
    ),
}


def build_tree(base: Path) -> list[str]:
    lines = [base.name]

    def walk(path: Path, prefix: str = "") -> None:
        entries = []
        for child in sorted(path.iterdir(), key=lambda item: (not item.is_dir(), item.name.lower())):
            if child.name in EXCLUDE_DIRS or child.name in EXCLUDE_FILES:
                continue
            entries.append(child)
        for index, child in enumerate(entries):
            branch = "\u2514\u2500\u2500 " if index == len(entries) - 1 else "\u251c\u2500\u2500 "
            lines.append(prefix + branch + child.name)
            if child.is_dir():
                extension = "    " if index == len(entries) - 1 else "\u2502   "
                walk(child, prefix + extension)

    walk(base)
    return lines


def count_items(base: Path) -> tuple[int, int]:
    dir_count = 0
    file_count = 0
    for path in base.rglob("*"):
        relative_parts = path.relative_to(base).parts
        if any(part in EXCLUDE_DIRS for part in relative_parts):
            continue
        if path.name in EXCLUDE_FILES:
            continue
        if path.is_dir():
            dir_count += 1
        else:
            file_count += 1
    return dir_count, file_count


def set_font(run, name: str, size: int) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def configure_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)


def add_paragraph_with_run(document: Document, body: str, bold: bool = False, size: int = 12) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(body)
    run.bold = bold
    set_font(run, "Times New Roman", size)


def main() -> None:
    root_items = []
    for child in sorted(ROOT.iterdir(), key=lambda item: (not item.is_dir(), item.name.lower())):
        if child.name in EXCLUDE_DIRS or child.name in EXCLUDE_FILES:
            continue
        root_items.append(child)

    folder_count, file_count = count_items(ROOT)
    tree_lines = build_tree(ROOT)

    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title_run = title.add_run(
        text(
            r"\u0421\u0442\u0440\u0443\u043a\u0442\u0443\u0440\u0430 ",
            r"\u043f\u0440\u043e\u0435\u043a\u0442\u0430 ",
            r"\u043c\u0430\u0433\u0438\u0441\u0442\u0435\u0440\u0441\u043a\u043e\u0439 ",
            r"\u0434\u0438\u0441\u0441\u0435\u0440\u0442\u0430\u0446\u0438\u0438",
        )
    )
    title_run.bold = True
    set_font(title_run, "Times New Roman", 16)

    add_paragraph_with_run(
        document,
        text(
            r"\u0422\u0435\u043c\u0430: \u00ab\u0420\u0430\u0437\u0440\u0430\u0431\u043e\u0442\u043a\u0430 ",
            r"\u0441\u0435\u0440\u0432\u0438\u0441\u0430 ",
            r"\u0438\u043d\u0442\u0435\u043b\u043b\u0435\u043a\u0442\u0443\u0430\u043b\u044c\u043d\u043e\u0433\u043e ",
            r"\u0430\u043d\u0430\u043b\u0438\u0437\u0430 \u0434\u0430\u043d\u043d\u044b\u0445 ",
            r"\u0434\u043b\u044f \u043f\u043e\u0434\u0434\u0435\u0440\u0436\u043a\u0438 ",
            r"\u043f\u0440\u0438\u043d\u044f\u0442\u0438\u044f \u0440\u0435\u0448\u0435\u043d\u0438\u0439 ",
            r"\u0432 \u043e\u0431\u043b\u0430\u0441\u0442\u0438 ",
            r"\u043f\u043e\u0436\u0430\u0440\u043d\u043e\u0439 ",
            r"\u0431\u0435\u0437\u043e\u043f\u0430\u0441\u043d\u043e\u0441\u0442\u0438 ",
            r"\u0441\u0435\u043b\u044c\u0441\u043a\u0438\u0445 ",
            r"\u0442\u0435\u0440\u0440\u0438\u0442\u043e\u0440\u0438\u0439\u00bb.",
        ),
    )
    add_paragraph_with_run(
        document,
        text(
            r"\u041a\u043e\u0440\u043d\u0435\u0432\u043e\u0439 ",
            r"\u043a\u0430\u0442\u0430\u043b\u043e\u0433 ",
            r"\u043f\u0440\u043e\u0435\u043a\u0442\u0430: F:\\filesFires\\base_import",
        ),
    )
    add_paragraph_with_run(
        document,
        text(
            r"\u0414\u0430\u0442\u0430 ",
            r"\u0444\u043e\u0440\u043c\u0438\u0440\u043e\u0432\u0430\u043d\u0438\u044f ",
            r"\u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442\u0430: 20.04.2026",
        ),
    )

    summary = document.add_paragraph()
    summary_head = summary.add_run(text(r"\u041a\u0440\u0430\u0442\u043a\u0430\u044f \u0441\u0432\u043e\u0434\u043a\u0430. "))
    summary_head.bold = True
    set_font(summary_head, "Times New Roman", 12)
    summary_text = summary.add_run(
        text(
            r"\u0412 \u0441\u0442\u0440\u0443\u043a\u0442\u0443\u0440\u0435 \u043f\u0440\u043e\u0435\u043a\u0442\u0430 ",
            rf"\u0437\u0430\u0444\u0438\u043a\u0441\u0438\u0440\u043e\u0432\u0430\u043d\u043e {folder_count} ",
            r"\u043a\u0430\u0442\u0430\u043b\u043e\u0433\u043e\u0432 \u0438 ",
            rf"{file_count} \u0444\u0430\u0439\u043b\u043e\u0432 ",
            r"\u0431\u0435\u0437 \u0443\u0447\u0435\u0442\u0430 \u0441\u043b\u0443\u0436\u0435\u0431\u043d\u043e\u0439 ",
            r"\u043f\u0430\u043f\u043a\u0438 .git.",
        )
    )
    set_font(summary_text, "Times New Roman", 12)

    add_paragraph_with_run(
        document,
        text(r"\u0421\u043e\u0441\u0442\u0430\u0432 \u043a\u043e\u0440\u043d\u0435\u0432\u043e\u0433\u043e \u0443\u0440\u043e\u0432\u043d\u044f:"),
        bold=True,
    )
    for item in root_items:
        label = text(r"\u043a\u0430\u0442\u0430\u043b\u043e\u0433") if item.is_dir() else text(r"\u0444\u0430\u0439\u043b")
        document.add_paragraph(f"{item.name} ({label})", style="List Bullet")

    add_paragraph_with_run(
        document,
        text(r"\u041d\u0430\u0437\u043d\u0430\u0447\u0435\u043d\u0438\u0435 \u043e\u0441\u043d\u043e\u0432\u043d\u044b\u0445 \u043a\u0430\u0442\u0430\u043b\u043e\u0433\u043e\u0432:"),
        bold=True,
    )
    for item in root_items:
        if item.is_dir() and item.name in TOP_LEVEL_NOTES:
            document.add_paragraph(f"{item.name} - {TOP_LEVEL_NOTES[item.name]}.", style="List Bullet")

    add_paragraph_with_run(
        document,
        text(r"\u041f\u043e\u043b\u043d\u0430\u044f \u0438\u0435\u0440\u0430\u0440\u0445\u0438\u044f \u043f\u0440\u043e\u0435\u043a\u0442\u0430:"),
        bold=True,
    )
    for line in tree_lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        set_font(run, "Consolas", 9)

    note = document.add_paragraph()
    note_head = note.add_run(text(r"\u041f\u0440\u0438\u043c\u0435\u0447\u0430\u043d\u0438\u0435. "))
    note_head.bold = True
    set_font(note_head, "Times New Roman", 12)
    note_text = note.add_run(
        text(
            r"\u0421\u0442\u0440\u0443\u043a\u0442\u0443\u0440\u0430 \u0441\u0444\u043e\u0440\u043c\u0438\u0440\u043e\u0432\u0430\u043d\u0430 ",
            r"\u0430\u0432\u0442\u043e\u043c\u0430\u0442\u0438\u0447\u0435\u0441\u043a\u0438 ",
            r"\u043f\u043e \u0442\u0435\u043a\u0443\u0449\u0435\u043c\u0443 ",
            r"\u0441\u043e\u0441\u0442\u043e\u044f\u043d\u0438\u044e \u0440\u0430\u0431\u043e\u0447\u0435\u0439 ",
            r"\u0434\u0438\u0440\u0435\u043a\u0442\u043e\u0440\u0438\u0438 \u043f\u0440\u043e\u0435\u043a\u0442\u0430. ",
            r"\u041f\u0430\u043f\u043a\u0430 .git \u0438\u0441\u043a\u043b\u044e\u0447\u0435\u043d\u0430 ",
            r"\u0438\u0437 \u043f\u0440\u0435\u0434\u0441\u0442\u0430\u0432\u043b\u0435\u043d\u0438\u044f ",
            r"\u043a\u0430\u043a \u0441\u043b\u0443\u0436\u0435\u0431\u043d\u0430\u044f.",
        )
    )
    set_font(note_text, "Times New Roman", 12)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
