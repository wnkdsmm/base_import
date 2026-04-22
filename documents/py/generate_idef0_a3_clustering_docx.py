from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUTPUT_PATH = Path("documents/IDEF0_A3_Пространственная_кластеризация_пожаров.docx")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_par(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_formula(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"


def count_question_marks_in_docx(path: Path) -> int:
    if not path.exists():
        return -1
    with ZipFile(path, "r") as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    text_chunks = re.findall(r">([^<]*)<", xml)
    text = "".join(text_chunks)
    return text.count("?")


def build_doc() -> Document:
    doc = Document()
    add_heading(doc, "IDEF0 A3: Пространственная кластеризация пожаров", level=0)
    add_par(
        doc,
        (
            "Проанализирован модуль app/services/clustering/. "
            "Описание основано на фактической реализации в файлах core_runner.py, core_algorithms.py, "
            "analysis_stats.py, analysis_features.py, data_impl.py, quality_*.py, jobs.py."
        ),
    )

    add_heading(doc, "1. IDEF0 — Диаграмма A3 и декомпозиция A3.1–A3.5", level=1)

    add_heading(doc, "A3.1 Подготовка и нормализация данных", level=2)
    add_bullet(doc, "Вход: table_name, sample_limit, sampling_strategy, сырые записи инцидентов из _collect_risk_inputs().")
    add_bullet(
        doc,
        (
            "Управление: CLUSTER_COUNT_OPTIONS, SAMPLE_LIMIT_OPTIONS, SAMPLING_STRATEGY_VALUES, "
            "RATE_SMOOTHING_PRIOR_STRENGTH, MEAN_SMOOTHING_PRIOR_STRENGTH."
        ),
    )
    add_bullet(
        doc,
        (
            "Выход: ClusteringDatasetBundle (entity_frame, feature_frame, candidate_features, support_summary), "
            "sampling_note, notes."
        ),
    )
    add_bullet(
        doc,
        (
            "Механизм: _load_territory_dataset(), _aggregate_territory_frame(), _sample_territory_frame(), "
            "_prepare_cluster_frame(), pandas, numpy, StandardScaler (на следующем шаге)."
        ),
    )

    add_heading(doc, "A3.2 Выбор алгоритма и параметров", level=2)
    add_bullet(doc, "Вход: cluster_frame/entity_frame, selected_features, requested_cluster_count, cluster_count_is_explicit.")
    add_bullet(
        doc,
        (
            "Управление: FEATURE_SELECTION_MIN_IMPROVEMENT, PROFILE_MODE_SCORE_TOLERANCE, "
            "PROFILE_MODE_SILHOUETTE_TOLERANCE, WEIGHTING_STRATEGY_*."
        ),
    )
    add_bullet(
        doc,
        (
            "Выход: diagnostics (rows, method_rows_by_cluster_count, best_quality_k, best_silhouette_k, best_gap_k, elbow_k), "
            "render_configuration (algorithm_key/method_key/weighting_strategy/cluster_count)."
        ),
    )
    add_bullet(
        doc,
        (
            "Механизм: _evaluate_cluster_counts(), _compare_clustering_methods(), _select_render_configuration(), "
            "_estimate_best_k_gap(), _estimate_elbow_k()."
        ),
    )

    add_heading(doc, "A3.3 Выполнение кластеризации", level=2)
    add_bullet(doc, "Вход: normalized feature matrix, sample_weights, выбранные algorithm_key и cluster_count.")
    add_bullet(doc, "Управление: MODEL_N_INIT, STABILITY_RANDOM_SEEDS, STABILITY_RESAMPLE_RATIO, HOPKINS_MIN_CLUSTERABLE.")
    add_bullet(
        doc,
        (
            "Выход: labels, scaled_centers/raw_centers, inertia, pca_projection, stability_ari, initialization_ari, "
            "method_comparison."
        ),
    )
    add_bullet(
        doc,
        (
            "Механизм: _run_clustering(), _fit_weighted_kmeans(), _fit_clustering_labels(), "
            "_compute_pca_projection(), _estimate_kmeans_initialization_stability(), _estimate_resampled_stability()."
        ),
    )

    add_heading(doc, "A3.4 Оценка качества кластеров", level=2)
    add_bullet(doc, "Вход: labels, scaled_points, diagnostics rows, support_summary, selected_features.")
    add_bullet(
        doc,
        (
            "Управление: пороги интерпретации silhouette/davies-bouldin/balance/stability в compute_segmentation_strength(), "
            "LOW_SUPPORT_TERRITORY_THRESHOLD."
        ),
    )
    add_bullet(
        doc,
        "Выход: quality_assessment (metric_cards, methodology_items, comparison_rows, dissertation_points), cluster_count_guidance.",
    )
    add_bullet(
        doc,
        (
            "Механизм: compute_clustering_metrics(), _cluster_quality_score(), _build_clustering_quality_assessment(), "
            "_build_cluster_count_guidance()."
        ),
    )

    add_heading(doc, "A3.5 Географическая интерпретация результатов", level=2)
    add_bullet(
        doc,
        (
            "Вход: labels + агрегированные территориальные сущности (Территория, Район, Тип территории), "
            "cluster centers и профили."
        ),
    )
    add_bullet(
        doc,
        "Управление: предметные поля территорий, справочники feature metadata, правила текста интерпретации в _build_notes().",
    )
    add_bullet(
        doc,
        (
            "Выход: cluster_profiles, centroid_rows, representative_rows, cluster_risk, charts (scatter/radar/distribution/diagnostics), "
            "summary + notes."
        ),
    )
    add_bullet(
        doc,
        (
            "Механизм: _build_cluster_profiles(), _build_centroid_table(), _build_representative_rows(), "
            "compute_cluster_risk_scores(), charts.py."
        ),
    )

    add_heading(doc, "Стрелки между подфункциями (выход → вход)", level=2)
    add_bullet(doc, "A3.1 → A3.2: candidate_features, feature_frame/entity_frame, support_summary.")
    add_bullet(doc, "A3.2 → A3.3: selected algorithm/method, weighting_strategy, cluster_count.")
    add_bullet(doc, "A3.3 → A3.4: labels, silhouette/davies/calinski/balance, stability, diagnostics rows.")
    add_bullet(doc, "A3.4 → A3.5: quality_assessment + cluster_count_guidance для интерпретации и финального описания.")
    add_bullet(doc, "A3.3 → A3.5: pca_projection, raw_centers, representative territories для географически осмысленной выдачи.")

    add_heading(doc, "2. Алгоритмы", level=1)

    add_heading(doc, "2.1 Алгоритм кластеризации", level=2)
    add_par(
        doc,
        (
            "Фактическая реализация: KMeans (основной), AgglomerativeClustering (Ward), Birch. "
            "DBSCAN в этом модуле не реализован: параметры eps/min_samples не используются."
        ),
    )
    add_bullet(doc, "Класс sklearn: sklearn.cluster.KMeans, sklearn.cluster.AgglomerativeClustering, sklearn.cluster.Birch.")
    add_bullet(
        doc,
        "Метрика расстояния в текущей реализации фактически евклидова в стандартизованном пространстве признаков (после StandardScaler).",
    )
    add_bullet(
        doc,
        (
            "Haversine/геодезическая дистанция не применяется: кластеризация идёт не по сырым координатам, "
            "а по агрегированным территориальным признакам риска."
        ),
    )
    add_bullet(
        doc,
        (
            "Весовая стратегия: uniform/not_applicable -> w_i = 1; incident_log -> "
            "w_i = log(1 + fire_count_i) / mean_j(log(1 + fire_count_j))."
        ),
    )

    add_formula(doc, "Целевая функция KMeans с весами:")
    add_formula(doc, "J = Σ_i w_i * ||x_i - μ_{c_i}||^2")

    add_par(doc, "Псевдокод (основной цикл):")
    add_code_block(
        doc,
        (
            "INPUT: cluster_frame, entity_frame, k, algorithm_key, weighting_strategy\n"
            "model_frame <- maybe_log_transform(cluster_frame)\n"
            "scaled_points <- StandardScaler.fit_transform(model_frame)\n"
            "sample_weights <- build_sample_weights(entity_frame, weighting_strategy)\n"
            "IF algorithm_key == 'kmeans':\n"
            "    model <- KMeans(n_clusters=k, n_init=MODEL_N_INIT, random_state=42)\n"
            "    model.fit(scaled_points, sample_weight=sample_weights)\n"
            "    labels <- model.labels_\n"
            "    centers_scaled <- model.cluster_centers_\n"
            "    inertia <- model.inertia_\n"
            "ELSE:\n"
            "    labels <- fit_predict(Agglomerative/Birch)\n"
            "    centers_scaled <- mean points per label\n"
            "    inertia <- sum squared distances to centers\n"
            "metrics <- compute_clustering_metrics(scaled_points, labels)\n"
            "stability_ari <- resampled_stability(scaled_points, labels)\n"
            "pca_projection <- PCA(n_components=2) for visualization\n"
            "RETURN labels, centers, metrics, inertia, stability, pca_projection\n"
        ),
    )

    add_heading(doc, "2.2 Алгоритм выбора оптимального числа кластеров", level=2)
    add_bullet(
        doc,
        (
            "Используются несколько критериев одновременно: silhouette, composite quality_score, "
            "Gap Statistic, Elbow (по инерции)."
        ),
    )
    add_bullet(doc, "Диапазон перебора: k из CLUSTER_COUNT_OPTIONS, ограниченный сверху len(cluster_frame)-1.")
    add_bullet(
        doc,
        (
            "Критерий выбора рабочего k: best_quality_k по _diagnostics_row_sort_key "
            "(quality_score -> silhouette -> -davies_bouldin -> balance -> -shape_penalty)."
        ),
    )
    add_bullet(
        doc,
        (
            "Критерий остановки: явной early-stop логики нет; выполняется полный перебор всех допустимых k, "
            "после чего выбирается лучший."
        ),
    )

    add_par(doc, "Псевдокод (подбор k):")
    add_code_block(
        doc,
        (
            "INPUT: cluster_frame, entity_frame, weighting_strategy\n"
            "available_ks <- {k in CLUSTER_COUNT_OPTIONS | 2 <= k <= min(max_k, n-1)}\n"
            "rows <- []\n"
            "FOR each k in available_ks:\n"
            "    method_rows <- compare_methods_at_k(kmeans/agglomerative/birch)\n"
            "    best_row_k <- recommended row by diagnostics sort key\n"
            "    rows.append(best_row_k with cluster_count=k)\n"
            "best_silhouette_k <- argmax_k silhouette(k)\n"
            "best_quality_k <- argmax_k diagnostics_sort_key(k)\n"
            "best_gap_k <- estimate_best_k_gap(gap_scores)\n"
            "elbow_k <- estimate_elbow_k(rows by inertia)\n"
            "IF cluster_count_is_explicit == False and best_quality_k exists:\n"
            "    selected_k <- best_quality_k\n"
            "ELSE:\n"
            "    selected_k <- requested_k\n"
            "RETURN diagnostics, selected_k\n"
        ),
    )

    add_heading(doc, "2.3 Алгоритм оценки качества (силуэт и сопутствующие метрики)", level=2)
    add_formula(doc, "Формула силуэта (для объекта i):")
    add_formula(doc, "s(i) = (b(i) - a(i)) / max(a(i), b(i))")
    add_formula(doc, "a(i): среднее расстояние от i до объектов своего кластера")
    add_formula(doc, "b(i): минимум средних расстояний от i до объектов соседних кластеров")
    add_par(doc, "Интерпретация в контексте пожарных территорий:")
    add_bullet(doc, "s(i) близко к 1: территория хорошо соответствует своему типу риска и отделена от соседних типов.")
    add_bullet(doc, "s(i) около 0: пограничная территория, профиль риска похож на несколько кластеров.")
    add_bullet(doc, "s(i) < 0: возможно некорректное отнесение территории к кластеру.")
    add_bullet(
        doc,
        (
            "В коде дополнительно считаются davies_bouldin, calinski_harabasz, "
            "cluster_balance_ratio и shape_penalty для устойчивого выбора."
        ),
    )

    add_par(doc, "Псевдокод (оценка качества):")
    add_code_block(
        doc,
        (
            "INPUT: scaled_points, labels\n"
            "IF number_of_clusters < 2 OR n_points <= number_of_clusters:\n"
            "    silhouette <- None\n"
            "    davies_bouldin <- None\n"
            "    calinski_harabasz <- None\n"
            "ELSE:\n"
            "    silhouette <- silhouette_score(scaled_points, labels)\n"
            "    davies_bouldin <- davies_bouldin_score(scaled_points, labels)\n"
            "    calinski_harabasz <- calinski_harabasz_score(scaled_points, labels)\n"
            "balance_ratio <- min_cluster_size / max_cluster_size\n"
            "shape_penalty <- penalty_for_microclusters_and_imbalance\n"
            "quality_score <- 0.55*silhouette + 0.20*(1/(1+DB)) + 0.15*scaled_CH + 0.10*balance_ratio - shape_penalty\n"
            "RETURN all metrics\n"
        ),
    )

    add_heading(doc, "2.4 Алгоритм взвешивания признаков (WeightingStrategy)", level=2)
    add_par(
        doc,
        (
            "В текущем коде WeightingStrategy влияет прежде всего на веса наблюдений (территорий), "
            "а не на явные ручные коэффициенты для каждого признака."
        ),
    )
    add_bullet(doc, "incident_log: усиливает вклад территорий с большим числом пожаров через log1p-нормированные sample weights.")
    add_bullet(doc, "uniform: все территории имеют одинаковый вес.")
    add_bullet(doc, "not_applicable: используется для алгоритмов, где sample weights не применяются.")
    add_formula(doc, "w_i(incident_log) = log(1 + c_i) / mean_j(log(1 + c_j)), где c_i — число пожаров территории i.")
    add_formula(doc, "Геометрический эффект: увеличенные w_i смещают центры кластеров в сторону территорий с высокой нагрузкой.")
    add_par(
        doc,
        (
            "Пользователь в UI влияет на геометрию косвенно: через выбор набора признаков, режима (profile/load) "
            "и параметров выборки; прямой ручной ввод вектора feature-weights в модуле не реализован."
        ),
    )

    add_heading(doc, "2.5 Алгоритм диагностики устойчивости", level=2)
    add_bullet(
        doc,
        (
            "Метод: perturbation/subsampling (без возвращения), а не классический bootstrap с replacement."
        ),
    )
    add_bullet(
        doc,
        (
            "Initialization stability: ARI между запусками KMeans с разными random_state на одном и том же датасете "
            "(_estimate_kmeans_initialization_stability)."
        ),
    )
    add_bullet(
        doc,
        (
            "Resampled stability: для каждого seed берётся подвыборка размера ~STABILITY_RESAMPLE_RATIO (0.8), "
            "строятся кластеры, затем ARI считается на пересечении индексов двух подвыборок "
            "(_estimate_resampled_stability)."
        ),
    )
    add_formula(doc, "Stability = mean(ARI(labeling_a, labeling_b)) по парам запусков.")
    add_bullet(doc, "Чем выше ARI, тем воспроизводимее сегментация при изменении состава данных.")

    add_heading(doc, "3. Примечания по соответствию запросу", level=1)
    add_bullet(
        doc,
        (
            "DBSCAN (eps, min_samples) и haversine-дистанция в текущем модуле app/services/clustering/ не используются. "
            "Если нужно, их можно добавить как отдельную ветку algorithm_key='dbscan' с метрическим режимом для координат."
        ),
    )
    add_bullet(
        doc,
        (
            "Географическая интерпретация в модуле выполняется через территориальные агрегаты (территория/район/тип), "
            "а не через геометрию точек на сфере."
        ),
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    qmarks = count_question_marks_in_docx(OUTPUT_PATH)
    print(f"DOCX_CREATED={OUTPUT_PATH.resolve()}")
    print(f"TEXT_QUESTION_MARKS={qmarks}")


if __name__ == "__main__":
    main()
