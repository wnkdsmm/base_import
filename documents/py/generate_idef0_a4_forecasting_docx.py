from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document


OUTPUT_PATH = Path("documents/IDEF0_A4_Прогнозирование_пожарной_активности.docx")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_par(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"


def count_text_question_marks(path: Path) -> int:
    if not path.exists():
        return -1
    with ZipFile(path, "r") as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    text = "".join(re.findall(r">([^<]*)<", xml))
    return text.count("?")


def build_doc() -> Document:
    doc = Document()
    add_heading(doc, "IDEF0 A4: Прогнозирование пожарной активности", level=0)
    add_par(
        doc,
        (
            "Анализ выполнен по модулю app/services/forecasting/: forecasting_pipeline.py, "
            "assembly_input.py, assembly_output.py, sql_sources_registry.py, sql_sources_query.py, "
            "sql_payload.py, sql_aggregations.py, shaping.py, quality.py, inputs.py."
        ),
    )

    add_heading(doc, "1. IDEF0 — A4 и декомпозиция A4.1–A4.5", level=1)

    add_heading(doc, "A4.1 Формирование выборки источников данных", level=2)
    add_bullet(
        doc,
        "Вход: table_name, district, cause, object_category, history_window; список пользовательских таблиц.",
    )
    add_bullet(
        doc,
        (
            "Управление: _build_forecasting_table_options(), _resolve_forecasting_selection(), "
            "_selected_source_tables(), _canonicalize_source_tables(), _history_window_year_span()."
        ),
    )
    add_bullet(
        doc,
        (
            "Выход: selected_table, source_tables, source_table_notes, metadata_items "
            "(resolved_columns, column_quality)."
        ),
    )
    add_bullet(
        doc,
        (
            "Механизм: selection.py, sources.py (_collect_forecasting_metadata/_collect_forecasting_inputs), "
            "get_table_columns_cached(), SQL-кэш в sql.py."
        ),
    )

    add_heading(doc, "A4.2 SQL-агрегация временных рядов", level=2)
    add_bullet(
        doc,
        "Вход: source_tables + metadata_items + фильтры (district/cause/object_category) + min_year.",
    )
    add_bullet(
        doc,
        (
            "Управление: SourceQueryRegistryMixin/SourceQuerySqlMixin, _build_scope_conditions(), "
            "_build_materialized_scope_conditions(), history_window."
        ),
    )
    add_bullet(
        doc,
        (
            "Выход: dense daily_history (date, count, avg_temperature) и filtered_records_count; "
            "option_catalog (districts/causes/object_categories)."
        ),
    )
    add_bullet(
        doc,
        (
            "Механизм: sql_payload.py (_build_daily_history_sql), sql_sources_registry.py "
            "(union fast-path + fallback), sql_aggregations.py (GROUP BY, merge, densification), SQLAlchemy engine."
        ),
    )

    add_heading(doc, "A4.3 Построение прогноза", level=2)
    add_bullet(doc, "Вход: daily_history, forecast_days, scenario temperature.")
    add_bullet(
        doc,
        (
            "Управление: _build_forecast_history_stats(), weekday/month seasonality factors, "
            "trend_ratio, robust_ceiling, temperature_slope."
        ),
    )
    add_bullet(
        doc,
        (
            "Выход: forecast_rows по дням (forecast_value, fire_probability, lower_bound, upper_bound, "
            "scenario_label, hints)."
        ),
    )
    add_bullet(
        doc,
        "Механизм: shaping.py (_build_forecast_rows, _forecast_event_probability, _forecast_temperature_effect).",
    )

    add_heading(doc, "A4.4 Расчёт доверительных интервалов (bootstrap)", level=2)
    add_bullet(
        doc,
        (
            "Вход: точечный прогноз и оценка вариативности ряда."
        ),
    )
    add_bullet(
        doc,
        (
            "Управление: в текущем коде нет статистического bootstrap; используется эвристический spread "
            "из volatility и горизонта шага."
        ),
    )
    add_bullet(
        doc,
        (
            "Выход: lower_bound/upper_bound для каждого дня."
        ),
    )
    add_bullet(
        doc,
        (
            "Механизм: shaping.py (_build_forecast_rows): "
            "spread = max(0.75, volatility * (0.9 + 0.15 * sqrt(step)))."
        ),
    )

    add_heading(doc, "A4.5 Оценка качества и формирование ответа", level=2)
    add_bullet(doc, "Вход: daily_history, forecast_rows, metadata, feature coverage, backtest windows.")
    add_bullet(
        doc,
        (
            "Управление: _run_scenario_backtesting(), _build_scenario_quality_assessment(), "
            "compute_count_metrics() (MAE/RMSE/SMAPE)."
        ),
    )
    add_bullet(
        doc,
        (
            "Выход: summary, quality_assessment, insights, charts, executive_brief, "
            "risk_prediction (через decision support), итоговый ForecastingPayload."
        ),
    )
    add_bullet(
        doc,
        (
            "Механизм: assembly_input.py/assembly_output.py, presentation.py, quality.py, "
            "cache _FORECASTING_CACHE (TTL 120 сек), input/sql cache (TTL 120 сек)."
        ),
    )

    add_heading(doc, "Стрелки между A4.x", level=2)
    add_bullet(doc, "A4.1 -> A4.2: source_tables, metadata_items, resolved_columns, source_table_notes.")
    add_bullet(doc, "A4.2 -> A4.3: daily_history + filtered_records_count + option_catalog.")
    add_bullet(doc, "A4.3 -> A4.4: point forecast (estimate) + volatility + horizon step.")
    add_bullet(doc, "A4.4 -> A4.5: интервальные границы lower/upper для UI и пояснений.")
    add_bullet(doc, "A4.2 + A4.3 + A4.4 -> A4.5: сводка, метрики, карточки признаков, графики, итоговый payload.")

    add_heading(doc, "Связь A4 -> A5 (оценка риска)", level=2)
    add_bullet(
        doc,
        (
            "Из forecasting в A5 передаются параметры среза: source_tables, selected_district, selected_cause, "
            "selected_object_category, history_window, planning_horizon_days."
        ),
    )
    add_bullet(
        doc,
        (
            "Передача выполняется через _build_decision_support_block() -> build_decision_support_payload() "
            "из app/services/forecast_risk/core.py."
        ),
    )
    add_bullet(
        doc,
        (
            "A5 возвращает risk_prediction/geo_prediction; A4 встраивает это в payload "
            "(risk_prediction, charts['geo'], executive_brief, decision_support_status)."
        ),
    )
    add_bullet(
        doc,
        (
            "Примечание: сами forecast_rows (дневной календарь) не являются прямым входом A5; "
            "A5 пересчитывает территориальный риск из исходных данных с теми же фильтрами и горизонтом."
        ),
    )

    add_heading(doc, "2. Алгоритмы", level=1)

    add_heading(doc, "2.1 Алгоритм отбора источников данных (SourceQueryRegistry)", level=2)
    add_bullet(
        doc,
        (
            "Регистрация источников: SourceQueryBuilder наследует SourceQueryRegistryMixin + SourceQuerySqlMixin; "
            "источники формируются из metadata_items и resolved_columns."
        ),
    )
    add_bullet(
        doc,
        (
            "Приоритет источника данных по таблице: materialized view (mv_forecasting_daily_*) "
            "используется, если существует; иначе прямой source query."
        ),
    )
    add_bullet(
        doc,
        (
            "При конфликте raw/clean таблиц (_canonicalize_source_tables): clean_* имеет приоритет, "
            "raw-дубликат исключается с note."
        ),
    )
    add_bullet(
        doc,
        (
            "При конфликте фильтра и схемы: если выбран фильтр (district/cause/object_category), "
            "но в таблице нет нужной колонки, scope_is_valid=False и таблица исключается из этого запроса."
        ),
    )
    add_par(doc, "Псевдокод построения итогового union-запроса:")
    add_code(
        doc,
        (
            "metadata_items <- collect_metadata(source_tables)\n"
            "view_status <- daily_aggregate_view_status_map(table_names)\n"
            "query_parts <- []\n"
            "for metadata in metadata_items:\n"
            "    table <- metadata.table_name\n"
            "    cols <- metadata.resolved_columns\n"
            "    if view_status[table] == True:\n"
            "        part <- materialized_part_sql(table, cols, filters, min_year)\n"
            "    else:\n"
            "        part <- source_part_sql(table, cols, filters, min_year)\n"
            "    if part is not None:\n"
            "        query_parts.append(part)\n"
            "final_sql <- SELECT fire_date, SUM(incident_count), weighted_avg_temp FROM (UNION ALL query_parts)\n"
            "return execute(final_sql)\n"
        ),
    )

    add_heading(doc, "2.2 Алгоритм агрегации временного ряда", level=2)
    add_bullet(
        doc,
        (
            "SQL-паттерн: агрегирование по дню (GROUP BY fire_date), опциональные фильтры по district/cause/object_category, "
            "ограничение по min_year для history_window."
        ),
    )
    add_bullet(
        doc,
        (
            "Для температуры используется взвешенное объединение по temperature_samples "
            "при слиянии источников."
        ),
    )
    add_bullet(
        doc,
        (
            "Пропуски по датам: после SQL выполняется densification "
            "(_dense_daily_history_from_merged_rows) с заполнением count=0 и avg_temperature=None."
        ),
    )
    add_bullet(
        doc,
        "Оконные функции SQL (OVER, LAG/LEAD) в модуле forecasting не используются.",
    )
    add_bullet(
        doc,
        "Нормализация по населению/площади в forecasting не применяется (используются абсолютные counts по дням).",
    )
    add_par(doc, "Псевдокод (алгоритм 2):")
    add_code(
        doc,
        (
            "INPUT: source_tables, filters, history_window\n"
            "metadata_items, min_year <- load_metadata_and_min_year(source_tables, history_window)\n"
            "try:\n"
            "    rows <- load_daily_history_rows_union(metadata_items, filters, min_year)\n"
            "except:\n"
            "    rows <- []\n"
            "if rows empty:\n"
            "    for each table in metadata_items:\n"
            "        table_rows <- load_daily_history_rows(table, filters, min_year)\n"
            "        if failed: table_rows <- record_fallback(table, filters, min_year)\n"
            "        merge_by_date(merged, table_rows)\n"
            "else:\n"
            "    merge_by_date(merged, rows)\n"
            "history <- dense_daily_history(min_date..max_date, merged)\n"
            "cache(history)\n"
            "return history\n"
        ),
    )

    add_heading(doc, "2.3 Алгоритм bootstrap доверительных интервалов", level=2)
    add_par(
        doc,
        (
            "Фактическое состояние модуля: статистический residual bootstrap (B=1000) не реализован. "
            "Поля bootstrap_mode в payload отражают этапы загрузки данных/блока decision support, "
            "а не bootstrap-ресемплинг."
        ),
    )
    add_par(
        doc,
        (
            "Реально используемый интервал в shaping.py: "
            "lower = max(0, estimate - spread), upper = min(robust_ceiling + spread, estimate + spread), "
            "где spread = max(0.75, volatility * (0.9 + 0.15 * sqrt(step)))."
        ),
    )
    add_par(
        doc,
        (
            "Параметр α для bootstrap-интервала в текущем коде не задаётся, потому что квантильный bootstrap отсутствует."
        ),
    )
    add_par(doc, "Формула интервальной оценки для bootstrap-подхода (целевая, если внедрять):")
    add_par(doc, "I_(1-α)(t+h) = [Q_(α/2)({ŷ*_b(t+h)}), Q_(1-α/2)({ŷ*_b(t+h)})],  b = 1..B.")
    add_par(doc, "Псевдокод (алгоритм 3, целевой bootstrap из ТЗ):")
    add_code(
        doc,
        (
            "INPUT: history, horizon H, bootstrap size B=1000, alpha\n"
            "fit model on history -> y_hat, residuals e_t = y_t - y_hat_t\n"
            "for h in 1..H:\n"
            "    samples <- []\n"
            "    for b in 1..B:\n"
            "        e_star <- sample_with_replacement(residuals)\n"
            "        y_star_pred <- model_point_forecast(h) + e_star\n"
            "        samples.append(y_star_pred)\n"
            "    lower[h] <- quantile(samples, alpha/2)\n"
            "    upper[h] <- quantile(samples, 1 - alpha/2)\n"
            "return lower, upper\n"
        ),
    )

    add_heading(doc, "2.4 Алгоритм оценки качества прогноза", level=2)
    add_bullet(
        doc,
        (
            "Backtesting: _run_scenario_backtesting() делает rolling one-step validation "
            "с минимальным train-окном min_train_days и минимумом 8 окон."
        ),
    )
    add_bullet(
        doc,
        "Метрики: MAE, RMSE, SMAPE (через compute_count_metrics). Сравнение с baseline и delta_vs_baseline.",
    )
    add_bullet(
        doc,
        (
            "Порогов pass/fail по значениям метрик в forecasting нет; "
            "есть пороги готовности по объёму данных (например, <30 дней => quality not ready)."
        ),
    )
    add_bullet(
        doc,
        (
            "Влияние на UI: quality_assessment cards/comparison_rows/dissertation_points; "
            "при недостатке данных показываются предупреждающие сообщения вместо метрик."
        ),
    )
    add_par(doc, "Формулы метрик:")
    add_par(doc, "MAE = (1/n) * Σ_i |y_i - ŷ_i|")
    add_par(doc, "RMSE = sqrt((1/n) * Σ_i (y_i - ŷ_i)^2)")
    add_par(doc, "SMAPE = (100/n) * Σ_i 2|ŷ_i - y_i| / (|y_i| + |ŷ_i|)")

    add_heading(doc, "3. Ключевые технические выводы", level=1)
    add_bullet(
        doc,
        (
            "SourceQueryRegistry реализован как mixin-архитектура (SourceQueryRegistryMixin/SourceQuerySqlMixin), "
            "а не как динамический DI-реестр."
        ),
    )
    add_bullet(
        doc,
        "Forecasting использует SQLAlchemy + SQL text queries и двухуровневый кэш на 120 секунд (inputs + SQL + payload).",
    )
    add_bullet(
        doc,
        (
            "Прогнозный движок в модуле — эвристический сценарный, без statsmodels/ARIMA и без статистического bootstrap."
        ),
    )
    add_bullet(
        doc,
        "Интеграция с A5 выполняется через вызов build_decision_support_payload() и встраивание risk payload в ответ forecasting.",
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"DOCX_CREATED={OUTPUT_PATH.resolve()}")
    print(f"TEXT_QUESTION_MARKS={count_text_question_marks(OUTPUT_PATH)}")


if __name__ == "__main__":
    main()
