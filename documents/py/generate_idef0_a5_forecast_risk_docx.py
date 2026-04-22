from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document


OUTPUT_PATH = Path("documents/IDEF0_A5_Оценка_территориального_пожарного_риска.docx")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_par(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_code(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Consolas"


def count_text_question_marks(path: Path) -> int:
    if not path.exists():
        return -1
    with ZipFile(path, "r") as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    text = "".join(re.findall(r">([^<]*)<", xml))
    return text.count("?")


def add_icom_block(
    doc: Document,
    title: str,
    input_text: str,
    control_text: str,
    output_text: str,
    mechanism_text: str,
) -> None:
    add_heading(doc, title, level=2)
    add_bullet(doc, f"Вход: {input_text}")
    add_bullet(doc, f"Управление: {control_text}")
    add_bullet(doc, f"Выход: {output_text}")
    add_bullet(doc, f"Механизм: {mechanism_text}")


def build_doc() -> Document:
    doc = Document()
    add_heading(doc, "IDEF0 A5: Оценка территориального пожарного риска", level=0)
    add_par(
        doc,
        (
            "Анализ выполнен по модулю app/services/forecast_risk/: core.py, data.py, data_impl.py, geo.py, "
            "profiles.py, profile_resolution.py, scoring_compute.py, scoring_history.py, scoring_ranking.py, "
            "reliability.py, validation.py, presentation.py, notes.py, types.py."
        ),
    )

    add_heading(doc, "1. IDEF0 — декомпозиция A5.1-A5.5", level=1)

    add_icom_block(
        doc,
        "A5.1 Геопривязка событий к территориям",
        (
            "records с полями date, district, territory_label, cause, object_category и опционально latitude/longitude; "
            "planning_horizon_days."
        ),
        (
            "GEO_LOOKBACK_DAYS, MAX_GEO_CHART_POINTS, MAX_GEO_HOTSPOTS; правила _derive_geo_cell_size; "
            "фильтрация записей без координат."
        ),
        (
            "geo_prediction: points, hotspots, districts, legend, top_zone_label, top_explanation; "
            "если координат нет, payload has_coordinates=False."
        ),
        (
            "geo.py::_build_geo_prediction, _derive_geo_cell_size, Counter, math.floor, timedelta; "
            "доминантная административная единица определяется как dominant_district внутри ячейки."
        ),
    )

    add_icom_block(
        doc,
        "A5.2 Расчёт базового риска по территории",
        (
            "исторические записи пожаров по выбранному срезу, агрегированные TerritoryBucket, горизонт planning_horizon_days."
        ),
        (
            "формулы _history_weight_for_record, _normalized_risk_fields, _logistics_fields, _water_fields; "
            "порог LONG_RESPONSE_THRESHOLD_MINUTES; defaults и thresholds профиля."
        ),
        (
            "для каждой территории: risk_score, component_scores, fire_probability, severe_probability, arrival_probability, "
            "water_deficit_probability и объяснимые факторы."
        ),
        (
            "scoring_history.py::_collect_territory_buckets/_horizon_context; "
            "scoring_compute.py::_normalized_risk_fields/_component_score_bundle/_build_territory_rows; "
            "app.services.explainable_logistics.build_explainable_logistics_profile."
        ),
    )

    add_icom_block(
        doc,
        "A5.3 Применение профиля весов",
        (
            "территориальные признаки и component_scores; параметр weight_mode из внешнего запроса."
        ),
        (
            "RISK_WEIGHT_PROFILES: adaptive, expert, calibratable; "
            "ограничения MIN_COMPONENT_WEIGHT, MAX_COMPONENT_WEIGHT; "
            "критерии калибровки MIN_CALIBRATION_WINDOWS, MIN_CALIBRATION_IMPROVEMENT."
        ),
        (
            "resolved profile с component_weights, calibration_summary, comparison с экспертным профилем, "
            "snapshot для UI (available_modes, metric_cards)."
        ),
        (
            "profiles.py::get_risk_weight_profile/resolve_component_weights/build_weight_profile_snapshot; "
            "profile_resolution.py::resolve_weight_profile_for_records и внутренние builders fallback/retained/calibrated."
        ),
    )

    add_icom_block(
        doc,
        "A5.4 Ранжирование территорий",
        (
            "список территорий с итоговым risk_score, history_pressure и component_scores."
        ),
        (
            "правило сортировки (risk_score, history_pressure) по убыванию; "
            "thresholds для risk_class и priority."
        ),
        (
            "упорядоченный список с ranking_position, ranking_gap_to_next, ranking_gap_to_top, ranking_reason, "
            "priority_label и action_hint."
        ),
        (
            "scoring_compute.py::_build_territory_rows; "
            "scoring_ranking.py::_attach_ranking_context/_priority_label/_risk_class/_recommended_action."
        ),
    )

    add_icom_block(
        doc,
        "A5.5 Оценка надёжности и формирование рекомендаций",
        (
            "ранжированные территории, quality_passport, historical_validation metrics_raw."
        ),
        (
            "validation_ready по условию windows_count >= 3; "
            "формулы confidence_norm в reliability.py; "
            "пороговые состояния _ranking_confidence_state."
        ),
        (
            "ranking_confidence_score/label/tone/note на каждую территорию, summary_cards, historical_validation блок, "
            "notes для принятия решения."
        ),
        (
            "reliability.py::_attach_ranking_reliability/_ranking_confidence_state/_build_summary_cards; "
            "validation.py::build_historical_validation_payload; notes.py::_build_risk_notes."
        ),
    )

    add_heading(doc, "2. Связи A4 -> A5 и A5 -> A7", level=1)
    add_heading(doc, "Входящие связи из A4 (прогнозирование)", level=2)
    add_bullet(
        doc,
        (
            "app/services/forecasting/assembly_input.py::_build_decision_support_section вызывает "
            "build_decision_support_payload(source_tables, selected_district, selected_cause, "
            "selected_object_category, history_window, planning_horizon_days)."
        ),
    )
    add_bullet(
        doc,
        (
            "A4 передаёт в A5 контекст среза и горизонт планирования. Прямой массив forecast_rows "
            "в базовом пути A4 -> A5 не передаётся."
        ),
    )
    add_bullet(
        doc,
        (
            "Совместимость с forecast_rows реализована отдельным адаптером "
            "build_risk_forecast_payload: planning_horizon_days = len(forecast_rows)."
        ),
    )

    add_heading(doc, "Исходящие связи в A7 (карта)", level=2)
    add_bullet(
        doc,
        (
            "A5 формирует geo_prediction и geo_summary; в A4 это попадает в charts['geo'] через "
            "forecasting/charts.py::_build_geo_chart."
        ),
    )
    add_bullet(
        doc,
        (
            "app/services/fire_map_service.py использует build_decision_support_payload для "
            "карточек приоритетов и executive brief на карте."
        ),
    )
    add_bullet(
        doc,
        (
            "core/mapping/mixins/analytics_hotspots.py импортирует app.services.forecast_risk.geo::_build_geo_prediction, "
            "то есть пространственная логика A5 переиспользуется в картографическом контуре A7."
        ),
    )

    add_heading(doc, "3. Алгоритмы", level=1)

    add_heading(doc, "3.1 Алгоритм геопривязки (geo.py)", level=2)
    add_bullet(
        doc,
        (
            "Метод сопоставления: грид-привязка по ячейкам, key = (floor(latitude/cell_size), floor(longitude/cell_size)). "
            "Это не point-in-polygon к официальным границам."
        ),
    )
    add_bullet(
        doc,
        (
            "Административная единица берётся как dominant_district из Counter по событиям в ячейке "
            "с fallback значением «Без района»."
        ),
    )
    add_bullet(
        doc,
        (
            "Outliers: явной проверки «вне административных границ» нет. "
            "Записи без координат полностью исключаются из geo_records."
        ),
    )
    add_par(doc, "Псевдокод:")
    add_code(
        doc,
        (
            "geo_records <- filter(records, latitude is not None and longitude is not None)\n"
            "if geo_records is empty:\n"
            "    return payload(has_coordinates=False)\n"
            "cell_size <- derive_cell_size(span(latitudes, longitudes))\n"
            "for record in geo_records:\n"
            "    cell_key <- (floor(lat/cell_size), floor(lon/cell_size))\n"
            "    recency_weight <- max(0.2, 1 - age_days/GEO_LOOKBACK_DAYS)\n"
            "    month_weight <- 1 + 0.35 * future_month_share(record.month)\n"
            "    weekday_weight <- 1 + 0.20 * future_weekday_share(record.weekday)\n"
            "    score <- recency_weight * month_weight * weekday_weight\n"
            "    accumulate cell score, incidents, centroid sums, district/cause counters\n"
            "for each cell:\n"
            "    raw_risk <- score * (1 + ln(1 + incidents) * 0.22) * (0.85 + 0.15 * freshness)\n"
            "rank cells by (raw_risk, incidents)\n"
            "normalize to risk_score 0..100 and build points/hotspots/district summaries\n"
            "return geo_prediction\n"
        ),
    )

    add_heading(doc, "3.2 Алгоритм расчёта риска территории", level=2)
    add_bullet(
        doc,
        (
            "Факторы частоты и сезонности: history_pressure, recency_pressure, seasonal_alignment, heating_pressure, "
            "fire_probability."
        ),
    )
    add_bullet(
        doc,
        (
            "Факторы последствий и типа объектов: severe_rate, casualty_pressure, damage_pressure, risk_category_factor, "
            "dominant_object_category."
        ),
    )
    add_bullet(
        doc,
        (
            "Логистика и доступ: long_arrival_rate, avg_response_pressure, distance_pressure, travel_time_pressure, "
            "service_coverage_gap, service_zone_pressure, water_gap_rate, tanker_dependency."
        ),
    )
    add_par(doc, "Ключевые формулы:")
    add_par(
        doc,
        "component_score_j = 100 * (Σ_k(signal_value_k * signal_weight_k) / Σ_k signal_weight_k)",
    )
    add_par(
        doc,
        "risk_score = clamp(Σ_j(component_score_j * component_weight_j), 1, 99)",
    )
    add_par(
        doc,
        "arrival_probability = clamp(0.24*long_arrival_rate + 0.18*response_pressure + 0.22*travel_time_pressure + 0.22*service_coverage_gap + 0.14*service_zone_pressure, 0.03, 0.98)",
    )
    add_par(
        doc,
        "water_deficit_probability = clamp(0.76*water_gap_rate + 0.24*tanker_dependency, 0.02, 0.99)",
    )
    add_bullet(
        doc,
        (
            "Роль A4: модуль forecasting задаёт фильтры среза и planning_horizon_days. "
            "Числа дневного прогноза не участвуют напрямую в формуле risk_score."
        ),
    )
    add_par(doc, "Псевдокод:")
    add_code(
        doc,
        (
            "horizon <- build_horizon_context(records, planning_horizon_days)\n"
            "territories <- aggregate_records_to_buckets(records, horizon)\n"
            "for each territory bucket:\n"
            "    compute history and seasonal signals\n"
            "    compute consequence signals\n"
            "    compute logistics and water signals\n"
            "    for each risk component j:\n"
            "        component_score_j <- weighted average of component signals\n"
            "    risk_score <- clamp(sum(component_score_j * component_weight_j), 1, 99)\n"
            "sort territories by (risk_score, history_pressure) desc\n"
            "attach ranking context and recommendations\n"
            "return ranked territories\n"
        ),
    )

    add_heading(doc, "3.3 Алгоритм применения профиля весов", level=2)
    add_par(doc, "Базовая формула профиля:")
    add_par(doc, "risk_score = Σ(w_i * f_i) / Σ(w_i)")
    add_bullet(
        doc,
        (
            "В коде пользовательские режимы: adaptive, expert, calibratable "
            "(profiles.py::RISK_WEIGHT_PROFILES)."
        ),
    )
    add_bullet(
        doc,
        (
            "Внутренние кандидаты калибровки: expert, balanced (равномерный), focus_fire_frequency, "
            "focus_consequence_severity, focus_long_arrival_risk, focus_water_supply_deficit и shift_*."
        ),
    )
    add_bullet(
        doc,
        (
            "Пользователь выбирает профиль через параметр weight_mode, который приходит в "
            "build_decision_support_payload и далее в resolve_weight_profile_for_records."
        ),
    )
    add_par(doc, "Псевдокод:")
    add_code(
        doc,
        (
            "requested_profile <- get_risk_weight_profile(weight_mode)\n"
            "if weight_mode == 'expert':\n"
            "    return requested_profile without calibration\n"
            "windows <- build_historical_windows(records, horizon)\n"
            "if windows not ready or count < MIN_CALIBRATION_WINDOWS:\n"
            "    return expert_fallback_profile\n"
            "candidates <- generate_weight_candidates(expert_profile)\n"
            "for each candidate:\n"
            "    metrics <- evaluate_profile_on_windows(candidate)\n"
            "    objective <- 0.24*top1 + 0.31*topk_capture + 0.20*precision + 0.25*ndcg\n"
            "    regularized <- objective - λ * L1_distance(candidate_weights, expert_weights)\n"
            "select best regularized candidate\n"
            "if improvement < MIN_CALIBRATION_IMPROVEMENT:\n"
            "    keep expert weights\n"
            "else:\n"
            "    return calibrated profile\n"
        ),
    )

    add_heading(doc, "3.4 Алгоритм оценки надёжности (reliability.py)", level=2)
    add_bullet(
        doc,
        (
            "Проблема малой выборки решается через fallback-ветку: validation_ready только если "
            "has_metrics=True и windows_count >= 3."
        ),
    )
    add_bullet(
        doc,
        (
            "Классический статистический доверительный интервал вида Wilson или Student в этом модуле не реализован."
        ),
    )
    add_par(doc, "Формулы, реально используемые в коде:")
    add_par(doc, "history_support = min(1, history_count / 8)")
    add_par(
        doc,
        "local_support = clamp(0.58*margin_support + 0.42*history_support, 0.15, 1.0)",
    )
    add_par(
        doc,
        "if validation_ready: confidence_norm = clamp(0.42*passport_score + 0.38*objective_score + 0.20*local_support, 0.18, 0.96)",
    )
    add_par(
        doc,
        "if validation not ready: confidence_norm = clamp(0.67*passport_score + 0.33*local_support, 0.16, 0.88)",
    )
    add_par(doc, "confidence_score = round(100 * confidence_norm)")
    add_bullet(
        doc,
        (
            "Предупреждение пользователю отображается через ranking_confidence_label/tone/note "
            "и статус historical_validation «Пока без проверки»."
        ),
    )
    add_par(doc, "Псевдокод:")
    add_code(
        doc,
        (
            "metrics <- historical_validation.metrics_raw\n"
            "validation_ready <- historical_validation.has_metrics and metrics.windows_count >= 3\n"
            "for each territory in ranking:\n"
            "    history_support <- min(1, history_count / 8)\n"
            "    margin_support <- gap-based score by ranking position\n"
            "    local_support <- clamp(0.58*margin_support + 0.42*history_support, 0.15, 1)\n"
            "    if validation_ready:\n"
            "        conf <- clamp(0.42*passport + 0.38*objective + 0.20*local_support, 0.18, 0.96)\n"
            "    else:\n"
            "        conf <- clamp(0.67*passport + 0.33*local_support, 0.16, 0.88)\n"
            "    apply position bias for top and tail\n"
            "    territory.ranking_confidence_score <- round(conf*100)\n"
            "    territory.ranking_confidence_note <- explanation with validation metrics or fallback note\n"
            "return annotated_ranking\n"
        ),
    )

    add_heading(doc, "3.5 Алгоритм ранжирования территорий", level=2)
    add_bullet(
        doc,
        (
            "Критерий сортировки: сначала risk_score по убыванию, затем history_pressure по убыванию."
        ),
    )
    add_bullet(
        doc,
        (
            "При одинаковом risk_score выше ставится территория с большим history_pressure; "
            "дополнительно считается gap_to_next и gap_to_top."
        ),
    )
    add_bullet(
        doc,
        (
            "Формат выхода: list[RiskScore] с полями label, risk_score, risk_class_label, priority_label, "
            "ranking_position, ranking_gap_to_next, ranking_gap_to_top, recommendations, explanation."
        ),
    )

    add_heading(doc, "4. Важные технические замечания по модулю", level=1)
    add_bullet(
        doc,
        (
            "В geo.py нет точного пространственного point-in-polygon по административным границам. "
            "Геопривязка для зоны риска выполняется через сетку и dominant_district из исходных записей."
        ),
    )
    add_bullet(
        doc,
        (
            "В reliability.py используется эвристическая оценка надёжности ранжирования, "
            "а не классический статистический доверительный интервал."
        ),
    )
    add_bullet(
        doc,
        (
            "Интеграция A4 -> A5 -> A7 реализована рабочими вызовами на уровне сервисов forecasting, "
            "forecast_risk и fire_map_service без дублирования бизнес-логики."
        ),
    )

    return doc


def main() -> None:
    document = build_doc()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"DOCX_CREATED={OUTPUT_PATH.resolve()}")
    print(f"TEXT_QUESTION_MARKS={count_text_question_marks(OUTPUT_PATH)}")


if __name__ == "__main__":
    main()
