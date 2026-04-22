from __future__ import annotations

from datetime import datetime
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"F:\filesFires\base_import")
OUTPUT_PATH = ROOT / "documents" / "Полный_анализ_архитектуры_Fire_Data_Pipeline.docx"


def set_font(run, name: str = "Times New Roman", size: int = 12, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def configure_document(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def add_paragraph(document: Document, text: str, *, bold: bool = False, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_font(run, bold=bold)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading("", level=level)
    run = heading.add_run(text)
    set_font(run, size=14 if level == 1 else 12, bold=True)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        add_paragraph(document, item, style="List Bullet")


def count_question_marks_in_docx_text_nodes(docx_path: Path) -> int:
    with ZipFile(docx_path, "r") as archive:
        xml_bytes = archive.read("word/document.xml")
    root = ET.fromstring(xml_bytes)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return sum((node.text or "").count("?") for node in root.findall(".//w:t", ns))


def add_icom_block(
    doc: Document,
    code: str,
    name: str,
    function_text: str,
    inputs: list[str],
    controls: list[str],
    outputs: list[str],
    mechanisms: list[str],
    links: list[str],
) -> None:
    add_heading(doc, f"{code}. {name}", level=2)
    add_paragraph(doc, "Функция:", bold=True)
    add_bullets(doc, [function_text])
    add_paragraph(doc, "Вход (Input):", bold=True)
    add_bullets(doc, inputs)
    add_paragraph(doc, "Управление (Control):", bold=True)
    add_bullets(doc, controls)
    add_paragraph(doc, "Выход (Output):", bold=True)
    add_bullets(doc, outputs)
    add_paragraph(doc, "Механизм (Mechanism):", bold=True)
    add_bullets(doc, mechanisms)
    add_paragraph(doc, "Связи с другими блоками:", bold=True)
    add_bullets(doc, links)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title_run = title.add_run("Комплексный системный анализ проекта Fire Data Pipeline")
    set_font(title_run, size=16, bold=True)

    add_paragraph(doc, "Дата формирования: " + datetime.now().strftime("%d.%m.%Y %H:%M"))
    add_paragraph(doc, "Основание: анализ модулей app/, core/, config/ текущего проекта.")

    add_heading(doc, "1. IDEF0 — Контекстная диаграмма (A-0)", level=1)
    add_paragraph(
        doc,
        "Единственный блок A-0: «Система поддержки принятия решений в области пожарной безопасности сельских территорий».",
        bold=True,
    )

    add_paragraph(doc, "ВХОД (Input):", bold=True)
    add_bullets(
        doc,
        [
            "Файлы пользователя: XLS/XLSX/CSV (endpoint POST /upload).",
            "Исторические таблицы пожаров в PostgreSQL (включая clean_* таблицы).",
            "Географические и атрибутивные признаки инцидентов: дата, район, причина, категория объекта, координаты, ущерб, логистика прибытия, вода.",
            "Параметры HTTP-запросов из браузера: фильтры, горизонты прогноза, выбор таблиц, feature_columns, k кластеров.",
        ],
    )

    add_paragraph(doc, "УПРАВЛЕНИЕ (Control):", bold=True)
    add_bullets(
        doc,
        [
            "Пороговые правила очистки и профилирования из config/constants.py: NULL_THRESHOLD, DOMINANT_VALUE_THRESHOLD, LOW_VARIANCE_THRESHOLD.",
            "Алгоритмические параметры: FORECASTING_FORECAST_DAY_OPTIONS, CLUSTER_COUNT_OPTIONS, ML_CACHE_LIMIT, модельные пороги и ограничения.",
            "Пользовательские настройки фильтров в API и UI: district/cause/object_category/year/history_window/temperature.",
            "Политики кэширования и инвалидации: TTL и invalidate_runtime_caches/invalidate_table_related_caches.",
            "Правила обязательных доменных колонок (KeepImportantColumnsStep + NatashaColumnMatcher).",
        ],
    )

    add_paragraph(doc, "ВЫХОД (Output):", bold=True)
    add_bullets(
        doc,
        [
            "REST JSON-ответы UTF-8 для dashboard/forecasting/ml/clustering/access_points/tables/ops.",
            "HTML-страницы FastAPI + Jinja2 (дашборд, прогноз, ML, кластеризация, проблемные точки, таблицы, fire-map).",
            "Интерактивные графики Plotly (фигуры в JSON и рендер в браузере).",
            "Интерактивная карта пожаров (fires_map.html) + экспорт spatial analytics (fires_map_analysis.json/.md).",
            "Промежуточные и итоговые артефакты обработки: CSV/XLSX-отчёты, clean_* таблицы в PostgreSQL.",
            "Статусы и логи фоновых задач (pending/running/completed/failed).",
        ],
    )

    add_paragraph(doc, "МЕХАНИЗМ (Mechanism):", bold=True)
    add_bullets(
        doc,
        [
            "Backend: FastAPI + APIRouter, Jinja2Templates, StaticFiles.",
            "Доступ к БД: SQLAlchemy Engine + psycopg2 (PostgreSQL).",
            "Аналитика данных: pandas, numpy.",
            "Математика и модели: scikit-learn, statsmodels (и косвенно scipy через экосистему sklearn).",
            "Визуализация: Plotly (графики), OpenLayers (карта), HTML/JS/CSS фронтенд.",
            "Внутрипроцессный кэш и синхронизация: CopyingTtlCache/CopyingLruCache, RLock.",
            "Инфраструктура фоновых задач: ThreadPoolExecutor + JobStore + polling API.",
        ],
    )

    add_heading(doc, "2. IDEF0 — Диаграмма верхнего уровня (A0)", level=1)
    add_paragraph(
        doc,
        "Система декомпозирована на 8 функциональных блоков A1–A8.",
    )

    add_icom_block(
        doc,
        "A1",
        "Приём и импорт данных",
        "Загрузка пользовательского файла и импорт в реляционную БД.",
        [
            "Multipart UploadFile (XLS/XLSX/CSV).",
            "session_id и job_id из cookie/запроса.",
        ],
        [
            "Ограничения форматов и правила чтения pandas.",
            "Настройки Settings(project_name/output_folder).",
        ],
        [
            "Таблица исходных данных в PostgreSQL (to_sql if_exists=replace).",
            "CSV-копия импортированного набора в result-папке.",
            "Job-логи импорта.",
        ],
        [
            "app/services/pipeline_service.py, ImportDataStep.",
            "pandas.read_excel/read_csv + DataFrame.to_sql.",
        ],
        [
            "Выход A1 (сырая таблица) поступает в A2 для очистки/профилирования.",
            "Выход A1 поступает в A3 для обновления метаданных и инвалидации кэшей.",
            "Выход A1 является источником данных для A4–A8.",
        ],
    )

    add_icom_block(
        doc,
        "A2",
        "Очистка, профилирование и формирование clean-слоя",
        "Оценка качества колонок, защита обязательных признаков и построение clean_* таблиц.",
        [
            "Сырые таблицы из A1.",
            "Пользовательские пороги очистки.",
        ],
        [
            "NULL_THRESHOLD/DOMINANT_VALUE_THRESHOLD/LOW_VARIANCE_THRESHOLD.",
            "Реестр обязательных признаков и правила сопоставления Natasha.",
        ],
        [
            "Профилинговые отчёты CSV/XLSX.",
            "Очищенная таблица clean_<table> и clean_<table>.xlsx.",
            "Сводка кандидатных/защищённых колонок.",
        ],
        [
            "FiresFeatureProfilingStep, KeepImportantColumnsStep, CreateCleanTableStep.",
            "NatashaColumnMatcher, pandas + SQLAlchemy text/DDL.",
        ],
        [
            "Выход A2 обновляет источники для A4–A8 (через prefer_clean=True в выборе таблиц).",
            "Факт создания clean-таблицы инициирует инвалидацию кэшей в A3.",
        ],
    )

    add_icom_block(
        doc,
        "A3",
        "Метаданные, кэш, сессии и управление заданиями",
        "Единый управляющий слой: метаданные БД, кэши сервисов, session/job состояние.",
        [
            "События изменения таблиц (после A1/A2/операций tables).",
            "Запросы на получение metadata/options/status.",
        ],
        [
            "TTL-политики и схемы ключей кэша.",
            "Правила безопасного доступа из многопоточного окружения.",
        ],
        [
            "Актуальные списки таблиц/колонок.",
            "Прогретые кэши и сброшенные stale-значения.",
            "Статусы и логи jobs для клиента.",
        ],
        [
            "app/db_metadata.py, app/cache.py, app/runtime_invalidation.py, app/state.py(JobStore).",
            "Startup warmup_runtime_caches в app/main.py.",
        ],
        [
            "A3 поставляет метаданные и кэшированные данные для A4, A5, A6, A7, A8.",
            "A3 принимает сигналы от A1/A2/A8 (операции с таблицами) для инвалидации.",
        ],
    )

    add_icom_block(
        doc,
        "A4",
        "Дашборд оперативной аналитики",
        "Формирование KPI, агрегатов ущерба/распределений/таймлайнов и управленческого summary.",
        [
            "Таблицы пожаров и фильтры пользователя (table/year/group/horizon).",
            "Метаданные из A3.",
        ],
        [
            "statistics_constants, domain analytics metadata, cache TTL 300/120.",
            "Правила fallback и deferred-mode для страницы.",
        ],
        [
            "Dashboard JSON payload + Plotly фигуры.",
            "HTML-контекст страницы и текстовые brief-выгрузки.",
        ],
        [
            "app/dashboard/service.py + cache.py + SQL-агрегаторы.",
            "Plotly bundle и frontend rendering.",
        ],
        [
            "A4 использует метаданные/кэш A3.",
            "Результаты A4 отображаются пользователю через A8 (веб-слой/API).",
        ],
    )

    add_icom_block(
        doc,
        "A5",
        "Сценарное прогнозирование и поддержка решений",
        "Построение краткосрочного прогноза и территориального риск-ранжирования.",
        [
            "Исторические ряды пожаров и фильтры (district/cause/object/history_window).",
            "Горизонт прогноза и сценарная температура.",
        ],
        [
            "Параметры forecast_days/history_window/weight_mode.",
            "Критерии качества и бизнес-правила приоритизации.",
        ],
        [
            "Прогнозные ряды, метрики качества, рекомендации по территориям.",
            "Geo-hotspot payload для карты и executive brief.",
        ],
        [
            "app/services/forecasting/*, app/services/forecast_risk/*.",
            "SQL-агрегаторы + stats-подходы + двуслойный кэш.",
        ],
        [
            "A5 использует метаданные/кэши A3.",
            "Geo/risk выход A5 поступает в A8 (карта и UI).",
            "Асинхронный режим A5 использует контур jobs из A3.",
        ],
    )

    add_icom_block(
        doc,
        "A6",
        "ML-прогноз и backtesting",
        "Обучение моделей счета/вероятности пожаров, расчёт интервалов и explainability.",
        [
            "Дневной исторический ряд и feature-набор.",
            "Параметры ML-фильтров и горизонта.",
        ],
        [
            "ML-константы обучения/валидации (min history, backtest windows, interval params).",
            "Схема ключей ML cache и лимиты LRU.",
        ],
        [
            "ML-прогноз, backtesting-метрики, importance признаков, интервалы предсказания.",
            "JSON payload и linked backtest job status.",
        ],
        [
            "app/services/ml_model/core.py + training/* + jobs.py.",
            "scikit-learn + statsmodels, LRU/TTL-кэши, ThreadPoolExecutor.",
        ],
        [
            "A6 потребляет данные/метаданные через A3 и forecasting data-layer.",
            "Выход A6 отображается в A8 (страница и API ml-model).",
        ],
    )

    add_icom_block(
        doc,
        "A7",
        "Кластеризация территорий и анализ проблемных точек",
        "Сегментация территорий и ranking точек доступа/риска с объяснением факторов.",
        [
            "Агрегированные признаки территорий/точек.",
            "Параметры k, sample_limit, sampling_strategy, feature_columns.",
        ],
        [
            "Правила отбора признаков, weighting strategy, stability thresholds.",
            "Ограничения минимальной поддержки и размерностей.",
        ],
        [
            "Профили кластеров, диагностические метрики, ranking проблемных точек.",
            "Графики scatter/radar/diagnostics и факторные декомпозиции.",
        ],
        [
            "app/services/clustering/* и app/services/access_points/*.",
            "sklearn алгоритмы + numpy/pandas + кэш + jobs.",
        ],
        [
            "A7 использует A3 для кэша/метаданных и A1/A2 как источники.",
            "Результаты A7 передаются в A8 для визуализации и API.",
        ],
    )

    add_icom_block(
        doc,
        "A8",
        "Веб-представление, API и геовизуализация",
        "Единая точка выдачи результатов: HTML UI, REST API, интерактивная карта.",
        [
            "Payload блоков A4–A7.",
            "Параметры страницы/фильтров и session cookie.",
        ],
        [
            "FastAPI маршруты и контракты ответов.",
            "Cache-Control для статики и versioning ассетов.",
        ],
        [
            "HTML-страницы, JSON API, downloads brief.",
            "fire_map HTML (OpenLayers) и слои hotspot/risk/clusters.",
        ],
        [
            "app/routes/*, app/main.py, page_common.py, fire_map_service.py, core/mapping/*.",
            "Jinja2 + JS/CSS + Plotly + OpenLayers.",
        ],
        [
            "A8 агрегирует и выдаёт наружу результаты A4/A5/A6/A7.",
            "A8 запускает фоновые задачи и polling через A3 (JobStore).",
        ],
    )

    add_paragraph(doc, "Ключевые межблочные связи (стрелки):", bold=True)
    add_bullets(
        doc,
        [
            "A1 → A2: импортированные таблицы становятся входом очистки и профилирования.",
            "A1/A2 → A3: изменение схемы и данных инициирует инвалидацию кэша и обновление metadata.",
            "A3 → A4/A5/A6/A7: метаданные таблиц, кэш-слой и session/job инфраструктура.",
            "A5 → A8: прогноз и территориальный риск передаются в UI/API и карту.",
            "A7 → A8: кластеры и проблемные точки передаются в визуализацию.",
            "A4/A5/A6/A7 → A8: все аналитические payload выдаются через маршруты FastAPI.",
        ],
    )

    add_heading(doc, "3. Архитектурная схема взаимодействия модулей", level=1)
    add_paragraph(doc, "Контур взаимодействий:", bold=True)
    add_bullets(
        doc,
        [
            "Клиент (браузер) ↔ FastAPI роутеры: app/routes/pages.py (HTML) и app/routes/api_*.py (JSON).",
            "Роутеры ↔ сервисы: dashboard, forecasting, ml_model, clustering, access_points, fire_map, table_workflows.",
            "Сервисы ↔ кэш: CopyingTtlCache/CopyingLruCache (локальная память процесса).",
            "Сервисы ↔ PostgreSQL: через config/db.py engine (SQLAlchemy) + SQL text/inspect/to_sql/read_sql.",
            "Сервисы ↔ core/processing: pipeline_service вызывает шаги ImportDataStep, FiresFeatureProfilingStep, KeepImportantColumnsStep, CreateCleanTableStep.",
            "Асинхронные jobs: jobs.py модулей запускают ThreadPoolExecutor и пишут прогресс в JobStore.",
            "JobStore ↔ клиент polling: POST запускает job, GET /api/*-jobs/{job_id} и GET /logs читают статус/логи.",
        ],
    )
    add_paragraph(doc, "Паттерн запроса и ответа:", bold=True)
    add_bullets(
        doc,
        [
            "Страницы открываются в shell/deferred-режиме, затем догружают данные через API.",
            "JSON формируется через utf8_json(..., ensure_ascii=False) с media type application/json; charset=utf-8.",
            "Session context единообразно пробрасывается через run_session_json_action / run_session_analytics_request.",
        ],
    )

    add_heading(doc, "4. Стек технологий с обоснованием", level=1)
    add_paragraph(doc, "FastAPI", bold=True)
    add_bullets(
        doc,
        [
            "Роль: HTTP API + сервер страниц + startup hooks + статика.",
            "Почему: высокая скорость разработки, typing-first, простая интеграция с Pydantic/Starlette и async-сценариями.",
            "Альтернативы: Flask (проще, но меньше встроенных async/typing возможностей), Django (тяжелее для данного сервиса).",
        ],
    )
    add_paragraph(doc, "SQLAlchemy", bold=True)
    add_bullets(
        doc,
        [
            "Роль: единый DB Engine, SQL execution, introspection и транзакционные контексты.",
            "Почему: зрелая ORM/SQL toolkit, поддержка pooling, inspect API, переносимость SQL-слоя.",
            "Альтернативы: psycopg2 raw SQL (больше ручного кода), Django ORM (жёстче связан с Django).",
        ],
    )
    add_paragraph(doc, "pandas", bold=True)
    add_bullets(
        doc,
        [
            "Роль: импорт табличных файлов, профилирование колонок, трансформации и экспорт CSV/XLSX.",
            "Почему: де-факто стандарт для ETL/аналитики в Python, удобный bridge между файлами и SQL.",
            "Альтернативы: polars (быстрее на некоторых задачах, но меньше совместимость со старым кодом), dask (сложнее для текущих объёмов).",
        ],
    )
    add_paragraph(doc, "scikit-learn", bold=True)
    add_bullets(
        doc,
        [
            "Роль: ML-модели и clustering (PoissonRegressor, LogisticRegression, KMeans, AgglomerativeClustering, Birch, PCA).",
            "Почему: стабильный API, богатый набор алгоритмов и встроенных метрик/препроцессинга.",
            "Альтернативы: xgboost/lightgbm (сильные бустинги, но выше сложность интерпретации и интеграции под текущую архитектуру).",
        ],
    )
    add_paragraph(doc, "numpy", bold=True)
    add_bullets(
        doc,
        [
            "Роль: базовая численная математика, векторизация, статистические вычисления, подготовка признаков.",
            "Почему: высокая производительность и совместимость почти со всеми ML/статистическими библиотеками.",
            "Альтернативы: чистый Python (слишком медленно), numba (ускорение, но требует дополнительной сложности поддержки).",
        ],
    )
    add_paragraph(doc, "scipy", bold=True)
    add_bullets(
        doc,
        [
            "Роль: в этом репозитории прямых import scipy нет; библиотека используется косвенно как вычислительная база экосистемы sklearn.",
            "Почему: промышленный набор численных алгоритмов, обычно необходим для расширенной статистики и оптимизации.",
            "Альтернативы: ручная реализация отдельных численных процедур на numpy (дольше и рискованнее по точности).",
        ],
    )
    add_paragraph(doc, "statsmodels", bold=True)
    add_bullets(
        doc,
        [
            "Роль: статистические count-модели в ML блоке (GLM NegativeBinomial).",
            "Почему: прозрачные статистические модели, интерпретируемые параметры, удобна для explainable-аналитики.",
            "Альтернативы: только sklearn (меньше специализированных статистических семейств), PyMC (сильнее для байесовщины, но сложнее эксплуатация).",
        ],
    )
    add_paragraph(doc, "plotly", bold=True)
    add_bullets(
        doc,
        [
            "Роль: интерактивные графики на страницах dashboard/forecasting/ml/clustering/access_points.",
            "Почему: удобный JSON-формат фигур, быстрый рендер в браузере, хорошая связка Python ↔ JS.",
            "Альтернативы: matplotlib (не web-native), ECharts/Highcharts (требуют другой стек интеграции на бэкенде).",
        ],
    )

    add_heading(doc, "5. Общая схема потока данных (от Excel до дашборда)", level=1)
    add_bullets(
        doc,
        [
            "1) Пользователь загружает файл через POST /upload, формат multipart/form-data; файл сохраняется на диск в data/uploads/<session>/<job>/.",
            "2) POST /import_data запускает ImportDataStep: XLS/XLSX читается pandas.read_excel, CSV — pandas.read_csv(encoding='utf-8-sig').",
            "3) Данные становятся pandas DataFrame в памяти процесса.",
            "4) Импортированный DataFrame сохраняется как промежуточный CSV (UTF-8-SIG) и пишется в PostgreSQL через DataFrame.to_sql(...).",
            "5) По запросу очистки POST /run_profiling строится профилинговый отчёт: DataFrame -> <table>_fires_profiling_report.csv/.xlsx.",
            "6) После правил KeepImportantColumnsStep + Natasha создаётся clean_<table> в PostgreSQL (CREATE TABLE AS SELECT ...).",
            "7) Очищенный набор дополнительно выгружается в clean_<table>.xlsx.",
            "8) runtime_invalidation сбрасывает кэши сервисов; db_metadata/dashboard/forecasting/ml/clustering/access_points/fire_map получают свежий контекст.",
            "9) При открытии страницы дашборда (GET /) сервер отдаёт HTML (Jinja2) и shell-контекст; frontend затем запрашивает GET /api/dashboard-data.",
            "10) Dashboard service выполняет SQL-агрегации по выбранным таблицам/фильтрам, строит Python payload + Plotly JSON-фигуры.",
            "11) Payload кэшируется (metadata TTL 300 сек, data TTL 120 сек) и возвращается клиенту как UTF-8 JSON.",
            "12) Браузер рендерит карточки и графики Plotly; пользователь видит итоговую аналитику и может перейти к прогнозу/ML/кластерам/карте.",
        ],
    )

    add_paragraph(doc, "Форматы данных по пути:", bold=True)
    add_bullets(
        doc,
        [
            "Вход: XLS/XLSX/CSV, multipart/form-data.",
            "Промежуточно: pandas DataFrame, Python dict/list payload, in-memory cache objects.",
            "Хранение: PostgreSQL tables, materialized views (для части forecasting SQL).",
            "Артефакты: CSV/XLSX, HTML карты, JSON/Markdown аналитики карты.",
            "Выдача: JSON UTF-8 API, HTML+JS/CSS страницы.",
        ],
    )

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    question_marks = count_question_marks_in_docx_text_nodes(OUTPUT_PATH)
    print(f"DOCX_CREATED={OUTPUT_PATH}")
    print(f"TEXT_QUESTION_MARKS={question_marks}")


if __name__ == "__main__":
    main()

