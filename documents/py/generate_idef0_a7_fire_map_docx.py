from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document


OUTPUT_PATH = Path("documents/IDEF0_A7_Интерактивная_картографическая_визуализация.docx")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_par(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_code(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Consolas"


def count_text_question_marks(path: Path) -> int:
    if not path.exists():
        return -1
    with ZipFile(path, "r") as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    text = "".join(re.findall(r">([^<]*)<", xml))
    return text.count("?")


def add_icom(
    doc: Document,
    title: str,
    input_text: str,
    control_text: str,
    output_text: str,
    mechanism_text: str,
) -> None:
    add_heading(doc, title, level=2)
    add_bullet(doc, f"Вход: {input_text}")
    add_bullet(doc, f"Управление: {control_text}")
    add_bullet(doc, f"Выход: {output_text}")
    add_bullet(doc, f"Механизм: {mechanism_text}")


def build_doc() -> Document:
    doc = Document()
    add_heading(doc, 'IDEF0 A7: "Интерактивная картографическая визуализация"', level=0)
    add_par(
        doc,
        (
            "Анализ выполнен по модулям: app/services/fire_map_service.py, core/mapping/*, "
            "core/processing/steps/create_fire_map.py, core/processing/steps/fire_map_loader.py, "
            "app/routes/pages.py, app/runtime_invalidation.py."
        ),
    )

    add_heading(doc, "1. IDEF0 — декомпозиция A7.1-A7.4", level=1)

    add_icom(
        doc,
        "A7.1 Сбор аналитических данных из сервисов",
        (
            "Выбранная таблица пожаров, строки с координатами/датами/причинами/территориями, "
            "параметры decision support (горизонт, фильтры)."
        ),
        (
            "selected_table, ограничение max_records_per_table, правила маппинга колонок Config, "
            "параметры build_decision_support_payload(..., planning_horizon_days=14)."
        ),
        (
            "Нормализованные spatial records, payload spatial_analytics, executive_brief и risk_prediction "
            "для страницы /fire-map."
        ),
        (
            "get_fire_map_page_context, build_decision_support_payload, load_fire_map_source, "
            "MapCreator._collect_spatial_records, MapCreator._build_spatial_analytics."
        ),
    )

    add_icom(
        doc,
        "A7.2 Построение слоёв карты",
        (
            "Подготовленные точки пожаров (GeoJSON Point), heatmap_points, hotspots, DBSCAN-кластеры, "
            "risk_zones (Polygon), priority_territories."
        ),
        (
            "layer_defaults (incidents/heatmap/hotspots/clusters/risk_zones/priorities), "
            "категории маркеров (deaths/injured/children/evacuated/other), "
            "heatmap radius=20 и blur=26."
        ),
        (
            "Набор GeoJSON-слоёв для клиентского OpenLayers: incidents, heatmap, hotspots, clusters, "
            "risk_zones, priorities."
        ),
        (
            "build_analytics_layer_geojsons, build_tab_script_lines, build_map_layer_script_lines, "
            "OpenLayers (ol.layer.Tile/Vector/Heatmap)."
        ),
    )

    add_icom(
        doc,
        "A7.3 Генерация HTML с встроенным JavaScript",
        (
            "Подготовленные таблицы prepared_tables с geojson + spatial_analytics + center/zoom."
        ),
        (
            "Структура HTML-страницы (head/style/body), порядок рендеринга вкладок и скриптов, "
            "экранирование JSON для JS через _json_for_script."
        ),
        (
            "Готовый файл fires_map.html (и сопутствующие fires_map_analysis.json / .md)."
        ),
        (
            "MapCreatorTemplateMixin._generate_html, generate_html, build_tab_outer_lines, "
            "build_tab_script_lines, MapCreator.create_map."
        ),
    )

    add_icom(
        doc,
        "A7.4 Кэширование и отдача результата",
        (
            "table_name запроса /fire-map/embed и /fire-map, ранее сохраненный HTML, "
            "кэш executive brief + risk prediction."
        ),
        (
            "TTL=300 секунд, ключ cache=normalized_table_name (str.strip()), "
            "централизованная инвалидация после импорта/очистки."
        ),
        (
            "Быстрый возврат HTML карты без повторной генерации; "
            "быстрый возврат контекста страницы с executive brief."
        ),
        (
            "CopyingTtlCache (_FIRE_MAP_CACHE, _FIRE_MAP_BRIEF_CACHE), "
            "clear_fire_map_cache, invalidate_runtime_caches."
        ),
    )

    add_heading(doc, "2. Входящие данные из A5 / A3 / A6", level=1)
    add_bullet(
        doc,
        (
            "Вход из A5 (forecast_risk) — прямой: "
            "app/services/fire_map_service.py -> build_decision_support_payload(...) формирует risk_prediction "
            "и executive brief для страницы карты."
        ),
    )
    add_bullet(
        doc,
        (
            "Вход из A5 (geo) — прямой: core/mapping/mixins/analytics_hotspots.py использует "
            "forecast_risk.geo._build_geo_prediction(...) для hotspot-кандидатов."
        ),
    )
    add_bullet(
        doc,
        (
            "Вход из A3 (clustering) — отдельного вызова app/services/clustering/* нет; "
            "в A7 кластеризация выполняется локально внутри core/mapping "
            "через sklearn.cluster.DBSCAN (независимый контур карты)."
        ),
    )
    add_bullet(
        doc,
        (
            "Вход из A6 (access_points) — отдельного вызова app/services/access_points/* нет; "
            "вместо этого используются признаки доступности из той же таблицы пожаров "
            "(fire_station_distance, report_time, arrival_time) и "
            "app/services/explainable_logistics.build_explainable_logistics_profile."
        ),
    )

    add_heading(doc, "3. Алгоритмы", level=1)

    add_heading(doc, "3.1 Алгоритм построения hotspot-зон", level=2)
    add_bullet(
        doc,
        (
            "Фактический метод для hotspot-центров: grid-based риск-оценка "
            "(forecast_risk.geo._build_geo_prediction), а не чистый server-side KDE."
        ),
    )
    add_bullet(
        doc,
        (
            "Параллельно в карте строится слой тепловой карты (KDE/heatmap) на клиенте OpenLayers "
            "по weighted points: radius=20, blur=26."
        ),
    )
    add_bullet(
        doc,
        (
            "Для кластерных зон используется DBSCAN: eps_km выбирается автоматически "
            "через 70-й перцентиль k-NN-дистанций * 1.15 (нижняя граница 0.9 км), "
            "min_samples = max(4, min(8, round(log2(n+1)+2)))."
        ),
    )
    add_par(doc, "Формула KDE (теоретическая форма для heatmap-слоя):")
    add_par(doc, "K(x) = (1 / (n * h)) * Σ k((x - x_i) / h)")
    add_par(
        doc,
        (
            "В реализации h интерпретируется параметрами визуального сглаживания "
            "radius/blur слоя ol.layer.Heatmap."
        ),
    )
    add_par(doc, "Визуальное представление:")
    add_bullet(doc, "heatmap: полупрозрачный heatmap-слой (интенсивность по weight).")
    add_bullet(doc, "risk_zones: полигоны (окружности) вокруг hotspot/DBSCAN центров.")
    add_bullet(doc, "hotspots/clusters/priorities: точечные векторные слои с popup.")
    add_par(doc, "Псевдокод:")
    add_code(
        doc,
        (
            "INPUT: records(lat, lon, date, severity, district, cause)\n"
            "dated_records <- filter records where date is not null\n"
            "if len(dated_records) >= 3:\n"
            "    geo_prediction <- _build_geo_prediction(dated_records, planning_horizon_days=30)\n"
            "    hotspots <- top geo_prediction.hotspots (<=8)\n"
            "else:\n"
            "    hotspots <- [] and add fallback notes\n"
            "\n"
            "heatmap_points <- normalize(record.weight / max_weight, clamp to [0.08, 1.0])\n"
            "dbscan <- DBSCAN(eps=auto_eps_km(records), min_samples=adaptive_min_samples(n))\n"
            "clusters <- aggregate DBSCAN labels >= 0 to centroids + risk_score\n"
            "\n"
            "risk_zone_candidates <- clusters + non-overlapping hotspots\n"
            "risk_zones <- top candidates with circle polygons\n"
            "if risk_zones is empty:\n"
            "    risk_zones <- fallback zones by priority territory centroids\n"
            "RETURN hotspots, heatmap_points, clusters, risk_zones\n"
        ),
    )

    add_heading(doc, "3.2 Алгоритм слоевой композиции карты (MapCreator + миксины)", level=2)
    add_bullet(
        doc,
        (
            "Паттерн миксинов: MapCreator(MapCreatorUtilityMixin, MapCreatorAnalyticsMixin, "
            "MapCreatorTemplateMixin, MapCreatorExportMixin)."
        ),
    )
    add_bullet(
        doc,
        (
            "Порядок разрешения конфликтов методов: стандартный Python MRO слева направо; "
            "при одинаковом имени берется первый найденный метод."
        ),
    )
    add_bullet(
        doc,
        (
            "Слои: 1) базовый OSM, 2) категории пожаров incidents, "
            "3) heatmap, 4) hotspots, 5) clusters, 6) risk_zones, 7) priorities."
        ),
    )
    add_par(doc, "Псевдокод:")
    add_code(
        doc,
        (
            "INPUT: tables_data\n"
            "prepared_tables <- []\n"
            "for each table in tables_data:\n"
            "    table_data <- _prepare_table_data(df)\n"
            "    # includes geojson incidents + spatial_analytics\n"
            "    if table_data has features:\n"
            "        prepared_tables.append(table_data)\n"
            "\n"
            "html <- generate_html(prepared_tables):\n"
            "    create OL map with OSM base layer\n"
            "    add incidents category vector layers\n"
            "    add analytics layers in order:\n"
            "        heatmap -> hotspots -> clusters -> risk_zones -> priorities\n"
            "    bind popup overlay + click handler\n"
            "    bind layer/category filter checkboxes\n"
            "save html to fires_map.html\n"
            "save analytics json/md\n"
            "RETURN html_path\n"
        ),
    )

    add_heading(doc, "3.3 Алгоритм определения ключа кэша", level=2)
    add_par(doc, "Ключ и правило:")
    add_code(
        doc,
        (
            "normalized_table_name = str(table_name or '').strip()\n"
            "cache_key = normalized_table_name\n"
        ),
    )
    add_bullet(
        doc,
        (
            "Одинаковый принцип для двух кэшей: "
            "_FIRE_MAP_CACHE (HTML) и _FIRE_MAP_BRIEF_CACHE (brief + risk_prediction)."
        ),
    )
    add_bullet(doc, "TTL для обоих кэшей: 300 секунд.")
    add_bullet(
        doc,
        (
            "Почему TTL=300 сек достаточен: генерация карты сравнительно тяжелая (чтение БД + аналитика + HTML), "
            "а оперативность в 5 минут приемлема для веб-аналитики; "
            "после реальных изменений данных кэш сбрасывается принудительно через invalidate_runtime_caches."
        ),
    )

    add_heading(doc, "3.4 Алгоритм шаблонизации (Jinja2 + встроенный JS)", level=2)
    add_bullet(
        doc,
        (
            "Важно: HTML карты в core/mapping строится не через Jinja2, а через Python-генераторы строк "
            "(PAGE_HEAD_LINES/PAGE_STYLE_LINES + script builders)."
        ),
    )
    add_bullet(
        doc,
        (
            "Jinja2 используется на уровне оболочки страницы app/templates/fire_map.html, "
            "которая встраивает итоговую карту через iframe /fire-map/embed."
        ),
    )
    add_bullet(
        doc,
        (
            "Передача данных Python -> JavaScript: _json_for_script() сериализует в JSON "
            "и экранирует <, >, &, U+2028, U+2029."
        ),
    )
    add_bullet(
        doc,
        (
            "Формат для клиентской стороны: GeoJSON FeatureCollection "
            "(Point для incidents/hotspots/clusters/priorities, Polygon для risk_zones) "
            "+ служебные JSON-объекты styles, layer defaults, heatmap config."
        ),
    )

    add_heading(doc, "4. Технические выводы по A7", level=1)
    add_bullet(
        doc,
        (
            "A7 интегрирует прогнозно-рисковую логику A5 напрямую, "
            "но A3/A6 подключены косвенно через данные и локальные алгоритмы карты."
        ),
    )
    add_bullet(
        doc,
        (
            "Контур построения карты полностью самодостаточный: "
            "load_fire_map_source -> MapCreator analytics -> HTML/OpenLayers -> TTL-кэш -> отдача через /fire-map/embed."
        ),
    )
    add_bullet(
        doc,
        (
            "Инвалидация кэша связана с пайплайном импорта/очистки, "
            "поэтому карта не «застаивается» после обновления таблиц."
        ),
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"DOCX_CREATED={OUTPUT_PATH.resolve()}")
    print(f"TEXT_QUESTION_MARKS={count_text_question_marks(OUTPUT_PATH)}")


if __name__ == "__main__":
    main()

