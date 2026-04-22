from __future__ import annotations

from datetime import datetime
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"F:\filesFires\base_import")
OUTPUT_PATH = ROOT / "documents" / "Анализ_pipeline_service_и_core_processing.docx"


def set_font(run, name: str = "Times New Roman", size: int = 12, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def configure_document(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def add_paragraph(document: Document, text: str, *, bold: bool = False, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_font(run, bold=bold)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading("", level=level)
    run = heading.add_run(text)
    set_font(run, size=14 if level == 1 else 12, bold=True)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        add_paragraph(document, item, style="List Bullet")


def count_question_marks_in_docx_text_nodes(docx_path: Path) -> int:
    with ZipFile(docx_path, "r") as archive:
        xml_bytes = archive.read("word/document.xml")
    root = ET.fromstring(xml_bytes)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return sum((node.text or "").count("?") for node in root.findall(".//w:t", ns))


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title_run = title.add_run(
        "Технический анализ app/services/pipeline_service.py и core/processing "
        "проекта Fire Data Pipeline"
    )
    set_font(title_run, size=16, bold=True)

    add_paragraph(doc, "Дата формирования: " + datetime.now().strftime("%d.%m.%Y %H:%M"))
    add_paragraph(doc, "Анализируемые модули:")
    add_bullets(
        doc,
        [
            "app/services/pipeline_service.py",
            "core/processing/pipeline.py",
            "core/processing/steps/import_data.py",
            "core/processing/steps/fires_feature_profiling.py",
            "core/processing/steps/keep_important_columns.py",
            "core/processing/steps/create_clean_table.py",
            "core/processing/steps/column_filter_match.py",
            "app/runtime_invalidation.py",
        ],
    )

    add_heading(doc, "1. Последовательность шагов пайплайна", level=1)
    add_paragraph(
        doc,
        "В приложении фактически используются два связанных конвейера: импорт данных и очистка с профилированием.",
    )
    add_paragraph(doc, "Поток 1: загрузка файла → импорт")
    add_bullets(
        doc,
        [
            "POST /upload → save_uploaded_file: файл сохраняется в data/uploads/<session>/<job>/.",
            "POST /import_data → import_uploaded_data: создаётся Settings(input_file=...), затем вызывается ImportDataStep.run(settings).",
            "После успешного импорта вызывается _invalidate_runtime_caches(...), статус job меняется на completed.",
        ],
    )
    add_paragraph(doc, "Поток 2: очистка и профилирование")
    add_bullets(
        doc,
        [
            "POST /run_profiling → run_profiling_for_table.",
            "Шаг 1: FiresFeatureProfilingStep(settings).run() строит отчёт по качеству колонок.",
            "Шаг 2: KeepImportantColumnsStep().run(...) защищает доменно важные колонки от удаления.",
            "Шаг 3: _load_profile_summary(...) формирует агрегированную сводку для API/UI.",
            "Шаг 4: CreateCleanTableStep().run(...) создаёт clean_<table> в PostgreSQL и выгружает clean_<table>.xlsx.",
            "После успеха также выполняется _invalidate_runtime_caches(...).",
        ],
    )

    add_heading(doc, "2. Базовый класс PipelineStep", level=1)
    add_bullets(
        doc,
        [
            "Класс core.processing.pipeline.PipelineStep задаёт интерфейс: run(self, settings, **kwargs).",
            "Базовая реализация run вызывает NotImplementedError; конкретная логика реализуется в подклассах шагов.",
            "Класс core.processing.pipeline.Pipeline управляет последовательным запуском self.steps.",
            "Pipeline.run() фиксирует elapsed_seconds по каждому шагу и возвращает structured result.",
            "При исключении Pipeline.run() возвращает success=False, failed_step и error, без автоматического rollback всего конвейера.",
        ],
    )

    add_heading(doc, "3. Шаг ImportDataStep", level=1)
    add_bullets(
        doc,
        [
            "Файл: core/processing/steps/import_data.py, класс ImportDataStep(PipelineStep).",
            "Поддерживаемые форматы входа: .xls/.xlsx через pandas.read_excel, .csv через pandas.read_csv(encoding='utf-8-sig').",
            "Неподдерживаемое расширение вызывает ValueError.",
            "Считанный DataFrame сохраняется в CSV: DataFrame.to_csv(..., encoding='utf-8-sig').",
            "Запись в PostgreSQL: DataFrame.to_sql(project_name, conn, if_exists='replace', index=False) внутри engine.begin().",
            "Артефакты шага кэшируются в Settings: _pipeline_source_df и _pipeline_import_csv.",
        ],
    )
    add_paragraph(doc, "Используемые SQLAlchemy-методы: engine.begin() как транзакционный контекст, connection из begin-контекста.")

    add_heading(doc, "4. Шаг CreateCleanTableStep", level=1)
    add_bullets(
        doc,
        [
            "Файл: core/processing/steps/create_clean_table.py, класс CreateCleanTableStep(PipelineStep).",
            "Источник решения об удалении колонок: profile_df с булевым признаком candidate_to_drop.",
            "Булево приведение выполняется через coerce_bool_series(...).",
            "keep_columns = profile_df[~candidate_mask]['column']; removed_columns = profile_df[candidate_mask]['column'].",
            "Создание таблицы в PostgreSQL: DROP TABLE IF EXISTS clean_<table>; CREATE TABLE clean_<table> AS SELECT <keep_columns> FROM <source_table>.",
            "SQL выполняется через conn.execute(sqlalchemy.text(...)) в блоке engine.begin().",
            "Экспорт результата: либо срез in-memory DataFrame, либо pandas.read_sql('SELECT ... FROM clean_<table>', engine), затем to_excel(..., engine='openpyxl').",
        ],
    )
    add_paragraph(
        doc,
        "Важно: этот шаг не выполняет value-level очистку дат и не нормализует значения колонок. "
        "Он фильтрует только набор столбцов на основе profiling/keep-правил.",
    )

    add_heading(doc, "5. Шаг FiresFeatureProfilingStep", level=1)
    add_bullets(
        doc,
        [
            "Файл: core/processing/steps/fires_feature_profiling.py, класс FiresFeatureProfilingStep(PipelineStep).",
            "Вход: DataFrame из памяти (предпочтительно) или SELECT * FROM <table> через pandas.read_sql.",
            "Профилируемые признаки по каждой колонке: dtype, null_ratio, unique_count, unique_ratio, dominant_ratio, variance.",
            "Для строковых колонок используется нормализация: np.char.strip + np.char.lower, плюс учёт MISSING_LIKE_VALUES.",
            "Флаги исключения: drop_null, drop_constant, low_variance, almost_constant.",
            "Итоговый флаг: candidate_to_drop = drop_null OR drop_constant OR low_variance OR almost_constant.",
            "Отчёты сохраняются в два формата: <table>_fires_profiling_report.csv и <table>_fires_profiling_report.xlsx.",
        ],
    )
    add_paragraph(doc, "Ключевые формулы:")
    add_bullets(
        doc,
        [
            "null_ratio = доля пропусков (для строк берётся максимум между isna() и missing-like).",
            "unique_ratio = unique_count / n_rows.",
            "dominant_ratio = max(value_counts(normalize=True)).",
            "low_variance = variance < low_variance_threshold.",
        ],
    )

    add_heading(doc, "6. Важные колонки: KeepImportantColumnsStep и NatashaColumnMatcher", level=1)
    add_bullets(
        doc,
        [
            "Файл шага: core/processing/steps/keep_important_columns.py, класс KeepImportantColumnsStep.",
            "Сопоставление выполняет core/processing/steps/column_filter_match.py, класс NatashaColumnMatcher.",
            "Алгоритм защиты колонки от удаления: mandatory registry → legacy explicit map → keyword rules.",
            "Нормализация имён: regex-очистка, lower-case, токенизация, лемматизация (Natasha: Segmenter, NewsMorphTagger, MorphVocab).",
            "apply_match_results(...) ставит protected_from_drop=True и candidate_to_drop=False для защищённых колонок.",
            "Формируются файлы: *_updated_fires_profiling_report.csv/.xlsx и *_protected_columns_report.csv/.xlsx.",
        ],
    )

    add_heading(doc, "7. Асинхронность и real-time логи", level=1)
    add_bullets(
        doc,
        [
            "Тяжёлые шаги pipeline_service выполняются синхронно в запросе import_data/run_profiling.",
            "Отдельного запуска в фоне через asyncio.to_thread или ThreadPoolExecutor для этих шагов в модуле нет.",
            "Маршрут /upload объявлен async, но импорт и profiling маршруты синхронные.",
            "Real-time логирование реализовано через app.state.job_store и _LiveLogStream.",
            "run_profiling_for_table оборачивает выполнение шагов в redirect_stdout(_LiveLogStream), и каждая строка пишет add_log(...).",
            "Клиент получает обновления через polling endpoint /logs; в import.js частота обновления 2000 мс.",
        ],
    )

    add_heading(doc, "8. Инвалидация кэшей после успешного импорта", level=1)
    add_bullets(
        doc,
        [
            "В import_uploaded_data после успешного ImportDataStep.run вызывается _invalidate_runtime_caches(session_id, job_id).",
            "В run_profiling_for_table после успешного создания clean-таблицы вызывается тот же механизм.",
            "Функция делегирует в app.runtime_invalidation.invalidate_runtime_caches(on_warning=...).",
            "Сбрасываются кэши модулей: db_metadata, dashboard, ml_model, forecasting, clustering, access_points, fire_map.",
            "Если какой-либо invalidator упал, ошибка логируется как warning, основной поток не прерывается.",
        ],
    )

    add_heading(doc, "9. Обработка ошибок и откат", level=1)
    add_bullets(
        doc,
        [
            "run_profiling_for_table перехватывает FileNotFoundError, ValueError и общий Exception; возвращает status='error'.",
            "В finally всегда выполняется log_stream.flush() и mark_job_status(session, job, final_status).",
            "import_uploaded_data в finally очищает текущий файл job_store.clear_current_file(...) и выставляет финальный статус.",
            "SQL-операции внутри engine.begin() используют транзакционный контекст SQLAlchemy: при исключении выполняется rollback транзакции этого шага.",
            "Глобального компенсационного rollback между шагами нет: артефакты предыдущих успешных шагов могут остаться.",
        ],
    )

    add_heading(doc, "Приложение: конкретные pandas и SQLAlchemy методы", level=1)
    add_paragraph(doc, "pandas:")
    add_bullets(
        doc,
        [
            "read_excel, read_csv, read_sql, to_csv, to_excel, to_sql",
            "select_dtypes, isna, value_counts, nunique, var, astype, fillna, to_numeric",
            "sort_values, loc, dropna, to_dict(orient='records')",
        ],
    )
    add_paragraph(doc, "SQLAlchemy:")
    add_bullets(
        doc,
        [
            "config.db.engine как процессный singleton Engine",
            "engine.begin() для транзакционных блоков",
            "sqlalchemy.text(...) + conn.execute(...) для DDL/DML SQL",
        ],
    )
    add_paragraph(doc, "Форматы данных в пайплайне: вход .xls/.xlsx/.csv, промежуточные DataFrame, выход .csv/.xlsx и таблицы PostgreSQL.")

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    question_marks = count_question_marks_in_docx_text_nodes(OUTPUT_PATH)
    print(f"DOCX_CREATED={OUTPUT_PATH}")
    print(f"TEXT_QUESTION_MARKS={question_marks}")


if __name__ == "__main__":
    main()

