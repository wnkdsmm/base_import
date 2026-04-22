from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document


OUTPUT_PATH = Path("documents/IDEF0_A8_Управление_кэшем_и_состоянием_сессий.docx")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_par(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_code(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Consolas"


def count_text_question_marks(path: Path) -> int:
    if not path.exists():
        return -1
    with ZipFile(path, "r") as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    text = "".join(re.findall(r">([^<]*)<", xml))
    return text.count("?")


def add_icom(
    doc: Document,
    title: str,
    input_text: str,
    control_text: str,
    output_text: str,
    mechanism_text: str,
) -> None:
    add_heading(doc, title, level=2)
    add_bullet(doc, f"Вход: {input_text}")
    add_bullet(doc, f"Управление: {control_text}")
    add_bullet(doc, f"Выход: {output_text}")
    add_bullet(doc, f"Механизм: {mechanism_text}")


def build_doc() -> Document:
    doc = Document()
    add_heading(doc, 'IDEF0 A8: "Управление кэшем и состоянием сессий"', level=0)
    add_par(
        doc,
        (
            "Анализ выполнен по модулям: app/cache.py, app/state.py, app/runtime_invalidation.py. "
            "Для контекста потоков данных использованы вызовы из app/routes/api_common.py, app/routes/api_ops.py, "
            "app/services/*/jobs.py и модулей инициализации кэшей."
        ),
    )

    add_heading(doc, "1. IDEF0 — декомпозиция A8.1-A8.4", level=1)

    add_icom(
        doc,
        "A8.1 Управление TTL-кэшем аналитических данных",
        (
            "Ключи запросов аналитики и метаданных, результаты SQL-агрегаций и подготовленные payload."
        ),
        (
            "TTL-правила сервисов, нормализация cache key, копирование и заморозка payload, "
            "правила thread-safety на уровне экземпляра кэша."
        ),
        (
            "Быстрый ответ из RAM-кэша, снижение повторных SQL и пересчета аналитики."
        ),
        (
            "CopyingTtlCache, freeze_mutable_payload, clone_mutable_payload, "
            "build_immutable_payload_ttl_cache."
        ),
    )

    add_icom(
        doc,
        "A8.2 Управление LRU-кэшем ML-артефактов",
        (
            "Ключи ML-запросов и тренировочных артефактов, вычисленные ML payload и промежуточные модели."
        ),
        (
            "Ограничение максимального размера, правило вытеснения Least Recently Used, "
            "политика иммутабельного хранения."
        ),
        (
            "Переиспользование последних ML-результатов и предсказаний, "
            "контроль потребления памяти процесса."
        ),
        (
            "CopyingLruCache, OrderedDict.move_to_end/popitem(last=False), "
            "build_immutable_payload_lru_cache, MLModelCaches.ml_cache."
        ),
    )

    add_icom(
        doc,
        "A8.3 Инвалидация при изменении данных",
        (
            "Событие изменения таблиц после импорта, очистки, удаления, изменения структуры."
        ),
        (
            "Каскадный сценарий invalidation, порядок очистки кэшей, "
            "флаг include_metadata, обработка предупреждений."
        ),
        (
            "Согласованное состояние кэшей метаданных и сервисов, "
            "исключение устаревших аналитических ответов."
        ),
        (
            "invalidate_table_related_caches, invalidate_runtime_caches, "
            "importlib.import_module, таблица _INVALIDATORS."
        ),
    )

    add_icom(
        doc,
        "A8.4 Управление состоянием асинхронных задач (JobStore)",
        (
            "session_id, job_id, параметры фоновых задач, статусы и сообщения прогресса."
        ),
        (
            "Правила жизненного цикла статусов, хранение логов и результата, "
            "политика thread-safety и snapshot-выдачи."
        ),
        (
            "Актуальный статус задачи для клиента, логи в режиме polling, "
            "результат или ошибка в рамках сессии."
        ),
        (
            "JobStore, SessionState, JobState, методы add_log/get_logs/get_job_snapshot/"
            "complete_job/fail_job/mark_job_status."
        ),
    )

    add_heading(doc, "2. Алгоритмы", level=1)

    add_heading(doc, "2.1 Алгоритм TTL-кэша (CopyingTtlCache)", level=2)
    add_bullet(
        doc,
        "Структура хранилища: self._items: dict[K, dict[str, object]] с полями value и expires_at.",
    )
    add_bullet(
        doc,
        (
            "Thread-safety: self._lock = RLock, блокировка на уровне экземпляра кэша; "
            "все операции чтения и изменения self._items выполняются внутри with self._lock."
        ),
    )
    add_bullet(
        doc,
        (
            "Особенность: в set операция _store_value(value) выполняется до входа в lock, "
            "что уменьшает длительность критической секции."
        ),
    )
    add_par(doc, "Псевдокод:")
    add_code(
        doc,
        (
            "CLASS CopyingTtlCache:\n"
            "    state:\n"
            "        ttl_seconds\n"
            "        lock: RLock\n"
            "        items: dict[key -> {value, expires_at}]\n"
            "\n"
            "    GET(key):\n"
            "        now = current_time()\n"
            "        lock.acquire()\n"
            "        item = items.get(key)\n"
            "        if item is None:\n"
            "            lock.release()\n"
            "            return None\n"
            "        if item.expires_at <= now:\n"
            "            items.pop(key, None)\n"
            "            lock.release()\n"
            "            return None\n"
            "        value = load_value(item.value)\n"
            "        lock.release()\n"
            "        return value\n"
            "\n"
            "    SET(key, value):\n"
            "        frozen = store_value(value)\n"
            "        expires = current_time() + ttl_seconds\n"
            "        lock.acquire()\n"
            "        items[key] = {value: frozen, expires_at: expires}\n"
            "        lock.release()\n"
            "        return load_value(frozen)\n"
        ),
    )

    add_heading(doc, "2.2 Алгоритм LRU-кэша (CopyingLruCache)", level=2)
    add_bullet(
        doc,
        "Структура данных: OrderedDict[K, object], где порядок соответствует recency доступа.",
    )
    add_bullet(
        doc,
        (
            "Алгоритм вытеснения: при переполнении max_size вызывается popitem(last=False), "
            "удаляется наименее недавно использованный элемент."
        ),
    )
    add_bullet(
        doc,
        (
            "Сложность: get и set с move_to_end и popitem работают за O(1) амортизированно."
        ),
    )
    add_bullet(
        doc,
        "Thread-safety: RLock вокруг операций get/set/clear/delete.",
    )
    add_par(doc, "Псевдокод:")
    add_code(
        doc,
        (
            "CLASS CopyingLruCache:\n"
            "    state:\n"
            "        max_size\n"
            "        lock: RLock\n"
            "        items: OrderedDict[key -> value]\n"
            "\n"
            "    GET(key):\n"
            "        lock.acquire()\n"
            "        item = items.get(key)\n"
            "        if item is None:\n"
            "            lock.release()\n"
            "            return None\n"
            "        items.move_to_end(key)\n"
            "        value = load_value(item)\n"
            "        lock.release()\n"
            "        return value\n"
            "\n"
            "    SET(key, value):\n"
            "        stored = store_value(value)\n"
            "        lock.acquire()\n"
            "        items[key] = stored\n"
            "        items.move_to_end(key)\n"
            "        while len(items) > max_size:\n"
            "            items.popitem(last=False)\n"
            "        lock.release()\n"
            "        return load_value(stored)\n"
        ),
    )

    add_heading(doc, "2.3 Алгоритм заморозки payload (freeze_mutable_payload)", level=2)
    add_bullet(
        doc,
        (
            "Рекурсивный обход: dict, list, tuple, set, frozenset преобразуются в иммутабельные обертки "
            "_FrozenDict, _FrozenList, _FrozenTuple, _FrozenSet."
        ),
    )
    add_bullet(
        doc,
        (
            "Неизменяемость нужна для двух целей: "
            "исключить случайную модификацию кэшируемого объекта и обеспечить безопасную выдачу "
            "через clone_mutable_payload."
        ),
    )
    add_bullet(
        doc,
        (
            "Для редких пользовательских типов применяется fallback "
            "_FrozenLeaf(copy.deepcopy(value))."
        ),
    )
    add_bullet(
        doc,
        "Сложность: O(n) по числу узлов структуры плюс стоимость deepcopy для нестандартных leaf.",
    )
    add_par(doc, "Псевдокод:")
    add_code(
        doc,
        (
            "FUNCTION freeze_mutable_payload(value):\n"
            "    if value is immutable_leaf_type:\n"
            "        return value\n"
            "    if value is already frozen_wrapper:\n"
            "        return value\n"
            "    if value is dict:\n"
            "        return FrozenDict(tuple((k, freeze_mutable_payload(v)) for each (k,v)))\n"
            "    if value is list:\n"
            "        return FrozenList(tuple(freeze_mutable_payload(x) for x in value))\n"
            "    if value is tuple:\n"
            "        return FrozenTuple(tuple(freeze_mutable_payload(x) for x in value))\n"
            "    if value is set:\n"
            "        return FrozenSet(frozenset(freeze_mutable_payload(x) for x in value))\n"
            "    if value is frozenset:\n"
            "        return FrozenSet(frozenset(freeze_mutable_payload(x) for x in value), preserve_frozenset=True)\n"
            "    return FrozenLeaf(deepcopy(value))\n"
        ),
    )

    add_heading(doc, "2.4 Алгоритм каскадной инвалидации (invalidate_table_related_caches)", level=2)
    add_par(doc, "Сценарий при изменении таблицы X:")
    add_code(
        doc,
        (
            "invalidate_table_related_caches(table_name=X):\n"
            "    db_metadata = import_module('app.db_metadata')\n"
            "    db_metadata.invalidate_db_metadata_cache(table_name=X)\n"
            "    invalidate_service_caches(include_metadata=False)\n"
            "\n"
            "invalidate_service_caches():\n"
            "    invalidate_runtime_caches(include_metadata=False)\n"
            "\n"
            "invalidate_runtime_caches():\n"
            "    for invalidator in _INVALIDATORS in declared order:\n"
            "        skip db_metadata when include_metadata=False\n"
            "        import target module lazily\n"
            "        call clear_* function\n"
        ),
    )
    add_bullet(
        doc,
        (
            "Почему порядок важен: сначала метаданные БД, затем dashboard, затем сервисы. "
            "Dashboard зависит от актуальной сигнатуры таблиц и не должен пересобраться на старом metadata snapshot."
        ),
    )
    add_bullet(
        doc,
        (
            "Lazy import через importlib.import_module уменьшает связанность модулей и "
            "снижает риск circular import между слоями маршрутов, сервисов и кэшей."
        ),
    )

    add_heading(doc, "2.5 Алгоритм управления JobStore", level=2)
    add_bullet(
        doc,
        (
            "Базовый жизненный цикл в проекте: created -> pending -> running -> completed или failed. "
            "Для upload-контура дополнительно используется статус uploaded."
        ),
    )
    add_bullet(
        doc,
        (
            "Логи накапливаются в job.logs через add_log; чтение идет через get_logs и get_job_snapshot. "
            "Клиент читает состояние polling-эндпоинтами: /logs, /api/ml-model-jobs/{job_id}, "
            "/api/clustering-jobs/{job_id}, /api/forecasting-decision-support-jobs/{job_id}."
        ),
    )
    add_bullet(
        doc,
        (
            "Thread-safety: единый RLock в JobStore защищает _sessions, jobs, logs, result и meta. "
            "RLock выбран потому, что внутри методов есть вложенные вызовы (например ensure_session из create_or_reset_job)."
        ),
    )
    add_bullet(
        doc,
        (
            "Результаты и meta замораживаются при записи (freeze_mutable_payload) "
            "и клонируются при выдаче (clone_mutable_payload), чтобы внешние модификации не ломали state."
        ),
    )

    add_heading(doc, "3. Технические детали по модулям", level=1)
    add_heading(doc, "3.1 app/cache.py", level=2)
    add_bullet(
        doc,
        "CopyingTtlCache и CopyingLruCache универсальны и параметризуются storer/loader.",
    )
    add_bullet(
        doc,
        (
            "Параметр skip_freeze переводит storer/loader в identity-функции. "
            "Это полезно для кэширования уже неизменяемых структур."
        ),
    )
    add_bullet(
        doc,
        (
            "Фабрики build_immutable_payload_ttl_cache и build_immutable_payload_lru_cache "
            "закрепляют безопасный паттерн хранения mutable payload."
        ),
    )

    add_heading(doc, "3.2 app/state.py", level=2)
    add_bullet(
        doc,
        "Хранилище сессий: _sessions: dict[session_id -> SessionState], внутри jobs и индексы latest_job_ids.",
    )
    add_bullet(
        doc,
        (
            "Job snapshot содержит logs, result, error_message, meta, created_at, updated_at и "
            "используется как канонический ответ для job status API."
        ),
    )
    add_bullet(
        doc,
        (
            "Метод prune_job_if_idle удаляет только полностью пустые завершенные job, "
            "не затрагивая задачи с результатом, логами или активным статусом."
        ),
    )

    add_heading(doc, "3.3 app/runtime_invalidation.py", level=2)
    add_bullet(
        doc,
        (
            "Таблица _INVALIDATORS задает единый реестр инвалидации: "
            "db_metadata, dashboard, ml_model, forecasting, clustering, access_points, fire_map."
        ),
    )
    add_bullet(
        doc,
        (
            "warmup_runtime_caches прогревает metadata и dashboard при старте приложения, "
            "ошибки прогрева не валят startup и пишутся как warning."
        ),
    )

    add_heading(doc, "4. Привязка к реальным TTL и LRU в проекте", level=1)
    add_bullet(doc, "db_metadata: TTL 60 секунд для имен таблиц и колонок.")
    add_bullet(doc, "dashboard metadata cache: TTL 300 секунд.")
    add_bullet(doc, "dashboard data cache: TTL 120 секунд.")
    add_bullet(doc, "forecasting metadata/base/sql caches: TTL 120 секунд.")
    add_bullet(doc, "fire_map HTML и brief caches: TTL 300 секунд.")
    add_bullet(doc, "ML payload LRU cache: max_size = ML_CACHE_LIMIT = 128.")
    add_bullet(doc, "ML training artifact OrderedDict LRU: лимит 32 артефакта.")

    add_heading(doc, "5. Вывод", level=1)
    add_par(
        doc,
        (
            "A8 реализует единый in-memory контур управления производительностью и состоянием: "
            "безопасные TTL и LRU кэши для аналитики, каскадная инвалидация при изменении данных, "
            "и потокобезопасный JobStore для асинхронных задач с polling-выдачей статусов и логов."
        ),
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"DOCX_CREATED={OUTPUT_PATH.resolve()}")
    print(f"TEXT_QUESTION_MARKS={count_text_question_marks(OUTPUT_PATH)}")


if __name__ == "__main__":
    main()

