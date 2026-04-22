from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document


OUTPUT_PATH = Path("documents/IDEF0_A6_Анализ_доступности_для_пожарных_подразделений.docx")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_par(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_code(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Consolas"


def count_text_question_marks(path: Path) -> int:
    if not path.exists():
        return -1
    with ZipFile(path, "r") as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    text = "".join(re.findall(r">([^<]*)<", xml))
    return text.count("?")


def add_icom(
    doc: Document,
    title: str,
    input_text: str,
    control_text: str,
    output_text: str,
    mechanism_text: str,
) -> None:
    add_heading(doc, title, level=2)
    add_bullet(doc, f"Вход: {input_text}")
    add_bullet(doc, f"Управление: {control_text}")
    add_bullet(doc, f"Выход: {output_text}")
    add_bullet(doc, f"Механизм: {mechanism_text}")


def build_doc() -> Document:
    doc = Document()
    add_heading(doc, "IDEF0 A6: Анализ доступности для пожарных подразделений", level=0)
    add_par(
        doc,
        (
            "Анализ выполнен по модулю app/services/access_points/: core.py, data.py, data_impl.py, point_data.py, "
            "analysis_factors.py, analysis_output.py, analysis_output_context.py, analysis_output_types.py, "
            "analysis_ranking.py, features.py, numeric.py, presentation.py, types.py."
        ),
    )

    add_heading(doc, "1. IDEF0 — декомпозиция A6.1-A6.4", level=1)

    add_icom(
        doc,
        "A6.1 Извлечение и нормализация данных о точках доступа",
        (
            "source_tables, SQL-таблицы пожаров, фильтры district/year, пользовательский набор feature_columns."
        ),
        (
            "карта кандидатов колонок (ADDRESS_*, DISTRICT_*, FIRE_STATION_DISTANCE_*, WATER_SUPPLY_*), "
            "LONG_RESPONSE_THRESHOLD_MINUTES=20.0, MIN_ACCESS_POINT_SUPPORT=3."
        ),
        (
            "нормализованные incident-level записи AccessPointInput, метаданные AccessPointMetadata, "
            "каталоги фильтров (доступные районы/годы), набор point records для последующей агрегации."
        ),
        (
            "data_impl.py::_load_table_metadata/_build_source_sql/_normalize_record/_record_to_access_point_input, "
            "app.db_metadata.get_table_columns_cached, SQLAlchemy engine.connect."
        ),
    )

    add_icom(
        doc,
        "A6.2 Географическая идентификация (адрес→район)",
        (
            "нормализованные записи с address/object/coords/settlement/territory/district."
        ),
        (
            "fallback-правила _resolve_point_identity; нормализация _normalize_match_text; фильтр generic-object токенов."
        ),
        (
            "PointIdentity (point_id, label, entity_type, entity_code, granularity_rank) и группировка в PointBucket."
        ),
        (
            "point_data.py::_resolve_point_identity/_aggregate_point_buckets/_update_point_bucket_from_record, Counter."
        ),
    )

    add_icom(
        doc,
        "A6.3 Расчёт балла доступности",
        (
            "entity_frame с агрегированными долями/средними: distance, response, long_arrival, water, consequences, recurrence, "
            "night, heating, completeness, support_weight."
        ),
        (
            "FACTOR_WEIGHTS, нормировка весов до 94 баллов, UNCERTAINTY_PENALTY_MAX=6.0, "
            "активные признаки selected_features."
        ),
        (
            "для каждой точки: access_score, water_score, severity_score, recurrence_score, data_gap_score, "
            "uncertainty_penalty, total_score, investigation_score, score_decomposition."
        ),
        (
            "analysis_factors.py::_build_access_point_score_series, "
            "analysis_output.py::_build_access_point_score_context/_score_total_and_uncertainty_penalty."
        ),
    )

    add_icom(
        doc,
        "A6.4 Ранжирование и формирование отчёта",
        (
            "scored rows + reason_details + типология + uncertainty flags."
        ),
        (
            "пороги severity: medium>=30, high>=55, critical>=75; "
            "правило сортировки total_score→severity_score→access_score→incident_count→granularity_rank."
        ),
        (
            "ranked points, top_points, summary_cards, score_distribution, reason_breakdown, incomplete_points, notes, charts."
        ),
        (
            "analysis_ranking.py::_build_access_point_rows_from_entity_frame/_select_top_points/_select_incomplete_points, "
            "presentation.py::_build_summary/_build_summary_cards/_build_notes, core.py::get_access_points_data."
        ),
    )

    add_heading(doc, "2. Алгоритмы", level=1)

    add_heading(doc, "2.1 Алгоритм географической идентификации (fallback-цепочка)", level=2)
    add_bullet(
        doc,
        (
            "В требуемой постановке: адрес -> населённый пункт -> территория -> район."
        ),
    )
    add_bullet(
        doc,
        (
            "В фактической реализации между адресом и населённым пунктом есть дополнительные уровни: "
            "объект и точные координаты (point_data.py::_resolve_point_identity)."
        ),
    )
    add_par(doc, "Псевдокод (с условиями перехода):")
    add_code(
        doc,
        (
            "INPUT: record(address, object_name, lat, lon, settlement, territory_label, district)\n"
            "normalize all text fields\n"
            "if address is not empty:\n"
            "    point_id <- 'address:' + normalized_address + optional_object_suffix\n"
            "    return entity_type='Объект/адрес', granularity_rank=5\n"
            "else if meaningful_object_name exists:\n"
            "    point_id <- 'object:' + normalized_object\n"
            "    return entity_type='Объект', granularity_rank=4\n"
            "else if coordinates are valid:\n"
            "    point_id <- 'coords:lat:lon'\n"
            "    return entity_type='Точная локация', granularity_rank=4\n"
            "else if settlement is not empty:                # Попытка 2\n"
            "    point_id <- 'settlement:' + normalized_settlement\n"
            "    return entity_type='Населённый пункт', granularity_rank=3\n"
            "else if territory_label is not empty:           # Попытка 3\n"
            "    point_id <- 'territory:' + normalized_territory\n"
            "    return entity_type='Territory label', granularity_rank=2\n"
            "else if district is not empty:                  # Попытка 4\n"
            "    point_id <- 'district:' + normalized_district\n"
            "    return entity_type='Район', granularity_rank=1\n"
            "else:\n"
            "    return point_id='unknown:unresolved', entity_type='Неуточнённая локация', granularity_rank=0\n"
        ),
    )
    add_bullet(
        doc,
        (
            "При неудаче всех уровней точка остаётся в выборке как unknown:unresolved; "
            "в UI это «Неуточнённая точка», с максимальной неопределённостью."
        ),
    )

    add_heading(doc, "2.2 Алгоритм расчёта балла доступности", level=2)
    add_par(doc, "Входные признаки и диапазоны:")
    add_bullet(doc, "incident_count: целое >= 1")
    add_bullet(doc, "incidents_per_year: >= 0")
    add_bullet(doc, "average_distance_km: >= 0 или NULL")
    add_bullet(doc, "average_response_minutes: >= 0 или NULL")
    add_bullet(doc, "long_arrival_share, no_water_share, severe_share, night_share, heating_share, rural_share: [0,1]")
    add_bullet(doc, "response_coverage_share, distance_coverage_share, water_unknown_share, completeness_share: [0,1]")
    add_bullet(doc, "support_weight: [0.4, 1.0], зависит от n_incidents и MIN_ACCESS_POINT_SUPPORT=3")

    add_par(doc, "Нормализация (в коде):")
    add_bullet(
        doc,
        "Не z-score; используется ratio-нормализация к верхней опоре (max/threshold) с clip [0,1].",
    )
    add_par(doc, "distance_norm = clip(average_distance / max(12, max_distance), 0, 1)")
    add_par(doc, "response_norm = clip(average_response / max(20, max_response), 0, 1)")
    add_par(
        doc,
        "recurrence_factor = clip(0.70 * frequency_norm + 0.30 * incidents_norm, 0, 1)",
    )
    add_par(
        doc,
        "severity_factor = clip(0.58 * severe_share + 0.24 * victim_share + 0.18 * major_damage_share, 0, 1)",
    )

    add_par(doc, "Формула агрегации в итоговый балл:")
    add_par(
        doc,
        "weight_i' = 94 * weight_i / Σ(weight_i по выбранным признакам)",
    )
    add_par(
        doc,
        "contribution_i = weight_i' * factor_i * support_weight",
    )
    add_par(doc, "pure_score = Σ(contribution_i по непенальным факторам)")
    add_par(doc, "uncertainty_penalty = 6 * uncertainty_factor")
    add_par(doc, "total_score = pure_score + uncertainty_penalty")
    add_par(
        doc,
        "investigation_score = min(100, 0.72 * pure_score + 0.28 * (uncertainty_penalty * 100 / 6))",
    )

    add_par(doc, "Штраф неопределённости δ = f(n_incidents):")
    add_par(
        doc,
        "support_weight(n) = 0.4 + 0.6 * min(1, n / 3)",
    )
    add_par(
        doc,
        "uncertainty_factor = 0.35*(1-response_coverage) + 0.30*water_unknown + 0.20*(1-distance_coverage) + 0.15*(1-support_weight(n))",
    )
    add_par(doc, "δ(n) = 6 * uncertainty_factor")

    add_par(doc, "Псевдокод:")
    add_code(
        doc,
        (
            "INPUT: entity_frame, selected_features\n"
            "resolve active feature codes and normalize factor weights to 94 points\n"
            "build base shares and normalized factors in [0,1]\n"
            "support_weight <- 0.4 + 0.6*min(1, incident_count/min_support)\n"
            "for each selected factor i:\n"
            "    contribution_i <- weight_i' * factor_i * support_weight\n"
            "pure_score <- sum(contribution_i)\n"
            "uncertainty_factor <- 0.35*(1-response_coverage) + 0.30*water_unknown + 0.20*(1-distance_coverage) + 0.15*(1-support_weight)\n"
            "uncertainty_penalty <- 6 * uncertainty_factor\n"
            "total_score <- pure_score + uncertainty_penalty\n"
            "investigation_score <- min(100, 0.72*pure_score + 0.28*(uncertainty_penalty*100/6))\n"
            "return scored_row\n"
        ),
    )

    add_heading(doc, "2.3 Алгоритм сглаживания при малом числе инцидентов", level=2)
    add_bullet(
        doc,
        (
            "Реализовано эмпирическое байесовское сглаживание долей (_smooth_share) с глобальным prior по всем точкам."
        ),
    )
    add_par(doc, "Формула:")
    add_par(
        doc,
        "posterior = (successes + prior_mean * α) / (n + α), где α = max(0, min_support - n)",
    )
    add_bullet(
        doc,
        "min_support = 3 (config.constants.MIN_ACCESS_POINT_SUPPORT).",
    )
    add_bullet(
        doc,
        (
            "Если n>=3, то α=0 и используется фактическая доля; если n<3, добавляется приорное «псевдонаблюдение»."
        ),
    )
    add_bullet(
        doc,
        (
            "При n<=0 в реализации возвращается 0.0 (защитное поведение для отсутствующих наблюдений)."
        ),
    )

    add_heading(doc, "2.4 Алгоритм ранжирования", level=2)
    add_bullet(
        doc,
        "Основная метрика сравнения: total_score (по убыванию).",
    )
    add_bullet(
        doc,
        "Tie-breaking: severity_score -> access_score -> incident_count -> granularity_rank (все по убыванию).",
    )
    add_bullet(
        doc,
        (
            "После сортировки назначаются rank и rank_display; точки с неопределённостью дополнительно попадают в "
            "список incomplete_points."
        ),
    )

    add_heading(doc, "3. Технические комментарии", level=1)
    add_bullet(
        doc,
        (
            "Модуль использует in-memory TTL-кэш на 120 секунд (core.py::_ACCESS_POINTS_CACHE), "
            "ключ учитывает таблицу, фильтры и выбранные признаки."
        ),
    )
    add_bullet(
        doc,
        (
            "Географическая идентификация сделана как fallback-приоритизация доступной сущности, а не как geocoder."
        ),
    )
    add_bullet(
        doc,
        (
            "Штраф неопределённости ограничен сверху 6 баллами и интерпретируется как управленческий приоритет верификации."
        ),
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"DOCX_CREATED={OUTPUT_PATH.resolve()}")
    print(f"TEXT_QUESTION_MARKS={count_text_question_marks(OUTPUT_PATH)}")


if __name__ == "__main__":
    main()
