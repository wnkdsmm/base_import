from __future__ import annotations

from datetime import datetime
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"F:\filesFires\base_import")
OUTPUT_PATH = ROOT / "documents" / "Анализ_routes_и_main.docx"


def set_font(run, name: str = "Times New Roman", size: int = 12, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def configure_document(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def add_paragraph(document: Document, text: str, *, bold: bool = False, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_font(run, bold=bold)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading("", level=level)
    run = heading.add_run(text)
    set_font(run, size=14 if level == 1 else 12, bold=True)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        add_paragraph(document, item, style="List Bullet")


def count_question_marks_in_docx_text_nodes(docx_path: Path) -> int:
    with ZipFile(docx_path, "r") as archive:
        xml_bytes = archive.read("word/document.xml")
    root = ET.fromstring(xml_bytes)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return sum((node.text or "").count("?") for node in root.findall(".//w:t", ns))


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title_run = title.add_run("Технический анализ app/routes/ и app/main.py проекта Fire Data Pipeline")
    set_font(title_run, size=16, bold=True)

    add_paragraph(doc, "Дата формирования: " + datetime.now().strftime("%d.%m.%Y %H:%M"))
    add_paragraph(doc, "Область анализа:")
    add_bullets(
        doc,
        [
            "app/main.py",
            "app/routes/api.py",
            "app/routes/api_common.py",
            "app/routes/api_dashboard.py",
            "app/routes/api_forecasting.py",
            "app/routes/api_ml_model.py",
            "app/routes/api_clustering.py",
            "app/routes/api_access_points.py",
            "app/routes/api_tables.py",
            "app/routes/api_column_search.py",
            "app/routes/api_ops.py",
            "app/routes/pages.py",
            "app/routes/page_common.py",
        ],
    )

    add_heading(doc, "1. FastAPI: создание приложения, роутеры, статика", level=1)
    add_bullets(
        doc,
        [
            "В app/main.py используется фабрика create_app() -> FastAPI.",
            "Конструктор: FastAPI(title='Fire Data Pipeline').",
            "Статика монтируется через application.mount('/static', StaticFiles(directory=str(STATIC_DIR)), name='static').",
            "Подключение роутеров: application.include_router(api_router) и application.include_router(pages_router).",
            "Экземпляр приложения создаётся как app = create_app().",
            "Startup-хук объявлен декоратором @app.on_event('startup').",
        ],
    )

    add_heading(doc, "2. Структура API: эндпоинты, методы, назначение", level=1)
    add_paragraph(doc, "Агрегация роутеров выполнена в app/routes/api.py через APIRouter.include_router(...).")

    add_paragraph(doc, "Dashboard:")
    add_bullets(
        doc,
        [
            "GET /api/dashboard-data — получить данные аналитической панели.",
        ],
    )

    add_paragraph(doc, "Forecasting:")
    add_bullets(
        doc,
        [
            "GET /api/forecasting-data — прогнозные данные (опционально с decision support).",
            "GET /api/forecasting-metadata — метаданные и фильтры для страницы прогноза.",
            "POST /api/forecasting-decision-support-jobs — старт фоновой задачи блока рекомендаций.",
            "GET /api/forecasting-decision-support-jobs/{job_id} — статус фоновой задачи.",
        ],
    )

    add_paragraph(doc, "ML:")
    add_bullets(
        doc,
        [
            "GET /api/ml-model-data — данные ML-анализа.",
            "POST /api/ml-model-jobs — запуск фоновой ML-задачи.",
            "GET /api/ml-model-jobs/{job_id} — статус ML-задачи (с linked backtest job).",
        ],
    )

    add_paragraph(doc, "Clustering:")
    add_bullets(
        doc,
        [
            "GET /api/clustering-data — данные кластеризации.",
            "POST /api/clustering-jobs — запуск фоновой clustering-задачи.",
            "GET /api/clustering-jobs/{job_id} — статус clustering-задачи.",
        ],
    )

    add_paragraph(doc, "Access Points:")
    add_bullets(
        doc,
        [
            "GET /api/access-points-data — данные рейтинга проблемных точек.",
        ],
    )

    add_paragraph(doc, "Tables:")
    add_bullets(
        doc,
        [
            "GET /api/tables/{table_name}/page — страница табличных данных (пагинация).",
            "DELETE /api/tables/{table_name} — удаление одной таблицы.",
            "POST /api/tables/delete — массовое удаление таблиц.",
        ],
    )

    add_paragraph(doc, "Column Search:")
    add_bullets(
        doc,
        [
            "GET /api/column-search — поиск колонок по запросу.",
            "POST /api/column-search/preview — предпросмотр модификации таблицы.",
            "POST /api/column-search/create-modify-table — создание/изменение таблицы по выбору колонок.",
        ],
    )

    add_paragraph(doc, "Ops (pipeline/import):")
    add_bullets(
        doc,
        [
            "POST /upload — загрузка входного файла.",
            "POST /import_data — импорт файла в БД.",
            "POST /run_profiling — запуск очистки/профилирования таблицы.",
            "GET /logs — получение логов операции по job_id.",
        ],
    )

    add_paragraph(doc, "Page endpoints (HTML, app/routes/pages.py):")
    add_bullets(
        doc,
        [
            "GET /, /forecasting, /ml-model, /backtesting, /clustering, /access-points, /column-search, /fire-map, /fire-map/embed, /tables, /tables/{table_name}, /select_table.",
            "GET /assets/plotly.js, /favicon.ico, /brief/dashboard.txt, /brief/forecasting.txt.",
        ],
    )

    add_heading(doc, "3. Управление сессиями и хранение состояния (JobStore)", level=1)
    add_bullets(
        doc,
        [
            "Cookie-ключ сессии: SESSION_COOKIE_NAME = 'fire_monitor_session_id' (app/state.py).",
            "ensure_session_id(request) берёт cookie и вызывает job_store.ensure_session(...).",
            "Если cookie нет, создаётся новый session_id (uuid4().hex).",
            "utf8_json(..., session_id=...) устанавливает cookie с параметрами httponly=True, samesite='lax', path='/'.",
            "Состояние хранится in-memory в JobStore._sessions: dict[session_id -> SessionState].",
            "Для потокобезопасности JobStore использует RLock.",
            "На уровне job хранятся: status, logs, result, error_message, meta, timestamps, current_file_path.",
        ],
    )

    add_heading(doc, "4. Паттерн run_session_json_action", level=1)
    add_bullets(
        doc,
        [
            "Функция: run_session_json_action(request, action, status_code=200).",
            "Шаг 1: resolve session_id через ensure_session_id(request).",
            "Шаг 2: выполнение action(session_id) внутри json_action_response.",
            "Шаг 3: возврат Response с JSON UTF-8 и cookie session_id.",
            "json_action_response оборачивает action в try/except и умеет обрабатывать ValueError/Exception через on_value_error/on_exception callbacks.",
            "Если обработчики ошибок не переданы — исключение пробрасывается выше.",
        ],
    )
    add_paragraph(doc, "Формат успешного ответа: произвольный dict payload, сериализованный через json.dumps(..., ensure_ascii=False, default=str).")

    add_heading(doc, "5. Фоновые задачи (jobs): запуск, polling, формат прогресса", level=1)
    add_bullets(
        doc,
        [
            "Запуск jobs в API: POST /api/forecasting-decision-support-jobs, POST /api/ml-model-jobs, POST /api/clustering-jobs.",
            "Проверка статуса: GET /api/...-jobs/{job_id} через job_status_response(...).",
            "job_status_response вычисляет session_id и вызывает loader(session_id=session_id, job_id=job_id).",
            "Если статус payload['status'] == 'missing', HTTP-код ответа = 404, иначе 200.",
            "Фоновые вычисления выполняются в ThreadPoolExecutor (в сервисах jobs.py, max_workers=2).",
            "Статусы в payload: pending, running, completed, failed, missing.",
            "Стандартный payload статуса: job_id, kind, status, logs, result, error_message, meta, reused, is_final.",
            "Для ML добавляется linked payload backtest_job (через build_linked_job_status_payload).",
            "Прогресс хранится в meta: stage_index, stage_label, stage_message; заполняется StageTrackingJobProgressReporter.",
        ],
    )
    add_paragraph(doc, "Для import/profiling (ops) используется job_store + GET /logs, где возвращается {job_id, status, logs}.")

    add_heading(doc, "6. Шаблонизация страниц и lazy-loading", level=1)
    add_bullets(
        doc,
        [
            "Шаблонизатор: Jinja2Templates(directory=str(TEMPLATES_DIR)) в page_common.py.",
            "Рендер через helpers: render_template_page(...) и render_context_page(...).",
            "В контекст всегда добавляется request + версионированные asset-параметры.",
            "Страницы объявлены декораторами @router.get(..., response_class=HTMLResponse).",
            "В pages.py реализован _lazy(module_path, attr): importlib.import_module(...) вызывается при первом обращении к сервису.",
            "resolve_page_mode_context(mode=...):",
            "mode != 'deferred' -> вызывается page_loader (полные данные).",
            "mode == 'deferred' -> вызывается shell_loader (облегчённый контекст, lazy подход к загрузке).",
        ],
    )

    add_heading(doc, "7. HTTP-кэширование статики", level=1)
    add_bullets(
        doc,
        [
            "В page_common.py задан PUBLIC_CACHE_HEADERS = {'Cache-Control': 'public, max-age=86400'}.",
            "cached_text_response(...) и empty_cached_response(...) добавляют этот заголовок.",
            "GET /assets/plotly.js возвращает cached_text_response(..., media_type='application/javascript; charset=utf-8').",
            "GET /favicon.ico возвращает empty_cached_response() с Cache-Control.",
            "Для /static/* применяется StaticFiles; в коде app/main.py нет отдельного middleware, который переопределяет cache headers.",
            "Версионирование ресурсов делается через asset_versions(): вычисляется st_mtime_ns файла.",
            "В шаблонах статика подключается с query-параметром версии v={{ *_version }}, что обеспечивает cache-busting при изменении файлов.",
        ],
    )

    add_heading(doc, "8. Startup hook: прогрев кэша и проверка БД", level=1)
    add_bullets(
        doc,
        [
            "Декоратор: @app.on_event('startup') в app/main.py.",
            "На старте вызывается check_connection() из config.db (тест SELECT 1).",
            "При success печатается [OK] + get_db_info().",
            "Далее запускается фоновый прогрев: asyncio.create_task(asyncio.to_thread(warmup_runtime_caches, ...)).",
            "То есть warmup выполняется в отдельном потоке и не блокирует event loop старта.",
            "При ошибке подключения печатается [ERROR] и подсказка по DATABASE_URL.",
        ],
    )

    add_heading(doc, "9. CORS, UTF-8 JSON и сериализация", level=1)
    add_bullets(
        doc,
        [
            "CORS в анализируемых модулях не настроен: отсутствуют CORSMiddleware и app.add_middleware(...).",
            "UTF-8 JSON формируется централизованно в utf8_json(...): media_type='application/json; charset=utf-8'.",
            "Сериализация выполняется json.dumps(payload, ensure_ascii=False, default=str).",
            "ensure_ascii=False сохраняет кириллицу без escape в ответах.",
            "default=str позволяет сериализовать несериализуемые типы (например datetime) в строку.",
            "Для текстовых download-ответов используется media_type='text/plain; charset=utf-8'.",
            "Ошибки в analytics-эндпоинтах упаковываются в унифицированный envelope: {ok: false, error: {code, message, status_code, error_id, detail (в debug-режиме)}}.",
        ],
    )

    add_heading(doc, "Использованные FastAPI-классы и декораторы", level=1)
    add_bullets(
        doc,
        [
            "Классы: FastAPI, APIRouter, Request, UploadFile, Body, Form, File, Query, StaticFiles.",
            "Классы ответов: Response, HTMLResponse, Jinja2Templates.TemplateResponse.",
            "Основные декораторы: @app.on_event('startup'), @router.get(...), @router.post(...), @router.delete(...).",
            "Паттерны: router composition (include_router), session wrapper (run_session_json_action), analytics wrapper (run_analytics_request / run_session_analytics_request), job status facade (job_status_response).",
        ],
    )

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    question_marks = count_question_marks_in_docx_text_nodes(OUTPUT_PATH)
    print(f"DOCX_CREATED={OUTPUT_PATH}")
    print(f"TEXT_QUESTION_MARKS={question_marks}")


if __name__ == "__main__":
    main()
