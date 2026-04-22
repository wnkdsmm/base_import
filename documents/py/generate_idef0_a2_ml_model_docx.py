from __future__ import annotations

from datetime import datetime
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"F:\filesFires\base_import")
OUTPUT_PATH = ROOT / "documents" / "IDEF0_A2_Машинное_обучение_и_прогнозирование.docx"


def set_font(run, name: str = "Times New Roman", size: int = 12, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def configure_document(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def add_paragraph(document: Document, text: str, *, bold: bool = False, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_font(run, bold=bold)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading("", level=level)
    run = heading.add_run(text)
    set_font(run, size=14 if level == 1 else 12, bold=True)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        add_paragraph(document, item, style="List Bullet")


def add_code_block(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(code)
    set_font(run, name="Consolas", size=10)


def count_question_marks_in_docx_text_nodes(docx_path: Path) -> int:
    with ZipFile(docx_path, "r") as archive:
        xml_bytes = archive.read("word/document.xml")
    root = ET.fromstring(xml_bytes)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return sum((node.text or "").count("?") for node in root.findall(".//w:t", ns))


def add_icom_block(
    document: Document,
    title: str,
    input_items: list[str],
    control_items: list[str],
    output_items: list[str],
    mechanism_items: list[str],
) -> None:
    add_heading(document, title, level=2)
    add_paragraph(document, "Вход:", bold=True)
    add_bullets(document, input_items)
    add_paragraph(document, "Управление:", bold=True)
    add_bullets(document, control_items)
    add_paragraph(document, "Выход:", bold=True)
    add_bullets(document, output_items)
    add_paragraph(document, "Механизм:", bold=True)
    add_bullets(document, mechanism_items)


def section_idef0_a2(document: Document) -> None:
    add_heading(document, 'IDEF0 — Диаграмма A2 "Машинное обучение и прогнозирование"', level=1)

    add_icom_block(
        document,
        "A2.1 Загрузка и агрегация обучающей выборки",
        [
            "Список таблиц-источников (`source_tables`) и выбранное окно истории (`selected_history_window`).",
            "Фильтры пользователя: `cause`, `object_category`, `table_name`, `forecast_days`, `temperature`.",
            "Метаданные таблиц прогнозного контура (`metadata_items`, `resolved_columns`).",
        ],
        [
            "Нормализация фильтров и таблиц: `_canonicalize_source_tables`, `_normalize_filter_value`.",
            "Политика кэша входных данных: TTL 120 секунд (`_ML_FILTER_BUNDLE_CACHE`, `_ML_AGGREGATION_INPUT_CACHE`).",
            "SQL-строители: `_build_daily_history_sql`, `_count_forecasting_records_sql`, `_build_option_catalog_sql`.",
        ],
        [
            "Дневной ряд `daily_history` (дата, число пожаров, средняя температура).",
            "Число записей после фильтров `filtered_records_count`.",
            "Каталоги значений фильтров (`option_catalog`) и preload-заметки.",
        ],
        [
            "`app/services/ml_model/core.py`: `_load_ml_filter_bundle`, `_load_ml_aggregation_inputs`.",
            "`app/services/ml_model/training/data_access.py`: `load_ml_filter_bundle`, `load_ml_aggregation_inputs`.",
            "`app/services/forecasting/sql_payload.py`: SQL-агрегации и merge/union fast path.",
        ],
    )

    add_icom_block(
        document,
        "A2.2 Отбор и конструирование признаков",
        [
            "Ряд `daily_history` из A2.1.",
            "Температурные значения (`avg_temperature`) и календарная ось дат.",
        ],
        [
            "Набор признаков `FEATURE_COLUMNS`: `temp_value`, `weekday`, `month`, `lag_1`, `lag_7`, `lag_14`, `rolling_7`, `rolling_28`, `trend_gap`.",
            "Порог применимости температуры: минимум 30 непустых дней и покрытие >= 20% (`MIN_TEMPERATURE_NON_NULL_DAYS`, `MIN_TEMPERATURE_COVERAGE`).",
            "Минимум строк после feature engineering: `MIN_FEATURE_ROWS=24`.",
        ],
        [
            "Матрица признаков для backtest (`dataset`) с отфильтрованными валидными строками.",
            "Цели: `count` (регрессия) и `event = (count > 0)` (классификация).",
            "Температурная статистика (`monthly`, `overall`, `usable`) и предупреждение о тренде.",
        ],
        [
            "`training_dataset.py`: `_build_history_frame`, `_feature_frame`, `_build_backtest_seed_dataset`, `_prepare_training_dataset`.",
            "`training_temperature.py`: `_fit_temperature_statistics`, `_apply_temperature_statistics`, `_temperature_feature_columns`.",
            "pandas/numpy: `shift`, `rolling.mean`, `get_dummies`, `notna().all(axis=1)`.",
        ],
    )

    add_icom_block(
        document,
        "A2.3 Обучение модели",
        [
            "Подготовленный набор признаков из A2.2.",
            "Выбранный рабочий ключ count-модели (`selected_count_model_key`) из A2.4.",
        ],
        [
            "Гиперпараметры Poisson: `alpha=0.40`, `max_iter=2000`.",
            "Гиперпараметры LogisticRegression: `solver='liblinear'`, `max_iter=500`, `class_weight='balanced'`, `random_state=42`.",
            "Порог пригодности Negative Binomial: `len(y)>=56` и `variance/mean >= 1.35`.",
            "Порог пригодности event-классификатора: минимум по 8 объектов каждого класса (`MIN_EVENT_CLASS_COUNT`).",
        ],
        [
            "Финальная count-модель (`poisson` или `negative_binomial`) на полной истории.",
            "Опциональная event-модель (`logistic_regression`) при валидности классификационного контура.",
            "Артефакты модели с колонками design matrix для инференса.",
        ],
        [
            "`training_models.py`: `_fit_count_model`, `_fit_negative_binomial_model_from_design`, `_fit_event_model_from_design`.",
            "scikit-learn: `PoissonRegressor`, `LogisticRegression`, `ColumnTransformer`, `StandardScaler`, `Pipeline`.",
            "statsmodels: `GLM(..., family=NegativeBinomial(alpha))`.",
        ],
    )

    add_icom_block(
        document,
        "A2.4 Backtest и оценка качества",
        [
            "Исторический frame и seed dataset из A2.2.",
            "Кандидаты методов: `seasonal_baseline`, `heuristic_forecast`, `poisson`, `negative_binomial`.",
        ],
        [
            "Схема: `rolling-origin backtesting (expanding window, lead-time-aware)`.",
            "Ограничения: `MIN_BACKTEST_POINTS=8`, `MAX_BACKTEST_POINTS=45`, `ROLLING_MIN_TRAIN_ROWS=28`.",
            "Правило выбора count-метода: сортировка по `(poisson_deviance, MAE, RMSE, sMAPE)` + tolerance 5% для explainability tie-break.",
            "Классификационный контур события: порог `CLASSIFICATION_THRESHOLD=0.5` и выбор по композитному event score.",
        ],
        [
            "Таблица окон backtest (`BacktestWindowRow`) и итоговая таблица оценки (`BacktestEvaluationRow`).",
            "Метрики качества count: MAE, RMSE, sMAPE, Poisson deviance; классификация: Brier, ROC-AUC, F1, Log-loss.",
            "Выбранный рабочий метод и обоснование (`selected_count_model_reason`).",
        ],
        [
            "`backtesting/training_backtesting_execution.py`: `_run_backtest`, `_select_backtest_count_model`.",
            "`backtesting/training_backtesting_horizons.py`: оценка по горизонтам и сбор payload.",
            "`app/services/model_quality.py`: `compute_count_metrics`, `compute_classification_metrics`.",
        ],
    )

    add_icom_block(
        document,
        "A2.5 Калибровка интервалов предсказания",
        [
            "Пары `(actual, predicted)` из backtest по каждому горизонту.",
            "Ось дат окон backtest для последовательной (time-aware) валидации.",
        ],
        [
            "Уровень интервала: `PREDICTION_INTERVAL_LEVEL=0.8` (alpha=0.2).",
            "Разделение калибровки/оценки: минимум 6/4 окна, доля калибровки 60%.",
            "Схемы валидации: фиксированная хронологическая, блочная, скользящая; приоритет rolling при равной стабильности.",
            "Адаптивные корзины по квантилям прогнозов (`PREDICTION_INTERVAL_TARGET_BINS=3`).",
        ],
        [
            "Калибровка интервала `PredictionIntervalCalibration` + `by_horizon`.",
            "Оценка покрытия интервала и пояснения по схеме валидации.",
            "Монотонно расширенные по горизонту интервалы (`_enforce_monotonic_horizon_interval_calibrations`).",
        ],
        [
            "`training/calibration_compute.py`: оценка квантилей остатков и кандидатов схем.",
            "`training/calibration_output.py`: выбор лучшей схемы и текст объяснения.",
            "`backtesting/training_backtesting_horizons.py` + `forecast_bounds.py`: пересчёт coverage и интервальных границ.",
        ],
    )

    add_icom_block(
        document,
        "A2.6 Формирование и кэширование результата",
        [
            "Артефакты обучения A2.3, метрики backtest A2.4, интервал-калибровка A2.5.",
            "Метаданные фильтров/источников и агрегаты A2.1.",
        ],
        [
            "Ключ верхнего ML-кэша: `(schema_version, table, cause, object_category, temperature, days_ahead, history_window)`.",
            "Лимиты: LRU-кэш payload `ML_CACHE_LIMIT=128`, артефакт-кэш обучения `limit=32`.",
            "Инвалидация: `clear_ml_model_cache()` очищает ML payload, input TTL cache, training artifact cache и forecasting SQL cache.",
        ],
        [
            "Итоговый payload для UI: summary, quality assessment, графики, таблицы прогноза, feature importance, заметки.",
            "Сериализованные backtest/horizon блоки для панели качества.",
            "Кэшированный результат для повторного вызова и фоновых jobs.",
        ],
        [
            "`payloads.py`: `_build_ml_payload`.",
            "`core.py`: `_cache_get/_cache_store`, `get_ml_model_data`.",
            "`training/training.py`: `_training_artifact_cache_store`.",
            "`jobs.py`: `ThreadPoolExecutor`, `start_ml_model_job`, статусные этапы через `job_store`.",
        ],
    )

    add_heading(document, "Стрелки-связи между подпроцессами", level=2)
    add_bullets(
        document,
        [
            "A2.1 -> A2.2: `daily_history`, `filtered_records_count`, `option_catalog`.",
            "A2.2 -> A2.4: `seed.frame` и `seed.dataset` для rolling-origin backtesting.",
            "A2.4 -> A2.3: `selected_count_model_key`, валидность event-контура, правила выбора рабочего метода.",
            "A2.4 -> A2.5: backtest-ряды факта/прогноза для калибровки интервалов и проверки coverage.",
            "A2.3 + A2.5 -> A2.6: финальные модели + интервальная калибровка для генерации future forecast rows.",
            "A2.1 + A2.6: фильтры/метаданные и UI-контекст дополняют итоговый payload.",
            "Техническая оговорка: в коде backtest (A2.4) выполняется до финального обучения модели (A2.3), чтобы сначала выбрать рабочий метод.",
        ],
    )


def section_algorithms(document: Document) -> None:
    add_heading(document, "Алгоритмы", level=1)

    add_heading(document, "1) Алгоритм обучения модели", level=2)
    add_paragraph(document, "Классы и типы моделей:", bold=True)
    add_bullets(
        document,
        [
            "Count-регрессия: `sklearn.linear_model.PoissonRegressor`.",
            "Альтернативный count-кандидат: `statsmodels.api.GLM(..., family=NegativeBinomial)`.",
            "Event-классификация: `sklearn.linear_model.LogisticRegression`.",
        ],
    )
    add_paragraph(document, "Гиперпараметры и подбор:", bold=True)
    add_bullets(
        document,
        [
            "PoissonRegressor: `alpha=0.40`, `max_iter=2000` (фиксированные значения).",
            "LogisticRegression: `solver='liblinear'`, `max_iter=500`, `class_weight='balanced'`, `random_state=42` (фиксированные значения).",
            "Grid/Random/Bayesian search не используется; вместо этого выполняется выбор метода по rolling-origin backtesting.",
            "Для Negative Binomial `alpha` оценивается из дисперсии ряда: `alpha = max((Var(y)-E[y]) / E[y]^2, 1e-4)` с верхней отсечкой 5.0.",
        ],
    )
    add_paragraph(document, "Целевая переменная и признаки:", bold=True)
    add_bullets(
        document,
        [
            "Цель регрессии: `y_count = count` (число пожаров в день).",
            "Цель классификации: `y_event = 1[count > 0]`.",
            "Признаки: `temp_value`, `weekday`, `month`, `lag_1`, `lag_7`, `lag_14`, `rolling_7`, `rolling_28`, `trend_gap`.",
        ],
    )
    add_paragraph(document, "Метрики качества:", bold=True)
    add_bullets(
        document,
        [
            "Count-контур: Poisson deviance, MAE, RMSE, sMAPE.",
            "Event-контур: Brier score, ROC-AUC, F1, Log-loss.",
            "Критично: рабочий метод выбирается по out-of-sample backtest, а не по train-loss.",
        ],
    )
    add_paragraph(document, "Псевдокод:", bold=True)
    add_code_block(
        document,
        """INPUT daily_history, forecast_days, scenario_temperature
frame = prepare_reference_frame(build_history_frame(daily_history))
dataset = build_backtest_seed_dataset(frame)
IF len(dataset) < MIN_FEATURE_ROWS: return empty_result

backtest = run_backtest(frame, dataset, validation_horizon=forecast_days)
IF not backtest.is_ready: return empty_result

selected_count_model_key = backtest.selected_count_model_key
final_frame, final_dataset, temperature_stats = prepare_training_dataset(frame)
feature_columns = temperature_feature_columns(temperature_stats)

IF selected_count_model_key == 'poisson':
    count_model = PoissonRegressor(alpha=0.40, max_iter=2000).fit(X, y_count)
ELIF selected_count_model_key == 'negative_binomial':
    alpha = estimate_alpha_from_dispersion(y_count)
    count_model = GLM(y_count, X, family=NegBin(alpha)).fit()
ELIF selected_count_model_key in {'seasonal_baseline','heuristic_forecast'}:
    count_model = None

IF logistic is validated and classes are sufficient:
    event_model = LogisticRegression(solver='liblinear', max_iter=500,
                                     class_weight='balanced', random_state=42).fit(X, y_event)
ELSE:
    event_model = None

feature_importance = build_feature_importance(count_model or explainable_fallback_model)
forecast_rows = build_future_forecast_rows(..., interval_calibration=backtest.interval_calibration)
RETURN assembled_ml_result""",
    )

    add_heading(document, "2) Алгоритм временного backtest", level=2)
    add_paragraph(document, "Схема разбиения ряда:", bold=True)
    add_bullets(
        document,
        [
            "В коде: `rolling-origin` с `expanding window` и lead-time-aware горизонтами.",
            "Train для окна `k`: все данные до origin `k`, Test: следующие `H` дней (`H=max_horizon_days`).",
            "Валидация ведётся по множеству горизонтов `1..H`, а итоговый выбор — на `validation_horizon_days`.",
        ],
    )
    add_paragraph(document, "Формулы метрик:", bold=True)
    add_bullets(
        document,
        [
            "MAE: `MAE = (1/n) * sum_i |y_i - y_hat_i|`.",
            "RMSE: `RMSE = sqrt((1/n) * sum_i (y_i - y_hat_i)^2)`.",
            "MAPE (классическая): `MAPE = (100/n) * sum_i |(y_i - y_hat_i) / y_i|`, при `y_i != 0`.",
            "В реализации сервиса используется `sMAPE`: `sMAPE = (100/n) * sum_i 2*|y_i-y_hat_i|/(|y_i|+|y_hat_i|)`.",
        ],
    )
    add_paragraph(document, "Критерий приемлемости:", bold=True)
    add_bullets(
        document,
        [
            "Backtest считается валидным, когда `backtest.is_ready=True` и есть минимум `MIN_BACKTEST_POINTS=8` сопоставимых окон.",
            "Модель-кандидат учитывается только при полном покрытии окон (`covered_window_count == window_count`).",
            "Выбор count-метода по правилу: минимизировать `(poisson_deviance, MAE, RMSE, sMAPE)` с `COUNT_MODEL_SELECTION_TOLERANCE=5%`.",
        ],
    )
    add_paragraph(document, "Псевдокод:", bold=True)
    add_code_block(
        document,
        """INPUT history_frame, dataset, validation_horizon_days, max_horizon_days
origin_dates = eligible_origins(dataset, min_train_rows, max_horizon_days)
origin_dates = last(MAX_BACKTEST_POINTS, origin_dates)

FOR origin in origin_dates:
    train = history_frame[:origin]
    future = history_frame[origin : origin + max_horizon_days]
    IF len(train_after_feature_filter) < min_train_rows: continue

    fit candidate models on train
    recursively simulate forecasts for candidates over horizons 1..H
    collect BacktestWindowRow for each horizon

evaluation_data = rows_for(validation_horizon_days)
FOR each candidate with full window coverage:
    metrics[candidate] = {
        MAE, RMSE, sMAPE, PoissonDeviance
    }
selected = select_count_method_with_tolerance(metrics, tolerance=0.05)

IF no valid rows or too few comparable windows: return not_ready
RETURN backtest_success(selected, rows, horizon_summaries, metrics)""",
    )

    add_heading(document, "3) Алгоритм калибровки интервалов предсказания", level=2)
    add_paragraph(document, "Типы интервалов и схемы валидации:", bold=True)
    add_bullets(
        document,
        [
            "Фиксированный хронологический (`fixed_chrono_split`) — базовая reference-схема.",
            "Блочный (`blocked_forward_cv`) — переоценка на блоках поздних окон.",
            "Скользящий (`rolling_split_conformal`) — рекалибровка на каждом следующем окне.",
            "В финальном выборе сравниваются blocked/rolling; fixed используется как эталон сравнения.",
        ],
    )
    add_paragraph(document, "Математика построения:", bold=True)
    add_bullets(
        document,
        [
            "Остатки: `r_i = |y_i - y_hat_i|`.",
            "Квантиль абсолютной ошибки: `delta = Q_level(r)` с рангом `ceil((n+1)*level)`.",
            "Адаптивные интервалы: прогнозы делятся на бин-квантили; для бина берётся собственный `delta_bin`.",
            "Покрытие: `coverage = (1/n) * sum_i 1[y_i in [y_hat_i-delta_i, y_hat_i+delta_i]]`.",
            "Цель: `P(y in [y_hat - delta, y_hat + delta]) >= 1 - alpha`, где `alpha = 1 - level`.",
        ],
    )
    add_paragraph(document, "Стабильность и отбор схемы:", bold=True)
    add_bullets(
        document,
        [
            "Стабильность кандидата: `stability_score = |coverage - level| + std(segment_coverages)`.",
            "Сегменты для stability: 2/3/4 блока в зависимости от числа evaluation окон.",
            "При близких score в сортировке предпочтение отдается rolling-схеме.",
            "После расчёта выполняется монотонное расширение интервала по горизонту (`absolute_error_quantile` не убывает).",
        ],
    )
    add_paragraph(document, "Псевдокод:", bold=True)
    add_code_block(
        document,
        """INPUT actuals, predictions, window_dates, level=0.8
split = split_windows(total_windows, calibration_fraction=0.6,
                      min_calibration=6, min_evaluation=4)
IF split is None:
    calibration = build_calibration_on_all_windows()
    mark coverage_validated=False
    RETURN calibration

fixed = evaluate_fixed_chrono(actuals, predictions, split)
blocked = evaluate_blocked_forward_cv(actuals, predictions, split)
rolling = evaluate_rolling_split(actuals, predictions, split)

candidate_score(c) = |coverage(c)-level| + std(segment_coverages(c))
selected = argmin_{c in [blocked, rolling]} candidate_score(c)

deployment_calibration = build_calibration(prefix=all_windows)
deployment_calibration.method_label = "validated by " + selected.scheme_label
deployment_calibration.validated_coverage = selected.coverage

FOR each horizon h:
    calibration_h = calibrate_horizon(h)
enforce_monotonic_non_decreasing_quantile(calibration_h by horizon)

RETURN selected calibration, coverage notes, by_horizon calibrations""",
    )

    add_heading(document, "4) Алгоритм оценки важности признаков", level=2)
    add_paragraph(document, "Метод:", bold=True)
    add_bullets(
        document,
        [
            "Основной: `sklearn.inspection.permutation_importance`.",
            "Score-функция при пермутации: `neg_mean_absolute_error`.",
            "Fallback: `feature_importances_` или `abs(coef_)` или `abs(params)`.",
        ],
    )
    add_paragraph(document, "Параметры расчёта:", bold=True)
    add_bullets(
        document,
        [
            "Holdout срез: хвост после train split `70/30` (`IMPORTANCE_TRAIN_SPLIT_RATIO=0.7`).",
            "Ограничение объёма: до 120 точек (`IMPORTANCE_MAX_SAMPLE_SIZE=120`).",
            "Повторы перестановки: `PERMUTATION_REPEATS=8`.",
        ],
    )
    add_paragraph(document, "Влияние на UI:", bold=True)
    add_bullets(
        document,
        [
            "Список `feature_importance` попадает в payload и таблицу факторов.",
            "График важности строится в `charts.importance` (`_build_importance_chart`).",
            "Лидирующий фактор выводится в summary (`top_feature_label`) и в блоке качества как факт explainability.",
        ],
    )
    add_paragraph(document, "Псевдокод:", bold=True)
    add_code_block(
        document,
        """INPUT model_bundle, dataset
X = build_design_matrix(dataset, expected_columns=model_bundle.columns)
y = dataset['count']

IF model backend is sklearn and permutation_importance available:
    holdout = tail_after_split(X, y, split_ratio=0.7)
    sample = tail(min(len(holdout), 120))
    perm = permutation_importance(model, sample.X, sample.y,
                                  n_repeats=8,
                                  scoring='neg_mean_absolute_error')
    score_by_column = max(0, perm.importances_mean[j])
ELSE:
    score_by_column = fallback_from_model_attributes(model)

group weekday_* -> weekday
group month_* -> month
total = sum(scores)
importance_share_i = score_i / total

RETURN sorted feature_importance rows for UI""",
    )


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title_run = title.add_run(
        "Анализ app/services/ml_model/\n"
        "IDEF0 A2 «Машинное обучение и прогнозирование» + алгоритмы"
    )
    set_font(title_run, size=16, bold=True)

    add_paragraph(doc, "Дата формирования: " + datetime.now().strftime("%d.%m.%Y %H:%M"))
    add_paragraph(doc, "Анализируемые модули:")
    add_bullets(
        doc,
        [
            "app/services/ml_model/core.py",
            "app/services/ml_model/jobs.py",
            "app/services/ml_model/ml_model_config_types.py",
            "app/services/ml_model/payloads.py",
            "app/services/ml_model/training/data_access.py",
            "app/services/ml_model/training/training.py",
            "app/services/ml_model/training/training_dataset.py",
            "app/services/ml_model/training/training_models.py",
            "app/services/ml_model/training/training_importance.py",
            "app/services/ml_model/training/calibration_compute.py",
            "app/services/ml_model/training/calibration_output.py",
            "app/services/ml_model/training/forecast_bounds.py",
            "app/services/ml_model/training/forecast_intervals.py",
            "app/services/ml_model/backtesting/training_backtesting_execution.py",
            "app/services/ml_model/backtesting/training_backtesting_horizons.py",
            "app/services/ml_model/backtesting/training_backtesting_results.py",
            "app/services/ml_model/backtesting/training_backtesting_events.py",
            "app/services/model_quality.py",
        ],
    )

    section_idef0_a2(doc)
    section_algorithms(doc)
    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    question_marks = count_question_marks_in_docx_text_nodes(OUTPUT_PATH)
    print(f"DOCX_CREATED={OUTPUT_PATH}")
    print(f"TEXT_QUESTION_MARKS={question_marks}")


if __name__ == "__main__":
    main()
