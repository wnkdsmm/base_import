from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"F:\filesFires\base_import")
OUTPUT_PATH = ROOT / "documents" / "ml_model_related_files.docx"


def ru(value: str) -> str:
    return value.encode("ascii").decode("unicode_escape")


def text(*parts: str) -> str:
    return ru("".join(parts))


def configure_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)


def set_font(run, name: str, size: int) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def add_paragraph(document: Document, body: str, *, bold: bool = False, size: int = 12) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(body)
    run.bold = bold
    set_font(run, "Times New Roman", size)


def add_file_list(document: Document, files: list[Path]) -> None:
    for path in files:
        relative = path.relative_to(ROOT)
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(str(relative))
        set_font(run, "Consolas", 10)


def resolve(paths: list[str]) -> list[Path]:
    result: list[Path] = []
    for item in paths:
        path = ROOT / item
        if path.exists():
            result.append(path)
    return result


def main() -> None:
    ml_module = sorted((ROOT / "app" / "services" / "ml_model").rglob("*.py"))
    ml_templates = sorted((ROOT / "app" / "templates" / "includes" / "ml_model").glob("*.html"))

    categories: list[tuple[str, list[Path], str]] = [
        (
            text(
                r"\u041c\u0430\u0440\u0448\u0440\u0443\u0442\u044b \u0438 ",
                r"\u0441\u0435\u0440\u0432\u0435\u0440\u043d\u0430\u044f \u0441\u0431\u043e\u0440\u043a\u0430 ",
                r"\u0441\u0442\u0440\u0430\u043d\u0438\u0446\u044b",
            ),
            resolve(
                [
                    "app/routes/pages.py",
                    "app/routes/page_common.py",
                    "app/routes/api_ml_model.py",
                ]
            ),
            text(
                r"\u0424\u0430\u0439\u043b\u044b, \u043a\u043e\u0442\u043e\u0440\u044b\u0435 ",
                r"\u043e\u0442\u0432\u0435\u0447\u0430\u044e\u0442 \u0437\u0430 URL ",
                r"\u0440\u0430\u0437\u0434\u0435\u043b\u0430 ML-\u043f\u0440\u043e\u0433\u043d\u043e\u0437\u0430, ",
                r"\u0432\u044b\u0434\u0430\u0447\u0443 \u043a\u043e\u043d\u0442\u0435\u043a\u0441\u0442\u0430 ",
                r"\u0441\u0442\u0440\u0430\u043d\u0438\u0446\u044b, API ",
                r"\u0440\u0430\u0441\u0447\u0435\u0442\u0430 \u0438 \u0444\u043e\u043d\u043e\u0432\u044b\u0445 ",
                r"\u0437\u0430\u0434\u0430\u0447.",
            ),
        ),
        (
            text(
                r"\u041c\u043e\u0434\u0443\u043b\u044c ML-\u043f\u0440\u043e\u0433\u043d\u043e\u0437\u0430",
            ),
            ml_module,
            text(
                r"\u041e\u0441\u043d\u043e\u0432\u043d\u0430\u044f ",
                r"\u0441\u0435\u0440\u0432\u0438\u0441\u043d\u0430\u044f \u043b\u043e\u0433\u0438\u043a\u0430 ",
                r"\u043a\u043e\u043b\u0438\u0447\u0435\u0441\u0442\u0432\u0435\u043d\u043d\u043e\u0433\u043e ML-\u043f\u0440\u043e\u0433\u043d\u043e\u0437\u0430: ",
                r"\u043a\u043e\u043d\u0444\u0438\u0433\u0443\u0440\u0430\u0446\u0438\u044f, payload-\u0441\u043b\u043e\u0439, ",
                r"\u043e\u0431\u0443\u0447\u0435\u043d\u0438\u0435, backtesting, ",
                r"\u043a\u0430\u043b\u0438\u0431\u0440\u043e\u0432\u043a\u0430, ",
                r"\u043f\u0440\u0435\u0434\u0441\u0442\u0430\u0432\u043b\u0435\u043d\u0438\u0435 ",
                r"\u0440\u0435\u0437\u0443\u043b\u044c\u0442\u0430\u0442\u043e\u0432 \u0438 \u0434\u0436\u043e\u0431\u044b.",
            ),
        ),
        (
            text(
                r"\u0428\u0430\u0431\u043b\u043e\u043d\u044b \u0438 ",
                r"\u043e\u0431\u0449\u0438\u0439 layout",
            ),
            resolve(
                [
                    "app/templates/base.html",
                    "app/templates/ml_model.html",
                    "app/templates/includes/sidebar_nav.html",
                ]
            )
            + ml_templates,
            text(
                r"\u0428\u0430\u0431\u043b\u043e\u043d\u044b, \u0438\u0437 ",
                r"\u043a\u043e\u0442\u043e\u0440\u044b\u0445 \u0441\u043e\u0431\u0438\u0440\u0430\u0435\u0442\u0441\u044f ",
                r"\u0441\u0442\u0440\u0430\u043d\u0438\u0446\u0430 ML-\u043f\u0440\u043e\u0433\u043d\u043e\u0437\u0430, ",
                r"\u0432\u043a\u043b\u044e\u0447\u0430\u044f hero-\u0431\u043b\u043e\u043a, ",
                r"\u043f\u0430\u043d\u0435\u043b\u044c \u043f\u0430\u0440\u0430\u043c\u0435\u0442\u0440\u043e\u0432, ",
                r"\u0431\u043b\u043e\u043a\u0438 \u043a\u0430\u0447\u0435\u0441\u0442\u0432\u0430 ",
                r"\u0438 \u0441\u0435\u043a\u0446\u0438\u0438 \u043f\u0440\u043e\u0433\u043d\u043e\u0437\u0430.",
            ),
        ),
        (
            text(r"\u0421\u0442\u0438\u043b\u0438 \u0438 JavaScript"),
            resolve(
                [
                    "app/static/css/base.css",
                    "app/static/css/layout.css",
                    "app/static/css/shared-components.css",
                    "app/static/css/analytics.css",
                    "app/static/css/dashboard.css",
                    "app/static/css/ml_model.css",
                    "app/static/js/sidebar.js",
                    "app/static/js/analytics_shared.js",
                    "app/static/js/ml_model.js",
                    "app/static/js/ml_model_api.js",
                    "app/static/js/ml_model_charts.js",
                    "app/static/js/ml_model_events.js",
                    "app/static/js/ml_model_render.js",
                    "app/static/js/ml_model_state.js",
                    "app/static/js/ml_model_ui.js",
                    "app/static/js/shared/api_client.js",
                    "app/static/js/shared/plotly_helpers.js",
                    "app/static/js/shared/state_factory.js",
                    "app/static/js/shared/ui_helpers.js",
                ]
            ),
            text(
                r"\u0421\u0442\u0438\u043b\u0438 \u0438 \u0441\u043a\u0440\u0438\u043f\u0442\u044b, ",
                r"\u043a\u043e\u0442\u043e\u0440\u044b\u0435 \u0437\u0430\u0433\u0440\u0443\u0436\u0430\u044e\u0442\u0441\u044f ",
                r"\u0434\u043b\u044f \u043e\u0442\u0440\u0438\u0441\u043e\u0432\u043a\u0438 ",
                r"\u043f\u0430\u043d\u0435\u043b\u0438 ML-\u043f\u0440\u043e\u0433\u043d\u043e\u0437\u0430, ",
                r"\u0430\u0441\u0438\u043d\u0445\u0440\u043e\u043d\u043d\u043e\u0433\u043e ",
                r"\u043e\u0431\u043d\u043e\u0432\u043b\u0435\u043d\u0438\u044f ",
                r"\u0438 \u0432\u0438\u0437\u0443\u0430\u043b\u0438\u0437\u0430\u0446\u0438\u0438 ",
                r"\u0433\u0440\u0430\u0444\u0438\u043a\u043e\u0432.",
            ),
        ),
        (
            text(
                r"\u0422\u0435\u0441\u0442\u044b, \u043e\u0442\u043d\u043e\u0441\u044f\u0449\u0438\u0435\u0441\u044f ",
                r"\u043a ML-\u043f\u0440\u043e\u0433\u043d\u043e\u0437\u0443",
            ),
            resolve(
                [
                    "tests/ml_model_comparison_model_cases.py",
                    "tests/ml_model_comparison_presentation_cases.py",
                    "tests/ml_model_payload_interval_cases.py",
                    "tests/ml_model_payload_variant_cases.py",
                    "tests/test_ml_model_comparison.py",
                    "tests/test_ml_model_jobs.py",
                    "tests/test_ml_model_payload.py",
                    "tests/test_ml_model_result_types_event_metrics.py",
                    "tests/test_ml_model_shell_context.py",
                    "tests/test_ml_temperature_quality.py",
                ]
            ),
            text(
                r"\u0424\u0430\u0439\u043b\u044b \u0442\u0435\u0441\u0442\u043e\u0432, ",
                r"\u043d\u0430\u043f\u0440\u044f\u043c\u0443\u044e \u043f\u0440\u043e\u0432\u0435\u0440\u044f\u044e\u0449\u0438\u0435 ",
                r"payload-\u0441\u043b\u043e\u0439, backtesting, ",
                r"\u0440\u0430\u0441\u0447\u0435\u0442 \u043c\u0435\u0442\u0440\u0438\u043a, ",
                r"\u0434\u0436\u043e\u0431\u044b \u0438 shell-\u043a\u043e\u043d\u0442\u0435\u043a\u0441\u0442 ",
                r"\u0440\u0430\u0437\u0434\u0435\u043b\u0430 ML-\u043f\u0440\u043e\u0433\u043d\u043e\u0437\u0430.",
            ),
        ),
    ]

    total_files = sum(len(files) for _, files, _ in categories)

    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title_run = title.add_run(
        text(
            r"\u0424\u0430\u0439\u043b\u044b, ",
            r"\u043e\u0442\u043d\u043e\u0441\u044f\u0449\u0438\u0435\u0441\u044f ",
            r"\u043a \u0431\u043b\u043e\u043a\u0443 \u00abML-\u043f\u0440\u043e\u0433\u043d\u043e\u0437\u00bb",
        )
    )
    title_run.bold = True
    set_font(title_run, "Times New Roman", 16)

    add_paragraph(
        document,
        text(
            r"\u041a\u043e\u0440\u043d\u0435\u0432\u043e\u0439 ",
            r"\u043a\u0430\u0442\u0430\u043b\u043e\u0433 ",
            r"\u043f\u0440\u043e\u0435\u043a\u0442\u0430: F:\\filesFires\\base_import",
        ),
    )
    add_paragraph(
        document,
        text(
            r"\u041f\u0440\u0438\u043d\u0446\u0438\u043f \u043e\u0442\u0431\u043e\u0440\u0430: ",
            r"\u0432 \u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442 \u0432\u043a\u043b\u044e\u0447\u0435\u043d\u044b ",
            r"\u043f\u0440\u044f\u043c\u044b\u0435 \u0444\u0430\u0439\u043b\u044b, ",
            r"\u043e\u043f\u0440\u0435\u0434\u0435\u043b\u044f\u044e\u0449\u0438\u0435 ",
            r"\u0441\u0442\u0440\u0430\u043d\u0438\u0446\u0443 ML-\u043f\u0440\u043e\u0433\u043d\u043e\u0437\u0430, ",
            r"\u0435\u0435 API, \u0448\u0430\u0431\u043b\u043e\u043d\u044b, ",
            r"\u0441\u0442\u0438\u043b\u0438, JavaScript, \u0441\u0435\u0440\u0432\u0438\u0441\u043d\u0443\u044e ",
            r"\u043b\u043e\u0433\u0438\u043a\u0443 \u043e\u0431\u0443\u0447\u0435\u043d\u0438\u044f, ",
            r"backtesting \u0438 \u0432\u044b\u0434\u0430\u0447\u0438 ",
            r"\u043f\u0440\u043e\u0433\u043d\u043e\u0437\u0430. ",
            r"\u041e\u0431\u0449\u0435\u0441\u0438\u0441\u0442\u0435\u043c\u043d\u044b\u0435 ",
            r"\u0444\u0430\u0439\u043b\u044b \u0431\u0435\u0437 \u043f\u0440\u044f\u043c\u043e\u0439 ",
            r"\u0441\u0432\u044f\u0437\u0438 \u0441 \u0431\u043b\u043e\u043a\u043e\u043c ",
            r"\u043d\u0435 \u0432\u043a\u043b\u044e\u0447\u0430\u043b\u0438\u0441\u044c.",
        ),
    )
    add_paragraph(
        document,
        text(
            r"\u0418\u0442\u043e\u0433\u043e \u043e\u0442\u043e\u0431\u0440\u0430\u043d\u043e ",
            str(total_files),
            r" \u0444\u0430\u0439\u043b\u043e\u0432.",
        ),
    )

    for heading, files, description in categories:
        add_paragraph(document, heading, bold=True)
        add_paragraph(document, description)
        add_paragraph(
            document,
            text(r"\u041a\u043e\u043b\u0438\u0447\u0435\u0441\u0442\u0432\u043e \u0444\u0430\u0439\u043b\u043e\u0432: ", str(len(files))),
        )
        add_file_list(document, files)

    note = document.add_paragraph()
    note_head = note.add_run(text(r"\u041f\u0440\u0438\u043c\u0435\u0447\u0430\u043d\u0438\u0435. "))
    note_head.bold = True
    set_font(note_head, "Times New Roman", 12)
    note_text = note.add_run(
        text(
            r"\u0415\u0441\u043b\u0438 \u043d\u0443\u0436\u043d\u043e, \u043d\u0430 ",
            r"\u043e\u0441\u043d\u043e\u0432\u0435 \u044d\u0442\u043e\u0433\u043e \u0441\u043f\u0438\u0441\u043a\u0430 ",
            r"\u043c\u043e\u0436\u043d\u043e \u043e\u0442\u0434\u0435\u043b\u044c\u043d\u043e ",
            r"\u0441\u043e\u0441\u0442\u0430\u0432\u0438\u0442\u044c \u0441\u0445\u0435\u043c\u0443 ",
            r"\u0432\u0437\u0430\u0438\u043c\u043e\u0434\u0435\u0439\u0441\u0442\u0432\u0438\u044f ",
            r"\u0444\u0430\u0439\u043b\u043e\u0432 \u0431\u043b\u043e\u043a\u0430 ",
            r"\u00abML-\u043f\u0440\u043e\u0433\u043d\u043e\u0437\u00bb ",
            r"\u0434\u043b\u044f \u0434\u0438\u0441\u0441\u0435\u0440\u0442\u0430\u0446\u0438\u0438.",
        )
    )
    set_font(note_text, "Times New Roman", 12)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
