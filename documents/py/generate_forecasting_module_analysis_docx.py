from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / "forecasting_module_analysis.docx"


def set_font(run, *, size: float = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def configure_styles(document: Document) -> None:
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)


def add_paragraph(document: Document, text: str, *, bold: bool = False, align=None) -> None:
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, bold=bold)
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_after = Pt(6)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_font(run)
    paragraph.paragraph_format.space_after = Pt(3)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_font(run, size=14 if level == 1 else 12, bold=True)
    paragraph.paragraph_format.space_after = Pt(6)


def build_document() -> Document:
    document = Document()
    configure_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        'Технический анализ модуля app/services/forecasting/ проекта "Fire Data Pipeline"'
    )
    set_font(run, size=14, bold=True)

    add_paragraph(
        document,
        "Ниже приведён анализ фактической реализации модуля прогнозирования пожаров по данным исходного кода. "
        "В документе сознательно отделены реальные механизмы, которые есть в текущей версии проекта, от "
        "ожидаемых, но отсутствующих механизмов. Это важно для корректного технического описания в магистерской диссертации.",
    )

    add_heading(document, "1. Метод прогнозирования", level=1)
    add_paragraph(
        document,
        "Основной расчёт прогноза реализован функцией "
        "`app/services/forecasting/shaping.py::_build_forecast_rows`. "
        "В данном модуле не используется готовая статистическая модель из sklearn или statsmodels. "
        "Вместо этого применена собственная эвристико-статистическая модель сценарного прогноза по дневному временному ряду пожаров.",
    )
    add_bullet(
        document,
        "Прогнозируемая величина: ожидаемое количество пожаров в день (`forecast_value`).",
    )
    add_bullet(
        document,
        "Дополнительная величина: вероятность наличия хотя бы одного пожара в день (`fire_probability`).",
    )
    add_bullet(
        document,
        "Горизонт прогноза задаётся пользователем через `forecast_days`; допустимые значения берутся из "
        "`config/constants.py::FORECASTING_FORECAST_DAY_OPTIONS = [7, 14, 30, 60]`.",
    )
    add_bullet(
        document,
        "Единица измерения прогноза: количество пожаров в сутки по выбранному срезу данных.",
    )
    add_paragraph(
        document,
        "Расчёт строится на комбинации нескольких компонентов: недавнего среднего уровня ряда, дневной сезонности "
        "по дням недели, месячной сезонности, ограниченного трендового поправочного коэффициента и температурного эффекта. "
        "В `_build_forecast_history_stats` вычисляются `overall_average`, `recent_average`, `very_recent_average`, "
        "`trend_ratio`, `volatility`, `robust_ceiling`. Далее `_build_weekday_forecast_factors` и "
        "`_build_month_forecast_factors` формируют сглаженные сезонные множители.",
    )
    add_paragraph(
        document,
        "Базовая ожидаемая интенсивность для дня строится как `usual_for_day = base_recent_level * seasonal_factor`, "
        "где `seasonal_factor = weekday_factor + month_factor - 1`. Затем добавляются `trend_effect` и "
        "`temperature_effect`. Итоговый прогноз ограничивается сверху через `robust_ceiling`, чтобы модель "
        "не выдавала нереалистично большие значения на редких всплесках.",
    )

    add_heading(document, "2. Источники данных и SQL-агрегация", level=1)
    add_paragraph(
        document,
        "Слой данных разделён между модулями `sources.py`, `sql_sources_query.py`, `sql_sources_registry.py`, "
        "`sql_aggregations.py` и `sql_payload.py`. Сначала `sources.py::_load_table_metadata` запрашивает "
        "список колонок через `app.db_metadata.get_table_columns_cached(table_name)`, после чего сопоставляет "
        "логические поля проекта с реальными колонками таблицы.",
    )
    add_bullet(document, "Обязательная логическая колонка: дата пожара (`date`).")
    add_bullet(document, "Опциональные колонки: район (`district`), причина (`cause`), категория объекта (`object_category`).")
    add_bullet(document, "Дополнительные аналитические колонки: температура (`temperature`), широта (`latitude`), долгота (`longitude`).")
    add_paragraph(
        document,
        "Для сырых таблиц SQL строится в `SourceQuerySqlMixin::_build_scope_conditions` и `_load_forecasting_records`. "
        "Запрос имеет паттерн `SELECT ... FROM <table> WHERE <conditions> ORDER BY fire_date`. "
        "В условия входят: `date IS NOT NULL`, опционально `EXTRACT(YEAR FROM date_expr) >= :min_year`, "
        "а также параметризованные фильтры по району, причине и категории объекта.",
    )
    add_paragraph(
        document,
        "Для построения дневной истории используется два паттерна. "
        "Первый паттерн, если материализованного представления нет: "
        "`SELECT fire_date, COUNT(*) AS incident_count, AVG(temperature) AS avg_temperature, COUNT(temperature) AS temperature_samples "
        "FROM table WHERE ... GROUP BY fire_date ORDER BY fire_date`. "
        "Второй паттерн, если представление есть: чтение из `mv_forecasting_daily_<suffix>` с уже подготовленными полями "
        "`fire_date`, `district_value`, `cause_value`, `object_category_value`, `incident_count`, `avg_temperature`, `temperature_samples`.",
    )
    add_paragraph(
        document,
        "Материализованные представления создаются в `AggregationQueryBuilder::prepare_forecasting_materialized_views`. "
        "Для PostgreSQL генерируется `CREATE MATERIALIZED VIEW`, а затем индексы по `fire_date` и по набору "
        "`(fire_date, district_value, cause_value, object_category_value)`. Это ускоряет повторяющиеся агрегации "
        "по календарю и фильтрам UI.",
    )
    add_paragraph(
        document,
        "При ограничении истории по окну (`all`, `recent_3`, `recent_5`) функция "
        "`sources.py::_resolve_history_window_min_year` строит UNION-запрос вида "
        "`SELECT MAX(EXTRACT(YEAR FROM date_expr)) AS max_year FROM table WHERE date_expr IS NOT NULL` "
        "для всех источников и затем вычисляет минимальный допустимый год как "
        "`max(latest_years) - (year_span - 1)`.",
    )
    add_paragraph(
        document,
        "Справочники фильтров для UI формируются в `PayloadQueryBuilder::_build_option_catalog_sql`. "
        "Там либо суммируются `incident_count` из материализованных представлений, либо выполняются прямые "
        "`COUNT(*) GROUP BY option_value` по сырым таблицам. Результат агрегируется в `Counter` и переводится "
        "в списки `available_districts`, `available_causes`, `available_object_categories`.",
    )

    add_heading(document, "3. Интервалы прогноза: фактическая реализация", level=1)
    add_paragraph(
        document,
        "В текущей реализации модуля bootstrap-интервалы не используются. Поиск по `app/services/forecasting/*.py` "
        "не показывает функций бутстрэп-ресемплинга, квантильной калибровки или повторных случайных прогонов. "
        "Название `bootstrap.py` в этом модуле относится не к статистическому bootstrap, а к bootstrap-подгрузке "
        "UI payload при ленивом заполнении страницы.",
    )
    add_paragraph(
        document,
        "Фактические интервалы прогноза строятся эвристически в `shaping.py::_build_forecast_rows`. "
        "После вычисления точечной оценки `estimate` определяется ширина интервала "
        "`spread = max(0.75, volatility * (0.9 + 0.15 * sqrt(step)))`. "
        "Далее рассчитываются границы: "
        "`lower_bound = max(0.0, estimate - spread)` и "
        "`upper_bound = min(robust_ceiling + spread, estimate + spread)`.",
    )
    add_paragraph(
        document,
        "Следовательно, в текущем коде это не доверительный интервал в строгом статистическом смысле, "
        "а инженерный диапазон неопределённости, зависящий от волатильности недавней истории и удалённости шага прогноза. "
        "При академическом описании корректно указывать, что интервал формируется по эвристическому правилу, "
        "а не бутстрэпом.",
    )

    add_heading(document, "4. SourceQueryRegistryMixin и SourceQuerySqlMixin", level=1)
    add_paragraph(
        document,
        "Класс `SourceQueryBuilder` в `sql_sources.py` построен через множественное наследование: "
        "`class SourceQueryBuilder(SourceQueryRegistryMixin, SourceQuerySqlMixin, QueryBuilder)`. "
        "Здесь реализовано разделение обязанностей по паттерну mixin.",
    )
    add_bullet(
        document,
        "`SourceQuerySqlMixin` отвечает за построение SQL-фрагментов и выполнение запросов: "
        "`_build_scope_conditions`, `_load_forecasting_records`, `_load_option_counts`, "
        "`_load_source_daily_history_rows`, `_daily_history_union_source_part_sql` и др.",
    )
    add_bullet(
        document,
        "`SourceQueryRegistryMixin` отвечает за маршрутизацию источника: "
        "использовать ли материализованное представление, прямую таблицу, UNION fast path или fallback на построчное чтение.",
    )
    add_paragraph(
        document,
        "Такое разделение полезно тем, что SQL-логика не смешивается с логикой выбора источника. "
        "Например, `_load_daily_history_rows` сначала проверяет наличие materialized view через "
        "`self._aggregations._daily_aggregate_view_exists(table_name)`. "
        "Если представление есть, используется `_load_materialized_daily_history_rows`; если нет, вызывается "
        "`_load_source_daily_history_rows`. Для нескольких таблиц применяется `_load_daily_history_rows_union`, "
        "который собирает общий UNION-запрос. Если fast path завершается ошибкой, `SourceQueryRegistryMixin` "
        "переходит к посттабличному fallback-режиму и затем сливает результаты через `_merge_daily_history_rows`.",
    )

    add_heading(document, "5. Метаданные прогноза и пользовательские параметры", level=1)
    add_paragraph(
        document,
        "Параметры пользовательского запроса проходят через `forecasting_pipeline.py`, `forecasting_bootstrap.py`, "
        "`assembly_input.py` и `assembly_output.py`. Ключевые фильтры и параметры содержатся в payload `filters`.",
    )
    add_bullet(document, "Источник данных: `table_name` или нормализованный список `source_tables`.")
    add_bullet(document, "Территориальный срез: `district`.")
    add_bullet(document, "Причина пожара: `cause`.")
    add_bullet(document, "Категория объекта: `object_category`.")
    add_bullet(document, "Горизонт прогноза: `forecast_days`.")
    add_bullet(document, "Окно истории: `history_window` (`all`, `recent_3`, `recent_5`).")
    add_bullet(document, "Температурный сценарий: `temperature` как строковый ввод пользователя, далее нормализуемый в число.")
    add_paragraph(
        document,
        "Функция `build_forecasting_metadata_payload` сначала загружает только метаданные и наборы опций для фильтров, "
        "а уже потом базовый прогноз и блок поддержки решений. Таким образом, UI может быстро отрисовать форму и "
        "показать пользователю доступные срезы ещё до завершения тяжёлых расчётов.",
    )

    add_heading(document, "6. Метрики качества и критерии приемлемости", level=1)
    add_paragraph(
        document,
        "Качество прогноза оценивается в `quality.py::_run_scenario_backtesting` и "
        "`quality.py::_build_scenario_quality_assessment`. Проверка построена как one-step-ahead rolling backtest, "
        "то есть скользящая одношаговая валидация по истории без использования будущих наблюдений.",
    )
    add_paragraph(
        document,
        "Сначала формируется сезонная базовая модель `_scenario_baseline_expected_count`: "
        "`baseline = 0.6 * mean(last_same_weekday) + 0.4 * recent_mean`, где recent_mean считается по последним 28 дням. "
        "Затем рабочий прогноз сравнивается с этой базой на одинаковых исторических окнах.",
    )
    add_paragraph(
        document,
        "Численные метрики реализованы в `app/services/model_quality.py::compute_count_metrics`.",
    )
    add_bullet(document, "MAE: средняя абсолютная ошибка, `MAE = mean(|ŷ - y|)`.")
    add_bullet(document, "RMSE: корень из средней квадратичной ошибки, `RMSE = sqrt(mean((ŷ - y)^2))`.")
    add_bullet(
        document,
        "sMAPE: симметричная процентная ошибка, `sMAPE = mean(2 * |ŷ - y| / (|y| + |ŷ|)) * 100%`.",
    )
    add_bullet(
        document,
        "Mean Poisson deviance: средняя пуассоновская девиансная функция, полезная для счётных рядов.",
    )
    add_bullet(
        document,
        "Дополнительно считаются относительные изменения метрик к baseline: "
        "`mae_delta_vs_baseline`, `rmse_delta_vs_baseline`, `smape_delta_vs_baseline`.",
    )
    add_paragraph(
        document,
        "Жёстких фиксированных порогов вида «RMSE < X» в `quality.py` нет. "
        "Модуль использует сравнительную стратегию: рабочий прогноз считается лучше, если его ошибка ниже ошибки "
        "сезонной базовой модели. При этом есть пороги готовности самой оценки качества: "
        "не менее 30 дней непрерывной дневной истории, не менее 8 доступных точек для валидации и не менее 8 "
        "собранных validation rows после прокрутки окон. Размер обучающего окна задаётся как "
        "`min_train_days = min(28, max(14, len(history) // 2))`.",
    )

    add_heading(document, "7. Кэширование: фактическая многоуровневая схема", level=1)
    add_paragraph(
        document,
        "В запросе была указана схема «метаданные 300 сек и SQL 120 сек», однако текущая версия кода показывает "
        "другую конфигурацию. В модуле `forecasting` все обнаруженные TTL-кэши настроены на 120 секунд.",
    )
    add_bullet(
        document,
        "`forecasting_pipeline.py::_FORECASTING_CACHE = build_immutable_payload_ttl_cache(ttl_seconds=120.0)` "
        "хранит итоговый payload прогноза.",
    )
    add_bullet(
        document,
        "`inputs.py::_FORECASTING_METADATA_BUNDLE_CACHE` с TTL 120 секунд хранит метаданные, наборы опций и feature cards.",
    )
    add_bullet(
        document,
        "`inputs.py::_FORECASTING_BASE_INPUT_CACHE` с TTL 120 секунд хранит базовые входы: дневную историю, "
        "число записей, temperature quality и др.",
    )
    add_bullet(
        document,
        "`sql_aggregations.py::_FORECASTING_SQL_CACHE = CopyingTtlCache(ttl_seconds=120.0)` "
        "кэширует результаты SQL-слоя.",
    )
    add_paragraph(
        document,
        "Логика всё равно двухуровневая по смыслу. Верхний уровень кэширует уже собранные Python-payload и "
        "готовые структуры для UI. Нижний уровень кэширует результаты SQL-операций и промежуточные выборки, "
        "чтобы повторно не ходить в БД за одинаковыми агрегатами.",
    )
    add_paragraph(
        document,
        "Структура ключей задаётся через tuple-ключи. Примеры: "
        "`('option_catalog', normalized_tables, history_window, ...)`, "
        "`('daily_history', normalized_tables, history_window, district, cause, object_category)`, "
        "`('filtered_record_count', ...)`, "
        "`('forecasting_records', [table_name], district, cause, object_category, history_window)`, "
        "`('temperature_quality', [table_name], date_column, temperature_column)`, "
        "`('history_window_year', source_tables, history_window)`.",
    )
    add_paragraph(
        document,
        "Инвалидация выполняется централизованно через `forecasting_pipeline.py::clear_forecasting_cache`, "
        "которая вызывает `clear_forecasting_input_cache()` и `clear_forecasting_sql_cache()`. "
        "Поэтому при изменении данных или структуры источников можно снять сразу и верхний, и нижний слой кэша.",
    )

    add_heading(document, "8. Взаимодействие с forecast_risk", level=1)
    add_paragraph(
        document,
        "Модуль сценарного прогноза не ограничивается выдачей временного ряда. Он напрямую интегрирован с "
        "`app/services/forecast_risk/core.py`. В `forecasting_pipeline.py` импортируется функция "
        "`build_decision_support_payload`, после чего она внедряется в зависимости forecasting-слоя.",
    )
    add_paragraph(
        document,
        "Интеграция происходит в `assembly_input.py::_build_decision_support_block` и "
        "`assembly_output.py::complete_forecasting_decision_support_payload`. "
        "В `forecast_risk` передаются фильтры текущего сценария: `source_tables`, `selected_district`, "
        "`selected_cause`, `selected_object_category`, `history_window`, `planning_horizon_days`.",
    )
    add_paragraph(
        document,
        "Возвращаемый объект включает `risk_prediction`, а также может содержать `geo_prediction`. "
        "После этого блок сценарного прогноза показывает приоритеты территорий, паспорт качества, краткое "
        "управленческое резюме и рекомендации, уже согласованные с выбранным временным горизонтом прогноза.",
    )
    add_paragraph(
        document,
        "Для тяжёлого decision-support расчёта применяется отдельный фоновой механизм: "
        "`app/services/forecasting/jobs.py` создаёт `ThreadPoolExecutor(max_workers=2)` и "
        "поддерживает статусы задачи через job store. Это позволяет сначала показать базовый сценарный прогноз, "
        "а затем догрузить приоритеты территорий и рекомендации без блокировки ответа.",
    )

    add_heading(document, "9. Использование statsmodels", level=1)
    add_paragraph(
        document,
        "Поиск по файлам `app/services/forecasting/*.py` не выявляет импортов `statsmodels`, а также классов "
        "`ARIMA`, `SARIMAX`, `ExponentialSmoothing`, `ETSModel` или функций декомпозиции. "
        "Следовательно, в текущем модуле `forecasting` библиотека `statsmodels` не используется.",
    )
    add_paragraph(
        document,
        "Фактически расчёты выполняются на базе стандартного Python-модуля `statistics` (`mean`, `pstdev`), "
        "ручных правил формирования сезонных факторов, SQLAlchemy для доступа к данным и собственных сервисных "
        "утилит проекта. Если в тексте диссертации требуется упомянуть `statsmodels`, то это будет неверно "
        "для текущей версии именно этого модуля.",
    )

    add_heading(document, "10. Вывод", level=1)
    add_paragraph(
        document,
        "Модуль `app/services/forecasting/` представляет собой не классическую библиотечную time-series модель, "
        "а объяснимый сценарный сервис поверх SQL-агрегаций, кэширования и эвристико-статистических правил. "
        "Его сильная сторона состоит в прозрачности вычислений, привязке к доступным атрибутам пожарных данных, "
        "быстрой реакции UI за счёт многоуровневого TTL-кэширования и интеграции с блоком `forecast_risk` "
        "для формирования рекомендаций. При этом важно корректно зафиксировать ограничения: отсутствуют "
        "bootstrap-интервалы в статистическом смысле и отсутствует использование `statsmodels`.",
    )

    return document


def count_question_marks(path: Path) -> int:
    with ZipFile(path, "r") as archive:
        xml_parts = []
        for name in archive.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                xml_parts.append(archive.read(name).decode("utf-8", errors="ignore"))
    return "".join(xml_parts).count("?")


def main() -> None:
    document = build_document()
    document.save(OUTPUT_PATH)
    print(f"saved={OUTPUT_PATH}")
    print(f"question_marks={count_question_marks(OUTPUT_PATH)}")


if __name__ == "__main__":
    main()
