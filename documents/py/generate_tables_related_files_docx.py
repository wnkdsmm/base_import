from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parents[1]
DOCUMENTS_DIR = BASE_DIR / "documents"
OUTPUT_PATH = DOCUMENTS_DIR / "tables_related_files.docx"


def ru(value: str) -> str:
    return value.encode("ascii").decode("unicode_escape")


def set_font(run, *, name: str = "Times New Roman", size: int = 12, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_paragraph(document: Document, text: str, *, bold: bool = False, size: int = 12, align=None) -> None:
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_font(run)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)


def main() -> None:
    categories = [
        (
            ru("\\u0412\\u0445\\u043e\\u0434\\u043d\\u044b\\u0435 \\u0442\\u043e\\u0447\\u043a\\u0438 \\u0438 \\u043c\\u0430\\u0440\\u0448\\u0440\\u0443\\u0442\\u044b"),
            [
                "app/main.py",
                "app/routes/api.py",
                "app/routes/pages.py",
                "app/routes/page_common.py",
                "app/routes/api_tables.py",
            ],
        ),
        (
            ru("\\u0428\\u0430\\u0431\\u043b\\u043e\\u043d\\u044b \\u0438 layout"),
            [
                "app/templates/base.html",
                "app/templates/includes/sidebar_nav.html",
                "app/templates/tables.html",
                "app/templates/table_view.html",
            ],
        ),
        (
            ru("\\u0421\\u0442\\u0438\\u043b\\u0438 \\u0438 frontend-\\u0441\\u043a\\u0440\\u0438\\u043f\\u0442\\u044b"),
            [
                "app/static/css/base.css",
                "app/static/css/layout.css",
                "app/static/css/shared-components.css",
                "app/static/css/analytics.css",
                "app/static/css/tables.css",
                "app/static/js/sidebar.js",
                "app/static/js/analytics_shared.js",
                "app/static/js/import.js",
                "app/static/js/tables.js",
                "app/static/js/table_view.js",
                "app/static/js/shared/api_client.js",
                "app/static/js/shared/plotly_helpers.js",
                "app/static/js/shared/state_factory.js",
                "app/static/js/shared/ui_helpers.js",
            ],
        ),
        (
            ru("\\u0421\\u0435\\u0440\\u0432\\u0438\\u0441\\u043d\\u044b\\u0439 \\u0438 \\u0434\\u0430\\u0442\\u0430-\\u0441\\u043b\\u043e\\u0439"),
            [
                "app/services/table_workflows.py",
                "app/services/table_summary.py",
                "app/table_metadata.py",
                "app/table_operations.py",
                "app/db_views.py",
            ],
        ),
        (
            ru("\\u0422\\u0435\\u0441\\u0442\\u044b"),
            [
                "tests/test_table_pagination_db_views.py",
                "tests/test_table_pagination_order.py",
                "tests/test_table_pagination_routes.py",
                "tests/test_template_layout.py",
            ],
        ),
    ]

    all_files = [item for _title, files in categories for item in files]

    document = Document()
    configure_styles(document)

    add_paragraph(
        document,
        ru("\\u0424\\u0430\\u0439\\u043b\\u044b, \\u043e\\u0442\\u043d\\u043e\\u0441\\u044f\\u0449\\u0438\\u0435\\u0441\\u044f \\u043a \\u0431\\u043b\\u043e\\u043a\\u0443 \\xab\\u0422\\u0430\\u0431\\u043b\\u0438\\u0446\\u044b\\xbb"),
        bold=True,
        size=16,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        document,
        ru(
            "\\u0412 \\u0434\\u043e\\u043a\\u0443\\u043c\\u0435\\u043d\\u0442 \\u0432\\u043a\\u043b\\u044e\\u0447\\u0435\\u043d\\u044b \\u0444\\u0430\\u0439\\u043b\\u044b, "
            "\\u043a\\u043e\\u0442\\u043e\\u0440\\u044b\\u0435 \\u043e\\u0431\\u0435\\u0441\\u043f\\u0435\\u0447\\u0438\\u0432\\u0430\\u044e\\u0442 \\u0441\\u0442\\u0440\\u0430\\u043d\\u0438\\u0446\\u0443 "
            "\\u0441\\u043f\\u0438\\u0441\\u043a\\u0430 \\u0442\\u0430\\u0431\\u043b\\u0438\\u0446 `/tables`, \\u0441\\u0442\\u0440\\u0430\\u043d\\u0438\\u0446\\u0443 \\u043f\\u0440\\u043e\\u0441\\u043c\\u043e\\u0442\\u0440\\u0430 "
            "`/tables/{table_name}`, API \\u043f\\u0430\\u0433\\u0438\\u043d\\u0430\\u0446\\u0438\\u0438 \\u0438 \\u0443\\u0434\\u0430\\u043b\\u0435\\u043d\\u0438\\u044f, "
            "\\u0430 \\u0442\\u0430\\u043a\\u0436\\u0435 \\u0441\\u0432\\u044f\\u0437\\u0430\\u043d\\u043d\\u044b\\u0435 \\u0441\\u0435\\u0440\\u0432\\u0438\\u0441\\u044b \\u0438 \\u0442\\u0435\\u0441\\u0442\\u044b."
        ),
    )
    add_paragraph(
        document,
        ru(f"\\u0412\\u0441\\u0435\\u0433\\u043e \\u0432 \\u043f\\u0435\\u0440\\u0435\\u0447\\u043d\\u0435: {len(all_files)} \\u0444\\u0430\\u0439\\u043b\\u0430(\\u043e\\u0432)."),
        bold=True,
    )

    for title, files in categories:
        add_paragraph(document, title, bold=True, size=14)
        for path in files:
            add_bullet(document, path)

    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
