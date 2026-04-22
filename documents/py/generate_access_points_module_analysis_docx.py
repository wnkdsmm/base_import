from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / "access_points_module_analysis.docx"


def set_font(run, *, size: float = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def configure_styles(document: Document) -> None:
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "List Bullet"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)


def add_paragraph(document: Document, text: str, *, bold: bool = False, align=None) -> None:
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, bold=bold)
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_after = Pt(6)


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_font(run, size=14, bold=True)
    paragraph.paragraph_format.space_after = Pt(6)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_font(run)
    paragraph.paragraph_format.space_after = Pt(3)


def build_document() -> Document:
    document = Document()
    configure_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Технический анализ модуля app/services/access_points/ проекта "Fire Data Pipeline"')
    set_font(run, size=14, bold=True)

    add_paragraph(
        document,
        "Документ описывает фактическую реализацию блока «Проблемные точки» по исходному коду модуля "
        "`app/services/access_points/`. Этот блок не прогнозирует число пожаров и не выполняет кластеризацию; "
        "он строит объяснимый рейтинг отдельных точек, где доступ пожарных подразделений, подтверждённость "
        "водоснабжения, тяжесть последствий и повторяемость инцидентов делают локацию приоритетной для проверки.",
    )

    add_heading(document, "1. Понятие точки доступа и структура исходных данных")
    add_paragraph(
        document,
        "В предметной области модуля «точка доступа» — это не административная территория, а наиболее детализированная "
        "локация, для которой можно агрегировать пожары и оценить проблемность доступа пожарных подразделений. "
        "Логика построения такой сущности реализована в `point_data.py::_resolve_point_identity`, а загрузка исходных "
        "полей — в `data_impl.py::_build_source_sql`, `_normalize_record` и `_record_to_access_point_input`.",
    )
    add_paragraph(
        document,
        "Источник данных — пользовательские таблицы пожаров в PostgreSQL. Для каждой таблицы модуль сначала получает "
        "метаданные через `app.db_metadata.get_table_columns_cached`, затем подбирает реальные имена колонок по "
        "набору candidate-полей из `constants.py`. SQL-выборка строится динамически: для отсутствующих колонок "
        "подставляются безопасные fallback-выражения `''` или `NULL` через "
        "`app.services.shared.sql_helpers.select_expression_or_fallback`.",
    )
    add_paragraph(
        document,
        "В результирующей строке инцидента модуль использует следующие группы полей: "
        "район `district`, метку территории `territory_label`, населённый пункт `settlement`, тип населённого пункта "
        "`settlement_type`, категорию объекта `object_category`, название объекта `object_name`, адрес `address`, "
        "комментарий к адресу `address_comment`, координаты `latitude` и `longitude`, времена `report_time`, "
        "`arrival_time`, `detection_time`, расстояние до пожарной части `distance_to_fire_station`, сведения о воде "
        "`water_supply_count` и `water_supply_details`, последствия `consequence`, погибших `deaths`, пострадавших "
        "`injuries`, признак жертв `casualty_flag`, площадь и строения `destroyed_area`, `destroyed_buildings`, "
        "а также денежный ущерб `registered_damage` и дату события `event_date`.",
    )
    add_paragraph(
        document,
        "После нормализации формируются производные признаки уровня инцидента: `response_minutes`, `long_arrival`, "
        "`has_water_supply`, `severe_consequence`, `victims_present`, `major_damage`, `night_incident`, "
        "`heating_season`, `source_table`, `year`. Если у записи отсутствуют координаты, функция "
        "`data_impl.py::_normalize_record` возвращает `None`, и такой инцидент не попадает в дальнейший анализ.",
    )

    add_heading(document, "2. Иерархия географической привязки точки")
    add_paragraph(
        document,
        "Fallback-цепочка идентификации точки реализована в `point_data.py::_resolve_point_identity`. Формально "
        "пользовательский запрос можно описать как «адрес -> населённый пункт -> территория -> район», но в коде "
        "цепочка более точная и содержит промежуточные уровни:",
    )
    add_bullet(document, "1) адрес: `point_id = \"address:<normalized_address>|<normalized_object>\"`, `granularity_rank = 5`.")
    add_bullet(document, "2) осмысленное название объекта: `point_id = \"object:<normalized_object>\"`, `granularity_rank = 4`.")
    add_bullet(document, "3) координатная точка при наличии `latitude` и `longitude`: `point_id = \"coords:<lat>:<lon>\"`, `granularity_rank = 4`.")
    add_bullet(document, "4) населённый пункт: `point_id = \"settlement:<normalized_settlement>\"`, `granularity_rank = 3`.")
    add_bullet(document, "5) территория: `point_id = \"territory:<normalized_territory>\"`, `granularity_rank = 2`.")
    add_bullet(document, "6) район: `point_id = \"district:<normalized_district>\"`, `granularity_rank = 1`.")
    add_bullet(document, "7) неизвестная локация: `point_id = \"unknown:unresolved\"`, `granularity_rank = 0`.")
    add_paragraph(
        document,
        "Именно `granularity_rank` затем участвует в сортировке рейтинга: при равном score выше ставится более "
        "детальная сущность. Это важно для практического использования, так как адресная точка полезнее для "
        "управленческого решения, чем агрегат по району.",
    )

    add_heading(document, "3. Агрегация инцидентов по точкам")
    add_paragraph(
        document,
        "Агрегация реализована в `point_data.py::_aggregate_point_buckets`, `_update_point_bucket_from_record` и "
        "`_build_point_bucket_row`. Каждый bucket накапливает счётчики инцидентов, суммарные времена реагирования, "
        "суммарные расстояния, количество случаев без воды, тяжёлых последствий, жертв, крупного ущерба, ночных "
        "пожаров, пожаров в отопительный сезон, а также `Counter`-структуры для районов, территорий, населённых "
        "пунктов, типов поселений, категорий объектов и исходных таблиц.",
    )
    add_paragraph(
        document,
        "Из bucket вычисляются агрегированные показатели точки: `incident_count`, `years_observed`, "
        "`incidents_per_year`, `average_response_minutes`, `average_distance_km`, `response_coverage_share`, "
        "`distance_coverage_share`, `water_coverage_share`, `water_unknown_share`, `long_arrival_share`, "
        "`no_water_share`, `severe_share`, `victim_share`, `major_damage_share`, `night_share`, `heating_share`, "
        "`rural_share`, а также координаты как среднее по всем валидным точкам пожаров.",
    )
    add_paragraph(
        document,
        "Ключевые базовые формулы находятся в `point_data.py` и `numeric.py`: "
        "`average_response = response_total / response_count`, "
        "`average_distance = distance_total / distance_count`, "
        "`response_coverage_share = response_count / incident_count`, "
        "`distance_coverage_share = distance_count / incident_count`, "
        "`water_coverage_share = known_water_count / incident_count`, "
        "`water_unknown_share = max(0, 1 - water_coverage_share)`, "
        "`incidents_per_year = incident_count / years_observed`.",
    )

    add_heading(document, "4. Сглаживание при малом числе инцидентов")
    add_paragraph(
        document,
        "Сглаживание реализовано в `point_data.py::_smooth_share` и применяется ко всем долевым признакам через "
        "`_build_smoothed_bucket_shares`. Это не простая средняя доля, а bayes-like сглаживание к глобальному prior, "
        "если по точке наблюдений меньше, чем `MIN_ACCESS_POINT_SUPPORT`.",
    )
    add_paragraph(
        document,
        "Формула имеет вид: "
        "`support_gap = max(0, minimum_support - observations)`; "
        "если `support_gap <= 0`, то `share = successes / observations`; "
        "иначе `share = (successes + prior_mean * support_gap) / (observations + support_gap)`.",
    )
    add_paragraph(
        document,
        "Глобальные priors строятся функцией `point_data.py::_build_point_priors` по всему набору bucket-ов: "
        "отдельно для `long_arrival`, `no_water`, `severe`, `victims`, `major_damage`, `night`, `heating`, `rural`. "
        "За счёт этого точка с 1-2 пожарами не получает экстремальные доли 0 или 1 только из-за случайности.",
    )
    add_paragraph(
        document,
        "Дополнительно используется коэффициент поддержки: "
        "`support_weight_raw = min(1, incident_count / resolved_support)`, "
        "`support_weight = 0.4 + 0.6 * support_weight_raw`. "
        "Он не зануляет вклад слабой точки, но ослабляет contribution основных факторов при малом числе наблюдений.",
    )

    add_heading(document, "5. Предобработка и pandas-операции")
    add_paragraph(
        document,
        "Предобработка признаков сосредоточена в `analysis_factors.py`, `features.py` и `numeric.py`. Модуль активно "
        "использует `pandas` и `numpy`; это один из немногих аналитических блоков проекта, где преобразования "
        "данных векторизованы на `DataFrame` и `Series`.",
    )
    add_paragraph(
        document,
        "Ключевые операции подготовки числовых полей: "
        "`pd.to_numeric(..., errors=\"coerce\")` для перевода строк в числа; "
        "`.astype(float)` для унификации типа; "
        "`.where(pd.notna(values) & np.isfinite(values))` для отсечения `NaN` и бесконечностей; "
        "`.fillna(...)` для подстановки безопасных значений; "
        "`.clip(lower=..., upper=...)` для ограничения долей диапазоном [0, 1]; "
        "`.reset_index(drop=True)` для выравнивания `entity_frame` и `feature_frame`; "
        "`.to_numpy(copy=False)` для подготовки массивов, которые затем читаются построчно без лишнего копирования.",
    )
    add_paragraph(
        document,
        "Функция `numeric.py::_finite_numeric_columns` формирует numeric-frame только по целевым колонкам. "
        "Это защищает scoring от ошибок из-за текстовых значений, пустых строк и нечисловых артефактов в исходной БД.",
    )

    add_heading(document, "6. Отбор признаков и обработка пропусков")
    add_paragraph(
        document,
        "Набор explainable-признаков задаётся в `app/domain/access_points_metadata.py` и экспортируется через "
        "`constants.py::DEFAULT_ACCESS_POINT_FEATURES`. В scoring используются следующие reason-коды: "
        "`DISTANCE_TO_STATION`, `RESPONSE_TIME`, `LONG_ARRIVAL_SHARE`, `NO_WATER`, `SEVERE_CONSEQUENCES`, "
        "`REPEAT_FIRES`, `NIGHT_PROFILE`, `HEATING_SEASON`.",
    )
    add_paragraph(
        document,
        "Функция `features.py::_build_access_point_candidate_features` для каждого признака считает покрытие "
        "`coverage = non_null_count / row_count`, дисперсию `variance = series.var(skipna=True)` и число уникальных "
        "значений `unique_count = series.nunique(dropna=True)`. Затем признаки сортируются по score, который по сути "
        "приоритизирует базовый explainable-набор и более полные колонки.",
    )
    add_paragraph(
        document,
        "Если пользователь запросил признаки, которых нет в доступных данных, "
        "`features.py::_resolve_selected_access_point_features` откатывается к базовому набору. Это защищает UI "
        "от ситуации, когда кастомный выбор делает scoring пустым или неконсистентным.",
    )
    add_paragraph(
        document,
        "Обработка пропусков зависит от природы признака. Для долей применяется `fillna(0)` и `clip(0, 1)`. "
        "Для `victim_share` и `major_damage_share` предусмотрен fallback: если готовой доли нет, она восстанавливается "
        "как `victims_count / incident_count` и `major_damage_count / incident_count`. Для `water_unknown_share` "
        "при отсутствии значения используется `1 - water_coverage_share`. Для `years_observed` и `incident_count` "
        "применяются нижние пороги через `.clip(lower=1)` и `.clip(lower=0)`.",
    )

    add_heading(document, "7. Формулы scoring и нормализация")
    add_paragraph(
        document,
        "Ядро scoring находится в `analysis_factors.py::_build_access_point_base_series`, "
        "`_build_access_point_factor_series`, `_resolve_access_point_weight_context` и "
        "`analysis_output_context.py::_build_access_point_score_decomposition`.",
    )
    add_paragraph(
        document,
        "Сначала строятся нормализованные факторы. Масштабы определяются по текущей выборке: "
        "`distance_scale = max(12.0, max(average_distance))`, "
        "`response_scale = max(LONG_RESPONSE_THRESHOLD_MINUTES, max(average_response))`, "
        "`max_incidents = max(1.0, max(incident_count))`, "
        "`max_incidents_per_year = max(1.0, max(incidents_per_year))`.",
    )
    add_bullet(document, "`distance_norm = clip(average_distance / distance_scale, 0, 1)`.")
    add_bullet(document, "`response_norm = clip(average_response / response_scale, 0, 1)`.")
    add_bullet(document, "`frequency_norm = clip(incidents_per_year / max_incidents_per_year, 0, 1)`.")
    add_bullet(document, "`incidents_norm = clip(incident_count / max_incidents, 0, 1)`.")
    add_bullet(document, "`severity_factor = clip(0.58 * severe_share + 0.24 * victim_share + 0.18 * major_damage_share, 0, 1)`.")
    add_bullet(document, "`recurrence_factor = clip(0.70 * frequency_norm + 0.30 * incidents_norm, 0, 1)`.")
    add_bullet(document, "`uncertainty_factor = clip(0.35 * arrival_missing_share + 0.30 * water_unknown_share + 0.20 * distance_missing_share + 0.15 * (1 - support_weight), 0, 1)`.")
    add_paragraph(
        document,
        "Базовые веса факторов заданы константой `FACTOR_WEIGHTS`: "
        "distance 16, response 14, long_arrival 10, water 12, severity 18, recurrence 14, night 6, heating 4. "
        "Функция `_resolve_access_point_weight_context` пересчитывает только выбранные пользователем признаки в "
        "шкалу 94 балла: `normalized_weight = 94 * factor_weight / sum(selected_weights)`. Оставшиеся 6 баллов "
        "резервируются под penalty неопределённости: `UNCERTAINTY_PENALTY_MAX = 6.0`.",
    )
    add_paragraph(
        document,
        "Компонентный score считается не как чистая сумма, а как объяснимый набор вкладов. Для каждого reason-кода "
        "формируется `contribution_points = normalized_factor_weight * factor_value * support_weight`. "
        "Именно `support_weight` уменьшает вклад плохо поддержанных точек. Затем "
        "`pure_score = sum(contribution_points по непенальтийным факторам)`.",
    )
    add_paragraph(
        document,
        "Подоценки верхнего уровня формируются так: "
        "`access_score` — нормированная смесь `distance_norm`, `response_norm`, `long_arrival_share`; "
        "`water_score = no_water_share * 100`; "
        "`severity_score = severity_factor * 100`; "
        "`recurrence_score` — нормированная смесь `recurrence_factor`, `night_share`, `heating_share`; "
        "`data_gap_score = uncertainty_factor * 100`.",
    )
    add_paragraph(
        document,
        "Итоговый риск: "
        "`uncertainty_penalty = 6 * uncertainty_factor`; "
        "`total_score = pure_score + uncertainty_penalty`. "
        "Дополнительно рассчитывается `investigation_score = min(100, 0.72 * pure_score + 0.28 * (uncertainty_penalty * 100 / 6))`. "
        "Он используется для выбора точек, где прежде всего нужна верификация данных.",
    )

    add_heading(document, "8. Штраф неопределённости")
    add_paragraph(
        document,
        "Штраф неопределённости оформлен как отдельный reason-код `DATA_UNCERTAINTY` в "
        "`analysis_output_context.py`. В decomposition он помечается флагом `is_penalty=True`, поэтому отделяется "
        "от основного risk-сигнала и трактуется отдельно. Это инженерно важное решение: пользователь видит, что "
        "часть score обусловлена не самой опасностью, а неполнотой данных.",
    )
    add_paragraph(
        document,
        "Penalty начисляется за четыре источника неопределённости: "
        "пропуск времени прибытия, пропуск статуса водоснабжения, пропуск расстояния до ПЧ и малую статистическую "
        "поддержку точки. Максимальный вклад ограничен шестью баллами, поэтому неполнота не может доминировать "
        "над реальным риском доступа.",
    )
    add_paragraph(
        document,
        "Флаг `uncertainty_flag` выставляется в `analysis_output.py::_build_access_point_payload_row`, если "
        "`uncertainty_penalty >= 2.5` или `low_support == True` или `completeness_share < 0.6`. "
        "При этом `missing_data_priority = uncertainty_flag and total_score < HIGH_THRESHOLD`. "
        "То есть модуль специально выделяет точки, где score ещё не очень высокий, но качество данных уже требует проверки.",
    )

    add_heading(document, "9. Ранжирование точек")
    add_paragraph(
        document,
        "Ранжирование находится в `analysis_ranking.py::_build_access_point_rows_from_entity_frame`. "
        "После построения payload-строк список сортируется по следующему ключу по убыванию: "
        "`(total_score, severity_score, access_score, incident_count, granularity_rank)`.",
    )
    add_paragraph(
        document,
        "Следовательно, основная метрика сравнения — итоговый `total_score`. При равенстве score приоритет получает "
        "точка с более тяжёлыми последствиями, затем с худшей доступностью, затем с большим числом инцидентов и, "
        "наконец, с более детальной географической гранулярностью. После сортировки каждой точке присваиваются "
        "`rank` и `rank_display`.",
    )
    add_paragraph(
        document,
        "Отдельно формируется выборка неполных точек в `analysis_ranking.py::_select_incomplete_points`. "
        "Туда попадают либо строки с `missing_data_priority`, либо строки, где `data_gap_score >= 50` и "
        "`investigation_score >= WATCH_RISK_THRESHOLD`. Эти кандидаты сортируются по "
        "`(investigation_score, data_gap_score, total_score)` по убыванию. Практический смысл — вынести наверх "
        "не просто опасные, а те локации, где управленческая польза от дообогащения данных максимальна.",
    )

    add_heading(document, "10. Кэширование и инвалидация")
    add_paragraph(
        document,
        "Кэш реализован в `core.py` как in-memory `CopyingTtlCache(ttl_seconds=120.0)` и живёт только внутри процесса "
        "сервера. Это однопроцессный runtime-кэш аналитического payload, а не распределённый кэш. Он ускоряет повторные "
        "запросы к тем же фильтрам и тем же наборам таблиц, не заставляя заново считывать инциденты из БД и пересчитывать "
        "ranking, summary и графики.",
    )
    add_paragraph(
        document,
        "Ключ кэша строится функцией `core.py::_build_access_points_cache_key` и имеет структуру: "
        "`(\"v6\", selected_table, *source_tables, district, year, limit, *feature_columns)`. "
        "Версия `v6` играет роль coarse-grained инвалидации: при изменении логики scoring достаточно поменять версию, "
        "и старые ключи станут недействительными.",
    )
    add_paragraph(
        document,
        "Повторное использование реализовано одинаково для shell-режима страницы и для API-ответа: "
        "`get_access_points_shell_context()` и `get_access_points_data()` сначала вызывают `_ACCESS_POINTS_CACHE.get(key)`, "
        "а при промахе строят payload и кладут его обратно через `_ACCESS_POINTS_CACHE.set(key, payload)`.",
    )
    add_paragraph(
        document,
        "Явная инвалидация в модуле — только `clear_access_points_cache()`, которая вызывает `clear()` у кэша. "
        "Автоматической подписки на изменение схемы или содержимого таблиц в самом `access_points` нет. Поэтому "
        "фактическая стратегия инвалидации такая: "
        "1) TTL-истечение через 120 секунд; "
        "2) ручной вызов `clear_access_points_cache()`; "
        "3) рестарт процесса; "
        "4) смена версии ключа `v6` в коде.",
    )

    add_heading(document, "11. Взаимодействие с остальными модулями")
    add_paragraph(
        document,
        "Внешние точки входа стандартные для аналитических сервисов проекта. Страница `/access-points` в "
        "`app/routes/pages.py` вызывает `app.services.access_points.core.get_access_points_shell_context`, а API "
        "`/api/access-points-data` в `app/routes/api_access_points.py` вызывает `core.get_access_points_data` через "
        "обёртку `run_analytics_request`.",
    )
    add_paragraph(
        document,
        "Прямой кодовой зависимости `access_points -> forecast_risk` в модуле не обнаружено. "
        "`app/services/access_points/` не импортирует `forecast_risk.core`, не вызывает его scoring и не передаёт туда "
        "свои результаты. Оба блока существуют рядом как самостоятельные аналитические сервисы. "
        "На уровне dashboard и общих страниц они могут отображаться в одной системе, но математически и программно "
        "их расчёты разделены.",
    )
    add_paragraph(
        document,
        "Аналогично отсутствует прямой вызов `fire_map_service` из `access_points`. Модуль сам строит свои визуализации "
        "в `charts.py`: scatter, histogram, factor bar chart и heatmap. В payload точек действительно присутствуют "
        "`latitude`, `longitude` и `coordinates_display`, однако это используется внутри самого блока и не передаётся "
        "напрямую в `app/services/fire_map_service.py`. Следовательно, связь с картой в текущем коде — косвенная: "
        "оба модуля работают с общей базой и координатами пожаров, но отдельного мостика API между ними нет.",
    )

    add_heading(document, "12. Итоговая техническая оценка")
    add_paragraph(
        document,
        "Модуль `access_points` реализует explainable-рейтинг проблемных локаций поверх сырых пожарных инцидентов. "
        "Сильная сторона реализации — прозрачные формулы, поддержка fallback-географии, сглаживание малых выборок, "
        "отдельный слой неопределённости и векторизованный расчёт на `pandas`. С инженерной точки зрения модуль хорошо "
        "подходит для диссертационного описания как сервис поддержки принятия решений, потому что каждый итоговый балл "
        "можно декомпозировать на измеримые причины: удалённость, время прибытия, воду, последствия, повторяемость "
        "и полноту данных.",
    )
    add_paragraph(
        document,
        "При этом необходимо корректно фиксировать ограничения текущей реализации: кэш локальный и краткоживущий, "
        "инвалидация данных не автоматизирована по событиям БД, а связка с `forecast_risk` и `fire_map_service` "
        "не реализована как прямой вычислительный пайплайн. Эти особенности лучше указывать в работе явно, "
        "чтобы описание совпадало с фактическим кодом проекта.",
    )

    return document


def verify_document_text(path: Path) -> tuple[int, int]:
    with ZipFile(path, "r") as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml)
    joined = "\n".join(texts)
    return joined.count("?"), len(texts)


def main() -> None:
    document = build_document()
    document.save(OUTPUT_PATH)
    question_marks, text_nodes = verify_document_text(OUTPUT_PATH)
    print(f"saved={OUTPUT_PATH}")
    print(f"text_question_marks={question_marks}")
    print(f"text_nodes={text_nodes}")


if __name__ == "__main__":
    main()
