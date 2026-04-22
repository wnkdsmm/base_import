from __future__ import annotations

from datetime import datetime
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"F:\filesFires\base_import")
OUTPUT_PATH = ROOT / "documents" / "Анализ_dashboard_модуля.docx"


def set_font(run, name: str = "Times New Roman", size: int = 12, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def configure_document(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def add_paragraph(document: Document, text: str, *, bold: bool = False, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_font(run, bold=bold)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading("", level=level)
    run = heading.add_run(text)
    set_font(run, size=14 if level == 1 else 12, bold=True)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        add_paragraph(document, item, style="List Bullet")


def count_question_marks_in_docx_text_nodes(docx_path: Path) -> int:
    with ZipFile(docx_path, "r") as archive:
        xml_bytes = archive.read("word/document.xml")
    root = ET.fromstring(xml_bytes)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return sum((node.text or "").count("?") for node in root.findall(".//w:t", ns))


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title_run = title.add_run("Технический анализ модуля app/dashboard проекта Fire Data Pipeline")
    set_font(title_run, size=16, bold=True)

    add_paragraph(doc, "Дата формирования: " + datetime.now().strftime("%d.%m.%Y %H:%M"))
    add_paragraph(doc, "Анализируемые файлы:")
    add_bullets(
        doc,
        [
            "app/dashboard/service.py",
            "app/dashboard/dashboard_service_data.py",
            "app/dashboard/dashboard_service_build.py",
            "app/dashboard/cache.py",
            "app/dashboard/metadata.py",
            "app/dashboard/data_access.py",
            "app/dashboard/impact.py",
            "app/dashboard/impact_fire_metrics.py",
            "app/dashboard/impact_forecast_metrics.py",
            "app/dashboard/distribution.py",
            "app/dashboard/distribution_logic.py",
            "app/dashboard/summary.py",
            "app/dashboard/summary_logic.py",
            "app/dashboard/management.py",
            "app/statistics_constants.py",
            "app/domain/analytics_metadata.py",
            "app/runtime_invalidation.py",
            "app/table_operations.py",
        ],
    )

    add_heading(doc, "1. Архитектура дашборда и сборка единого ответа", level=1)
    add_bullets(
        doc,
        [
            "Входная точка данных дашборда: get_dashboard_data в app/dashboard/dashboard_service_data.py.",
            "Слой service.py выступает как compatibility re-export и проксирует вызовы в специализированные модули.",
            "Подготовка запроса: нормализация horizon_days, загрузка метаданных, вычисление cache_key, попытка чтения из кэша по запрошенному и каноническому ключу.",
            "Агрегация выполняется в _build_dashboard_aggregation (app/dashboard/dashboard_service_build.py).",
            "Финальная упаковка JSON-подобного ответа выполняется _build_dashboard_payload.",
            "Поля итогового payload: generated_at, has_data, summary, scope, trend, management, highlights, rankings, widgets, charts, filters, notes.",
            "Состав charts: yearly_fires, yearly_area, distribution, monthly_heatmap, monthly_profile, area_buckets, cumulative_area.",
            "Состав widgets: causes, districts, seasons (SQL-виджеты).",
            "Состав management: блок управленческой сводки на основе build_decision_support_payload из forecast_risk.",
            "Типовые контракты данных описаны TypedDict-классами в app/dashboard/types.py: DashboardPayload, DashboardAggregation, DashboardGroupedCounts и др.",
        ],
    )

    add_heading(doc, "2. Оптимизация SQL: reuse grouped counts и устранение N+1", level=1)
    add_bullets(
        doc,
        [
            "Ключевая оптимизация реализована в _collect_dashboard_grouped_counts (app/dashboard/impact_fire_metrics.py).",
            "Вместо отдельных запросов по каждому измерению и каждой таблице строятся подзапросы на таблицу, которые объединяются через UNION ALL и исполняются одним roundtrip.",
            "Внутри каждого подзапроса используется GROUP BY GROUPING SETS для одновременного расчета причин, распределения, районов, месяцев, бакетов площади и таймлайна.",
            "metric_kind вычисляется CASE WHEN GROUPING(column)=0 THEN <kind> ... END, что позволяет вернуть мульти-метрику в одном результирующем наборе.",
            "В _build_dashboard_grouped_counts_query используется шаблон grouped_source -> grouped_counts_bundle_N, где grouped_source содержит нормализованные label/date/metric-колонки.",
            "Пример паттерна устранения N+1: ранее логика строила бы цикл таблица x измерение x запрос; теперь таблица x 1 подзапрос, затем единый UNION ALL.",
            "Повторное использование результатов включено в distribution_logic._build_standard_dashboard_charts: distribution_counts, month_counts, area_bucket_counts и impact_timeline_rows передаются без повторных SQL-запросов.",
            "Проверка возможности reuse для распределения делается через _can_reuse_distribution_counts: если выбранный group_column совместим с резолвингом колонки в таблицах, график берет уже посчитанные counts.",
        ],
    )
    add_paragraph(doc, "Итог: шаблон grouped counts сокращает число SQL-вызовов с N+1 до 1 агрегированного запроса на набор выбранных таблиц и фильтров.")

    add_heading(doc, "3. Виджеты ущерба damage: метрики и агрегация", level=1)
    add_bullets(
        doc,
        [
            "Группа damage активируется при selected_group_column == DAMAGE_GROUP_OPTION_VALUE (__group__:damage_overview).",
            "Список счетчиков ущерба собирается в _damage_count_columns как объединение DISTRIBUTION_GROUPS[2][1] и тематических колонок _DAMAGE_THEME_COLUMNS.",
            "Базовая агрегация выполняется _collect_positive_column_counts (app/dashboard/distribution.py).",
            "SQL-паттерн для каждой метрики: COALESCE(SUM(CASE WHEN COALESCE(<numeric_expression>,0) > 0 THEN 1 ELSE 0 END),0).",
            "numeric_expression строится через _numeric_expression_for_column: очистка пробелов и запятых, regex-проверка числа, безопасное приведение к double precision.",
            "Категориальные элементы ущерба строятся _build_damage_category_items: пары destroyed/damaged из DAMAGE_PAIR_COLUMNS плюс standalone-колонки.",
            "Тематические элементы строятся _build_damage_theme_items: агрегация по укрупненным темам (недвижимость, площадь, техника, сельхозпотери, животные, прямой ущерб).",
            "Графики damage-режима: _build_damage_overview_chart, _build_damage_pairs_chart, _build_damage_standalone_chart, _build_damage_share_chart.",
            "Сортировка категорий производится по убыванию value; используется top-N отфильтрованных элементов для визуализации.",
        ],
    )

    add_heading(doc, "4. Виджеты распределения distribution: измерения и алгоритм группировки", level=1)
    add_bullets(
        doc,
        [
            "Распределение строит _build_distribution_chart(selected_tables, selected_year, group_column, grouped_counts=None).",
            "Если grouped_counts передан из grouped bundle, SQL не выполняется повторно и используется уже готовая агрегация label -> count.",
            "Если grouped_counts не передан, выполняется SQL по таблицам: SELECT label, COUNT(*) FROM table WHERE <year filter> GROUP BY label ORDER BY fire_count DESC.",
            "Нормализация label: COALESCE(NULLIF(TRIM(CAST(column AS TEXT)), ''), 'Не указано').",
            "Фильтрация по году строится _build_year_filter_clause: либо TRUE, либо year_expression = :selected_year, либо исключение таблицы из расчета.",
            "Алгоритм построения результата: агрегирование в defaultdict(int) -> сортировка по убыванию -> top 12 элементов -> value_display формат.",
            "Для group_column == RISK_CATEGORY_COLUMN используется pie-вариант Plotly; для остальных бар-чарт.",
            "Ранжирование для блока rankings использует _build_rankings: top_distribution, top_tables, recent_years.",
        ],
    )

    add_heading(doc, "5. Impact timeline: построение временной шкалы и периодная агрегация", level=1)
    add_bullets(
        doc,
        [
            "SQL-шаблон таймлайна задает _build_impact_timeline_query (app/dashboard/data_access.py).",
            "Для таблиц с датой: date_value = parsed date expression, GROUP BY date_value.",
            "Для таблиц без даты, но с table_year: используется MAKE_DATE(table_year, 1, 1) как surrogate date.",
            "Суммируемые метрики таймлайна: deaths, injuries, evacuated, evacuated_children, rescued_children.",
            "Коллектор _collect_impact_timeline_rows объединяет запросы таблиц через UNION ALL.",
            "Пост-агрегация выполняется _build_combined_impact_timeline_chart: группировка по date_key, суммирование метрик по одной дате.",
            "Интегральное значение точки рассчитывается как deaths + injuries + evacuated_adults + evacuated_children + rescued_children.",
            "Сортировка точек производится по date_key в возрастающем порядке для корректной временной оси.",
        ],
    )

    add_heading(doc, "6. Двухуровневый кэш: metadata 300 сек и data 120 сек", level=1)
    add_bullets(
        doc,
        [
            "Кэш уровня метаданных: _DASHBOARD_METADATA_CACHE = CopyingTtlCache(ttl_seconds=METADATA_CACHE_TTL_SECONDS).",
            "Кэш уровня данных: _DASHBOARD_CACHE = CopyingTtlCache(ttl_seconds=DASHBOARD_CACHE_TTL_SECONDS).",
            "TTL-значения заданы в app/domain/analytics_metadata.py и реэкспортируются через app/statistics_constants.py: METADATA_CACHE_TTL_SECONDS=300, DASHBOARD_CACHE_TTL_SECONDS=120.",
            "Оба кэша используют freeze_mutable_payload/clone_mutable_payload, чтобы данные в кэше были иммутабельным снимком и не портились при внешней мутации.",
            "Ключ metadata-кэша: tuple(sorted(select_user_table_names(list(get_table_signature_cached())))).",
            "Ключ data-кэша: (_metadata_table_names(metadata), table_name, year, normalized_group_column, horizon_days).",
            "Два уровня кэша нужны для разделения стабильных и динамических слоев: каталог таблиц меняется реже, чем пользовательские комбинации фильтров.",
            "В get_dashboard_data реализована двойная проверка кэша: сначала raw cache_key из запроса, затем resolved_cache_key после канонизации фильтров.",
        ],
    )

    add_heading(doc, "7. Инвалидация кэша дашборда при смене таблиц БД", level=1)
    add_bullets(
        doc,
        [
            "Автопроверка выполняется в _collect_dashboard_metadata_cached: если table_signature в metadata не совпадает с текущими table_names, кэши metadata и data очищаются и метаданные пересобираются.",
            "Явная инвалидация: _invalidate_dashboard_caches вызывает invalidate_db_metadata_cache и clear для обоих dashboard-кэшей.",
            "Централизованный оркестратор app/runtime_invalidation.py включает invalidator для app.dashboard.cache._invalidate_dashboard_caches.",
            "Изменение таблиц через app/table_operations.py вызывает invalidate_table_related_caches, что приводит к сбросу metadata + dashboard + кэшей других сервисов.",
            "Импорт и очистка таблиц в app/services/pipeline_service.py вызывают invalidate_runtime_caches после успешного завершения шага.",
            "Таким образом, при создании/удалении/замене таблиц stale-данные дашборда не сохраняются и следующие запросы получают пересобранный payload.",
        ],
    )

    add_heading(doc, "8. Взаимодействие с statistics_constants", level=1)
    add_bullets(
        doc,
        [
            "TTL-константы: DASHBOARD_CACHE_TTL_SECONDS, METADATA_CACHE_TTL_SECONDS управляют временем жизни кэша.",
            "Группы распределений: DISTRIBUTION_GROUPS определяют доступные group_column и отдельный режим группы ущерба.",
            "Причинные колонки: CAUSE_COLUMNS используются в резолвере причин и в reuse-логике distribution.",
            "Дата и площадь: DATE_COLUMN и AREA_COLUMN участвуют в year/month/date и area-бакет выражениях.",
            "Конфиг impact-метрик: IMPACT_METRIC_CONFIG управляет поиском и агрегацией колонок deaths/injuries/evacuated и др.",
            "Словари подписей: COLUMN_LABELS, DAMAGE_OVERVIEW_LABELS формируют заголовки и подписи виджетов.",
            "MONTH_LABELS задает человекочитаемые подписи месяцев для seasonal/monthly визуализаций.",
            "DAMAGE_PAIR_COLUMNS и DAMAGE_STANDALONE_COLUMNS определяют состав пар destroyed/damaged и standalone-метрик ущерба.",
        ],
    )

    add_heading(doc, "Приложение: SQL-паттерны и pandas-операции", level=1)
    add_paragraph(doc, "SQL-паттерны, используемые в app/dashboard:")
    add_bullets(
        doc,
        [
            "UNION ALL для объединения подзапросов по нескольким таблицам.",
            "GROUP BY GROUPING SETS для мульти-агрегации разных измерений за один проход.",
            "CASE WHEN GROUPING(...) = 0 THEN ... END для маркировки metric_kind.",
            "COALESCE + NULLIF + TRIM для нормализации строковых меток.",
            "Регулярные выражения и CAST в _numeric_expression_for_column для безопасного парсинга чисел.",
            "Параметризованный фильтр года :selected_year.",
            "MAKE_DATE(year, 1, 1) как fallback для таблиц без поля даты.",
        ],
    )
    add_paragraph(doc, "pandas-операции:")
    add_bullets(
        doc,
        [
            "В модулях app/dashboard прямой импорт pandas отсутствует, расчеты выполнены SQLAlchemy + Python collections.",
            "Функциональные эквиваленты pandas groupby.size реализованы через SQL GROUP BY + COUNT.",
            "Эквивалент pandas fillna реализован через SQL COALESCE.",
            "Эквивалент pandas to_numeric(errors='coerce') реализован через regex + CASE + cast в _numeric_expression_for_column.",
            "Эквивалент pandas cumsum для накопленной площади реализован вручную в _collect_cumulative_area_rows через running_total по отсортированным дням.",
            "Форматирование и сортировка результата выполняются Python-операциями sorted, defaultdict, list comprehensions.",
        ],
    )

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    question_marks = count_question_marks_in_docx_text_nodes(OUTPUT_PATH)
    print(f"DOCX_CREATED={OUTPUT_PATH}")
    print(f"TEXT_QUESTION_MARKS={question_marks}")


if __name__ == "__main__":
    main()
