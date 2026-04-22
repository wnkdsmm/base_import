from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / "fire_map_module_analysis.docx"


def set_font(run, *, size: float = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def configure_styles(document: Document) -> None:
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "List Bullet"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)


def add_paragraph(document: Document, text: str, *, bold: bool = False, align=None) -> None:
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, bold=bold)
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_after = Pt(6)


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_font(run, size=14, bold=True)
    paragraph.paragraph_format.space_after = Pt(6)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_font(run)
    paragraph.paragraph_format.space_after = Pt(3)


def build_document() -> Document:
    document = Document()
    configure_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Технический анализ модулей app/services/fire_map_service.py и core/mapping/ проекта "Fire Data Pipeline"')
    set_font(run, size=14, bold=True)

    add_paragraph(
        document,
        "Документ описывает фактическую реализацию блока Fire-map по исходному коду проекта. Анализ охватывает "
        "сервисный слой `app/services/fire_map_service.py`, пакет `core/mapping/` и шаг пайплайна "
        "`core/processing/steps/create_fire_map.py`.",
    )

    add_heading(document, "1. Технология карты и структура слоёв")
    add_paragraph(
        document,
        "Интерактивная карта строится на OpenLayers (`ol` v8.2.0) с базовым тайловым слоем OpenStreetMap. "
        "Подключение выполняется в `core/mapping/mixins/template_layout.py` через CDN-ресурсы "
        "`ol.css` и `ol.js`. Bootstrap используется для UI-контейнеров и вкладок.",
    )
    add_paragraph(
        document,
        "Формат вывода — статический HTML-файл `fires_map.html`, который генерируется методом "
        "`MapCreator.create_map(...)` и сохраняется в output-папку выбранной таблицы. Помимо HTML, "
        "параллельно формируются аналитические артефакты: `fires_map_analysis.json` и `fires_map_analysis.md`.",
    )
    add_paragraph(
        document,
        "Основная структура слоёв в generated HTML (формируется в `template_scripts.py`) следующая:",
    )
    add_bullet(document, "Базовый слой: `new ol.layer.Tile({ source: new ol.source.OSM() })`.")
    add_bullet(document, "Инциденты-пожары по категориям: vector-слои `deaths`, `injured`, `children`, `evacuated`, `other`.")
    add_bullet(document, "Тепловой слой плотности (`heatmap`): `ol.layer.Heatmap`.")
    add_bullet(document, "Hotspot-слой (`hotspots`): point vector-layer с rank и risk-tone.")
    add_bullet(document, "DBSCAN-кластеры (`clusters`): point vector-layer.")
    add_bullet(document, "Зоны риска (`risk_zones`): polygon vector-layer.")
    add_bullet(document, "Приоритетные территории (`priorities`): point vector-layer.")
    add_paragraph(
        document,
        "Включение/выключение слоёв реализовано через checkboxes в фильтр-панели (`layer-filter`, `category-filter`). "
        "Логика переключения видимости находится в `build_filter_script_lines`.",
    )

    add_heading(document, "2. Какие данные реально отображаются на карте")
    add_paragraph(
        document,
        "На карте выводятся данные из таблицы пожаров после чистки координат и отбора валидных строк. "
        "Базовая сущность — GeoJSON `Feature` типа `Point`, где в `properties` размещаются popup-поля, "
        "категория инцидента и дополнительные атрибуты записи.",
    )
    add_paragraph(
        document,
        "Фактический набор отображаемых аналитических сущностей (из `SpatialAnalyticsPayload`):",
    )
    add_bullet(document, "`heatmap.points`: точки плотности с весом `weight`.")
    add_bullet(document, "`hotspots`: ранжированные hotspot-точки.")
    add_bullet(document, "`dbscan.clusters`: кластеры пространственной концентрации.")
    add_bullet(document, "`risk_zones`: полигоны зон повышенного риска.")
    add_bullet(document, "`priority_territories`: приоритетные территории (точки-центроиды).")
    add_paragraph(
        document,
        "Важно: отдельный модуль `app/services/access_points/` не подключается напрямую к Fire-map. "
        "В текущем коде слой «точек доступа» из этого модуля на карту не передаётся. "
        "Вместо него используются `priority_territories` из аналитики `core/mapping`.",
    )

    add_heading(document, "3. Архитектура MapCreator и миксинов")
    add_paragraph(
        document,
        "Класс `MapCreator` реализован в `core/mapping/creator.py` как композиция миксинов через множественное "
        "наследование: "
        "`MapCreatorUtilityMixin`, `MapCreatorAnalyticsMixin`, `MapCreatorTemplateMixin`, `MapCreatorExportMixin`. "
        "Базовый orchestrator-метод `create_map(...)` управляет pipeline и делегирует специализированные этапы миксинам.",
    )
    add_paragraph(
        document,
        "Паттерн по сути сочетает `Template Method` (единый сценарий в `create_map`) и модульную декомпозицию "
        "поведения через mixin-слои. Это уменьшает монолитность: аналитика, рендеринг HTML, вспомогательные "
        "геометрические/форматные функции и экспорт разделены по файлам.",
    )
    add_paragraph(
        document,
        "Роли ключевых миксинов:",
    )
    add_bullet(document, "`MapCreatorUtilityMixin`: парсинг дат/времени, расстояния, уровни риска, popup-формат, геометрия окружности.")
    add_bullet(document, "`MapCreatorAnalyticsMixin`: сбор `ProcessedRecord`, построение hotspot/DBSCAN/risk_zones/priority_territories.")
    add_bullet(document, "`MapCreatorTemplateMixin`: генерация HTML-страницы и табов, подготовка layer payload для JS.")
    add_bullet(document, "`MapCreatorExportMixin`: формирование `fires_map_analysis.json` и `fires_map_analysis.md`.")
    add_paragraph(
        document,
        "Дополнительно в конструктор `MapCreator` внедряются `ColumnFinder` и `DataCleaner` (из `data_access.py`), "
        "что упрощает подмену/тестирование этих зависимостей без изменения логики рендера.",
    )

    add_heading(document, "4. Алгоритм hotspot-зон и параметры плотности")
    add_paragraph(
        document,
        "Hotspot-вычисление в Fire-map выполняется через `build_hotspots_from_dated_records(...)` "
        "(`analytics_hotspots.py`), которая делегирует расчёт в `app/services/forecast_risk/geo.py::_build_geo_prediction` "
        "с горизонтом `planning_horizon_days=30`.",
    )
    add_paragraph(
        document,
        "Метод hotspot в текущей реализации — не KDE и не DBSCAN, а grid-based scoring по ячейкам. "
        "Координаты бьются на сетку: "
        "`key = (floor(latitude / cell_size), floor(longitude / cell_size))`. "
        "Размер ячейки `cell_size` адаптивный (`_derive_geo_cell_size`): "
        "0.05, 0.08, 0.12, 0.20 или до 0.60 градуса в зависимости от географического охвата.",
    )
    add_paragraph(
        document,
        "Для каждого инцидента в ячейке считается вклад: "
        "`recency_weight = max(0.2, 1 - min(age_days, GEO_LOOKBACK_DAYS) / GEO_LOOKBACK_DAYS)`, "
        "`month_weight = 1 + 0.35 * (future_month_freq)`, "
        "`weekday_weight = 1 + 0.20 * (future_weekday_freq)`, "
        "`score = recency_weight * month_weight * weekday_weight`. "
        "Константа `GEO_LOOKBACK_DAYS` = 180.",
    )
    add_paragraph(
        document,
        "Итоговый raw-risk ячейки: "
        "`raw_risk = cell_score * (1 + log1p(incidents) * 0.22) * (0.85 + 0.15 * freshness)`. "
        "Далее риск нормируется к максимуму до шкалы 0-100. "
        "В топ hotspot попадают первые `MAX_GEO_HOTSPOTS` (8) ячеек.",
    )
    add_paragraph(
        document,
        "Формирование зон риска (`risk_zones`) выполняется в `analytics_priority.py::build_spatial_risk_zones` "
        "комбинацией DBSCAN-кластеров и hotspot-кандидатов. DBSCAN параметры (`analytics_dbscan.py`):",
    )
    add_bullet(document, "`eps_km` оценивается как `max(0.9, percentile70(k-neighbors distance) * 1.15)`.")
    add_bullet(document, "`min_samples = max(4, min(8, round(log2(n+1)+2)))`.")
    add_bullet(document, "DBSCAN включается только при `len(records) >= 8` и доступном `sklearn`.")
    add_paragraph(
        document,
        "Кандидаты hotspot/DBSCAN дедуплицируются по расстоянию "
        "`distance < max(radius_a, radius_b) * 0.75`, сортируются по `(risk_score, support_count)`, "
        "и берутся первые 6 зон. Для каждой зоны строится polygon окружности через "
        "`_build_circle_polygon(...)` (36 сегментов).",
    )

    add_heading(document, "5. Шаблонизация карты и HTML/JS-структура")
    add_paragraph(
        document,
        "В архитектуре есть два уровня шаблонизации:",
    )
    add_bullet(document, "Jinja2-уровень приложения: `app/templates/fire_map.html` и `fire_map_error.html`.")
    add_bullet(document, "Генерация HTML карты в `core/mapping`: строковые template-фрагменты без Jinja2.")
    add_paragraph(
        document,
        "Jinja2-шаблон `fire_map.html` строит страницу вокруг `iframe` на `/fire-map/embed` с параметром `table_name`, "
        "показывает executive brief и территориальные карточки. Сам интерактивный map-canvas внутри iframe "
        "возвращается как готовый HTML (`HTMLResponse(map_html)` в `pages.py::fire_map_embed`).",
    )
    add_paragraph(
        document,
        "Внутренний HTML карты формируется функцией `generate_html(...)` (`template_layout.py`) и содержит:",
    )
    add_bullet(document, "HEAD: Bootstrap + OpenLayers CDN, встроенный CSS.")
    add_bullet(document, "BODY: tab-pane контейнеры, map div, filter-panel, analytics-panel.")
    add_bullet(document, "Inline JavaScript: инициализация `ol.Map`, создание всех слоёв, popup overlay, обработчики фильтров.")
    add_paragraph(
        document,
        "Данные в JS передаются как JSON-константы, сериализованные через безопасную функцию "
        "`_json_for_script` (экранирует `<`, `>`, `&`, `U+2028`, `U+2029`). "
        "Слои передаются как GeoJSON `FeatureCollection` (`heatmap`, `hotspots`, `clusters`, `risk_zones`, `priorities`).",
    )

    add_heading(document, "6. Кэширование HTML и поведение при смене данных")
    add_paragraph(
        document,
        "В `app/services/fire_map_service.py` используются два TTL-кэша (`CopyingTtlCache`):",
    )
    add_bullet(document, "`_FIRE_MAP_CACHE` (300 сек): хранит HTML-контент карты (`build_fire_map_html`).")
    add_bullet(document, "`_FIRE_MAP_BRIEF_CACHE` (300 сек): хранит `brief` и `risk_prediction` (`get_fire_map_page_context`).")
    add_paragraph(
        document,
        "Ключ кэша для полной карты — строка `normalized_table_name`. "
        "Ключ для brief-кэша — `selected_table`. Версионирования ключей и контроля изменения данных таблицы нет; "
        "поэтому при обновлении данных в БД кэш будет отдавать старую версию до истечения TTL или ручной очистки.",
    )
    add_paragraph(
        document,
        "Ручная инвалидация выполняется через `clear_fire_map_cache()`, которая очищает оба кэша. "
        "При cache miss выполняется полный pipeline: загрузка данных -> `CreateFireMapStep` -> генерация HTML -> чтение файла -> set в кэш.",
    )

    add_heading(document, "7. Краткая версия карты (brief) и отличие от полной")
    add_paragraph(
        document,
        "Краткая версия реализована как executive brief, а не как урезанный HTML map-файл. "
        "Она строится в `get_fire_map_page_context(...)` через вызов "
        "`build_decision_support_payload(...)` (`forecast_risk`) и `build_executive_brief_from_risk_payload(...)`.",
    )
    add_paragraph(
        document,
        "Содержание brief: `lead`, `top_territory_label`, `priority_reason`, `action_label`, "
        "`confidence_score_display`, список top-территорий и заметки качества данных. "
        "Этот блок используется на странице `/fire-map` для руководительского резюме до/вместе с просмотром карты.",
    )
    add_paragraph(
        document,
        "Полная версия — это `fires_map.html` с OpenLayers-слоями, фильтрами и popups, доступная через endpoint "
        "`/fire-map/embed`. Таким образом, brief отвечает за narrative и decision-support summary, а full map — "
        "за геовизуализацию и интерактивную аналитику.",
    )

    add_heading(document, "8. Взаимодействие с create_fire_map.py")
    add_paragraph(
        document,
        "Связка сервиса с пайплайном шагов выглядит так:",
    )
    add_bullet(document, "`fire_map_service.build_fire_map_html(table_name)` создаёт `Settings` и вызывает `CreateFireMapStep().run(...)`.")
    add_bullet(document, "`CreateFireMapStep.run(...)` формирует `Config(output_dir=...)` и загружает данные через `load_fire_map_source(...)`.")
    add_bullet(document, "`fire_map_loader.load_fire_map_source(...)` определяет реальные колонки, делает SQL `SELECT ... LIMIT :limit` и возвращает DataFrame.")
    add_bullet(document, "`MapCreator.create_map({table_name: df})` готовит таблицу, строит spatial analytics и генерирует `fires_map.html`.")
    add_bullet(document, "После генерации шаг возвращает путь к HTML; сервис читает файл и кэширует строку HTML.")
    add_paragraph(
        document,
        "Ключевые классы/методы взаимодействия:",
    )
    add_bullet(document, "`app/services/fire_map_service.py`: `build_fire_map_html`, `get_fire_map_page_context`, `clear_fire_map_cache`.")
    add_bullet(document, "`core/processing/steps/create_fire_map.py`: `CreateFireMapStep.run`.")
    add_bullet(document, "`core/processing/steps/fire_map_loader.py`: `load_fire_map_source`.")
    add_bullet(document, "`core/mapping/creator.py`: `MapCreator._prepare_table_data`, `MapCreator.create_map`.")
    add_bullet(document, "`core/mapping/mixins/*`: analytics/template/export utilities для расчёта и рендера.")
    add_paragraph(
        document,
        "Форматы данных в связке:",
    )
    add_bullet(document, "Input: `pandas.DataFrame` инцидентов из PostgreSQL.")
    add_bullet(document, "Intermediate: TypedDict payload (`ProcessedRecord`, `SpatialAnalyticsPayload`, `MapTablePayload`).")
    add_bullet(document, "Layer payload: GeoJSON `FeatureCollection`.")
    add_bullet(document, "Output: HTML (`fires_map.html`) + JSON (`fires_map_analysis.json`) + Markdown (`fires_map_analysis.md`).")

    add_heading(document, "Итог")
    add_paragraph(
        document,
        "Fire-map в текущей реализации — это OpenLayers-визуализация инцидентов и пространственной аналитики "
        "(heatmap, hotspots, DBSCAN, зоны риска, приоритетные территории) с двухуровневой подачей: executive brief "
        "и интерактивный embed. Кэширование организовано по таблице с TTL 300 секунд без контроля версии данных, "
        "поэтому при частых изменениях таблиц требуется либо ожидание TTL, либо вызов `clear_fire_map_cache()`.",
    )

    return document


def verify_document_text(path: Path) -> tuple[int, int]:
    with ZipFile(path, "r") as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml)
    joined = "\n".join(texts)
    return joined.count("?"), len(texts)


def main() -> None:
    document = build_document()
    document.save(OUTPUT_PATH)
    question_marks, text_nodes = verify_document_text(OUTPUT_PATH)
    print(f"saved={OUTPUT_PATH}")
    print(f"text_question_marks={question_marks}")
    print(f"text_nodes={text_nodes}")


if __name__ == "__main__":
    main()
