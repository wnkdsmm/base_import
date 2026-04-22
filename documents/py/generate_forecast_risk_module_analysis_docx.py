from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / "forecast_risk_module_analysis.docx"


def set_font(run, *, size: float = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def configure_styles(document: Document) -> None:
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)


def add_paragraph(document: Document, text: str, *, bold: bool = False, align=None) -> None:
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, bold=bold)
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_after = Pt(6)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_font(run)
    paragraph.paragraph_format.space_after = Pt(3)


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_font(run, size=14, bold=True)
    paragraph.paragraph_format.space_after = Pt(6)


def build_document() -> Document:
    document = Document()
    configure_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Технический анализ модуля app/services/forecast_risk/ проекта "Fire Data Pipeline"')
    set_font(run, size=14, bold=True)

    add_paragraph(
        document,
        "Документ описывает фактическую реализацию модуля `forecast_risk` по исходному коду проекта. "
        "Здесь важно различать территориальный ranking-риск и классический прогноз временного ряда: "
        "`forecast_risk` не предсказывает число пожаров по дням, а ранжирует территории по приоритету действий.",
    )

    add_heading(document, "1. Алгоритм расчёта территориального риска")
    add_paragraph(
        document,
        "Ядро расчёта реализовано в `app/services/forecast_risk/scoring_compute.py::_build_territory_rows`. "
        "Сначала записи группируются по территории в `scoring_history.py::_collect_territory_buckets`, затем для каждого "
        "bucket рассчитываются нормализованные факторы риска, логистики и водоснабжения, после чего строятся "
        "компонентные баллы и итоговый `risk_score` в диапазоне 1-99.",
    )
    add_paragraph(
        document,
        "Входные данные формируются в `data_impl.py::_load_risk_records`. Из БД извлекаются и нормализуются: "
        "дата, район, причина, категория объекта, метка территории, тип поселения, расстояние до пожарной части, "
        "факты наличия воды, времена сообщения/обнаружения/прибытия, последствия, ущерб, погибшие, пострадавшие, "
        "категория риска объекта. На их основе строятся производные признаки: `response_minutes`, `long_arrival`, "
        "`heating_season`, `night_incident`, `victims_present`, `major_damage`, `severe_consequence`, `risk_category_score`.",
    )
    add_paragraph(
        document,
        "На уровне bucket по территории накапливаются: число пожаров `incidents`, взвешенная история `weighted_history`, "
        "seasonal sums по месяцам и дням недели, число тяжёлых последствий, пострадавших, случаев крупного ущерба, "
        "доля ночных пожаров, отопительный контур, данные по расстоянию, воде и фактическому времени прибытия.",
    )
    add_paragraph(
        document,
        "Основные формулы `scoring_compute.py::_normalized_risk_fields` имеют вид:",
    )
    add_bullet(
        document,
        "Историческое давление: `history_pressure = incidents / max_incidents`.",
    )
    add_bullet(
        document,
        "Давление по свежести: `recency_pressure = weighted_history / max_weighted`.",
    )
    add_bullet(
        document,
        "Сезонное совпадение: `seasonal_alignment = clamp(0.62 * month_component + 0.38 * weekday_component, 0, 1)`, "
        "где month_component и weekday_component получаются из сумм сезонных совпадений, делённых на число событий.",
    )
    add_bullet(
        document,
        "Базовый fire signal: `base_fire_signal = clamp(1 - exp(-(recent_incidents / recent_window_days)), 0.08, 0.72)`.",
    )
    add_bullet(
        document,
        "Ожидаемая интенсивность на горизонте: "
        "`expected_value = (weighted_history / history_days) * horizon_days * (0.72 + base_fire_signal)`.",
    )
    add_bullet(
        document,
        "Вероятность пожара: "
        "`fire_probability = clamp(max(1 - exp(-expected_value), base_fire_signal * min(0.94, 0.22 + history_pressure * 0.52)), 0.02, 0.995)`.",
    )
    add_bullet(
        document,
        "Вероятность тяжёлых последствий: "
        "`severe_probability = clamp(0.46 * severe_rate + 0.26 * casualty_pressure + 0.18 * damage_pressure + 0.10 * risk_factor, 0.02, 0.98)`.",
    )
    add_bullet(
        document,
        "Вероятность проблем с прибытием: "
        "`arrival_probability = clamp(0.24 * long_arrival_rate + 0.18 * response_pressure + 0.22 * travel_time_pressure + "
        "0.22 * service_coverage_gap + 0.14 * service_zone_pressure, 0.03, 0.98)`.",
    )
    add_bullet(
        document,
        "Вероятность дефицита воды: "
        "`water_deficit_probability = clamp(0.76 * water_gap_rate + 0.24 * tanker_dependency, 0.02, 0.99)`.",
    )
    add_paragraph(
        document,
        "Итоговый риск формируется как сумма вкладов четырёх компонентов. Для каждого компонента сначала считается "
        "`score = 100 * weighted_sum / total_signal_weight`, затем `contribution = score * component_weight`. "
        "Общий показатель: `risk_score = clamp(sum(component_contribution), 1, 99)`. "
        "То есть это прозрачная линейная аддитивная модель со шкалой 0-100, а не ML-классификатор.",
    )

    add_heading(document, "2. Профили весов")
    add_paragraph(
        document,
        "Профили весов описаны в `profiles.py`. В коде есть три режима:",
    )
    add_bullet(
        document,
        "`expert` — базовый экспертный профиль. Компонентные веса: "
        "`fire_frequency = 0.34`, `consequence_severity = 0.24`, "
        "`long_arrival_risk = 0.24`, `water_supply_deficit = 0.18`.",
    )
    add_bullet(
        document,
        "`adaptive` — тот же объяснимый профиль, но с возможностью автоматической калибровки по historical windows. "
        "Это режим по умолчанию: `DEFAULT_RISK_WEIGHT_MODE = \"adaptive\"`.",
    )
    add_bullet(
        document,
        "`calibratable` — шаблон для ручной настройки и экспериментов с весами без изменения структуры компонентов.",
    )
    add_paragraph(
        document,
        "Компоненты у всех профилей одинаковы: "
        "`fire_frequency`, `consequence_severity`, `long_arrival_risk`, `water_supply_deficit`. "
        "Меняются только веса компонентов и их статус. Состав сигналов внутри компонента остаётся фиксированным, "
        "что сохраняет объяснимость модели.",
    )
    add_paragraph(
        document,
        "Для сельских территорий применяется `rural_weight_shift`. В экспертном профиле он равен: "
        "`fire_frequency -0.03`, `consequence_severity -0.02`, `long_arrival_risk +0.03`, "
        "`water_supply_deficit +0.02`. После сдвига веса нормализуются через `resolve_component_weights`, "
        "то есть для сельской местности логистика и вода автоматически усиливаются.",
    )
    add_paragraph(
        document,
        "Автокалибровка профиля реализована в `profile_resolution.py::resolve_weight_profile_for_records`. "
        "Она запускается только если есть хотя бы `MIN_CALIBRATION_WINDOWS = 4` исторических окна. "
        "Генератор кандидатов `_generate_weight_candidates` строит: экспертный профиль, сбалансированный профиль, "
        "фокус-профили на каждый компонент и пары `shift donor -> receiver` со сдвигами `0.04` и `0.08`.",
    )
    add_paragraph(
        document,
        "Кандидат оценивается на historical windows. Для него вычисляется "
        "`objective = ranking_objective(aggregate)` и регуляризованная целевая функция "
        "`regularized_objective = objective - 0.08 * weight_distance(candidate, expert)`."
        "Если выигрыш лучшего кандидата над экспертным профилем меньше `MIN_CALIBRATION_IMPROVEMENT = 0.015`, "
        "экспертные веса удерживаются. Иначе профиль получает статус `Калиброван по истории`.",
    )

    add_heading(document, "3. Геопривязка риска")
    add_paragraph(
        document,
        "Геослой реализован в `geo.py::_build_geo_prediction`. Здесь нет пространственного join с внешними "
        "административными полигонами. Геопривязка строится по точкам пожаров, у которых есть `latitude` и `longitude`, "
        "а затем по этим точкам формируются регулярные ячейки сетки.",
    )
    add_paragraph(
        document,
        "Размер ячейки выбирается эвристически функцией `_derive_geo_cell_size(latitudes, longitudes)` по ширине охвата координат. "
        "Например, при малом span используется `0.05`, при среднем `0.08` или `0.12`, при большом охвате вплоть до `0.60` градуса.",
    )
    add_paragraph(
        document,
        "Точка относится к ячейке по ключу "
        "`key = (floor(latitude / cell_size), floor(longitude / cell_size))`. "
        "То есть административная территория напрямую не вычисляется из геометрии; вместо этого в ячейке накапливаются "
        "счётчики `districts`, `causes`, `object_categories`, и затем район определяется как наиболее частый "
        "`dominant_district` через `Counter.most_common(1)`.",
    )
    add_paragraph(
        document,
        "Следовательно, связь с административной территорией реализована не геометрическим overlay, а по доминирующему районному "
        "атрибуту среди пожаров, попавших в ячейку. Это более лёгкая, но и более грубая схема пространственной привязки.",
    )

    add_heading(document, "4. Ранжирование территорий")
    add_paragraph(
        document,
        "После расчёта `risk_score` список территорий сортируется в `scoring_compute.py` правилом "
        "`territory_rows.sort(key=lambda item: (item['risk_score'], item['history_pressure']), reverse=True)`. "
        "Основной критерий сравнения — итоговый риск, вторичный — историческое давление.",
    )
    add_paragraph(
        document,
        "Для карты ranking ячеек строится отдельно. В `geo.py` сначала считается `raw_risk`, затем ячейки сортируются по "
        "`(raw_risk, incidents)` по убыванию. После этого риск приводится к шкале 0-100 через "
        "`risk_score = (raw_risk / max_risk) * 100`.",
    )
    add_paragraph(
        document,
        "На UI ranking дополняется поясняющим контекстом в `scoring_ranking.py::_attach_ranking_context`: "
        "позиция в рейтинге, разрыв с соседней территорией, отставание от лидера и два самых сильных компонента. "
        "Это делает ранжирование управленчески интерпретируемым, а не просто сортировкой по числу.",
    )

    add_heading(document, "5. Надёжность оценки")
    add_paragraph(
        document,
        "Оценка надёжности реализована в `reliability.py::_attach_ranking_reliability`. "
        "Она комбинирует три источника уверенности: качество паспорта данных, результат historical validation "
        "и локальную устойчивость позиции конкретной территории в ранжировании.",
    )
    add_paragraph(
        document,
        "Локальная поддержка строится так:",
    )
    add_bullet(document, "`history_support = min(1, history_count / 8)`.")
    add_bullet(
        document,
        "Для лидера: `margin_support = min(1, ranking_gap_to_next / 8)`. "
        "Для остальных: `margin_support = 1 - min(1, ranking_gap_to_top / 12)`.",
    )
    add_bullet(
        document,
        "`local_support = clamp(0.58 * margin_support + 0.42 * history_support, 0.15, 1.0)`.",
    )
    add_paragraph(
        document,
        "Если historical validation готова (`windows_count >= 3`), то "
        "`confidence_norm = clamp(0.42 * passport_score + 0.38 * objective_score + 0.20 * local_support, 0.18, 0.96)`. "
        "Иначе используется упрощённая формула "
        "`confidence_norm = clamp(0.67 * passport_score + 0.33 * local_support, 0.16, 0.88)`.",
    )
    add_paragraph(
        document,
        "После этого для первого места добавляется бонус `+0.03`, а для позиций ниже четвёртой применяется штраф `-0.04`. "
        "Итоговая уверенность переводится в шкалу `0-100` и дискретные статусы: "
        "`Высокая` (от 82), `Рабочая` (от 64), `Умеренная` (от 46), `Ограниченная` ниже.",
    )
    add_paragraph(
        document,
        "При малом числе событий уверенность снижается сразу по двум каналам: уменьшается `history_support`, "
        "а также часто не набирается достаточное число historical windows для полноценной validation-проверки. "
        "Это корректный защитный механизм против переуверенного ranking-а на редких данных.",
    )

    add_heading(document, "6. Декомпозиция риска по факторам")
    add_paragraph(
        document,
        "Декомпозиция строится через четыре компонента риска:",
    )
    add_bullet(
        document,
        "`fire_frequency` — повторяемость и сезонная релевантность пожаров: "
        "`predicted_repeat_rate`, `history_pressure`, `recency_pressure`, `seasonal_alignment`, `heating_pressure`.",
    )
    add_bullet(
        document,
        "`consequence_severity` — тяжесть последствий: "
        "`severe_rate`, `casualty_pressure`, `damage_pressure`, `risk_category_factor`, `heating_pressure`.",
    )
    add_bullet(
        document,
        "`long_arrival_risk` — логистический риск: "
        "`long_arrival_rate`, `avg_response_pressure`, `travel_time_pressure`, `service_coverage_gap`, "
        "`service_zone_pressure`, `distance_pressure`.",
    )
    add_bullet(
        document,
        "`water_supply_deficit` — дефицит воды: "
        "`water_gap_rate`, `tanker_dependency`, `rural_context`, `damage_pressure`.",
    )
    add_paragraph(
        document,
        "Для каждого компонента функция `_score_component` формирует список сигналов, сортирует их по `weighted_value`, "
        "считает `score`, `contribution`, `tone`, а затем генерирует `summary`, `rationale` и `driver_text`. "
        "Именно эти поля затем используются в UI для объяснения, почему территория находится высоко в ranking-е.",
    )
    add_paragraph(
        document,
        "Формула, которую видит пользователь, собирается в `scoring_ranking.py::_build_formula_display` и имеет вид "
        "`Компонент A + Компонент B + Компонент C + Компонент D = итоговый риск`. "
        "Кроме того, модуль формирует `action_label`, `action_hint` и до трёх рекомендаций по доминирующим причинам риска.",
    )

    add_heading(document, "7. Взаимодействие с forecasting")
    add_paragraph(
        document,
        "Связь с `forecasting` двусторонняя, но достаточно простая. Модуль `forecasting` вызывает "
        "`forecast_risk.core.build_decision_support_payload(...)` и передаёт туда тот же фильтр: "
        "таблицы, район, причину, категорию объекта, окно истории и горизонт планирования.",
    )
    add_paragraph(
        document,
        "Сам `forecast_risk` не использует прогнозные значения по дням как отдельный числовой ряд. "
        "Он использует горизонт сценарного прогноза только для формирования `planning_horizon_days`. "
        "Именно этот горизонт участвует в `future_dates`, `future_months`, `future_weekdays`, "
        "`future_heating_share` и через них влияет на сезонные веса, `expected_value`, `fire_probability` и гео-оценку.",
    )
    add_paragraph(
        document,
        "Функция `build_risk_forecast_payload(...)` дополнительно подтверждает это: из `forecast_rows` берётся только "
        "длина ряда, то есть `planning_horizon_days = max(1, len(forecast_rows) or 14)`. "
        "Следовательно, влияние `forecasting` на `forecast_risk` идёт через длительность горизонта, а не через подстановку "
        "дневных предсказаний в формулу ranking-а.",
    )

    add_heading(document, "8. Передача результатов на карту")
    add_paragraph(
        document,
        "Прямого вызова `fire_map_service` внутри `app/services/forecast_risk/` не обнаружено. "
        "Модуль сам формирует `geo_prediction` и компактный `geo_summary`, которые затем попадают в итоговый payload "
        "decision-support блока.",
    )
    add_paragraph(
        document,
        "Дальше этот payload используется в `forecasting`: в `assembly_output.py` карта строится вызовом "
        "`deps['build_geo_chart'](geo_prediction)`. То есть фактический путь данных такой: "
        "`forecast_risk.geo -> risk payload -> forecasting assembly -> geo chart`. "
        "Если в диссертации нужно описывать интеграцию с картой, корректнее говорить о передаче в фронтенд-карту через "
        "гео-payload, а не о прямом взаимодействии с отдельным `fire_map_service`.",
    )

    add_heading(document, "9. Библиотеки и операции")
    add_paragraph(
        document,
        "В текущей реализации `forecast_risk` не использует `numpy` и `pandas`. По коду модуля нет импортов `numpy as np` "
        "или `pandas as pd`. Все вычисления выполнены на стандартных структурах Python: `dict`, `list`, `Counter`, "
        "`math.exp`, `math.log1p`, `math.floor`, сортировки списков, простые суммирования и деления.",
    )
    add_paragraph(
        document,
        "Если перечислять конкретные операции, реально присутствующие в коде, то это:",
    )
    add_bullet(document, "SQLAlchemy `text(...)` и `engine.connect()` для чтения данных из БД.")
    add_bullet(document, "`Counter(...)` для подсчёта месяцев, дней недели, районов, причин и категорий.")
    add_bullet(document, "`sum(...)`, `max(...)`, `min(...)`, `sorted(...)`, `sort(...)` для агрегирования и ranking-а.")
    add_bullet(document, "`math.exp(...)` в моделировании вероятности пожара.")
    add_bullet(document, "`math.log1p(...)` в гео-риске и размере маркера.")
    add_bullet(document, "`math.floor(...)` при разбиении координат на grid-cells.")
    add_paragraph(
        document,
        "Поэтому требование «указать numpy/pandas операции» для этого модуля корректнее интерпретировать как "
        "«зафиксировать их отсутствие в текущей реализации».",
    )

    add_heading(document, "10. Вывод")
    add_paragraph(
        document,
        "Модуль `forecast_risk` реализует объяснимую систему приоритизации территорий. "
        "Её основа — компонентный скоринг с прозрачными весами, адаптацией к сельскому контексту, "
        "historical validation ranking-качества и отдельным гео-слоем по координатам пожаров. "
        "Сильная сторона модуля — интерпретируемость и управленческая применимость. Ограничения тоже понятны: "
        "нет `numpy/pandas`, нет геометрического spatial join с административными полигонами и нет прямого вызова `fire_map_service`.",
    )

    return document


def count_text_question_marks(path: Path) -> int:
    with ZipFile(path, "r") as archive:
        xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
    texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml)
    return "".join(texts).count("?")


def main() -> None:
    document = build_document()
    document.save(OUTPUT_PATH)
    print(f"saved={OUTPUT_PATH}")
    print(f"text_question_marks={count_text_question_marks(OUTPUT_PATH)}")


if __name__ == "__main__":
    main()
