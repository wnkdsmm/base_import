# -*- coding: utf-8 -*-
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCUMENTS_DIR = PROJECT_ROOT / "documents"
OUTPUT_PATH = DOCUMENTS_DIR / "Полный_архитектурный_анализ_Fire_Data_Pipeline.docx"


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_section(document: Document, title: str) -> None:
    document.add_heading(title, level=1)


def add_subsection(document: Document, title: str) -> None:
    document.add_heading(title, level=2)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def section_1_idef0_context(document: Document) -> None:
    add_section(document, "1. IDEF0 — Контекстная диаграмма (A-0)")
    document.add_paragraph(
        "Единственный блок A-0: «Система поддержки принятия решений в области пожарной безопасности "
        "сельских территорий»."
    )

    add_subsection(document, "ВХОД (Input)")
    add_bullets(
        document,
        [
            "Пользовательские файлы XLS/XLSX/CSV, загружаемые через API (`/upload`).",
            "Данные о пожарах и последствиях из таблиц PostgreSQL (исходные и `clean_*`).",
            "Атрибутивные данные по территории: район, населённый пункт, координаты, тип объекта, причины пожаров.",
            "Временные ряды событий (дата/время пожара, сезонность, частота, дополнительные признаки для прогноза).",
            "Сигналы качества данных: пропуски, доминирующие значения, низкая вариативность, неполнота геоданных.",
        ],
    )

    add_subsection(document, "УПРАВЛЕНИЕ (Control)")
    add_bullets(
        document,
        [
            "Параметры пользователя из UI/API: таблица, горизонт прогноза, окно истории, фильтры территории/причины/категории.",
            "Пороги очистки данных (`NULL_THRESHOLD`, `DOMINANT_VALUE_THRESHOLD`, `LOW_VARIANCE_THRESHOLD`).",
            "Параметры аналитических блоков: `cluster_count`, `sample_limit`, `sampling_strategy`, `feature_columns`.",
            "Профили и режимы веса риска (adaptive/expert/calibratable) в модуле `forecast_risk`.",
            "Конфигурация кэшей и джобов (TTL, LRU, max_workers, стратегия повторного использования cache key).",
            "Доменные словари и правила сопоставления колонок (обязательные признаки, Natasha matcher, alias-кандидаты).",
        ],
    )

    add_subsection(document, "ВЫХОД (Output)")
    add_bullets(
        document,
        [
            "Очищенные таблицы и артефакты профилирования (`*_profiling.csv`, `*_profiling.xlsx`, `clean_*`, XLSX-экспорт).",
            "JSON-payload аналитики: dashboard, forecasting, ml_model, clustering, access_points, decision support.",
            "Ранжированные территории риска, проблемные точки доступа, кластеры, диагностические метрики качества.",
            "Интерактивная карта пожаров (HTML), а также spatial-экспорт (`fires_map_analysis.json`, `.md`).",
            "Управленческие brief-материалы и рекомендации для приоритизации профилактических мероприятий.",
            "Статусы фоновых задач, логи и результаты в `JobStore` для polling-клиента.",
        ],
    )

    add_subsection(document, "МЕХАНИЗМ (Mechanism)")
    add_bullets(
        document,
        [
            "Backend: FastAPI + Jinja2 (`app/main.py`, `app/routes/*`, `app/templates/*`).",
            "Слой данных: SQLAlchemy Engine (`config/db.py`) и PostgreSQL (URL `postgresql://.../fires`).",
            "Обработка данных: pandas + numpy (импорт, очистка, агрегирование, feature engineering).",
            "Аналитика/ML: scikit-learn, statsmodels (часть count-моделей), статистические метрики качества.",
            "Визуализация: Plotly (графики) и OpenLayers + Bootstrap (интерактивная карта).",
            "Кэширование и state-management: `CopyingTtlCache`, `CopyingLruCache`, in-memory `JobStore`.",
            "Асинхронность: `ThreadPoolExecutor` для тяжёлых аналитических jobs (forecasting/ml/clustering).",
        ],
    )


def section_2_idef0_top_level(document: Document) -> None:
    add_section(document, "2. IDEF0 — Диаграмма верхнего уровня (A0)")
    document.add_paragraph("Функциональная декомпозиция A0 (8 блоков A1–A8).")

    blocks = [
        {
            "id": "A1",
            "name": "Приём данных и импорт в хранилище",
            "input": [
                "Файлы пользователя XLS/XLSX/CSV.",
                "Команда импорта из UI/API.",
            ],
            "control": [
                "Правила чтения форматов (`ImportDataStep`).",
                "Параметры проекта (`Settings`, `project_name`, `output_folder`).",
            ],
            "output": [
                "Исходная таблица в PostgreSQL (`to_sql if_exists='replace'`).",
                "Промежуточный CSV в output-папке.",
                "Job-логи загрузки.",
            ],
            "mechanism": [
                "FastAPI `/upload`, `/import_data`.",
                "pandas `read_excel/read_csv`, SQLAlchemy engine.",
            ],
            "links": [
                "Выход A1 (таблица в БД) -> вход A2, A4, A5, A6, A7, A8.",
            ],
        },
        {
            "id": "A2",
            "name": "Очистка, профилирование и подготовка витрин",
            "input": [
                "Таблица из A1.",
                "Пороговые параметры качества колонок.",
            ],
            "control": [
                "Правила профилирования и защиты обязательных признаков.",
                "Сопоставление колонок через Natasha matcher и реестр mandatory features.",
            ],
            "output": [
                "Отчёты профилирования CSV/XLSX.",
                "Очищенная таблица `clean_*` + XLSX выгрузка.",
                "Список удалённых/сохранённых колонок и причины.",
            ],
            "mechanism": [
                "PipelineStep: `FiresFeatureProfilingStep`, `KeepImportantColumnsStep`, `CreateCleanTableStep`.",
                "pandas + SQL (`CREATE TABLE ... AS SELECT ...`).",
            ],
            "links": [
                "Выход A2 (`clean_*`, profile metadata) -> вход A4, A5, A6, A7, A8.",
            ],
        },
        {
            "id": "A3",
            "name": "Управление метаданными и кэшами",
            "input": [
                "Сигнатура таблиц/колонок БД.",
                "События изменения данных (импорт/очистка/удаление).",
            ],
            "control": [
                "TTL политик кэшей (metadata/service/dashboard/map).",
                "Стратегия invalidation и startup warmup.",
            ],
            "output": [
                "Актуальный metadata cache, dashboard cache, service caches.",
                "Снижение повторных SQL-запросов и времени отклика.",
            ],
            "mechanism": [
                "`app/db_metadata.py`, `app/cache.py`, `app/runtime_invalidation.py`.",
                "CopyingTtlCache/CopyingLruCache, SQLAlchemy inspect.",
            ],
            "links": [
                "A3 обслуживает A4–A8 (предоставляет быстрый доступ к метаданным/результатам).",
            ],
        },
        {
            "id": "A4",
            "name": "Оперативная аналитика и дашборд",
            "input": [
                "Подготовленные таблицы из A2.",
                "Фильтры пользователя (год, таблица, группировка).",
            ],
            "control": [
                "Константы распределений/метрик (`statistics_constants`).",
                "Пороговые правила формирования карточек ущерба и timeline.",
            ],
            "output": [
                "Dashboard payload (карточки, распределения, timeline, графики Plotly).",
                "Краткие управленческие выводы.",
            ],
            "mechanism": [
                "Модули `app/dashboard/*`.",
                "SQL + pandas + Plotly + TTL-кэш дашборда.",
            ],
            "links": [
                "Выход A4 используется как контекст для A5/A6 и интерфейса A8.",
            ],
        },
        {
            "id": "A5",
            "name": "Сценарный прогноз и рекомендации",
            "input": [
                "Исторический ряд пожаров по дням.",
                "Температурный и территориальный контекст, фильтры причин/категорий.",
            ],
            "control": [
                "Горизонт прогноза, окно истории, сценарные параметры пользователя.",
                "Метрики качества сценарного прогноза и backtesting.",
            ],
            "output": [
                "Forecast payload (сценарный прогноз, интервалы, quality, пояснения).",
                "Decision support payload для приоритизации территорий.",
            ],
            "mechanism": [
                "`app/services/forecasting/*` + `app/services/forecast_risk/*`.",
                "SQL-агрегации, bootstrap, stats-проверки, service/job cache.",
            ],
            "links": [
                "Выход A5 (прогноз и риск) -> вход A7 и A8.",
            ],
        },
        {
            "id": "A6",
            "name": "ML-прогноз и валидация моделей",
            "input": [
                "Дневной ряд и признаки из forecasting data-layer.",
                "Фильтры пользователя и горизонт прогноза.",
            ],
            "control": [
                "Выбор методов и стратегий обучения/калибровки интервалов.",
                "Backtesting-конфигурация (rolling windows, горизонты, метрики).",
            ],
            "output": [
                "ML-forecast, интервалы предсказания, feature importance, метрики качества.",
                "Статусы/логи фоновой ML-задачи.",
            ],
            "mechanism": [
                "`app/services/ml_model/*`.",
                "scikit-learn + (опционально) statsmodels + numpy/pandas + LRU cache.",
            ],
            "links": [
                "Выход A6 дополняет A5 и отображается в A8.",
            ],
        },
        {
            "id": "A7",
            "name": "Кластеризация и анализ проблемных точек",
            "input": [
                "Агрегированные признаки территорий и точек доступа.",
                "Пользовательские настройки кластеризации/feature selection.",
            ],
            "control": [
                "Выбор алгоритма/весов, ограничения на k и диагностика устойчивости.",
                "Пороги uncertainty и правила ранжирования access points.",
            ],
            "output": [
                "Кластеры территорий, профили, центроиды, representative points.",
                "Рейтинг проблемных точек доступа и объяснения score-компонентов.",
            ],
            "mechanism": [
                "`app/services/clustering/*`, `app/services/access_points/*`.",
                "KMeans/Agglomerative/Birch, метрики кластеров, pandas scoring pipeline.",
            ],
            "links": [
                "Выход A7 -> вход A8 (картографическая визуализация и аналитические панели).",
            ],
        },
        {
            "id": "A8",
            "name": "Визуализация, API-выдача и взаимодействие с пользователем",
            "input": [
                "Результаты A4, A5, A6, A7.",
                "Команды клиента (браузер, polling jobs, выбор страницы/фильтров).",
            ],
            "control": [
                "Маршрутизация FastAPI, формат JSON UTF-8, правила HTTP-кэширования статики.",
                "Session/job semantics (cookie session_id, статусные эндпоинты).",
            ],
            "output": [
                "HTML-страницы Jinja2, API JSON payload, интерактивная fire-map.",
                "Файлы brief и экспортные артефакты карты.",
            ],
            "mechanism": [
                "`app/main.py`, `app/routes/*`, `core/mapping/*`, `CreateFireMapStep`.",
                "FastAPI, Jinja2, OpenLayers, Plotly, JobStore polling.",
            ],
            "links": [
                "A8 потребляет результаты A4–A7 и отдаёт их пользователю как UI/API/документы.",
            ],
        },
    ]

    for block in blocks:
        add_subsection(document, f"{block['id']}. {block['name']}")
        document.add_paragraph("Входы:")
        add_bullets(document, block["input"])
        document.add_paragraph("Управление:")
        add_bullets(document, block["control"])
        document.add_paragraph("Выходы:")
        add_bullets(document, block["output"])
        document.add_paragraph("Механизм:")
        add_bullets(document, block["mechanism"])
        document.add_paragraph("Связи:")
        add_bullets(document, block["links"])


def section_3_architecture(document: Document) -> None:
    add_section(document, "3. Архитектурная схема взаимодействия модулей")

    add_subsection(document, "Клиент (браузер) ↔ FastAPI роутеры")
    add_bullets(
        document,
        [
            "Страницы: `app/routes/pages.py` рендерит Jinja2-шаблоны (`index`, `forecasting`, `ml_model`, `clustering`, `access_points`, `fire-map`).",
            "API: `app/routes/api*.py` отдаёт JSON-payload для ленивой загрузки данных в UI.",
            "Session: `run_session_json_action` создаёт/переиспользует `session_id` cookie (`fire_monitor_session_id`).",
        ],
    )

    add_subsection(document, "Роутеры ↔ сервисы аналитики")
    add_bullets(
        document,
        [
            "`api_dashboard` -> `app.dashboard.service`.",
            "`api_forecasting` -> `app.services.forecasting` + decision support job.",
            "`api_ml_model` -> `app.services.ml_model` + backtesting job.",
            "`api_clustering` -> `app.services.clustering`.",
            "`api_access_points` -> `app.services.access_points`.",
            "`api_ops` -> `app.services.pipeline_service` (upload/import/profiling).",
        ],
    )

    add_subsection(document, "Сервисы ↔ слой кэша")
    add_bullets(
        document,
        [
            "Общий кэш-примитив: `CopyingTtlCache`/`CopyingLruCache` (`app/cache.py`).",
            "Метаданные БД: `app/db_metadata.py`, TTL 60 сек.",
            "Dashboard cache: metadata TTL 300 сек + data TTL 120 сек.",
            "Forecasting/service caches: TTL 120 сек (payload, metadata bundle, SQL cache).",
            "Clustering/access_points caches: TTL 120 сек.",
            "Fire map HTML и brief caches: TTL 300 сек.",
        ],
    )

    add_subsection(document, "Сервисы ↔ PostgreSQL (через SQLAlchemy)")
    add_bullets(
        document,
        [
            "Единый `engine` создаётся в `config/db.py` через `create_engine(..., pool_pre_ping=True)`.",
            "Импорт данных: pandas `to_sql` в шаге `ImportDataStep`.",
            "Аналитические чтения: SQL через `sqlalchemy.text`, `pd.read_sql`, и query-builders сервисов.",
            "Метаданные таблиц/колонок: `sqlalchemy.inspect(engine)` с кэшированием.",
        ],
    )

    add_subsection(document, "Сервисы ↔ core/processing (пайплайн)")
    add_bullets(
        document,
        [
            "`app/services/pipeline_service.py` запускает шаги `core/processing/pipeline.py`.",
            "Конвейер: import -> profiling -> keep-important-columns -> create-clean-table -> fire-map (по запросу).",
            "После успешных изменений схемы/данных вызывается `invalidate_runtime_caches`.",
        ],
    )

    add_subsection(document, "Асинхронные задачи (jobs) ↔ JobStore ↔ клиент (polling)")
    add_bullets(
        document,
        [
            "Фоновые задачи стартуют через `ThreadPoolExecutor` в `forecasting/jobs.py`, `ml_model/jobs.py`, `clustering/jobs.py`.",
            "Состояние и логи хранятся в in-memory `JobStore` (`app/state.py`) с `RLock`.",
            "Клиент опрашивает `/api/*-jobs/{job_id}` и получает статус (`pending/running/completed/failed`), прогресс и результат.",
        ],
    )


def section_4_stack(document: Document) -> None:
    add_section(document, "4. Стек технологий с обоснованием")

    stack_rows = [
        (
            "FastAPI",
            "HTTP API, роутинг, startup-hook, интеграция с Jinja2/StaticFiles.",
            "Высокая скорость разработки и исполнения, строгая типизация, удобная архитектура для async/sync сервисов.",
            "Flask (+ручная сборка), Django REST Framework, Sanic.",
        ),
        (
            "SQLAlchemy",
            "Единый доступ к PostgreSQL: engine, transactions, text SQL, inspect metadata.",
            "Стандартизованный доступ к БД, контроль подключения, переносимость, зрелый pooling и инструментарий.",
            "psycopg2/psycopg (низкоуровнево), asyncpg, peewee.",
        ),
        (
            "pandas",
            "Импорт/очистка/агрегации табличных данных, профилирование признаков, подготовка фич.",
            "Де-факто стандарт для ETL и аналитики табличных данных; богатый API для Excel/CSV/SQL.",
            "Polars, Dask (для распределённых больших данных), pure SQL-only pipeline.",
        ),
        (
            "scikit-learn",
            "ML-модели, кластеризация, метрики качества, permutation importance.",
            "Унифицированный API, стабильные алгоритмы, хорошая интеграция с pandas/numpy.",
            "XGBoost/LightGBM/CatBoost, H2O, PyTorch/TF (при более сложных моделях).",
        ),
        (
            "numpy",
            "Численные операции, нормализация, линейная алгебра, расчёт метрик и весов.",
            "Высокая производительность и фундамент для большинства научных Python-библиотек.",
            "JAX, CuPy (GPU), pure Python (значительно медленнее).",
        ),
        (
            "scipy",
            "В проекте явно почти не используется напрямую; функциональность в основном закрывается sklearn/numpy.",
            "SciPy остаётся совместимым научным стеком и потенциальной базой для расширения статистических/оптимизационных блоков.",
            "Без SciPy (только numpy/sklearn), statsmodels, специализированные domain-пакеты.",
        ),
        (
            "statsmodels",
            "Опциональный backend count-моделей в ML (`Negative Binomial GLM`), статистическая интерпретируемость.",
            "Сильная статистическая база и интерпретируемые параметры для прикладной научной работы.",
            "scikit-learn GLM-only, PyMC/Stan (байесовские модели), prophet-like инструменты.",
        ),
        (
            "plotly",
            "Интерактивные графики в dashboard/forecasting/ml/clustering/access_points.",
            "Богатые интерактивные визуализации без тяжёлой ручной JS-разработки; JSON-сериализация фигуры в API.",
            "Matplotlib/Seaborn (менее интерактивно в web), ECharts, Vega-Lite.",
        ),
    ]

    for lib_name, role, rationale, alternatives in stack_rows:
        add_subsection(document, lib_name)
        document.add_paragraph(f"Роль: {role}")
        document.add_paragraph(f"Почему выбрана: {rationale}")
        document.add_paragraph(f"Альтернативы: {alternatives}")


def section_5_data_flow(document: Document) -> None:
    add_section(document, "5. Общая схема потока данных (от Excel до дашборда)")

    steps = [
        "1) Загрузка файла пользователем: браузер отправляет `multipart/form-data` на `/upload`; файл сохраняется в `uploads/<session>/<job>/`.",
        "2) Импорт в систему: `/import_data` запускает `ImportDataStep`, который читает XLS/XLSX/CSV в `pandas.DataFrame`.",
        "3) Первичное сохранение: DataFrame экспортируется в промежуточный CSV (`utf-8-sig`) и в таблицу PostgreSQL через `to_sql`.",
        "4) Профилирование признаков: `FiresFeatureProfilingStep` строит отчёты по качеству колонок и сохраняет их в CSV/XLSX.",
        "5) Защита обязательных колонок: `KeepImportantColumnsStep` корректирует кандидатов на удаление и формирует обновлённые отчёты.",
        "6) Создание очищенной витрины: `CreateCleanTableStep` делает SQL `CREATE TABLE clean_* AS SELECT ...` и экспорт `clean_*.xlsx`.",
        "7) Инвалидация и прогрев кэшей: сбрасываются metadata/service caches, далее startup/runtime warmup переинициализирует быстрые слои.",
        "8) SQL-агрегации сервисов: dashboard/forecasting/ml/clustering/access_points читают данные из PostgreSQL (часть результатов кэшируется в памяти процесса).",
        "9) Формирование аналитических payload: сервисы собирают `dict/list` структуры, графики Plotly (JSON) и статистические summary.",
        "10) API-выдача клиенту: роутеры отдают UTF-8 JSON (`application/json; charset=utf-8`) с текущим `session_id`.",
        "11) Отрисовка интерфейса: Jinja2 + frontend JS загружают payload, рендерят карточки/таблицы/графики/heatmap.",
        "12) Карта пожаров: `CreateFireMapStep` + `MapCreator` генерируют `fires_map.html` (OpenLayers), а также spatial-analysis JSON/Markdown.",
    ]

    add_bullets(document, steps)

    add_subsection(document, "Форматы данных по слоям")
    add_bullets(
        document,
        [
            "Вход: `.xlsx/.xls/.csv`.",
            "Внутри ETL: `pandas.DataFrame`.",
            "Хранилище: таблицы PostgreSQL (`raw`, `clean_*`, агрегаты/представления сервисов).",
            "Промежуточные артефакты: `.csv`, `.xlsx`, `.json`, `.md`.",
            "Транспорт API: UTF-8 JSON.",
            "Визуализация: HTML + JS (Plotly/OpenLayers), map HTML-файл.",
        ],
    )


def build_document() -> Path:
    document = Document()
    add_title(document, "Fire Data Pipeline — полный архитектурный анализ проекта")
    document.add_paragraph(
        f"Дата формирования: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    document.add_paragraph(
        "Основание: анализ модулей `app/`, `core/`, `config/` рабочего проекта."
    )

    section_1_idef0_context(document)
    section_2_idef0_top_level(document)
    section_3_architecture(document)
    section_4_stack(document)
    section_5_data_flow(document)

    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


def main() -> None:
    path = build_document()
    print(path)


if __name__ == "__main__":
    main()

