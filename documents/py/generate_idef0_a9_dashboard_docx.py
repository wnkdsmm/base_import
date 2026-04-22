from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document


OUTPUT_PATH = Path("documents/IDEF0_A9_Аналитический_дашборд.docx")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_par(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_code(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Consolas"


def count_text_question_marks(path: Path) -> int:
    if not path.exists():
        return -1
    with ZipFile(path, "r") as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    text = "".join(re.findall(r">([^<]*)<", xml))
    return text.count("?")


def add_icom(
    doc: Document,
    title: str,
    input_text: str,
    control_text: str,
    output_text: str,
    mechanism_text: str,
) -> None:
    add_heading(doc, title, level=2)
    add_bullet(doc, f"Вход: {input_text}")
    add_bullet(doc, f"Управление: {control_text}")
    add_bullet(doc, f"Выход: {output_text}")
    add_bullet(doc, f"Механизм: {mechanism_text}")


def build_doc() -> Document:
    doc = Document()

    add_heading(doc, 'IDEF0 A9: "Аналитический дашборд"', level=0)
    add_par(
        doc,
        (
            "Анализ выполнен по модулю app/dashboard с опорой на функции в файлах: "
            "cache.py, metadata.py, dashboard_service_data.py, dashboard_service_build.py, "
            "impact_fire_metrics.py, impact_forecast_metrics.py, distribution.py, "
            "distribution_logic.py, summary.py, data_access.py. "
            "Для контекста вызовов использован app/routes/api_dashboard.py и app/runtime_invalidation.py."
        ),
    )

    add_heading(doc, "1. IDEF0 — декомпозиция A9.1-A9.5", level=1)

    add_icom(
        doc,
        "A9.1 Сбор метаданных таблиц",
        (
            "Сигнатура таблиц БД, имена таблиц пользователя, списки колонок, годы, "
            "параметры запроса table_name/year/group_column."
        ),
        (
            "Правила выбора таблиц (_prefer_original_source_tables), "
            "правила выбора доступных группировок (_collect_group_column_options), "
            "правила валидации фильтров (_resolve_dashboard_filters), "
            "TTL метаданных 300 секунд."
        ),
        (
            "DashboardMetadata: tables, table_signature, table_options, default_group_column, errors; "
            "DashboardRequestState: selected_tables, selected_year, selected_group_column, available_years."
        ),
        (
            "app.dashboard.cache._collect_dashboard_metadata_cached, "
            "app.dashboard.metadata._collect_dashboard_metadata, "
            "app.db_metadata.get_table_signature_cached/get_table_columns_cached."
        ),
    )

    add_icom(
        doc,
        "A9.2 SQL-агрегация с переиспользованием группировок",
        (
            "selected_tables и selected_year из A9.1, selected_group_column, "
            "конфигурация метрик и колонок ущерба."
        ),
        (
            "SQL-шаблон GROUPING SETS в _build_dashboard_grouped_counts_query, "
            "фильтр года _build_year_filter_clause, "
            "флаги include_area_buckets/include_impact_timeline/positive_count_columns."
        ),
        (
            "Единый grouped_counts_bundle: cause_counts, distribution_counts, district_counts, "
            "month_counts, area_bucket_counts, impact_timeline_rows, positive_column_counts; "
            "summary_bundle с summary_rows и yearly_grouped."
        ),
        (
            "app.dashboard.impact_fire_metrics._collect_dashboard_grouped_counts, "
            "app.dashboard.summary._collect_dashboard_summary_bundle, SQLAlchemy text + engine.connect()."
        ),
    )

    add_icom(
        doc,
        "A9.3 Формирование виджетов ущерба и распределения",
        (
            "grouped_counts_bundle из A9.2, summary_bundle, выбранная группировка selected_group_column."
        ),
        (
            "Логика ветвления damage vs standard (_is_damage_group_selection), "
            "правило reuse distribution counts (_can_reuse_distribution_counts), "
            "ограничение top-N для виджетов и графиков."
        ),
        (
            "Готовые dashboard widgets и chart blocks: distribution, damages, districts, seasons, causes, "
            "rankings, highlights, table breakdown."
        ),
        (
            "app.dashboard.distribution_logic._build_damage_dashboard_charts/_build_standard_dashboard_charts/"
            "_build_dashboard_widgets, app.dashboard.distribution._build_distribution_chart."
        ),
    )

    add_icom(
        doc,
        "A9.4 Формирование временной шкалы воздействия",
        (
            "impact_timeline_rows из A9.2 и метрики воздействия deaths, injuries, evacuated, "
            "evacuated_children, rescued_children."
        ),
        (
            "Агрегация по date_value и сортировка по времени; "
            "форматирование дат _format_chart_date; "
            "в текущем backend нет отдельного шага сглаживания скользящим средним."
        ),
        (
            "График impact timeline с суммарным вкладом по датам и сериями по категориям воздействия."
        ),
        (
            "app.dashboard.impact_forecast_metrics._build_combined_impact_timeline_chart и "
            "_build_combined_impact_timeline_plotly."
        ),
    )

    add_icom(
        doc,
        "A9.5 Сборка итогового ответа и кэширование",
        (
            "Агрегаты A9.2-A9.4, summary/scope/trend/management блоки, фильтры пользователя."
        ),
        (
            "Ключи кэша: metadata key и resolved data key; TTL metadata 300 сек, data 120 сек; "
            "fallback-политика allow_fallback; инвалидация при смене table signature."
        ),
        (
            "Финальный DashboardPayload для API и шаблона страницы: summary, charts, widgets, "
            "filters, notes, management.export_text."
        ),
        (
            "app.dashboard.dashboard_service_build._build_dashboard_payload, "
            "app.dashboard.dashboard_service_data.get_dashboard_data, "
            "app.dashboard.cache._get_dashboard_cache/_set_dashboard_cache."
        ),
    )

    add_heading(doc, "1.1 Как выход A9.1 управляет всеми последующими функциями", level=2)
    add_bullet(
        doc,
        (
            "A9.1 -> A9.2: selected_tables, selected_year, selected_group_column из "
            "_resolve_dashboard_filters задают WHERE, GROUPING SETS и состав SQL-запросов."
        ),
    )
    add_bullet(
        doc,
        (
            "A9.1 -> A9.3: список доступных group_column определяет ветку расчета "
            "(ущерб или стандартное распределение) и возможность reuse grouped counts."
        ),
    )
    add_bullet(
        doc,
        (
            "A9.1 -> A9.4: table metadata определяет наличие date-колонки и, следовательно, "
            "доступность временной шкалы воздействия для конкретных таблиц."
        ),
    )
    add_bullet(
        doc,
        (
            "A9.1 -> A9.5: table_signature входит в ключи кэша. При изменении сигнатуры "
            "_collect_dashboard_metadata_cached очищает metadata cache и data cache."
        ),
    )

    add_heading(doc, "2. Алгоритмы", level=1)

    add_heading(doc, "2.1 Алгоритм оптимизации SQL (grouped counts reuse)", level=2)
    add_bullet(
        doc,
        (
            "Проблема: при независимом построении каждого виджета возникает N отдельных SQL-запросов "
            "для causes, districts, seasons, distribution, area buckets, impact timeline и др."
        ),
    )
    add_bullet(
        doc,
        (
            "Решение в модуле: _collect_dashboard_grouped_counts формирует один объединенный SQL "
            "через GROUPING SETS + UNION ALL по выбранным таблицам и возвращает пакет агрегатов."
        ),
    )
    add_bullet(
        doc,
        (
            "Дальше каждый виджет берет свой срез в памяти из grouped_counts_bundle, "
            "без дополнительных запросов к БД."
        ),
    )
    add_bullet(
        doc,
        (
            "Оценка экономии: для стандартного дашборда минимум 6 тематических агрегатов "
            "получаются одним SQL round-trip вместо 6 отдельных запросов. "
            "Экономия по группе = N - 1, где N — число агрегатов, которые переиспользуются."
        ),
    )
    add_par(doc, "Псевдокод:")
    add_code(
        doc,
        (
            "INPUT: selected_tables, selected_year, selected_group_column\n"
            "OUTPUT: grouped_counts_bundle\n"
            "\n"
            "subqueries = []\n"
            "FOR each table IN selected_tables:\n"
            "    query = build_grouped_counts_query(\n"
            "        table, selected_year, selected_group_column,\n"
            "        dimensions=[cause, distribution, district, month, area_bucket, impact_timeline]\n"
            "    )\n"
            "    IF query is not None:\n"
            "        subqueries.append(query)\n"
            "\n"
            "rows = execute_sql(UNION_ALL(subqueries))\n"
            "\n"
            "bundle = init_empty_bundle()\n"
            "FOR each row IN rows:\n"
            "    metric_kind = row.metric_kind\n"
            "    IF metric_kind == 'cause':\n"
            "        bundle.cause_counts[row.label] += row.fire_count\n"
            "    ELIF metric_kind == 'distribution':\n"
            "        bundle.distribution_counts[row.label] += row.fire_count\n"
            "    ELIF metric_kind == 'district':\n"
            "        bundle.district_counts[row.label] += row.fire_count\n"
            "    ELIF metric_kind == 'month':\n"
            "        bundle.month_counts[int(row.label)] += row.fire_count\n"
            "    ELIF metric_kind == 'area_bucket':\n"
            "        bundle.area_bucket_counts[row.label] += row.fire_count\n"
            "    ELIF metric_kind == 'impact_timeline':\n"
            "        bundle.impact_timeline_rows.append(row)\n"
            "\n"
            "return bundle\n"
            "\n"
            "FOR each widget IN dashboard_widgets:\n"
            "    widget_data = take_slice_from(bundle)\n"
            "    render_widget(widget_data)\n"
        ),
    )

    add_heading(doc, "2.2 Алгоритм двухуровневого кэша дашборда", level=2)
    add_bullet(
        doc,
        "Уровень 1 metadata cache: ключ = tuple(sorted(table_names)), TTL = 300 секунд."
    )
    add_bullet(
        doc,
        (
            "Уровень 2 data cache: ключ = "
            "(metadata_table_signature, table_name, year, normalized_group_column, horizon_days), TTL = 120 секунд."
        ),
    )
    add_bullet(
        doc,
        (
            "Почему TTL разные: структура таблиц меняется реже, чем аналитические срезы по фильтрам. "
            "Поэтому метаданные хранятся дольше, а результат виджетов обновляется чаще."
        ),
    )
    add_bullet(
        doc,
        (
            "Инвалидация: при несовпадении table_signature или при runtime invalidation "
            "очищаются оба уровня кэша."
        ),
    )
    add_par(doc, "Псевдокод:")
    add_code(
        doc,
        (
            "FUNCTION collect_dashboard_metadata_cached():\n"
            "    current_key = tuple(sorted(select_user_table_names(get_table_signature_cached())))\n"
            "    cached_meta = metadata_cache.get(current_key)\n"
            "    IF cached_meta exists AND metadata_table_names(cached_meta) == current_key:\n"
            "        return cached_meta\n"
            "\n"
            "    metadata = collect_dashboard_metadata(current_key)\n"
            "    metadata_cache.clear()\n"
            "    data_cache.clear()\n"
            "    metadata_cache.set(current_key, metadata)\n"
            "    return metadata\n"
            "\n"
            "FUNCTION get_dashboard_data(params):\n"
            "    metadata = collect_dashboard_metadata_cached()\n"
            "    raw_key = (metadata_table_names(metadata), params.table, params.year, params.group, params.horizon)\n"
            "    cached = data_cache.get(raw_key)\n"
            "    IF cached exists:\n"
            "        return cached\n"
            "\n"
            "    resolved_state = resolve_dashboard_filters(metadata, params)\n"
            "    resolved_key = (\n"
            "        metadata_table_names(metadata),\n"
            "        resolved_state.selected_table,\n"
            "        resolved_state.selected_year_or_all,\n"
            "        resolved_state.selected_group,\n"
            "        params.horizon\n"
            "    )\n"
            "    cached2 = data_cache.get(resolved_key)\n"
            "    IF cached2 exists:\n"
            "        return cached2\n"
            "\n"
            "    aggregation = build_dashboard_aggregation(resolved_state)\n"
            "    payload = build_dashboard_payload(metadata, aggregation, resolved_state)\n"
            "    data_cache.set(resolved_key, payload)\n"
            "    return payload\n"
        ),
    )

    add_heading(doc, "2.3 Алгоритм построения виджета распределения", level=2)
    add_bullet(
        doc,
        (
            "Группировка делается по выбранной категории (например, причина, тип объекта, категория риска) "
            "через SQL GROUP BY label либо через reuse уже собранных distribution_counts."
        ),
    )
    add_bullet(
        doc,
        "Формируются элементы items = [{label, value, value_display}] с сортировкой по убыванию и ограничением top-12."
    )
    add_bullet(
        doc,
        (
            "Нормализация долей для аналитики: p_i = count_i / Σcount. "
            "В backend хранится абсолютный count, доля вычисляется на уровне визуализации Plotly."
        ),
    )
    add_bullet(
        doc,
        (
            "Формат для Plotly: series labels, values и hovertemplate. "
            "Для risk-category используется pie-представление, иначе bar."
        ),
    )

    add_heading(doc, "2.4 Алгоритм временной шкалы ущерба (impact timeline)", level=2)
    add_bullet(
        doc,
        (
            "В текущей реализации impact timeline агрегируется по date_value "
            "в _build_combined_impact_timeline_chart на основе impact_timeline_rows."
        ),
    )
    add_bullet(
        doc,
        (
            "Агрегации по неделям, месяцам и кварталам как отдельного backend-шага для этого графика нет. "
            "Месячная агрегированная аналитика выводится отдельными виджетами monthly_profile и monthly_heatmap."
        ),
    )
    add_bullet(
        doc,
        (
            "Скользящее среднее для impact timeline в коде сейчас не применяется, "
            "фактически окно сглаживания W = 1. Базовая формула для расширения: "
            "S_t = (1/W) * Σ_{k=0..W-1} y_{t-k}."
        ),
    )

    add_heading(doc, "3. Ключевые функции и классы", level=1)
    add_bullet(doc, "Кэш: CopyingTtlCache в app/dashboard/cache.py (metadata и data уровни).")
    add_bullet(doc, "Метаданные: _collect_dashboard_metadata, _resolve_dashboard_filters.")
    add_bullet(doc, "Агрегация: _collect_dashboard_grouped_counts, _collect_dashboard_summary_bundle.")
    add_bullet(doc, "Сборка: _build_dashboard_aggregation, _build_dashboard_payload, get_dashboard_data.")
    add_bullet(doc, "Инвалидация: _invalidate_dashboard_caches и invalidate_table_related_caches.")

    add_heading(doc, "4. Вывод", level=1)
    add_par(
        doc,
        (
            "A9 реализован как конвейер: метаданные -> групповые SQL-агрегаты -> "
            "визуальные блоки -> финальный payload в кэш. "
            "Ключевая оптимизация — единый grouped counts запрос с последующим reuse в памяти, "
            "что уменьшает число SQL-вызовов и ускоряет отдачу дашборда."
        ),
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    qm = count_text_question_marks(OUTPUT_PATH)
    print(f"DOCX_CREATED={OUTPUT_PATH}")
    print(f"TEXT_QUESTION_MARKS={qm}")


if __name__ == "__main__":
    main()

