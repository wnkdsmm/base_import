from __future__ import annotations

from datetime import datetime
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"F:\filesFires\base_import")
OUTPUT_PATH = ROOT / "documents" / "IDEF0_A1_Обработка_и_импорт_данных.docx"


def set_font(run, name: str = "Times New Roman", size: int = 12, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def configure_document(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def add_paragraph(document: Document, text: str, *, bold: bool = False, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_font(run, bold=bold)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading("", level=level)
    run = heading.add_run(text)
    set_font(run, size=14 if level == 1 else 12, bold=True)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        add_paragraph(document, item, style="List Bullet")


def add_code_block(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(code)
    set_font(run, name="Consolas", size=10)


def count_question_marks_in_docx_text_nodes(docx_path: Path) -> int:
    with ZipFile(docx_path, "r") as archive:
        xml_bytes = archive.read("word/document.xml")
    root = ET.fromstring(xml_bytes)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return sum((node.text or "").count("?") for node in root.findall(".//w:t", ns))


def section_idef0_a1(document: Document) -> None:
    add_heading(document, 'IDEF0 — Диаграмма A1 "Обработка и импорт данных"', level=1)

    add_heading(document, "A1.1 Приём и валидация файла", level=2)
    add_paragraph(document, "Вход:", bold=True)
    add_bullets(
        document,
        [
            "UploadFile из HTTP-запроса (`save_uploaded_file`).",
            "session_id и job_id (создаются/резолвятся через `job_store.create_or_reset_job`).",
        ],
    )
    add_paragraph(document, "Управление:", bold=True)
    add_bullets(
        document,
        [
            "Санитизация имени: `Path(file.filename).name` в `_build_upload_file_path`.",
            "Правило размещения: `UPLOAD_FOLDER / session_id / job_id`.",
            "Базовая проверка доступности файла на этапе импорта: `os.path.exists(input_file)` в `ImportDataStep.run`.",
        ],
    )
    add_paragraph(document, "Выход:", bold=True)
    add_bullets(
        document,
        [
            "Файл сохранён на диск в папке загрузок.",
            "В `JobStore` записаны путь файла и статус загрузки (`set_uploaded_file`, лог события).",
            "Служебный ответ API: `status`, `filename`, `path`, `job_id`.",
        ],
    )
    add_paragraph(document, "Механизм:", bold=True)
    add_bullets(
        document,
        [
            "`app/services/pipeline_service.py`: `_build_upload_file_path`, `save_uploaded_file`.",
            "FastAPI `UploadFile`, `pathlib.Path`, `shutil.copyfileobj`, `app.state.job_store`.",
        ],
    )

    add_heading(document, "A1.2 Импорт данных в СУБД", level=2)
    add_paragraph(document, "Вход:", bold=True)
    add_bullets(
        document,
        [
            "Путь к загруженному файлу (`job.current_file_path`).",
            "Параметры `Settings(input_file, output_folder, project_name)`.",
        ],
    )
    add_paragraph(document, "Управление:", bold=True)
    add_bullets(
        document,
        [
            "Поддерживаемые расширения: `.xls`, `.xlsx`, `.csv` (`ImportDataStep.run`).",
            "Кодировка CSV: `utf-8-sig` при чтении/выгрузке.",
            "Стратегия записи в БД: `to_sql(..., if_exists='replace', index=False)`.",
        ],
    )
    add_paragraph(document, "Выход:", bold=True)
    add_bullets(
        document,
        [
            "Таблица в PostgreSQL с именем `project_name`.",
            "Промежуточный CSV-файл `<project_name>.csv` в `output_folder`.",
            "Кэш in-memory в `settings`: `_pipeline_source_df`, `_pipeline_import_csv`.",
        ],
    )
    add_paragraph(document, "Механизм:", bold=True)
    add_bullets(
        document,
        [
            "`core/processing/steps/import_data.py`: класс `ImportDataStep`, метод `run`.",
            "pandas: `read_excel`, `read_csv`, `to_csv`, `to_sql`.",
            "SQLAlchemy Engine из `config.db`: транзакционный контекст `engine.begin()`.",
        ],
    )

    add_heading(document, "A1.3 Очистка и нормализация", level=2)
    add_paragraph(document, "Вход:", bold=True)
    add_bullets(
        document,
        [
            "Профиль колонок (`profile_df`) с флагом `candidate_to_drop`.",
            "Исходный DataFrame (`source_df`) или таблица-источник в БД.",
        ],
    )
    add_paragraph(document, "Управление:", bold=True)
    add_bullets(
        document,
        [
            "Если есть `<table>_updated_fires_profiling_report.csv`, используется он (приоритет после `KeepImportantColumnsStep`).",
            "Правило отбора: сохраняются только колонки с `candidate_to_drop=False`.",
            "Контроль целостности: если список `keep_columns` пуст, выбрасывается `ValueError`.",
        ],
    )
    add_paragraph(document, "Выход:", bold=True)
    add_bullets(
        document,
        [
            "Новая таблица `clean_<table>` в PostgreSQL.",
            "Экспорт `clean_<table>.xlsx`.",
            "Списки `kept_columns`, `removed_columns`, размерность результата.",
        ],
    )
    add_paragraph(document, "Механизм:", bold=True)
    add_bullets(
        document,
        [
            "`core/processing/steps/create_clean_table.py`: `CreateCleanTableStep.run`.",
            "`coerce_bool_series` из `column_transforms.py` для приведения булевых флагов.",
            "SQL DDL через `sqlalchemy.text` + `conn.execute`: `DROP TABLE IF EXISTS`, `CREATE TABLE ... AS SELECT ...`.",
            "pandas: `read_sql`, `to_excel(engine='openpyxl')`.",
        ],
    )
    add_paragraph(
        document,
        "Техническое уточнение: в `CreateCleanTableStep` нет value-level нормализации дат, иммутации пропусков или фильтрации строк; шаг выполняет колонко-ориентированную очистку (schema pruning).",
    )

    add_heading(document, "A1.4 Профилирование признаков", level=2)
    add_paragraph(document, "Вход:", bold=True)
    add_bullets(
        document,
        [
            "Таблица БД (`selected_table`/`project_name`) или переданный DataFrame.",
            "Пороговые параметры качества (`null_threshold`, `dominant_value_threshold`, `low_variance_threshold`).",
        ],
    )
    add_paragraph(document, "Управление:", bold=True)
    add_bullets(
        document,
        [
            "Дефолтные константы из `config/constants.py`: `NULL_THRESHOLD=0.9`, `DOMINANT_VALUE_THRESHOLD=0.85`, `LOW_VARIANCE_THRESHOLD=0.0001`.",
            "Список квазипустых строк `MISSING_LIKE_VALUES` для строковых полей.",
            "Состав правил исключения: `drop_null`, `drop_constant`, `low_variance`, `almost_constant`.",
        ],
    )
    add_paragraph(document, "Выход:", bold=True)
    add_bullets(
        document,
        [
            "Отчёты `<table>_fires_profiling_report.csv` и `.xlsx`.",
            "Матрица причин исключения и `reason_summary`.",
            "Список кандидатов на исключение и список сохраняемых колонок.",
        ],
    )
    add_paragraph(document, "Механизм:", bold=True)
    add_bullets(
        document,
        [
            "`core/processing/steps/fires_feature_profiling.py`: `FiresFeatureProfilingStep.run`.",
            "pandas + numpy: `select_dtypes`, `isna`, `value_counts`, `nunique`, `var`, `to_csv`, `to_excel`.",
        ],
    )

    add_heading(document, "A1.5 Инвалидация кэшей", level=2)
    add_paragraph(document, "Вход:", bold=True)
    add_bullets(
        document,
        [
            "Событие успешного завершения импорта (`import_uploaded_data`) или очистки/профилирования (`run_profiling_for_table`).",
            "Контекст `session_id` и `job_id` для логирования предупреждений.",
        ],
    )
    add_paragraph(document, "Управление:", bold=True)
    add_bullets(
        document,
        [
            "Реестр инвалидаторов `_INVALIDATORS` в `app/runtime_invalidation.py`.",
            "Флаг `include_metadata` определяет, сбрасывать ли кэш метаданных БД.",
            "Ошибки отдельных инвалидаторов не прерывают общий процесс (warning-стратегия).",
        ],
    )
    add_paragraph(document, "Выход:", bold=True)
    add_bullets(
        document,
        [
            "Сброшены runtime-кэши модулей: `db_metadata`, `dashboard`, `ml_model`, `forecasting`, `clustering`, `access_points`, `fire_map`.",
            "Предупреждения по сбою отдельных инвалидаторов добавлены в job-логи.",
        ],
    )
    add_paragraph(document, "Механизм:", bold=True)
    add_bullets(
        document,
        [
            "`pipeline_service._invalidate_runtime_caches` -> `app.runtime_invalidation.invalidate_runtime_caches`.",
            "`importlib.import_module` + вызов целевых функций-инвалидаторов.",
        ],
    )


def section_algorithms(document: Document) -> None:
    add_heading(document, "Алгоритмы", level=1)

    add_heading(document, "1) Алгоритм определения важных колонок (NatashaColumnMatcher)", level=2)
    add_paragraph(document, "Тип алгоритма:", bold=True)
    add_bullets(
        document,
        [
            "Гибридный rule-based + NLP-подход.",
            "Точное совпадение нормализованного имени, токен-сопоставление, проверка token_set-правил, лемматизация Natasha.",
        ],
    )
    add_paragraph(document, "Словари и правила приоритета:", bold=True)
    add_bullets(
        document,
        [
            "`MANDATORY_FEATURE_REGISTRY` (обязательные признаки, synonym/token_set/exclude_tokens).",
            "`LEGACY_EXPLICIT_IMPORTANT_COLUMNS` (legacy exact map).",
            "`KEYWORD_IMPORTANCE_RULES` (keyword include_all/include_any).",
            "Приоритет в `_match_column_payload_metadata`: `mandatory` -> `legacy explicit` -> `keyword`.",
        ],
    )
    add_paragraph(document, "Fallback-цепочка:", bold=True)
    add_bullets(
        document,
        [
            "Если нет совпадения в mandatory/legacy/keyword, колонка не получает защиту (`match=None`).",
            "В `KeepImportantColumnsStep.apply_match_results` такая колонка остаётся с исходным `candidate_to_drop` из profiling.",
            "Для поисковых запросов используется доп. fallback-лексика `FALLBACK_IMPORTANT_PATTERNS` при сборке query terms.",
        ],
    )
    add_paragraph(document, "Сложность:", bold=True)
    add_bullets(
        document,
        [
            "Для C колонок: O(C * R * T), где R — число правил/признаков, T — средняя стоимость token/lemma-проверок.",
            "Дополнительно есть амортизационное ускорение за счёт кэшей `_terms_cache` (до 4096) и `_group_catalog_cache` (до 32).",
        ],
    )
    add_paragraph(document, "Псевдокод:", bold=True)
    add_code_block(
        document,
        """INPUT: profile_df.columns, mandatory_registry, legacy_map, keyword_rules
FOR each column_name in profile_df.columns:
    payload = _column_terms(column_name)   # normalize + tokens + lemmas

    match = _match_mandatory_feature(payload)
    IF match is None:
        match = _match_legacy_explicit(payload)
    IF match is None:
        match = _match_keyword_rule(payload)

    IF match exists:
        write protection metadata into profile_df row
        IF profiling_candidate_to_drop == True:
            candidate_to_drop = False
            protected_from_drop = True
    ELSE:
        keep original candidate_to_drop from profiling

RETURN updated profile_df, protected_columns_report""",
    )

    add_heading(document, "2) Алгоритм очистки данных (CreateCleanTableStep)", level=2)
    add_paragraph(document, "Правила нормализации дат:", bold=True)
    add_bullets(
        document,
        [
            "На уровне `CreateCleanTableStep` нормализация дат не выполняется.",
            "Шаг работает с уже рассчитанным профилем колонок и делает структурную очистку по `candidate_to_drop`.",
        ],
    )
    add_paragraph(document, "Обработка пропущенных значений:", bold=True)
    add_bullets(
        document,
        [
            "Прямой иммутации пропусков в `CreateCleanTableStep` нет.",
            "Решение о качестве колонки (включая высокий `null_ratio`) принимается ранее в `FiresFeatureProfilingStep`.",
        ],
    )
    add_paragraph(document, "Фильтрация строк:", bold=True)
    add_bullets(
        document,
        [
            "Построчная фильтрация не выполняется; количество строк сохраняется.",
            "Фильтрация выполняется по измерению колонок: `keep_columns` vs `removed_columns`.",
        ],
    )
    add_paragraph(document, "SQL-операции для PostgreSQL:", bold=True)
    add_bullets(
        document,
        [
            "`DROP TABLE IF EXISTS \"clean_<table>\"`.",
            "`CREATE TABLE \"clean_<table>\" AS SELECT <kept_columns> FROM \"<source_table>\"`.",
            "При выгрузке (fallback): `SELECT <kept_columns> FROM \"clean_<table>\"` через `pd.read_sql`.",
        ],
    )
    add_paragraph(document, "Псевдокод:", bold=True)
    add_code_block(
        document,
        """INPUT: settings, profile_df(optional), source_df(optional)
profile_df = resolve_profile_df(profile_df, settings._pipeline_profile_df, profile_csv)
ASSERT 'candidate_to_drop' in profile_df

candidate_mask = coerce_bool_series(profile_df['candidate_to_drop'])
keep_columns = profile_df[~candidate_mask]['column']
removed_columns = profile_df[candidate_mask]['column']
ASSERT keep_columns is not empty

new_table = 'clean_' + source_table
EXEC SQL: DROP TABLE IF EXISTS new_table
EXEC SQL: CREATE TABLE new_table AS SELECT keep_columns FROM source_table

IF source_df exists AND all keep_columns in source_df:
    clean_df = source_df[keep_columns]
ELSE:
    clean_df = read_sql(SELECT keep_columns FROM new_table)

write clean_df to Excel (openpyxl)
RETURN clean_table metadata + kept/removed columns""",
    )

    add_heading(document, "3) Алгоритм профилирования признаков (FiresFeatureProfilingStep)", level=2)
    add_paragraph(document, "Какие статистики считаются по колонке:", bold=True)
    add_bullets(
        document,
        [
            "`dtype`",
            "`null_ratio`",
            "`unique_count`",
            "`unique_ratio`",
            "`dominant_ratio`",
            "`variance` (для numeric, иначе 1.0)",
        ],
    )
    add_paragraph(document, "Формат выходного отчёта:", bold=True)
    add_bullets(
        document,
        [
            "`<table>_fires_profiling_report.csv`",
            "`<table>_fires_profiling_report.xlsx`",
        ],
    )
    add_paragraph(document, "Как определяется качество данных:", bold=True)
    add_bullets(
        document,
        [
            "Колонка помечается `candidate_to_drop=True`, если срабатывает хотя бы одно правило: `drop_null`, `drop_constant`, `low_variance`, `almost_constant`.",
            "Формула: `candidate_to_drop = drop_null OR drop_constant OR low_variance OR almost_constant`.",
            "Отдельно сохраняется декомпозиция причин (`drop_reasons`, `reason_summary`) для объяснимости в UI/API.",
        ],
    )
    add_paragraph(document, "Псевдокод:", bold=True)
    add_code_block(
        document,
        """INPUT: table_name or source_df, thresholds
df = resolve_source_df(source_df, settings._pipeline_source_df, SELECT * FROM table)
n_rows = len(df)
string_cols = df.select_dtypes(object|string)
df_norm = lower(strip(df[string_cols]))

FOR each column in df.columns:
    col_data = df[column]
    IF column in string_cols:
        null_ratio = max(isna(col_data).mean, isin(MISSING_LIKE_VALUES, df_norm[column]).mean)
        dominant_ratio = max(value_counts(df_norm[column], normalize=True))
        unique_count = nunique(df_norm[column], dropna=True)
    ELSE:
        null_ratio = isna(col_data).mean
        dominant_ratio = max(value_counts(col_data, normalize=True))
        unique_count = nunique(col_data, dropna=True)

    unique_ratio = unique_count / n_rows
    collect row stats

variance = var(numeric columns), map by column (else 1.0)

drop_null = null_ratio > null_threshold
drop_constant = unique_count <= 1
low_variance = variance < low_variance_threshold
almost_constant = dominant_ratio > dominant_value_threshold
candidate_to_drop = drop_null OR drop_constant OR low_variance OR almost_constant

build drop_reasons and reason_summary
sort by candidate_to_drop, null_ratio, dominant_ratio
write CSV + XLSX
RETURN profile_df + candidates + kept_columns + files + thresholds""",
    )


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title_run = title.add_run(
        "Анализ core/processing и app/services/pipeline_service.py\n"
        "IDEF0 A1 «Обработка и импорт данных» + алгоритмы"
    )
    set_font(title_run, size=16, bold=True)

    add_paragraph(doc, "Дата формирования: " + datetime.now().strftime("%d.%m.%Y %H:%M"))
    add_paragraph(doc, "Анализируемые модули:")
    add_bullets(
        doc,
        [
            "app/services/pipeline_service.py",
            "core/processing/pipeline.py",
            "core/processing/steps/import_data.py",
            "core/processing/steps/fires_feature_profiling.py",
            "core/processing/steps/keep_important_columns.py",
            "core/processing/steps/create_clean_table.py",
            "core/processing/steps/column_filter_match.py",
            "core/processing/steps/column_filter_payload.py",
            "core/processing/steps/column_filter_text.py",
            "core/processing/steps/column_transforms.py",
            "app/domain/column_matching.py",
            "app/runtime_invalidation.py",
            "config/constants.py",
        ],
    )

    section_idef0_a1(doc)
    section_algorithms(doc)
    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    question_marks = count_question_marks_in_docx_text_nodes(OUTPUT_PATH)
    print(f"DOCX_CREATED={OUTPUT_PATH}")
    print(f"TEXT_QUESTION_MARKS={question_marks}")


if __name__ == "__main__":
    main()
